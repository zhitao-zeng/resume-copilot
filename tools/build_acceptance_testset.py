"""Build acceptance fixtures for the resume-copilot product flow."""

from __future__ import annotations

import csv
import json
import shutil
import textwrap
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "acceptance_testset"
FILES = OUT / "files"


def _font_path() -> str | None:
    candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/wqy-zenhei/wqy-zenhei.ttc",
        "/usr/share/fonts/wqy-microhei/wqy-microhei.ttc",
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return None


FONT_PATH = _font_path()
PDF_FONT = "Helvetica"
if FONT_PATH:
    try:
        pdfmetrics.registerFont(TTFont("FixtureCJK", FONT_PATH))
        PDF_FONT = "FixtureCJK"
    except Exception:
        PDF_FONT = "Helvetica"


def _reset_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def _docx(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(30, 64, 120)

    for heading, lines in sections:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(3)
        hr = h.add_run(heading)
        hr.bold = True
        hr.font.size = Pt(12)
        hr.font.color.rgb = RGBColor(30, 64, 120)
        for line in lines:
            para = doc.add_paragraph(style=None)
            para.paragraph_format.left_indent = Inches(0.12)
            para.paragraph_format.space_after = Pt(2)
            if line.startswith("- "):
                para.style = doc.styles["List Bullet"]
                para.add_run(line[2:])
            else:
                para.add_run(line)
    doc.save(path)


def _pdf(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    styles = getSampleStyleSheet()
    styles["Title"].fontName = PDF_FONT
    styles["Heading2"].fontName = PDF_FONT
    styles["BodyText"].fontName = PDF_FONT
    styles["BodyText"].fontSize = 9
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=16 * mm, bottomMargin=16 * mm)
    story: list[Any] = [Paragraph(title, styles["Title"]), Spacer(1, 5 * mm)]
    for heading, lines in sections:
        story.append(Paragraph(heading, styles["Heading2"]))
        for line in lines:
            text = line[2:] if line.startswith("- ") else line
            story.append(Paragraph(text, styles["BodyText"]))
        story.append(Spacer(1, 3 * mm))
    doc.build(story)


def _image(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    font_title = ImageFont.truetype(FONT_PATH, 30) if FONT_PATH else ImageFont.load_default()
    font_head = ImageFont.truetype(FONT_PATH, 22) if FONT_PATH else ImageFont.load_default()
    font_body = ImageFont.truetype(FONT_PATH, 18) if FONT_PATH else ImageFont.load_default()
    width, height = 1300, 1800
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    y = 44
    draw.text((48, y), title, font=font_title, fill=(27, 57, 106))
    y += 58
    for heading, lines in sections:
        draw.rectangle((48, y, width - 48, y + 34), fill=(232, 240, 250))
        draw.text((60, y + 4), heading, font=font_head, fill=(27, 57, 106))
        y += 48
        for line in lines:
            wrapped = textwrap.wrap(line, width=52)
            for chunk in wrapped[:3]:
                draw.text((68, y), chunk, font=font_body, fill=(25, 25, 25))
                y += 28
            y += 4
        y += 10
    img.save(path)


def _sections(name: str, base: str, education: str, experiences: list[str], skills: str) -> list[tuple[str, list[str]]]:
    return [
        ("个人信息", [base]),
        ("个人总结", ["候选人具备清晰的问题拆解和执行闭环能力，过往经历以真实岗位职责和结果为准。"]),
        ("教育背景", [education]),
        ("经历", experiences),
        ("技能", [skills]),
    ]


def _write_fixture_files() -> None:
    for sub in ("cv", "jd", "templates"):
        (FILES / sub).mkdir(parents=True, exist_ok=True)

    cv_defs = {
        "product_pm": _sections(
            "张晨",
            "张晨 | 13810001001 | zhangchen@example.com | 4年经验 | 男",
            "复旦大学 本科 计算机科学 09-2016 - 06-2020",
            [
                "07-2022 - 05-2026 第四范式 产品经理，负责企业数据平台需求调研、版本规划与跨团队推进，推动报表配置效率提升30%。",
                "07-2020 - 06-2022 星河科技 产品助理，负责客户反馈整理、竞品分析和需求文档维护。",
            ],
            "SQL、Axure、Figma、数据分析、项目管理",
        ),
        "finance_bank": _sections(
            "李慧",
            "李慧 | 13910002002 | lihui@example.com | 5年经验 | 女",
            "中央财经大学 学士 金融学 09-2015 - 06-2019",
            [
                "03-2020 - 06-2025 中国银行 贷款管理岗，负责对公大客户贷款资料审核、授信流程跟进和风险材料归档。",
                "07-2019 - 02-2020 中信证券 实习生，参与行业资料整理和客户台账维护。",
            ],
            "Excel、信贷审核、客户管理、风险识别",
        ),
        "teacher": _sections(
            "王宁",
            "王宁 | 13710003003 | wangning@example.com | 3年经验 | 女",
            "华东师范大学 硕士 教育学 09-2017 - 06-2020",
            [
                "09-2020 - 06-2025 上海市第一实验学校 语文教师，负责初中语文授课、班级管理和校本教研。",
                "03-2019 - 06-2019 上海市育才中学 教育实习，参与课堂观察和作业批改。",
            ],
            "课程设计、班级管理、教研、家校沟通",
        ),
        "doctor": _sections(
            "陈雪",
            "陈雪 | 13610004004 | chenxue@example.com | 6年经验 | 女",
            "北京大学医学部 硕士 临床医学 09-2012 - 06-2019",
            [
                "07-2019 - 05-2026 北京协和医院 内科医师，负责门诊接诊、住院患者病程记录和常见慢病诊疗。",
                "09-2017 - 06-2019 北京大学第三医院 规培医师，参与病例讨论和科研资料整理。",
            ],
            "临床诊疗、病例记录、科研资料整理、执业医师资格",
        ),
        "operations": _sections(
            "赵可",
            "赵可 | 13510005005 | zhaoke@example.com | 4年经验 | 男",
            "武汉大学 本科 市场营销 09-2016 - 06-2020",
            [
                "08-2021 - 05-2026 云舟科技 用户运营，负责活动策划、用户分层和数据复盘，推动老客复购率提升12%。",
                "07-2020 - 07-2021 南方电商 内容运营，负责内容选题、社群维护和投放素材跟进。",
            ],
            "Excel、SQL、用户增长、活动运营、数据复盘",
        ),
        "sales": _sections(
            "刘洋",
            "刘洋 | 13410006006 | liuyang@example.com | 7年经验 | 男",
            "南京大学 本科 工商管理 09-2011 - 06-2015",
            [
                "01-2021 - 05-2026 明源云 售前顾问，负责大客户需求澄清、方案宣讲和招投标材料支持。",
                "07-2015 - 12-2020 东华软件 销售经理，负责客户开发、商机推进和合同回款跟进。",
            ],
            "客户开发、方案宣讲、招投标、CRM",
        ),
        "design": _sections(
            "周然",
            "周然 | 13310007007 | zhouran@example.com | 4年经验 | 女",
            "中国美术学院 本科 视觉传达 09-2016 - 06-2020",
            [
                "09-2021 - 05-2026 蓝湖科技 UI设计师，负责SaaS后台界面、组件规范和可用性走查。",
                "07-2020 - 08-2021 风起互动 视觉设计师，负责活动页视觉和品牌物料设计。",
            ],
            "Figma、Sketch、设计系统、交互走查",
        ),
    }
    for key, sections in cv_defs.items():
        _docx(FILES / "cv" / f"cv_{key}.docx", f"{sections[0][1][0].split('|')[0].strip()} 简历", sections)
        _pdf(FILES / "cv" / f"cv_{key}.pdf", f"{sections[0][1][0].split('|')[0].strip()} 简历", sections)
        _image(FILES / "cv" / f"cv_{key}.png", f"{sections[0][1][0].split('|')[0].strip()} 简历", sections)

    jd_defs = {
        "product_pm": [
            ("岗位职责", ["负责B端产品需求分析、PRD撰写、跨团队推进和上线后数据复盘。"]),
            ("任职要求", ["熟悉数据产品或企业软件，具备SQL、指标分析和项目管理能力。"]),
        ],
        "operations": [
            ("岗位职责", ["负责用户增长、活动策划、社群运营和转化漏斗分析。"]),
            ("任职要求", ["具备数据复盘、内容策划和跨团队协同能力。"]),
        ],
        "finance_risk": [
            ("岗位职责", ["负责对公信贷资料审核、风险识别、授信流程跟进和客户材料归档。"]),
            ("任职要求", ["熟悉银行信贷流程，具备审慎细致的风险意识。"]),
        ],
        "teacher": [
            ("岗位职责", ["负责课程设计、课堂授课、作业反馈、班级管理和家校沟通。"]),
            ("任职要求", ["具备教师资格证，表达清晰，重视教学效果沉淀。"]),
        ],
        "doctor": [
            ("岗位职责", ["负责门诊接诊、病历书写、慢病管理和科室科研资料整理。"]),
            ("任职要求", ["具备执业医师资格，诊疗表达合规、严谨。"]),
        ],
        "sales": [
            ("岗位职责", ["负责客户开发、方案宣讲、商机推进、招投标支持和回款跟进。"]),
            ("任职要求", ["具备大客户沟通、需求澄清和解决方案表达能力。"]),
        ],
        "design": [
            ("岗位职责", ["负责产品界面设计、组件规范、用户路径优化和设计交付。"]),
            ("任职要求", ["熟练使用Figma，具备B端产品设计和可用性分析经验。"]),
        ],
        "education_ops": [
            ("岗位职责", ["负责教务流程、课程排期、学员服务和培训项目交付。"]),
            ("任职要求", ["具备教育服务意识、流程优化和沟通协调能力。"]),
        ],
    }
    for key, sections in jd_defs.items():
        text = "\n".join([key, *[f"{h}\n" + "\n".join(lines) for h, lines in sections]])
        (FILES / "jd" / f"jd_{key}.txt").write_text(text, encoding="utf-8")
        _docx(FILES / "jd" / f"jd_{key}.docx", f"{key} JD", sections)
        _pdf(FILES / "jd" / f"jd_{key}.pdf", f"{key} JD", sections)
        _image(FILES / "jd" / f"jd_{key}.png", f"{key} JD", sections)

    _docx(FILES / "templates" / "template_compact.docx", "标准简历模板", [("版式要求", ["单栏、紧凑、深蓝标题、经历 bullet 展示。"])])
    _docx(FILES / "templates" / "template_modern.docx", "现代简历模板", [("版式要求", ["姓名居中、分区清晰、浅灰分隔线、最多三页。"])])
    _pdf(FILES / "templates" / "template_minimal.pdf", "简历模板参考", [("版式偏好", ["黑白、单栏、紧凑间距。"])])
    _image(FILES / "templates" / "template_blue.png", "蓝色简历模板参考", [("版式偏好", ["深蓝标题、单栏、头像区域可忽略。"])])


def _case(
    case_id: str,
    scenario: str,
    industry: str,
    user_stage: str,
    query: str,
    *,
    cv: str | None = None,
    jd_text: str | None = None,
    jd_file: str | None = None,
    template: str | None = None,
    missing: list[str] | None = None,
    conflicts: list[str] | None = None,
    forbidden: list[str] | None = None,
) -> dict[str, Any]:
    return {
        "id": case_id,
        "scenario": scenario,
        "industry": industry,
        "user_stage": user_stage,
        "query": query,
        "cv_path": cv,
        "target_jd": jd_text,
        "target_jd_file_path": jd_file,
        "cv_template_path": template,
        "expected_missing_fields": missing or [],
        "expected_conflicts": conflicts or [],
        "forbidden_fabrication": forbidden or [],
        "expected_output": {
            "format": "docx",
            "max_pages": 3,
            "reply_text_must_include": ["生成方向", "补充", "确认"],
            "score_target_after_llm": 90,
        },
    }


def _cases() -> list[dict[str, Any]]:
    product_jd = (FILES / "jd" / "jd_product_pm.txt").read_text(encoding="utf-8")
    ops_jd = (FILES / "jd" / "jd_operations.txt").read_text(encoding="utf-8")
    finance_jd = (FILES / "jd" / "jd_finance_risk.txt").read_text(encoding="utf-8")
    teacher_jd = (FILES / "jd" / "jd_teacher.txt").read_text(encoding="utf-8")
    doctor_jd = (FILES / "jd" / "jd_doctor.txt").read_text(encoding="utf-8")
    sales_jd = (FILES / "jd" / "jd_sales.txt").read_text(encoding="utf-8")
    design_jd = (FILES / "jd" / "jd_design.txt").read_text(encoding="utf-8")
    education_jd = (FILES / "jd" / "jd_education_ops.txt").read_text(encoding="utf-8")

    cases = [
        _case("S1-PRD-DOCX-JD-TEXT-001", "scenario1", "product_research", "experienced", "请根据目标JD优化我的简历，突出B端产品和数据分析能力。", cv="files/cv/cv_product_pm.docx", jd_text=product_jd, template="files/templates/template_compact.docx", missing=[], forbidden=["新增算法工程师经历"]),
        _case("S1-FIN-PDF-JD-PNG-002", "scenario1", "finance", "experienced", "请针对金融风控岗优化，不要编造贷款金额。", cv="files/cv/cv_finance_bank.pdf", jd_file="files/jd/jd_finance_risk.png", missing=[], forbidden=["10亿贷款", "新增工商银行"]),
        _case("S1-TEA-PNG-JD-DOCX-003", "scenario1", "teacher", "experienced", "请优化为初中语文老师岗位，保留原学校。", cv="files/cv/cv_teacher.png", jd_file="files/jd/jd_teacher.docx", template="files/templates/template_blue.png", forbidden=["新增教师资格证"]),
        _case("S1-DOC-PDF-JD-PDF-004", "scenario1", "doctor", "experienced", "请按内科医师JD优化，表达要合规。", cv="files/cv/cv_doctor.pdf", jd_file="files/jd/jd_doctor.pdf", forbidden=["主刀手术量"]),
        _case("S1-OPS-DOCX-JD-TEXT-005", "scenario1", "operations", "experienced", "优化为用户增长运营，强调转化和复盘。", cv="files/cv/cv_operations.docx", jd_text=ops_jd),
        _case("S1-SAL-DOCX-JD-TEXT-006", "scenario1", "sales_presale", "experienced", "优化为售前顾问，突出方案宣讲。", cv="files/cv/cv_sales.docx", jd_text=sales_jd),
        _case("S1-DES-PNG-JD-TEXT-007", "scenario1", "design", "experienced", "投B端UI设计师，请优化表达。", cv="files/cv/cv_design.png", jd_text=design_jd, template="files/templates/template_modern.docx"),
        _case("S1-EDU-DOCX-JD-TEXT-008", "scenario1", "education", "experienced", "想转教务运营，请优化简历方向。", cv="files/cv/cv_teacher.docx", jd_text=education_jd, forbidden=["新增教务主管"]),
        _case("S1-OTHER-DOCX-JD-TEXT-009", "scenario1", "other", "experienced", "按综合管理岗优化，避免技术化表述。", cv="files/cv/cv_operations.docx", jd_text="综合管理岗：流程梳理、部门协同、会议纪要、项目跟进。"),
        _case("S1-FIN-DOCX-JD-DOCX-010", "scenario1", "finance", "experienced", "突出资料审核、风险意识和合规表达。", cv="files/cv/cv_finance_bank.docx", jd_file="files/jd/jd_finance_risk.docx"),
        _case("S1-PRD-PNG-JD-PDF-011", "scenario1", "product_research", "experienced", "请优化为数据产品经理，保留原始时间。", cv="files/cv/cv_product_pm.png", jd_file="files/jd/jd_product_pm.pdf"),
        _case("S1-DOC-DOCX-JD-TEXT-012", "scenario1", "doctor", "experienced", "请优化为门诊医生，不能写未提供的资质。", cv="files/cv/cv_doctor.docx", jd_text=doctor_jd),
        _case("S2-STU-OPS-QUERY-013", "scenario2", "operations", "student", "姓名林一，电话13210008008，邮箱linyi@example.com，复旦大学本科市场营销，09-2022到06-2026，在校期间做过社群活动项目，负责报名转化统计和活动复盘，想投运营实习。技能Excel、PowerPoint。"),
        _case("S2-STU-PRD-QUERY-014", "scenario2", "product_research", "student", "我叫何然，电话13210008009，邮箱heran@example.com，同济大学软件工程本科，09-2021到06-2025，做过课程选课系统项目，负责需求梳理、原型和测试，想做产品经理。技能Axure、SQL。"),
        _case("S2-EXP-FIN-QUERY-015", "scenario2", "finance", "experienced", "姓名孙倩，电话13210008010，邮箱sunqian@example.com，上海财经大学金融学本科09-2014到06-2018，2018年7月到2025年6月在招商银行柜面和信贷资料审核岗位，想投风控岗。", missing=["skills"]),
        _case("S2-TEA-MISSING-016", "scenario2", "teacher", "student", "我想生成一份老师简历。我是师范专业应届生，做过小学数学试讲和班级活动组织。", missing=["name", "phone", "email", "education.school", "education.period", "skills"]),
        _case("S2-DOC-MISSING-017", "scenario2", "doctor", "experienced", "本人临床医学硕士，2019年后在三甲医院内科工作，负责门诊接诊和病历书写，想投内科医生。", missing=["name", "phone", "email", "education.school", "experience.period", "skills"]),
        _case("S2-DES-QUERY-018", "scenario2", "design", "student", "姓名唐心，电话13210008011，邮箱tangxin@example.com，中国美术学院视觉传达本科09-2022到06-2026，有校园App改版项目，负责界面设计和组件整理，技能Figma。"),
        _case("S2-SAL-QUERY-019", "scenario2", "sales_presale", "experienced", "姓名马腾，电话13210008012，邮箱mateng@example.com，南京财经大学本科09-2013到06-2017，2017年7月至2025年5月做企业软件销售，负责客户开发、方案演示和回款。"),
        _case("S2-EDU-QUERY-020", "scenario2", "education", "experienced", "姓名许悦，电话13210008013，邮箱xuyue@example.com，华南师范大学本科09-2015到06-2019，2019年7月至2025年5月做培训机构教务，负责排课、学员服务和续费提醒。"),
        _case("S2-OPS-CONFLICT-021", "scenario2", "operations", "experienced", "姓名邓明，电话13210008014，邮箱dengming@example.com，武汉大学本科09-2012到06-2016。2020年3月到2025年6月在A公司做运营，2024年1月到2025年5月在B公司做运营主管。", conflicts=["experience_time_overlap"], missing=["skills"]),
        _case("S2-PRD-NO-PHONE-022", "scenario2", "product_research", "experienced", "姓名宋青，邮箱songqing@example.com，浙江大学本科计算机09-2014到06-2018，2018年7月到2025年6月负责企业后台产品需求、PRD和项目推进。", missing=["phone", "skills"]),
        _case("S2-FIN-NO-EMAIL-023", "scenario2", "finance", "experienced", "姓名高璐，电话13210008015，中央财经大学本科金融09-2015到06-2019，在银行负责贷款资料审核和客户材料归档，想做风控。", missing=["email", "experience.period", "skills"]),
        _case("S2-OTHER-QUERY-024", "scenario2", "other", "experienced", "姓名秦远，电话13210008016，邮箱qinyuan@example.com，北京工商大学本科09-2014到06-2018，2018年7月至2025年5月做行政项目协调，负责流程跟进、会议纪要和供应商沟通。"),
        _case("S3-PRD-CV-DOCX-025", "scenario3", "product_research", "experienced", "我想投数据产品经理，请优化简历说明，不要增加新公司。", cv="files/cv/cv_product_pm.docx"),
        _case("S3-OPS-CV-PDF-026", "scenario3", "operations", "experienced", "我想投用户增长运营，请把经历改得更有目标、动作、复盘。", cv="files/cv/cv_operations.pdf"),
        _case("S3-TEA-CV-PNG-027", "scenario3", "teacher", "experienced", "想投班主任兼语文老师，请优化教育表达。", cv="files/cv/cv_teacher.png"),
        _case("S3-DOC-CV-DOCX-028", "scenario3", "doctor", "experienced", "想投门诊内科医生，请优化但不要写没有的科研论文。", cv="files/cv/cv_doctor.docx", forbidden=["SCI论文"]),
        _case("S3-FIN-CV-DOCX-029", "scenario3", "finance", "experienced", "想投银行风控审核岗，请突出资料审核、风险意识。", cv="files/cv/cv_finance_bank.docx"),
        _case("S3-SAL-CV-PDF-030", "scenario3", "sales_presale", "experienced", "想投售前解决方案顾问，请突出客户需求澄清。", cv="files/cv/cv_sales.pdf"),
        _case("S3-DES-CV-PNG-031", "scenario3", "design", "experienced", "想投SaaS UI设计，请突出设计系统和交付。", cv="files/cv/cv_design.png"),
        _case("S3-EDU-CV-DOCX-032", "scenario3", "education", "experienced", "想投教务运营，请把教学经历转成教育服务和流程管理方向。", cv="files/cv/cv_teacher.docx", forbidden=["新增教务经理经历"]),
        _case("S3-PRD-TEMPLATE-033", "scenario3", "product_research", "experienced", "请按这个模板优化产品经理简历。", cv="files/cv/cv_product_pm.pdf", template="files/templates/template_modern.docx"),
        _case("S3-FIN-TEMPLATE-PDF-034", "scenario3", "finance", "experienced", "参考模板风格优化，不要改变事实。", cv="files/cv/cv_finance_bank.png", template="files/templates/template_minimal.pdf"),
        _case("S3-TEA-MISSING-035", "scenario3", "teacher", "experienced", "请优化老师岗位简历，如果缺信息请提示。", cv="files/cv/cv_teacher.docx"),
        _case("S3-OTHER-CV-036", "scenario3", "other", "experienced", "想投项目协调岗位，请弱化过度运营术语。", cv="files/cv/cv_operations.docx"),
        _case("S4-PRD-JD-DOCX-037", "scenario4", "product_research", "student", "姓名程洛，电话13210008017，邮箱chengluo@example.com，上海交通大学软件工程本科09-2022到06-2026，做过课程选课系统项目，负责需求分析和原型设计，技能SQL、Axure。", jd_file="files/jd/jd_product_pm.docx"),
        _case("S4-OPS-JD-PNG-038", "scenario4", "operations", "student", "姓名陆佳，电话13210008018，邮箱lujia@example.com，厦门大学广告学本科09-2021到06-2025，做过校园社群活动，负责招募、内容发布和复盘。", jd_file="files/jd/jd_operations.png"),
        _case("S4-FIN-JD-TEXT-039", "scenario4", "finance", "experienced", "姓名吴越，电话13210008019，邮箱wuyue@example.com，江西财经大学本科09-2016到06-2020，2020年7月至2025年6月在农商行做柜面和贷款资料审核。", jd_text=finance_jd),
        _case("S4-TEA-JD-PDF-040", "scenario4", "teacher", "student", "姓名韩青，电话13210008020，邮箱hanqing@example.com，首都师范大学本科09-2022到06-2026，有小学语文试讲和课程设计经历。", jd_file="files/jd/jd_teacher.pdf"),
        _case("S4-DOC-JD-TEXT-041", "scenario4", "doctor", "experienced", "姓名孟雨，电话13210008021，邮箱mengyu@example.com，临床医学硕士，2019年7月至2025年5月在社区医院内科负责门诊接诊和慢病随访。", jd_text=doctor_jd, missing=["education.school", "education.period"]),
        _case("S4-SAL-JD-DOCX-042", "scenario4", "sales_presale", "experienced", "姓名杜凯，电话13210008022，邮箱dukai@example.com，山东大学本科09-2012到06-2016，2016年7月至2025年5月做企业软件销售，负责客户开发和方案演示。", jd_file="files/jd/jd_sales.docx"),
        _case("S4-DES-JD-TEXT-043", "scenario4", "design", "student", "姓名叶岚，电话13210008023，邮箱yelan@example.com，广州美术学院本科09-2022到06-2026，做过校园服务App界面设计，技能Figma。", jd_text=design_jd),
        _case("S4-EDU-JD-PNG-044", "scenario4", "education", "experienced", "姓名严希，电话13210008024，邮箱yanxi@example.com，华南师范大学本科09-2015到06-2019，2019年7月至2025年5月负责培训机构排课、学员服务和数据台账。", jd_file="files/jd/jd_education_ops.png"),
        _case("S4-PRD-MISSING-045", "scenario4", "product_research", "student", "我想投产品实习，做过课程系统原型和用户访谈。", jd_text=product_jd, missing=["name", "phone", "email", "education.school", "education.period", "skills"]),
        _case("S4-FIN-CONFLICT-046", "scenario4", "finance", "experienced", "姓名白宁，电话13210008025，邮箱baining@example.com，中央财经大学本科09-2014到06-2018，2020年3月到2025年6月在A银行信贷审核，2024年6月到2025年5月在B银行风控审核。", jd_text=finance_jd, conflicts=["experience_time_overlap"]),
        _case("S4-TEA-TEMPLATE-047", "scenario4", "teacher", "student", "姓名罗晴，电话13210008026，邮箱luoqing@example.com，东北师范大学本科09-2022到06-2026，做过初中英语试讲、教案设计和作业反馈。", jd_file="files/jd/jd_teacher.png", template="files/templates/template_blue.png"),
        _case("S4-OTHER-JD-TEXT-048", "scenario4", "other", "experienced", "姓名顾川，电话13210008027，邮箱guchuan@example.com，北京工商大学本科09-2013到06-2017，2017年7月至2025年5月负责行政流程、供应商沟通和项目材料跟进。", jd_text="项目协调岗位：负责跨部门事项推进、材料整理、流程跟进和会议纪要。"),
    ]
    return cases


def _write_cases(cases: list[dict[str, Any]]) -> None:
    with (OUT / "cases.jsonl").open("w", encoding="utf-8") as handle:
        for case in cases:
            handle.write(json.dumps(case, ensure_ascii=False) + "\n")

    with (OUT / "cases.csv").open("w", encoding="utf-8", newline="") as handle:
        fieldnames = ["id", "scenario", "industry", "user_stage", "cv_path", "target_jd_file_path", "cv_template_path", "expected_missing_fields", "expected_conflicts"]
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for case in cases:
            writer.writerow({key: json.dumps(case.get(key), ensure_ascii=False) if isinstance(case.get(key), list) else case.get(key, "") for key in fieldnames})

    summary: dict[str, Any] = {"total": len(cases), "by_scenario": {}, "by_industry": {}, "by_stage": {}, "file_coverage": {}}
    for case in cases:
        for key, bucket in (("scenario", "by_scenario"), ("industry", "by_industry"), ("user_stage", "by_stage")):
            value = case[key]
            summary[bucket][value] = summary[bucket].get(value, 0) + 1
        for field in ("cv_path", "target_jd_file_path", "cv_template_path"):
            path = case.get(field)
            if not path:
                continue
            ext = Path(path).suffix.lower().lstrip(".")
            summary["file_coverage"][ext] = summary["file_coverage"].get(ext, 0) + 1
    (OUT / "manifest.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    _reset_dir(FILES)
    _write_fixture_files()
    cases = _cases()
    _write_cases(cases)
    print(json.dumps({"output": str(OUT), "case_count": len(cases)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
