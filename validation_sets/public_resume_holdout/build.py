"""Build frozen public resume holdout and shadow splits.

The source is the pinned Chinese JobResQA TSV. Source text is synthetic and
anonymized upstream. This builder intentionally does not call an LLM: split
membership, metadata, spans, and fixtures must remain deterministic.
"""

from __future__ import annotations

import csv
import hashlib
import io
import json
import math
import os
import random
import re
import shutil
import tempfile
import textwrap
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any, Iterable
from urllib.request import urlopen

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageEnhance, ImageFont


ROOT = Path(__file__).resolve().parent
REVISION = "1dbe6ffcf82cb06ec00f29bb0e4aeebde556addf"
SOURCE_URL = (
    "https://raw.githubusercontent.com/Avature/jobresqa-benchmark/"
    f"{REVISION}/data/jobresqa.zh.tsv"
)
SOURCE_SHA256 = "a749770f4064e3723fb5f7a712f66b88f28a2b5545c589a8d3998f83c8852814"
SUPPLEMENTAL_PATH = ROOT / "supplemental_profiles.json"
SUPPLEMENTAL_SHA256 = "e66ac5fb13ef61762a9e3fdebec2b5b4c655f2927462015ec82519713152b847"
SUPPLEMENTAL_REVISION = "178ab864dcad9910c5670d43e4bdbbb901a11f18"
SEED = 20260813
EXCLUDED_RESUME_IDS = {
    # Upstream contains two materially different end dates for this resume ID
    # ("至今" versus "2023年5月"). Keeping either would make the ground truth
    # ambiguous, so the complete root is excluded from both frozen splits.
    "01267": "conflicting source variants for one employment end date",
}

SPLIT_SIZES = {
    "holdout_v2": {"scenario1": 15, "scenario2": 15, "scenario3": 15, "scenario4": 15},
    "shadow_v3": {"scenario1": 6, "scenario2": 6, "scenario3": 6, "scenario4": 6},
}

INDUSTRY_KEYWORDS: dict[str, tuple[str, ...]] = {
    "doctor": (
        "医院", "医疗", "临床", "患者", "病人", "医生", "护士", "药房", "药剂", "健康信息",
        "medical", "clinical", "patient", "hospital", "nurse", "pharmacy",
    ),
    "teacher": (
        "教师", "老师", "授课", "课堂", "教学", "课程", "学校", "教研", "学生",
        "teacher", "teaching", "classroom", "school", "curriculum",
    ),
    "education": (
        "教育协调", "教育项目", "培训师", "培训项目", "学术项目", "招生",
        "education coordinator", "academic program", "training specialist",
    ),
    "finance": (
        "银行", "金融", "会计", "审计", "贷款", "信贷", "证券", "预算", "税务", "律师",
        "法律", "诉讼", "bank", "finance", "accounting", "audit", "loan", "legal", "attorney",
    ),
    "design": (
        "设计师", "平面设计", "视觉设计", "用户界面", "媒体制作", "艺术", "摄影",
        "designer", "graphic design", "visual design", "creative", "media production",
    ),
    "sales_presale": (
        "销售", "业务发展", "市场营销", "客户开发", "客户关系", "零售", "商机",
        "sales", "business development", "marketing", "retail", "account executive",
    ),
    "product_research": (
        "软件", "数据库", "数据分析", "信息技术", "程序", "网络工程", "系统工程", "研究科学",
        "software", "database", "data analyst", "information technology", "developer", "engineer",
        "research scientist", "technical consultant",
    ),
    "operations": (
        "运营", "行政", "人力资源", "项目协调", "办公室", "客户服务", "餐饮", "酒店", "物流",
        "采购", "维护", "现场服务", "operations", "administrative", "human resources", "coordinator",
        "customer service", "hospitality", "food service", "logistics", "maintenance",
    ),
}

SECTION_HEADINGS = {
    "摘要或目标", "专业经验", "工作经历", "教育背景", "教育", "技能", "专业技能", "核心技能",
    "个人简介", "联系信息", "职业目标", "资格证书", "证书", "项目经历", "荣誉", "语言",
}


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _sha256_text(text: str) -> str:
    return _sha256_bytes(text.encode("utf-8"))


def _atomic_write(path: Path, data: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd, temp_name = tempfile.mkstemp(prefix=f".{path.name}.", dir=path.parent)
    try:
        with os.fdopen(fd, "wb") as handle:
            handle.write(data)
            handle.flush()
            os.fsync(handle.fileno())
        os.replace(temp_name, path)
    except Exception:
        try:
            Path(temp_name).unlink(missing_ok=True)
        finally:
            raise


def _write_text(path: Path, text: str) -> None:
    _atomic_write(path, text.encode("utf-8"))


def _write_json(path: Path, payload: Any) -> None:
    _write_text(path, json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True) + "\n")


def _write_jsonl(path: Path, rows: Iterable[dict[str, Any]]) -> None:
    text = "".join(json.dumps(row, ensure_ascii=False, sort_keys=True) + "\n" for row in rows)
    _write_text(path, text)


def _download_source() -> bytes:
    with urlopen(SOURCE_URL, timeout=60) as response:
        data = response.read()
    actual = _sha256_bytes(data)
    if actual != SOURCE_SHA256:
        raise RuntimeError(f"JobResQA SHA256 mismatch: {actual} != {SOURCE_SHA256}")
    return data


def _parse_rows(data: bytes) -> list[dict[str, str]]:
    text = data.decode("utf-8-sig")
    rows = list(csv.DictReader(io.StringIO(text), delimiter="\t"))
    if len(rows) != 581:
        raise RuntimeError(f"Unexpected JobResQA row count: {len(rows)}")
    required = {"example_id", "resume_id", "resume", "jd_id", "jd", "language"}
    if not rows or not required.issubset(rows[0]):
        raise RuntimeError("JobResQA schema changed")
    return rows


def _consistent_by_id(
    rows: list[dict[str, str]],
    id_key: str,
    text_key: str,
    excluded_ids: set[str] | None = None,
) -> dict[str, str]:
    excluded_ids = excluded_ids or set()
    grouped: dict[str, set[str]] = defaultdict(set)
    for row in rows:
        if row[id_key] in excluded_ids:
            continue
        grouped[row[id_key]].add(row[text_key])
    inconsistent = {key: len(values) for key, values in grouped.items() if len(values) != 1}
    if inconsistent:
        raise RuntimeError(f"Inconsistent {text_key} for IDs: {inconsistent}")
    return {key: next(iter(values)) for key, values in grouped.items()}


def _industry(text: str) -> str:
    lowered = text.casefold()
    scores = {
        industry: sum(lowered.count(keyword.casefold()) for keyword in keywords)
        for industry, keywords in INDUSTRY_KEYWORDS.items()
    }
    best_score = max(scores.values(), default=0)
    if best_score == 0:
        return "other"
    priority = [
        "doctor", "teacher", "education", "finance", "design", "sales_presale",
        "product_research", "operations",
    ]
    return next(name for name in priority if scores[name] == best_score)


def _stage(text: str) -> str:
    lowered = text.casefold()
    explicit_years = [int(value) for value in re.findall(r"(?:超过|拥有|具备)?\s*(\d{1,2})\s*年", text)]
    if explicit_years and max(explicit_years) >= 6:
        return "experienced"
    if any(token in lowered for token in ("在校生", "应届", "在读", "学生简历", "seeking an internship")):
        return "student"
    year_tokens = {int(value) for value in re.findall(r"(?:19|20)\d{2}", text)}
    if len(year_tokens) >= 4 or len(re.findall(r"(?:至今|present)", lowered)) >= 1:
        return "experienced"
    return "job_seeker"


def _supplemental_industry(role: str) -> str:
    lowered = role.casefold()
    if any(token in lowered for token in ("engineer", "data analyst", "product manager")):
        return "product_research"
    if any(token in lowered for token in ("marketing", "sales", "business development", "account executive")):
        return "sales_presale"
    if any(token in lowered for token in ("financial", "accountant", "fp&a")):
        return "finance"
    if any(token in lowered for token in ("support", "project manager", "program coordinator", "operations")):
        return "operations"
    return "other"


def _supplemental_resume_text(profile: dict[str, Any]) -> str:
    skills = ", ".join(str(value) for value in profile["skills"])
    bullets = "\n".join(f"- {value}" for value in profile["experience_bullets"])
    return (
        "ANONYMIZED SYNTHETIC PROFILE\n"
        f"Professional Direction: {profile['role']}\n"
        f"Career Level: {profile['seniority']}\n"
        f"Years of Experience: {profile['years_experience']}\n"
        f"Industry: {profile['industry']}\n\n"
        "Summary\n"
        f"{profile['summary']}\n\n"
        "Education\n"
        f"{profile['education']}\n\n"
        "Skills\n"
        f"{skills}\n\n"
        "Experience Highlights\n"
        f"{bullets}"
    )


def _load_supplemental_rows() -> list[dict[str, Any]]:
    raw = SUPPLEMENTAL_PATH.read_bytes()
    actual = _sha256_bytes(raw)
    if actual != SUPPLEMENTAL_SHA256:
        raise RuntimeError(f"Supplemental profile SHA256 mismatch: {actual} != {SUPPLEMENTAL_SHA256}")
    payload = json.loads(raw)
    if payload.get("revision") != SUPPLEMENTAL_REVISION or payload.get("license") != "MIT":
        raise RuntimeError("Supplemental profile provenance changed")
    profiles = payload.get("profiles") or []
    if len(profiles) != 18 or len({item["resume_id"] for item in profiles}) != 18:
        raise RuntimeError("Expected 18 unique supplemental profiles")
    rows: list[dict[str, Any]] = []
    for profile in profiles:
        years = int(profile["years_experience"])
        rows.append(
            {
                "_source_dataset": payload["dataset"],
                "_source_revision": SUPPLEMENTAL_REVISION,
                "_source_license": "MIT",
                "_source_language": "en",
                "_source_slug": "candidate_matching",
                "_industry_label": _supplemental_industry(str(profile["role"])),
                "_source_domain": profile["industry"],
                "_user_stage": "job_seeker" if years == 0 else "experienced",
                "_career_level": profile["seniority"],
                "_years_experience": years,
                "resume_id": profile["resume_id"],
                "resume": _supplemental_resume_text(profile),
                "jd_id": "",
                "jd": "",
            }
        )
    return rows


def _stratified_ids(resumes: dict[str, str], count: int, seed: int) -> list[str]:
    groups: dict[str, list[str]] = defaultdict(list)
    for resume_id, text in sorted(resumes.items()):
        groups[_industry(text)].append(resume_id)
    rng = random.Random(seed)
    for ids in groups.values():
        rng.shuffle(ids)
    order = [
        "product_research", "operations", "doctor", "teacher", "education",
        "finance", "sales_presale", "design", "other",
    ]
    selected: list[str] = []
    while len(selected) < count:
        progressed = False
        for name in order:
            if groups[name]:
                selected.append(groups[name].pop())
                progressed = True
                if len(selected) == count:
                    break
        if not progressed:
            raise RuntimeError(f"Only found {len(selected)} resumes, need {count}")
    return selected


def _representative_rows(rows: list[dict[str, str]]) -> dict[str, dict[str, str]]:
    output: dict[str, dict[str, str]] = {}
    for row in sorted(rows, key=lambda item: (item["resume_id"], item["jd_id"], item["example_id"])):
        output.setdefault(row["resume_id"], row)
    return output


def _source_units(text: str, source_kind: str, candidate: bool, offset: int = 0) -> list[dict[str, Any]]:
    units: list[dict[str, Any]] = []
    section = "unknown"
    for line_number, match in enumerate(re.finditer(r"[^\r\n]+", text), start=1):
        raw = match.group(0)
        stripped = raw.strip()
        if not stripped:
            continue
        if stripped in SECTION_HEADINGS or (len(stripped) <= 16 and stripped.endswith("经历")):
            section = stripped
            fact_type = "section_heading"
            eligible = False
        elif re.fullmatch(r"[\[【].+[\]】][:：]?", stripped):
            fact_type = "placeholder"
            eligible = False
        elif any(token in stripped for token in ("[邮箱]", "[电话]", "[地址]", "[姓名]")):
            fact_type = "identity_or_contact"
            eligible = candidate
        elif re.search(r"(?:19|20)\d{2}|至今|Present", stripped, re.IGNORECASE):
            fact_type = "dated_record"
            eligible = candidate
        elif stripped.startswith(("•", "-", "*", "·", "▪")):
            fact_type = "experience_detail"
            eligible = candidate
        else:
            fact_type = "source_statement"
            eligible = candidate
        start = offset + match.start()
        end = offset + match.end()
        units.append(
            {
                "fact_id": f"{source_kind}-{len(units) + 1:04d}",
                "source_kind": source_kind,
                "source_span": [start, end],
                "line_number": line_number,
                "section": section,
                "fact_type": fact_type,
                "text": raw,
                "candidate_for_resume": bool(eligible),
                "confidence": 1.0,
            }
        )
    return units


def _font_path() -> str:
    candidates = [
        # Unifont preserves both standard CJK code points and the bullet glyph
        # during PDF text extraction. Micro Hei maps common characters such as
        # 力/年 to compatibility glyphs, while Zen Hei maps • to NUL.
        "/usr/share/fonts/truetype/unifont/unifont.ttf",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    raise RuntimeError("No Chinese font found")


def _wrapped_lines(text: str, width: int) -> list[str]:
    output: list[str] = []
    for line in text.splitlines():
        if not line.strip():
            output.append("")
            continue
        output.extend(textwrap.wrap(line, width=width, replace_whitespace=False) or [""])
    return output


def _render_docx(path: Path, text: str, two_column: bool) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(9)
    lines = text.splitlines()
    if two_column:
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        midpoint = math.ceil(len(lines) / 2)
        for cell, chunk in zip(table.rows[0].cells, (lines[:midpoint], lines[midpoint:])):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.width = Cm(8.6)
            for index, line in enumerate(chunk):
                paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(1.5)
                run = paragraph.add_run(line)
                run.font.name = "Microsoft YaHei"
                run.font.size = Pt(8.5)
    else:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("个人简历")
        title_run.bold = True
        title_run.font.size = Pt(17)
        title_run.font.color.rgb = RGBColor(35, 67, 110)
        for line in lines:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1.5)
            run = paragraph.add_run(line)
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(9)
    path.parent.mkdir(parents=True, exist_ok=True)
    temp = path.with_name(f".{path.name}.tmp.docx")
    document.save(temp)
    os.replace(temp, path)


def _render_pdf(path: Path, text: str) -> None:
    font_path = _font_path()
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    page.insert_font(fontname="holdoutcjk", fontfile=font_path)
    page.draw_rect(fitz.Rect(28, 28, 567, 814), color=(0.15, 0.27, 0.45), width=0.8)
    page.insert_text((44, 54), "个人简历", fontname="holdoutcjk", fontsize=18, color=(0.12, 0.25, 0.45))
    lines = _wrapped_lines(text, width=34)
    midpoint = math.ceil(len(lines) / 2)
    columns = (lines[:midpoint], lines[midpoint:])
    for left, column in zip((44, 310), columns):
        rect = fitz.Rect(left, 72, left + 235, 800)
        value = "\n".join(column)
        remaining = page.insert_textbox(
            rect,
            value,
            fontname="holdoutcjk",
            fontsize=7.4,
            lineheight=1.25,
            color=(0.08, 0.08, 0.08),
        )
        if remaining < 0:
            raise RuntimeError(f"PDF fixture overflow for {path.name}: {remaining}")
    path.parent.mkdir(parents=True, exist_ok=True)
    temp = path.with_name(f".{path.name}.tmp.pdf")
    document.save(temp, garbage=4, deflate=True)
    document.close()
    os.replace(temp, path)


def _render_png(path: Path, text: str) -> None:
    font_path = _font_path()
    body_font = ImageFont.truetype(font_path, 21)
    title_font = ImageFont.truetype(font_path, 40)
    # PIL does not wrap text to a pixel box. Keep the character width
    # conservative so CJK and mixed Latin lines cannot bleed into the other
    # column even when the chosen font has wide glyphs.
    lines = _wrapped_lines(text, width=25)
    midpoint = math.ceil(len(lines) / 2)
    max_column_lines = max(midpoint, len(lines) - midpoint)
    height = max(1980, 230 + max_column_lines * 29)
    image = Image.new("RGB", (1400, height), (247, 246, 242))
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((38, 38, 1362, height - 38), radius=18, fill="white", outline=(45, 76, 117), width=3)
    draw.text((72, 66), "个人简历", font=title_font, fill=(35, 67, 110))
    for x, column in zip((74, 735), (lines[:midpoint], lines[midpoint:])):
        y = 138
        for line in column:
            draw.text((x, y), line, font=body_font, fill=(24, 24, 24))
            y += 29
    image = ImageEnhance.Contrast(image.rotate(0.18, resample=Image.Resampling.BICUBIC, fillcolor=(242, 242, 238))).enhance(0.96)
    path.parent.mkdir(parents=True, exist_ok=True)
    temp = path.with_name(f".{path.name}.tmp.png")
    image.save(temp, format="PNG", optimize=True)
    os.replace(temp, path)


def _render_cv(path: Path, text: str, profile: str) -> None:
    if path.suffix == ".docx":
        _render_docx(path, text, two_column=profile == "two_column_docx")
    elif path.suffix == ".pdf":
        _render_pdf(path, text)
    elif path.suffix == ".png":
        _render_png(path, text)
    else:
        _write_text(path, text.rstrip() + "\n")


def _case_prefix(split: str) -> str:
    return "HV2" if split == "holdout_v2" else "SV3"


def _relative_for_runner(path: Path) -> str:
    # acceptance_testset/run_api_testset.py resolves fixture paths from its own directory.
    project_root = ROOT.parents[1]
    acceptance_root = project_root / "acceptance_testset"
    return os.path.relpath(path, acceptance_root)


def _base_case(case_id: str, scenario: str, industry: str, stage: str) -> dict[str, Any]:
    return {
        "id": case_id,
        "scenario": scenario,
        "industry": industry,
        "user_stage": stage,
        "cv_path": None,
        "target_jd": None,
        "target_jd_file_path": None,
        "cv_template_path": None,
        "expected_missing_fields": [],
        "expected_conflicts": [],
        "forbidden_fabrication": [],
        "expected_output": {
            "format": "docx",
            "allow_multi_page": True,
            "reply_text_must_cover": ["生成方向", "缺失信息", "岗位建议", "冲突检查"],
        },
    }


def _build_split(
    split: str,
    assignments: dict[str, list[dict[str, Any]]],
    layout_budget: int,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    split_root = ROOT / split
    cases: list[dict[str, Any]] = []
    annotations: list[dict[str, Any]] = []
    layout_profiles = [
        ("single_column_docx", ".docx"),
        ("two_column_pdf", ".pdf"),
        ("scanned_two_column_png", ".png"),
    ]
    layout_index = 0
    prefix = _case_prefix(split)

    for scenario_index, scenario in enumerate(("scenario1", "scenario2", "scenario3", "scenario4"), start=1):
        for case_index, row in enumerate(assignments[scenario], start=1):
            resume_text = row.get("resume", "")
            jd_text = row.get("jd", "")
            industry = row.get("_industry_label") or _industry(resume_text if scenario != "scenario4" else jd_text)
            stage = row.get("_user_stage") or (_stage(resume_text) if scenario != "scenario4" else "unknown")
            case_id = f"{prefix}-S{scenario_index}-{case_index:03d}"
            case = _base_case(case_id, scenario, industry, stage)
            case["career_level"] = row.get("_career_level", "unspecified")
            case["source_domain"] = row.get("_source_domain", industry)
            case["years_experience"] = row.get("_years_experience")
            resume_id = row["resume_id"] if scenario != "scenario4" else None
            jd_id = row["jd_id"] if scenario in {"scenario1", "scenario4"} else None
            source_slug = row.get("_source_slug", "jobresqa")
            annotation_sources: list[dict[str, Any]] = []

            if scenario in {"scenario1", "scenario3"}:
                raw_cv_path = split_root / "sources" / "resumes" / f"{source_slug}_{resume_id}.txt"
                _write_text(raw_cv_path, resume_text.rstrip() + "\n")
                use_layout = split == "holdout_v2" and layout_index < layout_budget
                if use_layout:
                    profile, suffix = layout_profiles[layout_index % len(layout_profiles)]
                    cv_path = split_root / "files" / "cv" / f"{source_slug}_{resume_id}_{profile}{suffix}"
                    _render_cv(cv_path, resume_text, profile)
                    layout_index += 1
                else:
                    profile = "plain_text"
                    cv_path = split_root / "files" / "cv" / f"{source_slug}_{resume_id}.txt"
                    _render_cv(cv_path, resume_text, profile)
                case["cv_path"] = _relative_for_runner(cv_path)
                case["input_profile"] = profile
                annotation_sources.append(
                    {
                        "kind": "cv",
                        "canonical_text_path": str(raw_cv_path.relative_to(ROOT)),
                        "sha256": _sha256_text(resume_text.rstrip() + "\n"),
                        "candidate_for_resume": True,
                        "units": _source_units(resume_text, "cv", True),
                    }
                )

            if scenario == "scenario1":
                case["query"] = "请根据目标岗位优化这份简历；只能使用简历中的个人事实，不得把JD要求写成候选人经历。"
                jd_path = split_root / "files" / "jd" / f"jobresqa_{jd_id}.txt"
                _write_text(jd_path, jd_text.rstrip() + "\n")
                case["target_jd_file_path"] = _relative_for_runner(jd_path)
                annotation_sources.append(
                    {
                        "kind": "jd",
                        "canonical_text_path": str(jd_path.relative_to(ROOT)),
                        "sha256": _sha256_text(jd_text.rstrip() + "\n"),
                        "candidate_for_resume": False,
                        "units": _source_units(jd_text, "jd", False),
                    }
                )
            elif scenario == "scenario2":
                query_prefix = "以下是我提供的全部个人信息。请生成结构化简历，完整保留事实，不要补造：\n\n"
                query = query_prefix + resume_text
                case["query"] = query
                case["input_profile"] = "query_only_full_profile"
                raw_query_path = split_root / "sources" / "queries" / f"{source_slug}_{resume_id}.txt"
                _write_text(raw_query_path, query.rstrip() + "\n")
                annotation_sources.append(
                    {
                        "kind": "query",
                        "canonical_text_path": str(raw_query_path.relative_to(ROOT)),
                        "sha256": _sha256_text(query.rstrip() + "\n"),
                        "candidate_for_resume": True,
                        "units": _source_units(resume_text, "query", True, offset=len(query_prefix)),
                    }
                )
            elif scenario == "scenario3":
                case["query"] = "请优化这份简历的结构和表达；保留全部有价值的源信息，不要增加新公司、岗位、数字或资质。"
            else:
                case["query"] = "请根据目标岗位生成一份可填写的结构化简历框架。我没有提供任何个人信息，请勿编造。"
                jd_path = split_root / "files" / "jd" / f"jobresqa_{jd_id}.txt"
                _write_text(jd_path, jd_text.rstrip() + "\n")
                case["target_jd_file_path"] = _relative_for_runner(jd_path)
                case["expected_missing_fields"] = ["个人信息", "教育背景", "经历", "技能"]
                case["expected_output"]["framework_only"] = True
                case["input_profile"] = "jd_only_no_personal_facts"
                annotation_sources.append(
                    {
                        "kind": "jd",
                        "canonical_text_path": str(jd_path.relative_to(ROOT)),
                        "sha256": _sha256_text(jd_text.rstrip() + "\n"),
                        "candidate_for_resume": False,
                        "units": _source_units(jd_text, "jd", False),
                    }
                )

            case["provenance"] = {
                "dataset": row.get("_source_dataset", "JobResQA"),
                "revision": row.get("_source_revision", REVISION),
                "language": row.get("_source_language", "zh"),
                "resume_id": resume_id,
                "jd_id": jd_id,
                "license": row.get("_source_license", "CC-BY-SA-2.0"),
            }
            case["evaluation_policy"] = {
                "candidate_fact_sources": ["cv", "query"],
                "non_candidate_fact_sources": ["jd", "instruction"],
                "require_atomic_grounding": True,
                "require_ownership_integrity": True,
                "human_expression_review": True,
            }
            cases.append(case)
            annotations.append(
                {
                    "case_id": case_id,
                    "annotation_version": "1.0",
                    "status": "exact_spans_ready_semantic_human_review_pending",
                    "sources": annotation_sources,
                }
            )

    return cases, annotations


def _select_assignments(rows: list[dict[str, str]]) -> dict[str, dict[str, list[dict[str, Any]]]]:
    resumes = _consistent_by_id(rows, "resume_id", "resume", set(EXCLUDED_RESUME_IDS))
    _consistent_by_id(rows, "jd_id", "jd")
    representative = _representative_rows(rows)
    supplemental_rows = _load_supplemental_rows()
    supplemental_counts = {
        "holdout_v2": {"scenario2": 6, "scenario3": 6},
        "shadow_v3": {"scenario2": 3, "scenario3": 3},
    }
    total_profile_cases = sum(
        SPLIT_SIZES[split][scenario]
        for split in SPLIT_SIZES
        for scenario in ("scenario1", "scenario2", "scenario3")
    )
    resume_needed = total_profile_cases - len(supplemental_rows)
    selected_ids = _stratified_ids(resumes, resume_needed, SEED)
    cursor = 0
    supplemental_cursor = 0
    assignments: dict[str, dict[str, list[dict[str, Any]]]] = {}
    used_resume_ids: set[str] = set()
    used_jd_ids: set[str] = set()
    for split in ("holdout_v2", "shadow_v3"):
        split_assignments: dict[str, list[dict[str, Any]]] = {}
        for scenario in ("scenario1", "scenario2", "scenario3"):
            size = SPLIT_SIZES[split][scenario]
            supplemental_size = supplemental_counts.get(split, {}).get(scenario, 0)
            jobresqa_size = size - supplemental_size
            ids = selected_ids[cursor : cursor + jobresqa_size]
            cursor += jobresqa_size
            selected_rows = [representative[resume_id] for resume_id in ids]
            if supplemental_size:
                selected_rows.extend(
                    supplemental_rows[supplemental_cursor : supplemental_cursor + supplemental_size]
                )
                supplemental_cursor += supplemental_size
            split_assignments[scenario] = selected_rows
            used_resume_ids.update(ids)
            if scenario == "scenario1":
                used_jd_ids.update(row["jd_id"] for row in selected_rows)
        assignments[split] = split_assignments

    if cursor != len(selected_ids) or supplemental_cursor != len(supplemental_rows):
        raise RuntimeError("Profile assignment accounting mismatch")

    jd_candidates: list[dict[str, str]] = []
    seen_jds: set[str] = set()
    rng = random.Random(SEED + 99)
    shuffled_rows = list(rows)
    rng.shuffle(shuffled_rows)
    for row in shuffled_rows:
        if row["resume_id"] in used_resume_ids or row["jd_id"] in used_jd_ids or row["jd_id"] in seen_jds:
            continue
        seen_jds.add(row["jd_id"])
        jd_candidates.append(row)
    needed_jds = sum(SPLIT_SIZES[split]["scenario4"] for split in SPLIT_SIZES)
    if len(jd_candidates) < needed_jds:
        raise RuntimeError(f"Only {len(jd_candidates)} disjoint JDs available, need {needed_jds}")
    jd_cursor = 0
    for split in ("holdout_v2", "shadow_v3"):
        size = SPLIT_SIZES[split]["scenario4"]
        assignments[split]["scenario4"] = jd_candidates[jd_cursor : jd_cursor + size]
        jd_cursor += size
    return assignments


def _file_manifest(split_root: Path) -> list[dict[str, Any]]:
    output: list[dict[str, Any]] = []
    for path in sorted(item for item in split_root.rglob("*") if item.is_file()):
        if path.name in {"manifest.json"}:
            continue
        data = path.read_bytes()
        output.append(
            {
                "path": str(path.relative_to(split_root)),
                "bytes": len(data),
                "sha256": _sha256_bytes(data),
            }
        )
    return output


def main() -> None:
    source_data = _download_source()
    rows = _parse_rows(source_data)
    resumes = _consistent_by_id(rows, "resume_id", "resume", set(EXCLUDED_RESUME_IDS))
    jds = _consistent_by_id(rows, "jd_id", "jd")
    assignments = _select_assignments(rows)

    # These are generated-only directories owned by this builder. Resetting
    # them prevents stale fixtures from an older frozen split entering a new
    # manifest after the deterministic selection logic changes.
    for split in SPLIT_SIZES:
        split_root = (ROOT / split).resolve()
        if split_root.parent != ROOT.resolve():
            raise RuntimeError(f"Refusing to reset unexpected path: {split_root}")
        if split_root.exists():
            shutil.rmtree(split_root)

    source_manifest = {
        "name": "JobResQA Chinese split",
        "homepage": "https://github.com/Avature/jobresqa-benchmark",
        "source_url": SOURCE_URL,
        "revision": REVISION,
        "license": "CC-BY-SA-2.0",
        "source_sha256": SOURCE_SHA256,
        "source_bytes": len(source_data),
        "source_rows": len(rows),
        "unique_resumes": len(resumes),
        "unique_jds": len(jds),
        "excluded_resume_ids": EXCLUDED_RESUME_IDS,
        "selection_seed": SEED,
        "supplemental_source": {
            "name": "Candidate-Job Matching Synthetic Dataset",
            "homepage": "https://huggingface.co/datasets/michaelozon/candidate-matching-synthetic",
            "revision": SUPPLEMENTAL_REVISION,
            "license": "MIT",
            "selected_profile_file": SUPPLEMENTAL_PATH.name,
            "selected_profile_sha256": SUPPLEMENTAL_SHA256,
            "selected_profiles": 18,
        },
    }
    _write_json(ROOT / "source_manifest.json", source_manifest)

    split_summaries: dict[str, Any] = {}
    for split in ("holdout_v2", "shadow_v3"):
        layout_budget = 12 if split == "holdout_v2" else 0
        cases, annotations = _build_split(split, assignments[split], layout_budget)
        split_root = ROOT / split
        _write_jsonl(split_root / "cases.jsonl", cases)
        _write_jsonl(split_root / "annotations.jsonl", annotations)
        summary = {
            "split": split,
            "frozen": True,
            "case_count": len(cases),
            "scenario_counts": dict(sorted(Counter(case["scenario"] for case in cases).items())),
            "industry_counts": dict(sorted(Counter(case["industry"] for case in cases).items())),
            "stage_counts": dict(sorted(Counter(case["user_stage"] for case in cases).items())),
            "career_level_counts": dict(sorted(Counter(case["career_level"] for case in cases).items())),
            "source_domain_counts": dict(sorted(Counter(case["source_domain"] for case in cases).items())),
            "source_dataset_counts": dict(sorted(Counter(case["provenance"]["dataset"] for case in cases).items())),
            "input_profile_counts": dict(sorted(Counter(case["input_profile"] for case in cases).items())),
            "source_resume_ids": sorted(
                case["provenance"]["resume_id"] for case in cases if case["provenance"]["resume_id"]
            ),
            "source_jd_ids": sorted(
                case["provenance"]["jd_id"] for case in cases if case["provenance"]["jd_id"]
            ),
        }
        _write_json(split_root / "summary.json", summary)
        _write_json(split_root / "manifest.json", {**summary, "files": _file_manifest(split_root)})
        split_summaries[split] = summary

    _write_json(
        ROOT / "split_manifest.json",
        {
            "dataset_version": "public-resume-holdout-1.0",
            "source": source_manifest,
            "splits": split_summaries,
            "policy": {
                "holdout_for_promotion_only": True,
                "shadow_not_for_daily_development": True,
                "retire_case_after_fixing_against_it": True,
                "jd_is_not_candidate_evidence": True,
                "one_page_limit": False,
            },
        },
    )
    print(json.dumps(split_summaries, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
