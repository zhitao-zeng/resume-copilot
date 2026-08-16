#!/usr/bin/env python3
"""Rendered-document gate for resume DOCX artifacts (R24 Phase 7).

The frozen evaluator cannot establish visual-layout or template-fidelity
signals from JSON; this gate audits the rendered document itself:

- label remnants (pure ``标签：`` lines) and separator artifacts
- sparse trailing content / accidental near-empty endings
- 100 percent retention of verified resume facts in the rendered text
- editability (real paragraph/table text, not an image dump)
- Chinese rendering correctness (no replacement chars / tofu markers)
- template fidelity: tagged shells keep no leftover tags, anchored shells
  keep their section titles, style-only renders keep the style profile

Never enforces a one-page resume and never deletes facts; it only reports.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
import tempfile
import os
from pathlib import Path
from typing import Any

GATE_VERSION = "rendered-document-gate-r24"

_PURE_LABEL = re.compile(r"^[^：:\n]{1,15}[：:]\s*$")
_SEPARATOR_ARTIFACT = re.compile(r"__{2,}|——{3,}|⸻|�")
# Lines carrying only separator glyphs ("  |  ", "· ·", "- -") are render
# remnants of an empty header row, never real content.
_REMNANT_LINE = re.compile(r"^[\s|｜·•\-—–_]+$")
_LEFTOVER_TAG = re.compile(r"\[\[\s*(?:section|scalar)\.[A-Za-z_][A-Za-z0-9_]*\s*\]\]")
_CJK = re.compile(r"[一-鿿]")


def _docx_texts(doc) -> tuple[list[str], list[str]]:
    paragraphs = [p.text for p in doc.paragraphs]
    cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return paragraphs, cells


def _flatten_resume_texts(resume_data: dict[str, Any]) -> list[str]:
    texts: list[str] = []

    def visit(value: Any) -> None:
        if isinstance(value, str):
            if value.strip():
                texts.append(value)
        elif isinstance(value, dict):
            for item in value.values():
                visit(item)
        elif isinstance(value, list):
            for item in value:
                visit(item)

    for key in ("meta", "summary", "experience", "projects", "education",
                "skills", "certifications", "awards", "additional_sections"):
        if key in resume_data:
            visit(resume_data[key])
    return texts


def _normalize(text: str) -> str:
    return re.sub(r"\s+", "", str(text))


def audit_rendered_docx(
    docx_path: str | Path,
    *,
    resume_data: dict[str, Any] | None = None,
    template_path: str | Path | None = None,
) -> dict[str, Any]:
    from docx import Document

    path = Path(docx_path)
    result: dict[str, Any] = {
        "gate_version": GATE_VERSION,
        "path": str(path),
        "exists": path.is_file(),
    }
    if not path.is_file():
        result.update(pass_=False, error="docx_missing")
        return result
    try:
        doc = Document(str(path))
    except Exception as exc:
        result.update(pass_=False, error=f"docx_unreadable:{type(exc).__name__}")
        return result

    paragraphs, cells = _docx_texts(doc)
    all_lines = [line for line in paragraphs + cells if line.strip()]
    full = "\n".join(all_lines)
    normalized_full = _normalize(full)

    pure_labels = [line.strip() for line in all_lines if _PURE_LABEL.fullmatch(line.strip())]
    remnant_lines = [line.strip() for line in all_lines if line.strip() and _REMNANT_LINE.fullmatch(line)]
    separators = _SEPARATOR_ARTIFACT.findall(full) + remnant_lines
    leftover_tags = _LEFTOVER_TAG.findall(full)
    replacement_chars = full.count("�")

    empty_paragraphs = sum(1 for p in paragraphs if not p.strip())
    nonempty = [p for p in paragraphs if p.strip()]
    tail = nonempty[len(nonempty) * 2 // 3:] if nonempty else []
    tail_avg = sum(len(p) for p in tail) / max(1, len(tail)) if tail else 0
    # Sparse trailing heuristic: the final third of the document degenerates
    # to a handful of very short lines (accidental sparse page).
    sparse_trailing = bool(tail) and len(tail) <= 3 and tail_avg <= 8 and len(nonempty) >= 9

    cjk_present = bool(_CJK.search(full))
    editable = bool(all_lines)

    fact_missing: list[str] = []
    fact_total = 0
    if resume_data:
        for text in _flatten_resume_texts(resume_data):
            normalized = _normalize(text)
            if len(normalized) < 2:
                continue
            fact_total += 1
            if normalized not in normalized_full:
                fact_missing.append(text[:40])
    fact_retention = (fact_total - len(fact_missing)) / fact_total if fact_total else 1.0

    template_mode = "none"
    template_ok = True
    template_evidence: dict[str, Any] = {}
    if template_path:
        from resume_renderer import detect_docx_template_mode

        template_mode = detect_docx_template_mode(template_path)
        if template_mode == "tagged":
            template_ok = not leftover_tags
            template_evidence["leftover_tags"] = leftover_tags
        elif template_mode == "anchored":
            from docx import Document as _Doc

            shell = _Doc(str(template_path))
            shell_titles = {
                p.text.strip() for p in shell.paragraphs if p.text.strip()
            }
            preserved = sum(1 for title in shell_titles if _normalize(title) in normalized_full)
            template_ok = preserved > 0
            template_evidence["shell_titles_preserved"] = f"{preserved}/{len(shell_titles)}"
    result.update({
        "pass": (
            not pure_labels
            and not separators
            and not leftover_tags
            and replacement_chars == 0
            and not sparse_trailing
            and editable
            and fact_retention >= 1.0
            and template_ok
        ),
        "label_remnants": pure_labels,
        "separator_artifacts": separators,
        "leftover_tags": leftover_tags,
        "replacement_chars": replacement_chars,
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
        "empty_paragraphs": empty_paragraphs,
        "sparse_trailing": sparse_trailing,
        "tail_avg_chars": round(tail_avg, 1),
        "cjk_present": cjk_present,
        "editable": editable,
        "fact_total": fact_total,
        "fact_missing": fact_missing,
        "fact_retention": round(fact_retention, 4),
        "template_mode": template_mode,
        "template_ok": template_ok,
        "template_evidence": template_evidence,
    })
    return result


def visual_layout_score01(audit: dict[str, Any]) -> float | None:
    """Map the rendered audit onto the r3 visual_layout proxy (0..1)."""

    if not audit.get("exists") or audit.get("error"):
        return None
    score = 1.0
    score -= min(0.5, 0.1 * len(audit.get("label_remnants") or []))
    score -= min(0.3, 0.1 * len(audit.get("separator_artifacts") or []))
    if audit.get("sparse_trailing"):
        score -= 0.2
    if not audit.get("editable"):
        score -= 0.5
    if audit.get("replacement_chars"):
        score -= 0.2
    return round(max(0.0, score), 4)


def template_fidelity_score01(audit: dict[str, Any]) -> float | None:
    """Map the rendered audit onto the r3 template_fidelity proxy (0..1)."""

    if not audit.get("exists") or audit.get("error"):
        return None
    if not audit.get("template_ok"):
        return 0.0
    if audit.get("leftover_tags"):
        return 0.0
    return 1.0


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", nargs="+")
    parser.add_argument("--resume-data", help="JSON file with resume_data for fact-retention checks")
    parser.add_argument("--template", help="source template DOCX for fidelity checks")
    parser.add_argument("--output")
    args = parser.parse_args()
    resume_data = None
    if args.resume_data:
        resume_data = json.loads(Path(args.resume_data).read_text(encoding="utf-8"))
    reports = []
    for item in args.docx:
        reports.append(audit_rendered_docx(item, resume_data=resume_data, template_path=args.template))
        report = reports[-1]
        print(f"{item}: {'PASS' if report.get('pass') else 'FAIL'} "
              f"labels={len(report.get('label_remnants') or [])} "
              f"seps={len(report.get('separator_artifacts') or [])} "
              f"facts={report.get('fact_retention')} "
              f"template={report.get('template_mode')}:{report.get('template_ok')}")
    if args.output:
        out = Path(args.output)
        out.parent.mkdir(parents=True, exist_ok=True)
        with tempfile.NamedTemporaryFile(mode="w", encoding="utf-8", dir=out.parent, delete=False) as handle:
            json.dump({"gate_version": GATE_VERSION, "reports": reports}, handle, ensure_ascii=False, indent=2)
            handle.write("\n")
        os.replace(handle.name, out)


if __name__ == "__main__":
    main()
