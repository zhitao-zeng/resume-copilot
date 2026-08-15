"""Verify the frozen public resume holdout without running the product model."""

from __future__ import annotations

import hashlib
import json
from collections import Counter
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from PIL import Image, ImageStat


ROOT = Path(__file__).resolve().parent
EXPECTED = {
    "holdout_v2": {"cases": 60, "per_scenario": 15},
    "shadow_v3": {"cases": 24, "per_scenario": 6},
}


def _load_jsonl(path: Path) -> list[dict[str, Any]]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def _sha(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _verify_sources(annotation: dict[str, Any]) -> None:
    for source in annotation["sources"]:
        source_path = ROOT / source["canonical_text_path"]
        assert source_path.exists(), source_path
        assert _sha(source_path) == source["sha256"], source_path
        text = source_path.read_text(encoding="utf-8")
        for unit in source["units"]:
            start, end = unit["source_span"]
            assert 0 <= start < end <= len(text), (source_path, unit)
            assert text[start:end] == unit["text"], (source_path, unit)
            assert unit["candidate_for_resume"] is source["candidate_for_resume"] or not unit["candidate_for_resume"]


def _normalized_character_recall(source: str, rendered: str) -> float:
    source_counts = Counter("".join(source.split()))
    rendered_counts = Counter("".join(rendered.split()))
    matched = sum(min(count, rendered_counts[char]) for char, count in source_counts.items())
    return matched / max(sum(source_counts.values()), 1)


def _docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def _pdf_text(path: Path) -> str:
    with fitz.open(path) as document:
        return "\n".join(page.get_text() for page in document)


def _verify_rendered_cv(case: dict[str, Any], annotation: dict[str, Any]) -> float | None:
    value = case.get("cv_path")
    if not value:
        return None
    rendered_path = (ROOT.parents[1] / "acceptance_testset" / value).resolve()
    cv_source = next((source for source in annotation["sources"] if source["kind"] == "cv"), None)
    if cv_source is None:
        return None
    source_path = ROOT / cv_source["canonical_text_path"]
    source_text = source_path.read_text(encoding="utf-8")
    if rendered_path.suffix == ".txt":
        assert rendered_path.read_text(encoding="utf-8") == source_text, rendered_path
        return 1.0
    if rendered_path.suffix == ".docx":
        recall = _normalized_character_recall(source_text, _docx_text(rendered_path))
        assert recall == 1.0, (rendered_path, recall)
        return recall
    if rendered_path.suffix == ".pdf":
        recall = _normalized_character_recall(source_text, _pdf_text(rendered_path))
        assert recall == 1.0, (rendered_path, recall)
        return recall
    if rendered_path.suffix == ".png":
        with Image.open(rendered_path) as image:
            assert image.width >= 1000 and image.height >= 1600, rendered_path
            contrast = ImageStat.Stat(image.convert("L")).stddev[0]
            assert contrast >= 5.0, (rendered_path, contrast)
        # Pixel fixtures need OCR to compare text, so exact source recall is
        # intentionally not claimed here.
        return None
    raise AssertionError(f"Unsupported CV fixture: {rendered_path}")


def main() -> None:
    reports: dict[str, Any] = {}
    split_cases: dict[str, list[dict[str, Any]]] = {}
    split_annotations: dict[str, list[dict[str, Any]]] = {}

    for split, expected in EXPECTED.items():
        root = ROOT / split
        cases = _load_jsonl(root / "cases.jsonl")
        annotations = _load_jsonl(root / "annotations.jsonl")
        split_cases[split] = cases
        split_annotations[split] = annotations
        assert len(cases) == expected["cases"]
        assert len(annotations) == expected["cases"]
        assert len({case["id"] for case in cases}) == len(cases)
        assert {case["id"] for case in cases} == {item["case_id"] for item in annotations}
        annotation_by_id = {item["case_id"]: item for item in annotations}
        counts = Counter(case["scenario"] for case in cases)
        assert counts == Counter({f"scenario{i}": expected["per_scenario"] for i in range(1, 5)})

        for case in cases:
            for key in ("cv_path", "target_jd_file_path", "cv_template_path"):
                value = case.get(key)
                if value:
                    path = ROOT.parents[1] / "acceptance_testset" / value
                    assert path.resolve().exists(), (case["id"], path)
            if case["scenario"] == "scenario4":
                assert case["cv_path"] is None
                assert case["provenance"]["resume_id"] is None
                assert case["expected_output"]["framework_only"] is True
                assert "没有提供任何个人信息" in case["query"]
            if case["scenario"] == "scenario2":
                assert case["cv_path"] is None
                assert case["target_jd"] is None
                assert case["target_jd_file_path"] is None

        for annotation in annotations:
            _verify_sources(annotation)

        rendered_recalls = [
            recall
            for case in cases
            if (recall := _verify_rendered_cv(case, annotation_by_id[case["id"]])) is not None
        ]

        manifest = json.loads((root / "manifest.json").read_text(encoding="utf-8"))
        for entry in manifest["files"]:
            path = root / entry["path"]
            assert path.exists(), path
            assert path.stat().st_size == entry["bytes"], path
            assert _sha(path) == entry["sha256"], path

        reports[split] = {
            "case_count": len(cases),
            "scenario_counts": dict(sorted(counts.items())),
            "industry_counts": dict(sorted(Counter(case["industry"] for case in cases).items())),
            "input_profile_counts": dict(sorted(Counter(case["input_profile"] for case in cases).items())),
            "exact_span_annotations": sum(
                len(source["units"])
                for annotation in annotations
                for source in annotation["sources"]
            ),
            "rendered_text_recall_min": min(rendered_recalls, default=None),
            "rendered_text_verified_files": len(rendered_recalls),
        }

    holdout_resumes = {
        case["provenance"]["resume_id"] for case in split_cases["holdout_v2"]
        if case["provenance"]["resume_id"]
    }
    shadow_resumes = {
        case["provenance"]["resume_id"] for case in split_cases["shadow_v3"]
        if case["provenance"]["resume_id"]
    }
    holdout_jds = {
        case["provenance"]["jd_id"] for case in split_cases["holdout_v2"]
        if case["provenance"]["jd_id"]
    }
    shadow_jds = {
        case["provenance"]["jd_id"] for case in split_cases["shadow_v3"]
        if case["provenance"]["jd_id"]
    }
    assert holdout_resumes.isdisjoint(shadow_resumes)
    assert holdout_jds.isdisjoint(shadow_jds)

    output = {
        "ok": True,
        "reports": reports,
        "cross_split": {
            "resume_overlap": 0,
            "jd_overlap": 0,
            "holdout_unique_resume_roots": len(holdout_resumes),
            "shadow_unique_resume_roots": len(shadow_resumes),
        },
    }
    (ROOT / "verification_report.json").write_text(
        json.dumps(output, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    print(json.dumps(output, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
