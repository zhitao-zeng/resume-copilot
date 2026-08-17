"""Binary input adapters for the production V3 compiler.

The adapters preserve transport structure and exact offsets.  They do not
classify occupations or repair semantic mistakes: DOCX hierarchy, PDF
coordinates and PP-Structure metadata are evidence, while section/fact meaning
is assigned later through the frozen semantic schema.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
import os
from pathlib import Path
from typing import Any, Iterable

from experimental_model_candidates import OCR_LAYOUT_REGION_SEPARATOR

from .contracts import DocumentGraph, LayoutNode, SourceAsset, SourceSpan
from .document_graph import _line_kind, from_native_text, from_ppstructure_blocks
from .ocr_numeric_witness import witness_ppstructure_blocks


@dataclass
class _NodeDraft:
    text: str
    kind: str = "paragraph"
    page: int = 1
    bbox: tuple[float, float, float, float] | None = None
    order: int = 0
    confidence: float = 1.0
    label: str = ""
    column_id: str | None = None
    region_id: str | None = None
    parent_id: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


def _graph_from_drafts(
    asset: SourceAsset,
    drafts: Iterable[_NodeDraft],
    *,
    engine: str,
    metadata: dict[str, Any] | None = None,
) -> DocumentGraph:
    cleaned = [draft for draft in drafts if str(draft.text or "").strip()]
    source_text = "\n".join(str(draft.text).strip() for draft in cleaned)
    nodes: list[LayoutNode] = []
    cursor = 0
    for index, draft in enumerate(cleaned):
        text = str(draft.text).strip()
        start, end = cursor, cursor + len(text)
        cursor = end + (1 if index < len(cleaned) - 1 else 0)
        node_id = f"{asset.source_id}:{engine}:{index}"
        nodes.append(LayoutNode(
            node_id=node_id,
            source_id=asset.source_id,
            kind=draft.kind,  # type: ignore[arg-type]
            text=text,
            page=max(1, draft.page),
            bbox=draft.bbox,
            order=draft.order if draft.order >= 0 else index,
            confidence=max(0.0, min(1.0, draft.confidence)),
            label=draft.label,
            column_id=draft.column_id,
            region_id=draft.region_id,
            parent_id=draft.parent_id,
            source_spans=[SourceSpan(
                source_id=asset.source_id,
                char_start=start,
                char_end=end,
                page=max(1, draft.page),
                node_id=node_id,
            )],
            metadata=draft.metadata,
        ))
    return DocumentGraph(
        source_id=asset.source_id,
        source_type=asset.source_type,
        extraction_engine=engine,  # type: ignore[arg-type]
        source_text=source_text,
        nodes=nodes,
        metadata=metadata or {},
    )


def _docx_graph(content: bytes, asset: SourceAsset) -> DocumentGraph:
    from docx import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    document = DocxDocument(BytesIO(content))
    drafts: list[_NodeDraft] = []

    def add_paragraph(paragraph: Paragraph, *, parent_id: str | None = None, metadata: dict[str, Any] | None = None) -> None:
        text = paragraph.text.strip()
        if not text:
            return
        style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "")
        kind = "heading" if style_name.casefold().startswith("heading") or _line_kind(text) == "heading" else "paragraph"
        payload = dict(metadata or {})
        payload.update({"style_name": style_name, "xml_kind": "paragraph"})
        drafts.append(_NodeDraft(
            text=text,
            kind=kind,
            order=len(drafts),
            parent_id=parent_id,
            metadata=payload,
        ))

    table_index = 0
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            add_paragraph(Paragraph(child, document))
            continue
        if not isinstance(child, CT_Tbl):
            continue
        table = Table(child, document)
        table_id = f"table:{table_index}"
        table_index += 1
        seen_cells: set[int] = set()
        for row_index, row in enumerate(table.rows):
            row_id = f"{table_id}:row:{row_index}"
            for cell_index, cell in enumerate(row.cells):
                cell_identity = id(cell._tc)
                if cell_identity in seen_cells:
                    continue
                seen_cells.add(cell_identity)
                cell_id = f"{row_id}:cell:{cell_index}"
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    add_paragraph(
                        paragraph,
                        parent_id=row_id,
                        metadata={
                            "table_id": table_id,
                            "table_row_id": row_id,
                            "row_id": row_id,
                            "cell_id": cell_id,
                            "row_index": row_index,
                            "cell_index": cell_index,
                            "paragraph_index": paragraph_index,
                        },
                    )

    # Contact information is often stored in a header.  Read each distinct
    # header/footer part once and keep its structural kind; no template text is
    # involved because this adapter receives the candidate CV asset only.
    seen_parts: set[str] = set()
    for section_index, section in enumerate(document.sections):
        for kind, part in (("header", section.header), ("footer", section.footer)):
            part_name = str(getattr(getattr(part, "part", None), "partname", ""))
            if part_name in seen_parts:
                continue
            seen_parts.add(part_name)
            for paragraph_index, paragraph in enumerate(part.paragraphs):
                text = paragraph.text.strip()
                if not text:
                    continue
                drafts.append(_NodeDraft(
                    text=text,
                    kind=kind,
                    order=len(drafts),
                    parent_id=f"{kind}:{section_index}",
                    metadata={
                        "xml_kind": kind,
                        "section_index": section_index,
                        "paragraph_index": paragraph_index,
                        "part_name": part_name,
                    },
                ))
    return _graph_from_drafts(
        asset,
        drafts,
        engine="native_docx",
        metadata={"native": True, "paragraph_table_order": True, "style_hierarchy": True},
    )


def _pdf_native_drafts(page: Any, page_number: int) -> list[_NodeDraft]:
    payload = page.get_text("dict", sort=True)
    drafts: list[_NodeDraft] = []
    for block_index, block in enumerate(payload.get("blocks", [])):
        if not isinstance(block, dict) or block.get("type", 0) != 0:
            continue
        block_bbox = block.get("bbox")
        for line_index, line in enumerate(block.get("lines", [])):
            spans = line.get("spans", []) if isinstance(line, dict) else []
            text = "".join(str(span.get("text") or "") for span in spans if isinstance(span, dict)).strip()
            if not text:
                continue
            raw_bbox = line.get("bbox") or block_bbox
            bbox = tuple(float(value) for value in raw_bbox) if isinstance(raw_bbox, (list, tuple)) and len(raw_bbox) == 4 else None
            max_size = max((float(span.get("size") or 0.0) for span in spans if isinstance(span, dict)), default=0.0)
            drafts.append(_NodeDraft(
                text=text,
                kind=_line_kind(text),
                page=page_number,
                bbox=bbox,  # type: ignore[arg-type]
                order=len(drafts),
                label="native_pdf_line",
                region_id=f"p{page_number}:b{block_index}",
                parent_id=f"p{page_number}:b{block_index}",
                metadata={
                    "block_index": block_index,
                    "line_index": line_index,
                    "font_size_max": max_size,
                    "native_spans": spans,
                },
            ))
    return drafts


def _ppstructure_drafts(blocks: Iterable[dict[str, Any]], *, page_number: int) -> list[_NodeDraft]:
    drafts: list[_NodeDraft] = []
    for index, row in enumerate(blocks):
        text = str(row.get("block_content") or row.get("text") or "").strip()
        if not text:
            continue
        raw_bbox = row.get("block_bbox", row.get("bbox"))
        bbox = tuple(float(value) for value in raw_bbox) if isinstance(raw_bbox, (list, tuple)) and len(raw_bbox) == 4 else None
        confidence = row.get("confidence", row.get("score", 1.0))
        try:
            confidence_value = float(confidence)
        except (TypeError, ValueError):
            confidence_value = 1.0
        label = str(row.get("block_label", row.get("label", "")) or "")
        drafts.append(_NodeDraft(
            text=text,
            kind="heading" if label.casefold() in {"title", "heading", "doc_title", "paragraph_title"} or _line_kind(text) == "heading" else "region",
            page=page_number,
            bbox=bbox,  # type: ignore[arg-type]
            order=int(row.get("block_order", index)) if str(row.get("block_order", index)).lstrip("-").isdigit() else index,
            confidence=confidence_value,
            label=label,
            column_id=str(row.get("column_id")) if row.get("column_id") is not None else None,
            region_id=str(row.get("region_id", row.get("block_id"))) if row.get("region_id", row.get("block_id")) is not None else None,
            parent_id=str(row.get("parent_id")) if row.get("parent_id") is not None else None,
            metadata=dict(row),
        ))
    return drafts


def _pdf_graph(content: bytes, asset: SourceAsset) -> DocumentGraph:
    try:
        import fitz
    except ImportError as exc:  # pragma: no cover - production dependency
        raise RuntimeError("PyMuPDF is unavailable") from exc
    from ppstructure_runtime import extract_ppstructure_blocks

    native_min_chars = max(1, int(os.getenv("V3_PDF_NATIVE_MIN_CHARS", "40")))
    drafts: list[_NodeDraft] = []
    engines: list[str] = []
    document = fitz.open(stream=content, filetype="pdf")
    try:
        for page_index, page in enumerate(document, start=1):
            native = _pdf_native_drafts(page, page_index)
            native_chars = sum(len(item.text.replace(" ", "")) for item in native)
            if native_chars >= native_min_chars:
                drafts.extend(native)
                engines.append("native_pdf")
                continue
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
            png = pixmap.tobytes("png")
            blocks = extract_ppstructure_blocks(png, filename=f"page-{page_index}.png")
            blocks = witness_ppstructure_blocks(blocks, png)
            page_drafts = _ppstructure_drafts(blocks, page_number=page_index)
            if page_drafts:
                drafts.extend(page_drafts)
                engines.append("ppstructure")
            else:
                drafts.extend(native)
                engines.append("native_pdf_empty")
    finally:
        document.close()
    engine = "ppstructure" if "ppstructure" in engines else "native_pdf"
    return _graph_from_drafts(
        asset,
        drafts,
        engine=engine,
        metadata={"native": engine == "native_pdf", "page_engines": engines, "mixed_pdf": len(set(engines)) > 1},
    )


def build_input_document_graph(
    content: bytes,
    *,
    filename: str,
    fallback_text: str = "",
    source_id: str = "cv",
    source_type: str = "cv",
) -> DocumentGraph:
    """Build the strongest available graph and preserve a safe text fallback."""

    suffix = Path(filename or "").suffix.casefold()
    media_types = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".pdf": "application/pdf",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".markdown": "text/markdown",
    }
    media_type = media_types.get(suffix, "image/*" if suffix in {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".gif", ".tif", ".tiff"} else "application/octet-stream")
    asset = SourceAsset(
        source_id=source_id,
        source_type=source_type,  # type: ignore[arg-type]
        filename=filename,
        media_type=media_type,
        native=suffix in {".docx", ".pdf", ".txt", ".md", ".markdown"},
    )
    try:
        if suffix == ".docx":
            return _docx_graph(content, asset)
        if suffix == ".pdf":
            return _pdf_graph(content, asset)
        if suffix in {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".gif", ".tif", ".tiff"}:
            layout_engine = os.getenv(
                "LAYOUT_ORDER_ENGINE", "bbox",
            ).strip().casefold()
            if layout_engine != "ppstructure" and fallback_text.strip():
                if layout_engine == "ppstructure_hybrid":
                    group_index = 0
                    drafts: list[_NodeDraft] = []
                    for line in fallback_text.splitlines():
                        if line.strip() == OCR_LAYOUT_REGION_SEPARATOR:
                            group_index += 1
                            continue
                        text = line.strip()
                        if not text:
                            continue
                        drafts.append(_NodeDraft(
                            text=text,
                            kind=_line_kind(text),
                            order=len(drafts),
                            region_id=f"hybrid-region:{group_index}",
                            metadata={
                                "layout_source": "PP-StructureV3",
                                "text_source": "PP-OCRv6",
                                "hybrid_region_index": group_index,
                            },
                        ))
                    graph = _graph_from_drafts(
                        asset.model_copy(update={"native": False}),
                        drafts,
                        engine="ocr",
                    )
                else:
                    graph = from_native_text(
                        asset.model_copy(update={
                            "text": fallback_text,
                            "native": False,
                        }),
                        fallback_text,
                        engine="ocr",
                    )
                graph.metadata.update({
                    "layout_source": (
                        "PP-StructureV3 geometry + PP-OCRv6 text"
                        if layout_engine == "ppstructure_hybrid"
                        else "PP-OCRv6+BBOX"
                    ),
                    "hybrid_layout": layout_engine == "ppstructure_hybrid",
                    "text_source": "PP-OCRv6",
                })
                return graph
            if layout_engine != "ppstructure":
                raise ValueError(
                    "raster BBOX/hybrid route requires extracted PP-OCRv6 text"
                )
            from ppstructure_runtime import extract_ppstructure_blocks

            blocks = extract_ppstructure_blocks(content, filename=filename)
            blocks = witness_ppstructure_blocks(blocks, content)
            return from_ppstructure_blocks(asset, blocks)
        if suffix in {".txt", ".md", ".markdown"}:
            decoded = content.decode("utf-8", errors="replace")
            return from_native_text(asset.model_copy(update={"text": decoded}), decoded)
    except Exception as exc:
        if not fallback_text.strip():
            raise
        graph = from_native_text(
            asset.model_copy(update={"text": fallback_text, "native": False}),
            fallback_text,
            engine="ocr" if media_type.startswith("image/") or suffix == ".pdf" else "text",
        )
        graph.metadata.update({"adapter_fallback": True, "adapter_error": type(exc).__name__})
        return graph
    if fallback_text.strip():
        return from_native_text(asset.model_copy(update={"text": fallback_text}), fallback_text)
    raise ValueError(f"unsupported or empty V3 candidate asset: {filename}")


__all__ = ["build_input_document_graph"]
