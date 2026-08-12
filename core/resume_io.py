import io
import math
import multiprocessing as mp
import os
import re
import signal
import threading
import time
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Optional
from xml.etree import ElementTree as ET

# Native math libraries otherwise see the host's CPU count rather than the 2C
# cgroup quota and may try to create hundreds of threads during a spawn import.
for _native_thread_env in (
    "OPENBLAS_NUM_THREADS",
    "OMP_NUM_THREADS",
    "MKL_NUM_THREADS",
    "NUMEXPR_NUM_THREADS",
):
    os.environ.setdefault(_native_thread_env, "1")

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from http_compat import HTTPException

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import pytesseract
except ImportError:
    pytesseract = None

# RapidOCR. PP-OCRv6 small is the preferred document model on both CPU and
# CUDA. The legacy v5 mobile/server bundles remain bounded fallbacks for images
# built before the v6 model artifact was added.
_RAPID_OCR = None
_RAPID_OCR_PATH = None
_RAPID_OCR_INITED = False
_RAPID_OCR_VERSION = "v4"
_RAPID_OCR_MODEL_TYPE = "unknown"
_RAPID_OCR_PROVIDER = "unavailable"
_RAPID_OCR_INIT_LOCK = threading.Lock()
_RAPID_OCR_RUN_LOCK = threading.Lock()


def _positive_int_env(name: str, default: int) -> int:
    try:
        value = int(os.getenv(name, str(default)))
    except (TypeError, ValueError):
        return default
    return value if value > 0 else default


def _positive_float_env(name: str, default: float) -> float:
    try:
        value = float(os.getenv(name, str(default)))
    except (TypeError, ValueError):
        return default
    return value if value > 0 else default


def _cgroup_cpu_limit() -> Optional[int]:
    """Read cgroup v2/v1 CPU quota as a conservative whole-core count."""
    try:
        quota_text = Path("/sys/fs/cgroup/cpu.max").read_text().strip().split()
        if len(quota_text) == 2 and quota_text[0] != "max":
            quota, period = int(quota_text[0]), int(quota_text[1])
            if quota > 0 and period > 0:
                return max(1, quota // period)
    except (OSError, ValueError):
        pass

    for root in (Path("/sys/fs/cgroup/cpu"), Path("/sys/fs/cgroup")):
        try:
            quota = int((root / "cpu.cfs_quota_us").read_text().strip())
            period = int((root / "cpu.cfs_period_us").read_text().strip())
            if quota > 0 and period > 0:
                return max(1, quota // period)
        except (OSError, ValueError):
            continue
    return None


def _available_cpu_count() -> int:
    """Return the lower of CPU affinity and the container's CPU quota."""
    try:
        affinity_count = max(1, len(os.sched_getaffinity(0)))
    except (AttributeError, OSError):
        affinity_count = max(1, os.cpu_count() or 1)
    quota_count = _cgroup_cpu_limit()
    return min(affinity_count, quota_count) if quota_count else affinity_count


def _rapid_ocr_cuda_available() -> bool:
    # CPU/mobile is the safe default because this service normally shares its
    # GPU with vLLM.  Set RAPID_OCR_DEVICE=cuda (or auto) for a deployment with
    # enough dedicated GPU headroom.
    device = os.getenv("RAPID_OCR_DEVICE", "cpu").strip().lower()
    if device in {"cpu", "off", "0", "false"}:
        return False
    try:
        import onnxruntime as ort

        # ONNX Runtime installed next to PyTorch may need its CUDA/cuDNN wheels
        # preloaded from site-packages before a CUDA session can be created.
        preload_dlls = getattr(ort, "preload_dlls", None)
        if callable(preload_dlls):
            try:
                preload_dlls()
            except Exception as exc:
                logger.warning("RapidOCR could not preload CUDA libraries: %s", exc)
        return (
            ort.get_device() == "GPU"
            and "CUDAExecutionProvider" in ort.get_available_providers()
        )
    except Exception:
        return False


def _rapid_ocr_model_bundles(model_dir: Path) -> dict[tuple[str, str], dict[str, Any]]:
    return {
        ("v6", "small"): {
            "det": model_dir / "det.onnx",
            "cls": model_dir / "cls.onnx",
            "rec": model_dir / "rec.onnx",
            "keys": model_dir / "keys.txt",
            "cls_version": "v4",
            "cls_model_type": "mobile",
        },
        ("v5", "server"): {
            "det": model_dir / "ch_PP-OCRv5_det_server.onnx",
            "cls": model_dir / "ch_PP-LCNet_x1_0_textline_ori_cls_server.onnx",
            "rec": model_dir / "ch_PP-OCRv5_rec_server.onnx",
            "cls_image_shape": [3, 80, 160],
        },
        ("v5", "mobile"): {
            "det": model_dir / "ch_PP-OCRv5_det_mobile.onnx",
            "cls": model_dir / "ch_PP-LCNet_x0_25_textline_ori_cls_mobile.onnx",
            "rec": model_dir / "ch_PP-OCRv5_rec_mobile.onnx",
            "cls_image_shape": [3, 80, 160],
        },
        ("v4", "server"): {
            "det": model_dir / "ch_PP-OCRv4_det_server.onnx",
            "cls": model_dir / "ch_ppocr_mobile_v2.0_cls_mobile.onnx",
            "rec": model_dir / "ch_PP-OCRv4_rec_server.onnx",
        },
        ("v4", "mobile"): {
            "det": model_dir / "ch_PP-OCRv4_det_mobile.onnx",
            "cls": model_dir / "ch_ppocr_mobile_v2.0_cls_mobile.onnx",
            "rec": model_dir / "ch_PP-OCRv4_rec_mobile.onnx",
        },
    }


def _rapid_ocr_candidate_dirs() -> list[Path]:
    configured_v6 = os.getenv("PPOCRV6_MODEL_DIR", "").strip()
    candidates = [
        Path(__file__).parent / "models" / "ppocrv6-small-ort",
        Path(__file__).parent.parent / "models" / "ppocrv6-small-ort",
        Path("/mounted_model/ppocrv6-small-ort"),
        Path("/root/app/models/ppocrv6-small-ort"),
        Path(__file__).parent / "models" / "rapidocr_multilang",
        Path(__file__).parent.parent / "models" / "rapidocr_multilang",
        Path("/mounted_model/rapidocr_multilang"),
        Path("/root/app/models/rapidocr_multilang"),
    ]
    if configured_v6:
        candidates.insert(0, Path(configured_v6))
    return list(dict.fromkeys(candidates))


def _select_rapid_ocr_bundle(
    candidates: list[Path],
    *,
    prefer_server: bool,
    forced_model_type: str = "auto",
) -> Optional[dict[str, Any]]:
    """Select an explicit model bundle without reconstructing filenames later."""
    if forced_model_type not in {"auto", "small", "server", "mobile"}:
        forced_model_type = "auto"
    if forced_model_type == "auto":
        type_order = ["small", "server", "mobile"] if prefer_server else ["small", "mobile", "server"]
    else:
        type_order = [forced_model_type]

    for model_type in type_order:
        for version in ("v6", "v5", "v4"):
            for model_dir in candidates:
                files = _rapid_ocr_model_bundles(model_dir).get((version, model_type))
                if files is None:
                    continue
                model_paths = [value for value in files.values() if isinstance(value, Path)]
                try:
                    bundle_exists = bool(model_paths) and all(
                        path.is_file() for path in model_paths
                    )
                except OSError:
                    # A protected system path is just a non-candidate. Continue
                    # probing later configured/local bundles instead of aborting
                    # OCR initialization with PermissionError.
                    bundle_exists = False
                if bundle_exists:
                    return {
                        "model_dir": model_dir,
                        "version": version,
                        "model_type": model_type,
                        **files,
                    }
    return None


def _rapid_ocr_session_providers(ocr_instance: Any) -> list[str]:
    providers: list[str] = []
    for component_name in ("text_det", "text_cls", "text_rec"):
        component = getattr(ocr_instance, component_name, None)
        wrapper = getattr(component, "session", None)
        session = getattr(wrapper, "session", wrapper)
        get_providers = getattr(session, "get_providers", None)
        if callable(get_providers):
            try:
                for provider in get_providers() or []:
                    if provider not in providers:
                        providers.append(str(provider))
            except Exception:
                continue
    return providers


def _build_rapid_ocr(bundle: dict[str, Any], *, use_cuda: bool, cpu_threads: int):
    from rapidocr import RapidOCR
    from rapidocr.utils.typings import ModelType as _ModelType
    from rapidocr.utils.typings import OCRVersion as _OCRVersion

    version_map = {
        "v4": _OCRVersion.PPOCRV4,
        "v5": _OCRVersion.PPOCRV5,
        "v6": getattr(_OCRVersion, "PPOCRV6", None),
    }
    ocr_version = version_map.get(str(bundle["version"]))
    if ocr_version is None:
        raise RuntimeError(
            "PP-OCRv6 requires rapidocr>=3.9.1; installed RapidOCR has no PPOCRV6 support"
        )
    cls_version = version_map.get(str(bundle.get("cls_version", bundle["version"])))
    if cls_version is None:
        cls_version = _OCRVersion.PPOCRV4
    model_type_map = {
        "small": _ModelType.SMALL,
        "server": _ModelType.SERVER,
        "mobile": _ModelType.MOBILE,
    }
    model_type = model_type_map[str(bundle["model_type"])]
    cls_model_type = model_type_map[str(bundle.get("cls_model_type", bundle["model_type"]))]
    gpu_mem_limit = _positive_int_env("RAPID_OCR_GPU_MEM_LIMIT_MB", 4096) * 1024 * 1024
    params = {
        "Det.model_path": str(bundle["det"]),
        "Cls.model_path": str(bundle["cls"]),
        "Rec.model_path": str(bundle["rec"]),
        "Det.ocr_version": ocr_version,
        "Det.model_type": model_type,
        "Cls.ocr_version": cls_version,
        "Cls.model_type": cls_model_type,
        "Rec.ocr_version": ocr_version,
        "Rec.model_type": model_type,
        "Global.max_side_len": _positive_int_env("RAPID_OCR_ENGINE_MAX_SIDE", 1600),
        "EngineConfig.onnxruntime.intra_op_num_threads": cpu_threads,
        "EngineConfig.onnxruntime.inter_op_num_threads": 1,
        "EngineConfig.onnxruntime.use_cuda": use_cuda,
        "EngineConfig.onnxruntime.cuda_ep_cfg.gpu_mem_limit": gpu_mem_limit,
    }
    if bundle.get("keys"):
        params["Rec.rec_keys_path"] = str(bundle["keys"])
    if bundle.get("cls_image_shape"):
        params["Cls.cls_image_shape"] = list(bundle["cls_image_shape"])
    return RapidOCR(params=params)


def _switch_rapid_ocr_to_cpu_mobile(reason: str) -> bool:
    """Replace a failed CUDA engine with the bounded CPU v6-small path."""
    global _RAPID_OCR, _RAPID_OCR_PATH, _RAPID_OCR_VERSION
    global _RAPID_OCR_MODEL_TYPE, _RAPID_OCR_PROVIDER
    with _RAPID_OCR_INIT_LOCK:
        try:
            bundle = _select_rapid_ocr_bundle(
                _rapid_ocr_candidate_dirs(),
                prefer_server=False,
                forced_model_type="auto",
            )
            if bundle is None:
                return False
            cpu_threads = min(
                _positive_int_env("RAPID_OCR_CPU_THREADS", 4),
                _available_cpu_count(),
            )
            replacement = _build_rapid_ocr(
                bundle,
                use_cuda=False,
                cpu_threads=cpu_threads,
            )
            providers = _rapid_ocr_session_providers(replacement)
            with _RAPID_OCR_RUN_LOCK:
                _RAPID_OCR = replacement
                _RAPID_OCR_PATH = str(bundle["model_dir"])
                _RAPID_OCR_VERSION = str(bundle["version"])
                _RAPID_OCR_MODEL_TYPE = str(bundle["model_type"])
                _RAPID_OCR_PROVIDER = providers[0] if providers else "unknown"
            logger.warning(
                "RapidOCR switched to CPU %s after CUDA failure | reason=%s provider=%s",
                bundle["model_type"],
                reason,
                _RAPID_OCR_PROVIDER,
            )
            return True
        except Exception as exc:
            logger.warning("RapidOCR CPU failover could not initialize: %s", exc)
            return False

def _init_rapid_ocr():
    global _RAPID_OCR, _RAPID_OCR_PATH, _RAPID_OCR_INITED, _RAPID_OCR_VERSION
    global _RAPID_OCR_MODEL_TYPE, _RAPID_OCR_PROVIDER
    if _RAPID_OCR_INITED:
        return
    with _RAPID_OCR_INIT_LOCK:
        if _RAPID_OCR_INITED:
            return
        started = time.perf_counter()
        try:
            candidates = _rapid_ocr_candidate_dirs()
            use_cuda = _rapid_ocr_cuda_available()
            requested_device = os.getenv("RAPID_OCR_DEVICE", "cpu").strip().lower()
            if requested_device in {"cuda", "gpu"} and not use_cuda:
                logger.warning("RapidOCR CUDA requested but CUDAExecutionProvider is unavailable; using CPU")

            forced_model = os.getenv("RAPID_OCR_MODEL", "auto").strip().lower()
            bundle = _select_rapid_ocr_bundle(
                candidates,
                prefer_server=use_cuda,
                forced_model_type=forced_model,
            )
            if bundle is None:
                logger.warning("RapidOCR model files not found, falling back to pytesseract")
                return

            cpu_threads = min(
                _positive_int_env("RAPID_OCR_CPU_THREADS", 4),
                _available_cpu_count(),
            )
            ocr_instance = _build_rapid_ocr(bundle, use_cuda=use_cuda, cpu_threads=cpu_threads)
            providers = _rapid_ocr_session_providers(ocr_instance)

            # A registered CUDA provider can still fail to create a CUDA session
            # (for example due to a CUDA/cuDNN mismatch).  In auto mode, fall back
            # to the bounded CPU path instead of silently running a CUDA-selected
            # model through the CPU provider.
            if (
                use_cuda
                and "CUDAExecutionProvider" not in providers
                and forced_model == "auto"
            ):
                fallback_bundle = _select_rapid_ocr_bundle(
                    candidates,
                    prefer_server=False,
                    forced_model_type="auto",
                )
                if fallback_bundle is not None:
                    logger.warning(
                        "RapidOCR CUDA session unavailable; switching to CPU %s model",
                        fallback_bundle["model_type"],
                    )
                    bundle = fallback_bundle
                    use_cuda = False
                    ocr_instance = _build_rapid_ocr(
                        bundle,
                        use_cuda=False,
                        cpu_threads=cpu_threads,
                    )
                    providers = _rapid_ocr_session_providers(ocr_instance)

            _RAPID_OCR = ocr_instance
            _RAPID_OCR_PATH = str(bundle["model_dir"])
            _RAPID_OCR_VERSION = str(bundle["version"])
            _RAPID_OCR_MODEL_TYPE = str(bundle["model_type"])
            _RAPID_OCR_PROVIDER = providers[0] if providers else "unknown"
            logger.info(
                "RapidOCR initialized | version=%s model=%s provider=%s cpu_threads=%s init_s=%.3f path=%s",
                _RAPID_OCR_VERSION,
                _RAPID_OCR_MODEL_TYPE,
                _RAPID_OCR_PROVIDER,
                cpu_threads,
                time.perf_counter() - started,
                _RAPID_OCR_PATH,
            )
        except Exception as exc:
            logger.warning("RapidOCR init failed: %s, falling back to pytesseract", exc)
        finally:
            _RAPID_OCR_INITED = True


def _prepare_ocr_image(content: bytes):
    """Decode, orient, normalize and bound an image before any OCR engine sees it."""
    from PIL import ImageOps

    with Image.open(io.BytesIO(content)) as source:
        pil_img = ImageOps.exif_transpose(source)
        if pil_img.mode != "RGB":
            pil_img = pil_img.convert("RGB")
        else:
            pil_img = pil_img.copy()

    original_size = pil_img.size
    width, height = original_size
    max_pixels = _positive_int_env("RAPID_OCR_MAX_PIXELS", 6_000_000)
    max_long_edge = _positive_int_env("RAPID_OCR_MAX_LONG_EDGE", 3000)
    scale = min(
        1.0,
        max_long_edge / max(width, height),
        math.sqrt(max_pixels / max(1, width * height)),
    )
    if scale < 1.0:
        new_size = (
            max(1, int(round(width * scale))),
            max(1, int(round(height * scale))),
        )
        resampling = getattr(Image, "Resampling", Image).LANCZOS
        pil_img = pil_img.resize(new_size, resampling)
    return pil_img, original_size


def _rapid_result_to_text(result: Any, *, img_width: int, img_height: int) -> str:
    boxes = result.boxes if hasattr(result, "boxes") else (result[0] if isinstance(result, (tuple, list)) else None)
    txts = result.txts if hasattr(result, "txts") else (result[1] if isinstance(result, (tuple, list)) and len(result) > 1 else None)
    if boxes is None or len(boxes) == 0:
        return ""
    ordered_texts = _reconstruct_ocr_reading_order(
        boxes,
        txts,
        img_width=img_width,
        img_height=img_height,
    )
    return "\n".join(ordered_texts) if ordered_texts else ""


def _run_rapid_ocr(np_img):
    # RapidOCR/ONNX sessions are shared globally.  Serialize calls because the
    # service may later raise its task-worker count and not every backend is
    # guaranteed to be re-entrant.
    with _RAPID_OCR_RUN_LOCK:
        return _RAPID_OCR(np_img)


def _ocr_image_with_rapid(content: bytes, *, prepared_image=None) -> str:
    """Extract text from image using global RapidOCR (single engine call).

    PP-OCRv6 was trained for RGB document input; collapsing it to one channel
    can erase colored section headings even when body text remains readable.
    Legacy bundles retain the established green-channel transform for colored
    sidebars.  The shared engine is called exactly once in either case.
    """
    import numpy as np  # noqa: F811
    from PIL import ImageOps

    pil_img = prepared_image if prepared_image is not None else _prepare_ocr_image(content)[0]

    arr = np.array(pil_img)
    if _RAPID_OCR_VERSION == "v6":
        np_img = arr
    else:
        # Legacy recognition bundles benefited from this transform on white
        # text over a red/blue sidebar.
        green = arr[:, :, 1]
        green_img = Image.fromarray(green).convert("L")
        enhanced = ImageOps.autocontrast(green_img, cutoff=3)
        np_img = np.array(enhanced.convert("RGB"))

    try:
        result = _run_rapid_ocr(np_img)
    except Exception as exc:
        logger.warning("RapidOCR inference failed, falling back: %s", exc)
        if _RAPID_OCR_PROVIDER == "CUDAExecutionProvider" and _switch_rapid_ocr_to_cpu_mobile(
            type(exc).__name__
        ):
            try:
                result = _run_rapid_ocr(np_img)
            except Exception as retry_exc:
                logger.warning("RapidOCR CPU mobile retry failed: %s", retry_exc)
                return ""
        else:
            return ""
    return _rapid_result_to_text(
        result,
        img_width=np_img.shape[1],
        img_height=np_img.shape[0],
    )



# ── OCR Layout Reconstruction ─────────────────────────────────────────────────


def _reconstruct_ocr_reading_order(
    boxes,
    texts,
    *,
    img_width: int,
    img_height: int,
) -> list[str]:
    """Reconstruct deterministic reading order from OCR line polygons.

    Resume pages mix several layout grammars: a persistent sidebar, balanced
    columns, right-aligned date annotations and ordinary key/value rows.  A
    single x-coordinate threshold cannot distinguish them.  Use adaptive
    recursive XY cuts instead: large persistent gutters form columns, smaller
    horizontal gaps form bands, and ambiguous annotation gutters fall back to
    row-major order.  Every source block is returned exactly once.
    """
    import numpy as np  # noqa: F811

    if boxes is None or len(boxes) == 0:
        return []

    # ── 1. Build block list with structured bbox fields ──
    raw_blocks: list[dict] = []
    for i in range(len(boxes)):
        box = boxes[i]
        confidence = 1.0
        if texts and i < len(texts) and texts[i]:
            raw_val = str(texts[i][0]) if isinstance(texts[i], (tuple, list)) else str(texts[i])
            if isinstance(texts[i], (tuple, list)) and len(texts[i]) > 1:
                try:
                    confidence = float(texts[i][1])
                except (TypeError, ValueError):
                    confidence = 1.0
        else:
            raw_val = str(texts[i]) if texts and i < len(texts) and texts[i] is not None else ""
        text = raw_val.strip()
        if not text:
            continue

        pts = np.array(box)
        x_min, x_max = float(pts[:, 0].min()), float(pts[:, 0].max())
        y_min, y_max = float(pts[:, 1].min()), float(pts[:, 1].max())

        raw_blocks.append({
            "text": text,
            "x_min": x_min,
            "x_max": x_max,
            "y_min": y_min,
            "y_max": y_max,
            "x_center": (x_min + x_max) / 2.0,
            "y_center": (y_min + y_max) / 2.0,
            "width": x_max - x_min,
            "height": y_max - y_min,
            "confidence": confidence,
            "source_index": i,
        })

    if not raw_blocks:
        return []

    median_height = float(np.median([block["height"] for block in raw_blocks]))

    def _row_overlap(left: dict, right: dict) -> float:
        overlap = max(
            0.0,
            min(left["y_max"], right["y_max"])
            - max(left["y_min"], right["y_min"]),
        )
        return overlap / max(1.0, min(left["height"], right["height"]))

    def _row_major(blocks: list[dict]) -> list[dict]:
        ordered = sorted(
            blocks,
            key=lambda block: (
                block["y_center"], block["x_min"], block["source_index"]
            ),
        )
        rows: list[list[dict]] = []
        for block in ordered:
            if not rows:
                rows.append([block])
                continue
            current = rows[-1]
            row_center = sum(item["y_center"] for item in current) / len(current)
            row_height = max(item["height"] for item in current)
            same_row = any(_row_overlap(block, item) >= 0.35 for item in current)
            same_row = same_row or abs(block["y_center"] - row_center) <= max(
                2.0, 0.32 * max(row_height, block["height"])
            )
            if same_row:
                current.append(block)
            else:
                rows.append([block])
        result: list[dict] = []
        for row in rows:
            result.extend(sorted(
                row,
                key=lambda block: (block["x_min"], block["source_index"]),
            ))
        return result

    def _projection_gaps(
        blocks: list[dict],
        *,
        start_key: str,
        end_key: str,
    ) -> list[tuple[float, float, float]]:
        intervals = sorted(
            (float(block[start_key]), float(block[end_key]))
            for block in blocks
        )
        if len(intervals) < 2:
            return []
        merged: list[list[float]] = [[intervals[0][0], intervals[0][1]]]
        for start, end in intervals[1:]:
            if start <= merged[-1][1] + 1.0:
                merged[-1][1] = max(merged[-1][1], end)
            else:
                merged.append([start, end])
        return [
            (right[0] - left[1], left[1], right[0])
            for left, right in zip(merged, merged[1:])
            if right[0] > left[1]
        ]

    def _median_width(blocks: list[dict]) -> float:
        return float(np.median([block["width"] for block in blocks])) if blocks else 0.0

    def _vertical_cut_is_annotation(
        left: list[dict],
        right: list[dict],
        *,
        normalized_gap: float,
    ) -> bool:
        smaller, larger = (left, right) if len(left) <= len(right) else (right, left)
        if not smaller or not larger:
            return True
        paired = sum(
            any(_row_overlap(block, other) >= 0.50 for other in larger)
            for block in smaller
        ) / len(smaller)
        count_ratio = len(smaller) / max(1, len(larger))
        width_ratio = _median_width(smaller) / max(1.0, _median_width(larger))
        # A shallow gap between equally populated row fragments is normally a
        # key/value form, not two independent reading columns.
        # Detector boxes tightly wrap glyphs, so an ordinary form can expose a
        # whitespace gap close to 10% of the page even when the visual label
        # and value cells are adjacent.
        if normalized_gap < 0.11 and paired >= 0.75:
            return True
        # A narrow, smaller group aligned one-to-one with wider record headers
        # is normally a date/location annotation column. Read it in-row.
        return count_ratio < 0.67 and paired >= 0.75 and width_ratio < 0.65

    def _best_vertical_gap(
        blocks: list[dict],
        *,
        force_aligned_gutter: bool = False,
    ) -> Optional[tuple[float, float, bool, bool]]:
        minimum = max(8.0, img_width * 0.055, median_height * 1.35)
        gap_splits: list[tuple[float, float, list[dict], list[dict]]] = []
        for gap, start, end in _projection_gaps(
            blocks, start_key="x_min", end_key="x_max"
        ):
            if gap < minimum:
                continue
            cut = (start + end) / 2.0
            left = [block for block in blocks if block["x_center"] < cut]
            right = [block for block in blocks if block["x_center"] >= cut]
            if len(left) < 2 or len(right) < 2:
                continue
            normalized_gap = gap / max(1.0, float(img_width))
            gap_splits.append((normalized_gap, cut, left, right))

        # Multiple parallel gutters are strong evidence for an actual
        # three-column grid, even when every row happens to align.
        multi_gutter = len(gap_splits) >= 2
        candidates: list[tuple[float, float, bool, bool]] = []
        for normalized_gap, cut, left, right in gap_splits:
            if not (multi_gutter or force_aligned_gutter) and _vertical_cut_is_annotation(
                left, right, normalized_gap=normalized_gap
            ):
                continue
            candidates.append((
                normalized_gap,
                cut,
                multi_gutter and any(other[1] < cut for other in gap_splits),
                multi_gutter and any(other[1] > cut for other in gap_splits),
            ))
        return max(candidates, default=None)

    def _best_horizontal_gap(blocks: list[dict]) -> Optional[tuple[float, float]]:
        minimum = max(4.0, img_height * 0.006, median_height * 0.35)
        candidates: list[tuple[float, float]] = []
        for gap, start, end in _projection_gaps(
            blocks, start_key="y_min", end_key="y_max"
        ):
            if gap < minimum:
                continue
            cut = (start + end) / 2.0
            top = [block for block in blocks if block["y_center"] < cut]
            bottom = [block for block in blocks if block["y_center"] >= cut]
            if not top or not bottom:
                continue
            candidates.append((gap / max(1.0, float(img_height)), cut))
        return max(candidates, default=None)

    def _recursive_order(
        blocks: list[dict],
        depth: int = 0,
        *,
        force_aligned_gutter: bool = False,
    ) -> list[dict]:
        if len(blocks) <= 1:
            return list(blocks)
        if depth >= 16:
            return _row_major(blocks)

        vertical = _best_vertical_gap(
            blocks,
            force_aligned_gutter=force_aligned_gutter,
        )
        horizontal = _best_horizontal_gap(blocks)
        vertical_score = vertical[0] if vertical else -1.0
        horizontal_score = horizontal[0] if horizontal else -1.0

        # Prefer a persistent column gutter only when it is materially stronger
        # than the whitespace between ordinary text rows.
        if vertical and (
            not horizontal or vertical_score >= horizontal_score * 1.20
        ):
            cut = vertical[1]
            left = [block for block in blocks if block["x_center"] < cut]
            right = [block for block in blocks if block["x_center"] >= cut]
            return _recursive_order(
                left,
                depth + 1,
                force_aligned_gutter=vertical[2],
            ) + _recursive_order(
                right,
                depth + 1,
                force_aligned_gutter=vertical[3],
            )
        if horizontal:
            cut = horizontal[1]
            top = [block for block in blocks if block["y_center"] < cut]
            bottom = [block for block in blocks if block["y_center"] >= cut]
            return _recursive_order(
                top,
                depth + 1,
                force_aligned_gutter=force_aligned_gutter,
            ) + _recursive_order(
                bottom,
                depth + 1,
                force_aligned_gutter=force_aligned_gutter,
            )
        return _row_major(blocks)

    return [block["text"] for block in _recursive_order(raw_blocks)]


def _ocr_image_multicandidate(content: bytes, *, prepared_image=None) -> str:
    """Try a bounded number of alternate image transforms on the shared engine."""
    import numpy as np  # noqa: F811
    from PIL import ImageOps

    pil_img = prepared_image if prepared_image is not None else _prepare_ocr_image(content)[0]

    if _RAPID_OCR is None:
        return ""

    def _run(pil: Image.Image) -> str:
        try:
            np_img = np.array(pil.convert("RGB"))
            result = _run_rapid_ocr(np_img)
            return _rapid_result_to_text(
                result,
                img_width=np_img.shape[1],
                img_height=np_img.shape[0],
            )
        except Exception as exc:
            logger.warning("RapidOCR alternate preprocessing failed: %s", exc)
            return ""

    # The primary path uses the green channel for colored resume sidebars.
    # First retry a conventional grayscale image; this is the only retry by
    # default, keeping an empty/unsupported image from multiplying OCR cost.
    auto = ImageOps.autocontrast(pil_img.convert("L"), cutoff=3)
    text = _run(auto.convert("RGB"))
    if text.strip():
        return text

    if _positive_int_env("RAPID_OCR_FALLBACK_ATTEMPTS", 1) < 2:
        return ""

    # Optional second retry for white text on a dark left column.
    w, h = pil_img.size
    if w > 200:
        arr = np.array(pil_img)
        left_inv = 255 - arr[:, :int(w * 0.35), :]
        right_orig = arr[:, int(w * 0.35):, :]
        merged = np.concatenate([left_inv, right_orig], axis=1)
        return _run(Image.fromarray(merged).convert("RGB"))
    return ""

from server_runtime import (
    ALLOWED_UPLOAD_EXTENSIONS,
    AVATAR_DIR,
    ENABLE_AVATAR_EXTRACTION,
    ENABLE_AVATAR_FACE_SCORE,
    ENABLE_TEXT_LAYOUT_NORMALIZATION,
    PROJECT_SECTION_HEADERS,
    SECTION_HEADING_KEYWORDS,
    SUPPORTED_FILE_PATH_EXTENSIONS,
    _FACE_CASCADE,
    cv2,
    fitz,
    logger,
    np,
    pikepdf,
    remaining_request_seconds,
)
def _normalize_heading_line(line: str) -> str:
    text = str(line or "").strip()
    text = re.sub(r"\s+", "", text)
    return text.strip("：:;；|·•-—–")


def _looks_like_section_header(line: str, keywords: tuple[str, ...], max_len: int = 32) -> bool:
    normalized = _normalize_heading_line(line)
    if not normalized or len(normalized) > max_len:
        return False
    lowered = normalized.lower()
    for keyword in keywords:
        key = _normalize_heading_line(keyword)
        if key and (key in normalized or key.lower() in lowered):
            return True
    return False


def _normalize_extracted_resume_text(text: str) -> str:
    raw = str(text or "")
    if not raw:
        return ""
    if not ENABLE_TEXT_LAYOUT_NORMALIZATION:
        return raw.strip()

    normalized = raw.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\u3000", " ").replace("\xa0", " ")
    normalized = normalized.replace("\ufeff", "").replace("\u200b", "")
    # Remove spaces/tabs between Chinese chars (NOT newlines \u2014 those separate
    # OCR blocks and must be preserved for correct reading order).
    normalized = re.sub(r"(?<=[\u4e00-\u9fff])[ \t]+(?=[\u4e00-\u9fff])", "", normalized)

    heading_alt = "|".join(re.escape(item) for item in sorted(SECTION_HEADING_KEYWORDS, key=len, reverse=True))
    if heading_alt:
        normalized = re.sub(rf"\s*({heading_alt})\s*", r"\n\1\n", normalized, flags=re.IGNORECASE)

    normalized = re.sub(r"[•▪◦●○■□▶►▸▹◆◇\uF000-\uF8FF]+", "\n- ", normalized)
    normalized = re.sub(
        r"(?<!\n)\s*((?:19|20)\d{2}[./年]\d{1,2}(?:月)?\s*(?:[-–—~至到]+\s*(?:(?:19|20)\d{2}[./年]\d{1,2}(?:月)?|至今|Present|present)))",
        r"\n\1",
        normalized,
        flags=re.IGNORECASE,
    )

    lines: list[str] = []
    for line in normalized.splitlines():
        cleaned = re.sub(r"\s{2,}", " ", line).strip()
        if not cleaned:
            continue
        lines.append(cleaned)

    return "\n".join(lines).strip()


def _detect_extension(file_name: str) -> str:
    return Path(file_name or "").suffix.lower()


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def _collect_docx_table_text(table: Any) -> list[str]:
    lines: list[str] = []
    rows = getattr(table, "rows", None)
    if not rows:
        return lines

    for row in rows:
        for cell in getattr(row, "cells", []) or []:
            for paragraph in getattr(cell, "paragraphs", []) or []:
                text = str(getattr(paragraph, "text", "") or "").strip()
                if text:
                    lines.append(text)
            for nested in getattr(cell, "tables", []) or []:
                lines.extend(_collect_docx_table_text(nested))
    return lines


def _extract_text_from_docx_document(doc: Any) -> str:
    """Extract text from DOCX preserving paragraph/table XML body order.

    Traverses ``<w:body>`` child elements (``<w:p>`` and ``<w:tbl>``) in their
    original document order, rather than collecting all paragraphs then all
    tables separately.  This preserves the structure that python-docx's
    ``.paragraphs`` / ``.tables`` iterators discard.
    """
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    lines: list[str] = []
    seen: set[str] = set()

    body = doc.element.body
    for child in body:
        tag = child.tag
        if tag == qn("w:p"):
            parts: list[str] = []
            for node in child.iter(f"{{{ns_w}}}t"):
                if node.text:
                    parts.append(node.text)
            text = re.sub(r"\s+", " ", "".join(parts)).strip()
            if text and text not in seen:
                seen.add(text)
                lines.append(text)
        elif tag == qn("w:tbl"):
            _collect_docx_table_text_inline(child, ns_w, lines, seen)

    return "\n".join(lines).strip()


def _collect_docx_table_text_inline(
    tbl_elem: Any, ns_w: str, lines: list[str], seen: set[str]
) -> None:
    """Extract cell text from a single ``<w:tbl>`` element, in row order."""
    for row in tbl_elem.iter(f"{{{ns_w}}}tr"):
        for cell in row.iter(f"{{{ns_w}}}tc"):
            cell_texts: list[str] = []
            for p in cell.iter(f"{{{ns_w}}}p"):
                parts: list[str] = []
                for t_node in p.iter(f"{{{ns_w}}}t"):
                    if t_node.text:
                        parts.append(t_node.text)
                ct = re.sub(r"\s+", " ", "".join(parts)).strip()
                if ct:
                    cell_texts.append(ct)
            combined = " ".join(cell_texts)
            if combined and combined not in seen:
                seen.add(combined)
                lines.append(combined)


def _extract_text_from_docx_xml_bytes(content: bytes) -> str:
    lines: list[str] = []
    seen: set[str] = set()
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
    candidate_pattern = re.compile(r"^word/(document|header\d+|footer\d+|footnotes|endnotes)\.xml$")

    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            names = [name for name in zf.namelist() if candidate_pattern.match(name)]
            for name in sorted(names):
                try:
                    root = ET.fromstring(zf.read(name))
                except Exception:
                    continue
                for paragraph in root.iter(f"{{{ns_w}}}p"):
                    parts: list[str] = []
                    for node in paragraph.iter():
                        if node.tag in {f"{{{ns_w}}}t", f"{{{ns_a}}}t"}:
                            value = str(node.text or "")
                            if value:
                                parts.append(value)
                    text = re.sub(r"\s+", " ", "".join(parts)).strip()
                    if text and text not in seen:
                        seen.add(text)
                        lines.append(text)
    except Exception:
        return ""

    return "\n".join(lines).strip()


def _sanitize_pdf_bytes(content: bytes) -> bytes:
    """Normalize PDF bytes with pikepdf/qpdf when available."""
    if pikepdf is None:
        return content
    try:
        with pikepdf.open(io.BytesIO(content)) as pdf:
            output = io.BytesIO()
            pdf.save(output)
            sanitized = output.getvalue()
            if sanitized:
                return sanitized
    except Exception as exc:
        logger.warning("pikepdf sanitize failed; fallback to raw bytes | error=%s", exc)
    return content


def _normalize_image_ext(ext: str) -> str:
    value = str(ext or "").strip().lower().lstrip(".")
    if value == "jpeg":
        value = "jpg"
    if value not in {"jpg", "png", "webp", "bmp", "gif", "tif", "tiff"}:
        value = "png"
    return value


def _coerce_avatar_bytes_to_rgb_png(image_bytes: bytes, ext: str) -> tuple[bytes, str]:
    payload = bytes(image_bytes or b"")
    image_ext = _normalize_image_ext(ext)
    if not payload or fitz is None:
        return payload, image_ext

    try:
        image_doc = fitz.open(stream=payload, filetype=image_ext)
        if image_doc.page_count <= 0:
            image_doc.close()
            return payload, image_ext
        page = image_doc.load_page(0)
        pix = page.get_pixmap(alpha=False)
        png_bytes = pix.tobytes("png")
        image_doc.close()
        if png_bytes:
            return bytes(png_bytes), "png"
    except Exception:
        return payload, image_ext

    return payload, image_ext


def _persist_avatar_bytes(image_bytes: bytes, ext: str, source_name: str) -> Optional[str]:
    if not ENABLE_AVATAR_EXTRACTION:
        return None
    payload, image_ext = _coerce_avatar_bytes_to_rgb_png(image_bytes, ext)
    if len(payload) < 1024:
        return None
    stem = re.sub(r"[^\w\u4e00-\u9fff\-]+", "_", Path(source_name or "resume").stem).strip("_") or "resume"
    filename = f"{stem}_avatar_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}.{image_ext}"
    output_path = AVATAR_DIR / filename
    output_path.write_bytes(payload)
    return str(output_path)


def _extract_pdf_image_png_bytes(doc: Any, xref: int, smask: int, source_name: str) -> bytes:
    pix = None
    try:
        base_pix = fitz.Pixmap(doc, xref)
        pix = base_pix
        if smask > 0:
            try:
                mask_pix = fitz.Pixmap(doc, smask)
                pix = fitz.Pixmap(base_pix, mask_pix)
            except Exception:
                pix = base_pix
        if pix.alpha:
            pix = fitz.Pixmap(pix, 0)
        if pix.colorspace is not None and pix.colorspace.n > 3:
            pix = fitz.Pixmap(fitz.csRGB, pix)
        png_bytes = pix.tobytes("png")
        if png_bytes:
            return bytes(png_bytes)
    except Exception as exc:
        logger.warning("Avatar pixmap normalize failed | source=%s xref=%s error=%s", source_name, xref, exc)

    try:
        base = doc.extract_image(xref) or {}
        return bytes(base.get("image") or b"")
    except Exception:
        return b""


def _avatar_face_stats(image_bytes: bytes) -> tuple[int, float]:
    if (
        not ENABLE_AVATAR_EXTRACTION
        or not ENABLE_AVATAR_FACE_SCORE
        or cv2 is None
        or np is None
        or _FACE_CASCADE is None
    ):
        return 0, 0.0

    try:
        arr = np.frombuffer(image_bytes or b"", dtype=np.uint8)
        if arr.size == 0:
            return 0, 0.0
        img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
        if img is None:
            return 0, 0.0
        h, w = img.shape[:2]
        if w < 40 or h < 40:
            return 0, 0.0
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        faces = _FACE_CASCADE.detectMultiScale(
            gray,
            scaleFactor=1.1,
            minNeighbors=5,
            minSize=(40, 40),
        )
        if len(faces) == 0:
            return 0, 0.0
        image_area = max(float(w * h), 1.0)
        max_ratio = 0.0
        for face in faces:
            try:
                fw = float(face[2] or 0.0)
                fh = float(face[3] or 0.0)
                max_ratio = max(max_ratio, (fw * fh) / image_area)
            except Exception:
                continue
        return int(len(faces)), float(max_ratio)
    except Exception:
        return 0, 0.0


def _apply_avatar_face_bonus(base_score: float, image_bytes: bytes) -> float:
    face_count, max_ratio = _avatar_face_stats(image_bytes)
    if face_count <= 0:
        return base_score

    # Use face signal as a bonus only; keep layout/size heuristics as primary score.
    multiplier = 1.0 + min(2.8, face_count * 1.05 + max_ratio * 4.0)
    additive = 25000.0 * float(min(face_count, 3))
    return base_score * multiplier + additive


def _extract_avatar_from_pdf_bytes(content: bytes, source_name: str) -> Optional[str]:
    if fitz is None or not ENABLE_AVATAR_EXTRACTION:
        return None

    sanitized = _sanitize_pdf_bytes(content)
    candidates: list[bytes] = [sanitized] if sanitized != content else []
    candidates.append(content)

    for pdf_bytes in candidates:
        doc = None
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            if doc.page_count <= 0:
                continue
            page = doc.load_page(0)
            rect = page.rect
            page_w = float(rect.width or 0.0)
            page_h = float(rect.height or 0.0)
            page_area = max(1.0, page_w * page_h)

            scored: list[tuple[float, int, int, bytes]] = []
            for img in page.get_images(full=True) or []:
                if not isinstance(img, (list, tuple)) or not img:
                    continue
                xref = int(img[0] or 0)
                if xref <= 0:
                    continue
                smask = int(img[1] or 0) if len(img) > 1 else 0
                rects: list[Any] = []
                try:
                    rects = page.get_image_rects(xref) or []
                except Exception:
                    rects = []

                best_rect = None
                best_area = 0.0
                for r in rects:
                    try:
                        rw = float(r.width or 0.0)
                        rh = float(r.height or 0.0)
                        area = rw * rh
                    except Exception:
                        continue
                    if area > best_area:
                        best_area = area
                        best_rect = r

                if best_rect is None:
                    # No placement rect (rare) - still allow as low-confidence fallback.
                    png_bytes = _extract_pdf_image_png_bytes(doc, xref, smask, source_name)
                    if not png_bytes:
                        continue
                    score = _apply_avatar_face_bonus(1.0, png_bytes)
                    scored.append((score, xref, smask, png_bytes))
                    continue

                width = max(0.0, float(best_rect.width or 0.0))
                height = max(0.0, float(best_rect.height or 0.0))
                area = width * height
                if width < 36 or height < 36:
                    continue
                ratio = area / page_area
                if ratio < 0.0015 or ratio > 0.28:
                    continue

                x0 = float(best_rect.x0 or 0.0)
                y0 = float(best_rect.y0 or 0.0)
                aspect = width / max(height, 1.0)
                score = float(area)
                if 0.55 <= aspect <= 1.45:
                    score *= 1.2
                if y0 <= page_h * 0.6:
                    score *= 1.2
                if x0 >= page_w * 0.38:
                    score *= 1.1
                png_bytes = _extract_pdf_image_png_bytes(doc, xref, smask, source_name)
                if not png_bytes:
                    continue
                score = _apply_avatar_face_bonus(score, png_bytes)
                scored.append((score, xref, smask, png_bytes))

            if not scored:
                continue

            scored.sort(key=lambda item: item[0], reverse=True)
            best_score, best_xref, best_smask, best_png_bytes = scored[0]
            if best_score <= 0:
                continue

            saved = _persist_avatar_bytes(best_png_bytes, "png", source_name=source_name)
            if saved:
                logger.info(
                    "Extracted avatar from PDF | source=%s path=%s xref=%s smask=%s",
                    source_name,
                    saved,
                    best_xref,
                    best_smask,
                )
                return saved
        except Exception as exc:
            logger.warning("Avatar extraction from PDF failed | source=%s error=%s", source_name, exc)
        finally:
            if doc is not None:
                doc.close()

    return None


def _extract_avatar_from_docx_bytes(content: bytes, source_name: str) -> Optional[str]:
    if not ENABLE_AVATAR_EXTRACTION:
        return None

    # Prefer inline-shape order as it better reflects the original document flow.
    try:
        doc = DocxDocument(io.BytesIO(content))
        scored: list[tuple[float, bytes, str]] = []
        for idx, shape in enumerate(getattr(doc, "inline_shapes", [])):
            try:
                inline = shape._inline  # type: ignore[attr-defined]
                rid = inline.graphic.graphicData.pic.blipFill.blip.embed
                image_part = doc.part.related_parts.get(rid)
                if image_part is None:
                    continue
                payload = bytes(image_part.blob or b"")
                if len(payload) < 1024:
                    continue
                ext = _normalize_image_ext(Path(getattr(image_part, "filename", "avatar.png")).suffix)
                width = float(getattr(shape, "width", 0) or 0)
                height = float(getattr(shape, "height", 0) or 0)
                area = max(1.0, width * height)
                score = area * (1.25 if idx == 0 else 1.0 / (idx + 1))
                score = _apply_avatar_face_bonus(score, payload)
                scored.append((score, payload, ext))
            except Exception:
                continue
        if scored:
            scored.sort(key=lambda item: item[0], reverse=True)
            saved = _persist_avatar_bytes(scored[0][1], scored[0][2], source_name=source_name)
            if saved:
                logger.info("Extracted avatar from DOCX inline shape | source=%s path=%s", source_name, saved)
                return saved
    except Exception as exc:
        logger.warning("DOCX inline-shape avatar extraction failed | source=%s error=%s", source_name, exc)

    # Fallback: scan embedded media files.
    try:
        scored_zip: list[tuple[float, bytes, str]] = []
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            for info in zf.infolist():
                name = info.filename.lower()
                if not name.startswith("word/media/"):
                    continue
                ext = _normalize_image_ext(Path(name).suffix)
                payload = zf.read(info.filename)
                if len(payload) < 1024:
                    continue
                # Prefer earlier media entries and larger files.
                score = float(info.file_size) + max(0, 200000 - info.header_offset) * 0.1
                score = _apply_avatar_face_bonus(score, payload)
                scored_zip.append((score, payload, ext))
        if scored_zip:
            scored_zip.sort(key=lambda item: item[0], reverse=True)
            saved = _persist_avatar_bytes(scored_zip[0][1], scored_zip[0][2], source_name=source_name)
            if saved:
                logger.info("Extracted avatar from DOCX media | source=%s path=%s", source_name, saved)
                return saved
    except Exception as exc:
        logger.warning("DOCX media avatar extraction failed | source=%s error=%s", source_name, exc)

    return None


def _extract_avatar_from_upload_bytes(content: bytes, filename: str) -> Optional[str]:
    if not ENABLE_AVATAR_EXTRACTION:
        return None
    ext = _detect_extension(filename)
    if ext == ".pdf":
        return _extract_avatar_from_pdf_bytes(content, source_name=filename or "resume.pdf")
    if ext == ".docx":
        return _extract_avatar_from_docx_bytes(content, source_name=filename or "resume.docx")
    return None


def _extract_avatar_from_file_path(file_path: str) -> Optional[str]:
    if not ENABLE_AVATAR_EXTRACTION or not file_path:
        return None
    path = Path(file_path)
    if not path.exists() or not path.is_file():
        return None
    try:
        payload = path.read_bytes()
    except Exception:
        return None
    return _extract_avatar_from_upload_bytes(payload, path.name)


def _extract_native_pdf_page_text(page: Any) -> str:
    """Extract a PDF text layer using the same layout ordering as OCR.

    ``page.get_text()`` follows the PDF object's internal insertion order,
    which is often column-interleaved in resume templates.  PyMuPDF's dict
    representation retains line coordinates, so feed those coordinates into
    the common, text-preserving layout sorter.  Fall back to plain extraction
    for malformed or unusually shaped text dictionaries.
    """
    try:
        page_dict = page.get_text("dict", sort=False) or {}
        boxes: list[list[list[float]]] = []
        texts: list[str] = []
        for block in page_dict.get("blocks", []) or []:
            if block.get("type", 0) != 0:
                continue
            for line in block.get("lines", []) or []:
                line_text = "".join(
                    str(span.get("text", ""))
                    for span in (line.get("spans", []) or [])
                ).strip()
                bbox = line.get("bbox")
                if not line_text or not bbox or len(bbox) != 4:
                    continue
                x_min, y_min, x_max, y_max = map(float, bbox)
                if x_max <= x_min or y_max <= y_min:
                    continue
                boxes.append([
                    [x_min, y_min],
                    [x_max, y_min],
                    [x_max, y_max],
                    [x_min, y_max],
                ])
                texts.append(line_text)

        if boxes:
            rect = getattr(page, "rect", None)
            width = max(1, int(round(float(getattr(rect, "width", 0) or 0))))
            height = max(1, int(round(float(getattr(rect, "height", 0) or 0))))
            if width <= 1:
                width = max(1, int(math.ceil(max(point[0] for box in boxes for point in box))))
            if height <= 1:
                height = max(1, int(math.ceil(max(point[1] for box in boxes for point in box))))
            ordered = _reconstruct_ocr_reading_order(
                boxes,
                texts,
                img_width=width,
                img_height=height,
            )
            if ordered:
                return "\n".join(ordered).strip()
    except Exception as exc:
        logger.debug("Coordinate-aware PDF text extraction failed: %s", exc)

    return (page.get_text() or "").strip()


def _extract_text_from_pdf_bytes(content: bytes) -> str:
    if fitz is None:
        raise HTTPException(status_code=500, detail="pymupdf is required for PDF parsing")

    sanitized = _sanitize_pdf_bytes(content)
    candidates: list[tuple[str, bytes]] = []
    if sanitized != content:
        candidates.append(("sanitized", sanitized))
    candidates.append(("raw", content))

    last_exc: Optional[Exception] = None
    for source, pdf_bytes in candidates:
        doc = None
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            page_texts: list[str] = []
            max_ocr_pages = _positive_int_env("OCR_PDF_MAX_PAGES", 12)
            sparse_threshold = _positive_int_env("OCR_PDF_NATIVE_MIN_CHARS", 40)
            render_scale = max(1.0, min(3.0, _positive_float_env("OCR_PDF_RENDER_SCALE", 1.7)))
            ocr_pages = 0
            for page_index, page in enumerate(doc):
                native_text = _extract_native_pdf_page_text(page)
                native_chars = len(re.sub(r"\s+", "", native_text))
                if native_chars >= sparse_threshold or ocr_pages >= max_ocr_pages:
                    page_texts.append(native_text)
                    continue

                # Scanned and hybrid PDFs commonly contain an empty text layer.
                # Render only sparse pages; native pages remain cheap and exact.
                pixmap = page.get_pixmap(
                    matrix=fitz.Matrix(render_scale, render_scale),
                    alpha=False,
                )
                image_bytes = pixmap.tobytes("png")
                ocr_text = extract_text_from_image_bytes(
                    image_bytes,
                    f"page-{page_index + 1}.png",
                ).strip()
                ocr_pages += 1
                page_texts.append(ocr_text if len(ocr_text) > len(native_text) else native_text)
            if ocr_pages:
                logger.info(
                    "PDF sparse-page OCR finished | pages=%s ocr_pages=%s source=%s",
                    len(page_texts),
                    ocr_pages,
                    source,
                )
            return "\n".join(text for text in page_texts if text).strip()
        except Exception as exc:
            last_exc = exc
            logger.warning("PDF parse failed via %s bytes | error=%s", source, exc)
        finally:
            if doc is not None:
                doc.close()

    raise HTTPException(status_code=400, detail=f"Failed to parse PDF: {last_exc}") from last_exc


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".gif", ".tif", ".tiff"}


def _validate_upload_signature(content: bytes, filename: str) -> None:
    """Reject extension spoofing and compressed document bombs early."""
    ext = _detect_extension(filename)
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
    if ext == ".pdf" and not content.lstrip().startswith(b"%PDF-"):
        raise HTTPException(status_code=400, detail="File content is not a valid PDF")
    if ext == ".docx":
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                names = set(archive.namelist())
                if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                    raise ValueError("missing DOCX package entries")
                expanded = sum(info.file_size for info in archive.infolist())
                compressed = max(1, sum(info.compress_size for info in archive.infolist()))
                if expanded > 100 * 1024 * 1024 or expanded / compressed > 100:
                    raise ValueError("compressed document exceeds safety limits")
        except (zipfile.BadZipFile, ValueError) as exc:
            raise HTTPException(status_code=400, detail="File content is not a safe DOCX") from exc
    if ext in IMAGE_EXTENSIONS:
        try:
            with Image.open(io.BytesIO(content)) as image:
                width, height = image.size
                if width <= 0 or height <= 0 or width * height > 40_000_000:
                    raise ValueError("image dimensions exceed safety limits")
                image.verify()
        except Exception as exc:
            raise HTTPException(status_code=400, detail="File content is not a valid image") from exc
    if ext in {".txt", ".md"} and b"\x00" in content[:4096]:
        raise HTTPException(status_code=400, detail="Text upload contains binary data")


def extract_text_from_image_bytes(content: bytes, filename: str) -> str:
    total_started = time.perf_counter()
    prepared_image, original_size = _prepare_ocr_image(content)
    logger.info(
        "OCR image prepared | extension=%s original=%sx%s input=%sx%s resized=%s",
        Path(filename).suffix.lower(),
        original_size[0],
        original_size[1],
        prepared_image.width,
        prepared_image.height,
        prepared_image.size != original_size,
    )

    # Prefer RapidOCR (CUDA server or CPU mobile), then bounded fallbacks.
    if not _RAPID_OCR_INITED:
        _init_rapid_ocr()

    text = ""
    if _RAPID_OCR is not None:
        stage_started = time.perf_counter()
        text = _ocr_image_with_rapid(content, prepared_image=prepared_image)
        logger.info(
            "OCR primary finished | provider=%s model=%s elapsed_s=%.3f chars=%s",
            _RAPID_OCR_PROVIDER,
            _RAPID_OCR_MODEL_TYPE,
            time.perf_counter() - stage_started,
            len(text),
        )

    if not text.strip():
        stage_started = time.perf_counter()
        try:
            text = _ocr_image_multicandidate(content, prepared_image=prepared_image)
        except Exception as exc:
            logger.warning("RapidOCR preprocessing fallback failed: %s", exc)
        logger.info(
            "OCR alternate preprocessing finished | elapsed_s=%.3f chars=%s",
            time.perf_counter() - stage_started,
            len(text),
        )

    if not text.strip() and pytesseract is not None:
        stage_started = time.perf_counter()
        try:
            text = pytesseract.image_to_string(
                prepared_image,
                lang="chi_sim+eng",
                timeout=_positive_float_env("TESSERACT_TIMEOUT_SECONDS", 15.0),
            )
        except Exception as exc:
            logger.warning("pytesseract OCR also failed: %s", exc)
        logger.info(
            "OCR tesseract fallback finished | elapsed_s=%.3f chars=%s",
            time.perf_counter() - stage_started,
            len(text),
        )

    if not text.strip():
        logger.warning(
            "All OCR engines failed | extension=%s total_s=%.3f",
            Path(filename).suffix.lower(),
            time.perf_counter() - total_started,
        )
        return ""

    normalized = _normalize_extracted_resume_text(text)
    logger.info(
        "OCR extraction finished | total_s=%.3f normalized_chars=%s",
        time.perf_counter() - total_started,
        len(normalized),
    )
    return normalized if normalized else ""

def _extract_text_from_bytes_impl(content: bytes, filename: str) -> str:
    ext = _detect_extension(filename)
    if ext not in ALLOWED_UPLOAD_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Unsupported file extension: {ext}")
    _validate_upload_signature(content, filename)

    if ext in IMAGE_EXTENSIONS:
        return extract_text_from_image_bytes(content, filename)

    if ext in {".txt", ".md"}:
        text = content.decode("utf-8", errors="replace").strip()
    elif ext == ".pdf":
        text = _extract_text_from_pdf_bytes(content)
    else:
        try:
            doc = DocxDocument(io.BytesIO(content))
            text = _extract_text_from_docx_document(doc)
            if len(text) < 20:
                xml_text = _extract_text_from_docx_xml_bytes(content)
                if len(xml_text) > len(text):
                    logger.info(
                        "DOCX XML fallback improved text extraction | filename=%s before_chars=%s after_chars=%s",
                        filename,
                        len(text),
                        len(xml_text),
                    )
                    text = xml_text
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {exc}") from exc

    text = _normalize_extracted_resume_text(text)
    if not text:
        raise HTTPException(status_code=400, detail="No valid text extracted from file")
    return text


def _isolated_descendant_pids(root_pid: int) -> set[int]:
    if os.name != "posix" or not Path("/proc").is_dir():
        return set()
    descendants: set[int] = set()
    pending = [int(root_pid)]
    while pending:
        parent_pid = pending.pop()
        try:
            child_files = list(Path(f"/proc/{parent_pid}/task").glob("*/children"))
        except OSError:
            continue
        for child_file in child_files:
            try:
                child_ids = [int(value) for value in child_file.read_text().split()]
            except (OSError, ValueError):
                continue
            for child_pid in child_ids:
                if child_pid in descendants or child_pid == root_pid:
                    continue
                descendants.add(child_pid)
                pending.append(child_pid)
    return descendants


def _signal_isolated_descendants(descendants: set[int], sig: int) -> None:
    for child_pid in descendants:
        try:
            os.kill(child_pid, sig)
        except (OSError, ProcessLookupError):
            pass


def _isolated_extract_entry(connection, content: bytes, filename: str) -> None:
    """Child entry for native OCR libraries that cannot be cancelled safely."""

    process_group_ready = False
    try:
        nested_in_task_group = os.getenv("RESUME_TASK_PROCESS_GROUP", "").strip() == "1"
        if os.name == "posix" and not nested_in_task_group:
            os.setsid()
            process_group_ready = True
        connection.send(("ready", {"process_group": process_group_ready}))
        text = _extract_text_from_bytes_impl(content, filename)
        connection.send(("ok", text))
    except BaseException as exc:  # child must serialize failures, not traceback objects
        connection.send((
            "error",
            {
                "status_code": int(getattr(exc, "status_code", 400) or 400),
                "detail": str(getattr(exc, "detail", exc) or type(exc).__name__),
                "exception_type": type(exc).__name__,
            },
        ))
    finally:
        connection.close()


def _terminate_isolated_process(process: mp.Process, *, process_group_ready: bool) -> None:
    """Terminate an OCR child and any native subprocesses it spawned."""

    if process.pid is None:
        return
    if not process.is_alive():
        process.join(timeout=0.2)
        if os.name == "posix" and process_group_ready:
            try:
                os.killpg(process.pid, signal.SIGTERM)
            except (OSError, ProcessLookupError):
                pass
            time.sleep(0.05)
            try:
                os.killpg(process.pid, signal.SIGKILL)
            except (OSError, ProcessLookupError):
                pass
        return
    descendants = _isolated_descendant_pids(process.pid)
    _signal_isolated_descendants(descendants, signal.SIGTERM)
    try:
        if os.name == "posix" and process_group_ready:
            os.killpg(process.pid, signal.SIGTERM)
        else:
            process.terminate()
    except (OSError, ProcessLookupError):
        pass
    process.join(timeout=1.0)
    if not process.is_alive():
        _signal_isolated_descendants(descendants, signal.SIGKILL)
        return
    descendants.update(_isolated_descendant_pids(process.pid))
    _signal_isolated_descendants(descendants, signal.SIGKILL)
    try:
        if os.name == "posix" and process_group_ready:
            os.killpg(process.pid, signal.SIGKILL)
        elif hasattr(process, "kill"):
            process.kill()
        else:
            process.terminate()
    except (OSError, ProcessLookupError):
        pass
    process.join(timeout=1.0)


def _ocr_hard_timeout_seconds() -> float:
    configured = _positive_float_env("OCR_HARD_TIMEOUT_SECONDS", 60.0)
    remaining = remaining_request_seconds()
    if remaining is None:
        return configured
    # Leave a small budget for deterministic fallback/reporting after OCR.
    return max(0.1, min(configured, remaining - 2.0))


def _extract_text_isolated(content: bytes, filename: str, timeout_seconds: float) -> str:
    """Run OCR/PDF extraction in a spawn child with a real kill boundary."""

    context = mp.get_context(os.getenv("OCR_PROCESS_START_METHOD", "spawn"))
    parent_connection, child_connection = context.Pipe(duplex=False)
    process = context.Process(
        target=_isolated_extract_entry,
        args=(child_connection, content, filename),
        name="resume-ocr",
    )
    process.start()
    child_connection.close()
    deadline = time.monotonic() + max(0.1, float(timeout_seconds))
    process_group_ready = False
    result: Optional[tuple[str, Any]] = None
    try:
        while time.monotonic() < deadline:
            wait_seconds = min(0.1, max(0.0, deadline - time.monotonic()))
            if parent_connection.poll(wait_seconds):
                message = parent_connection.recv()
                if not isinstance(message, tuple) or len(message) != 2:
                    continue
                status, payload = message
                if status == "ready":
                    process_group_ready = bool(
                        isinstance(payload, dict) and payload.get("process_group")
                    )
                    continue
                result = (str(status), payload)
                break
            if not process.is_alive():
                break
    finally:
        if result is None and process.is_alive():
            _terminate_isolated_process(
                process,
                process_group_ready=process_group_ready,
            )
        else:
            process.join(timeout=1.0)
            if process.is_alive():
                _terminate_isolated_process(
                    process,
                    process_group_ready=process_group_ready,
                )
        parent_connection.close()

    if result is None:
        if time.monotonic() >= deadline:
            raise HTTPException(
                status_code=408,
                detail=f"OCR exceeded the {timeout_seconds:.1f}s hard timeout",
            )
        raise HTTPException(
            status_code=500,
            detail=f"OCR worker exited without a result (exit={process.exitcode})",
        )
    status, payload = result
    if status == "ok":
        return str(payload or "")
    if isinstance(payload, dict):
        raise HTTPException(
            status_code=int(payload.get("status_code", 400) or 400),
            detail=str(payload.get("detail", "OCR worker failed")),
        )
    raise HTTPException(status_code=400, detail=str(payload or "OCR worker failed"))


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """Extract text, isolating OCR-capable formats behind a hard timeout."""

    ext = _detect_extension(filename)
    needs_native_ocr = ext in IMAGE_EXTENSIONS or ext == ".pdf"
    isolation_enabled = os.getenv("OCR_PROCESS_ISOLATION", "1").strip().lower() not in {
        "0", "false", "no", "off",
    }
    # Daemonic workers cannot create children. The whole-task supervisor still
    # provides an outer kill boundary in that uncommon deployment mode.
    if needs_native_ocr and isolation_enabled and not mp.current_process().daemon:
        return _extract_text_isolated(
            content,
            filename,
            timeout_seconds=_ocr_hard_timeout_seconds(),
        )
    return _extract_text_from_bytes_impl(content, filename)


def extract_text_from_path(file_path: str) -> str:
    if not file_path:
        raise HTTPException(status_code=400, detail="file_path is required")

    path = Path(file_path)
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=400, detail=f"file_path not found: {file_path}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_FILE_PATH_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Unsupported file extension: {ext}")

    try:
        payload = path.read_bytes()
    except OSError as exc:
        raise HTTPException(status_code=400, detail=f"Failed to read file_path: {exc}") from exc
    # Keep path-based callers on the same signature validation, OCR hard timeout,
    # scanned-PDF fallback and normalization path as uploaded bytes.
    return extract_text_from_bytes(payload, path.name)


def resolve_resume_text(resume_content: Optional[str], file_path: Optional[str]) -> str:
    if resume_content and str(resume_content).strip():
        return _normalize_extracted_resume_text(str(resume_content))
    if file_path and str(file_path).strip():
        return extract_text_from_path(str(file_path))
    raise HTTPException(status_code=400, detail="Either resume_content or file_path is required")
