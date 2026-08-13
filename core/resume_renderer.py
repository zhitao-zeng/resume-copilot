import html as html_lib
import copy
import hashlib
import io
import logging
import os
import re
import shutil
import subprocess
import tempfile
import uuid
import zipfile
from pathlib import Path
from typing import Any, Optional

from resume_common import _normalize_project_name

from docx import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from lxml import etree

from template_layout import apply_template_style_profile, extract_template_style_profile

try:
    from fontTools import subset as font_subset
    from fontTools.ttLib import TTCollection, TTFont
except ImportError:
    font_subset = None
    TTCollection = None
    TTFont = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from weasyprint import HTML as WeasyHTML
except ImportError:
    WeasyHTML = None

SUPPORTED_TEMPLATES = {"classic", "modern", "minimal"}
DEFAULT_DOC_FONT = os.getenv("RESUME_DOC_FONT", "Noto Sans CJK SC")
DEFAULT_DOC_MONO_FONT = os.getenv("RESUME_DOC_MONO_FONT", "DejaVu Sans Mono")
DEFAULT_DOC_EAST_ASIA_FONT = os.getenv("RESUME_DOC_EAST_ASIA_FONT", "Noto Sans CJK SC")
DEFAULT_DOC_EMBED_CJK_FONT = os.getenv("RESUME_DOC_EMBED_CJK_FONT", "1").strip().lower() not in {
    "0", "false", "no",
}
DEFAULT_DOC_MAX_EMBEDDED_FONT_BYTES = 4 * 1024 * 1024
PDF_ACCENT_COLORS = {
    "classic": (0x0F, 0x34, 0x60),
    "modern": (0x00, 0x5A, 0x64),
    "minimal": (0x33, 0x33, 0x33),
}
PDF_BODY_COLOR = (0x22, 0x22, 0x22)
PDF_MUTED_COLOR = (0x66, 0x66, 0x66)
PDF_MARGIN_MM = os.getenv("RESUME_PDF_MARGIN_MM", "10mm")
logger = logging.getLogger(__name__)
logging.getLogger("fontTools.subset").setLevel(logging.WARNING)

_OOXML_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_OOXML_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_OOXML_PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_OOXML_CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_OOXML_FONT_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font"
)
_OOXML_OBFUSCATED_FONT_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.obfuscatedFont"
)
_DOCX_CJK_FONT_CANDIDATES = (
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf",
    "/usr/share/fonts/truetype/noto/NotoSansCJKsc-Regular.otf",
)

_RESUME_SECTION_TITLES = {
    "个人简介", "个人总结", "基本信息", "求职信息", "工作经历", "工作/实习经历", "实习经历",
    "科研经历", "校园与志愿经历", "项目经历", "教育经历", "专业技能",
    "论文成果", "荣誉与奖项", "个人技能", "教育背景", "培训与进修", "教学经历",
    "PERSONAL INFORMATION", "CONTACT", "PROFILE", "PROFESSIONAL SUMMARY", "SUMMARY",
    "WORK EXPERIENCE", "EMPLOYMENT HISTORY", "PROJECT EXPERIENCE", "RESEARCH EXPERIENCE",
    "EDUCATION", "SKILLS", "CERTIFICATIONS", "AWARDS", "PUBLICATIONS", "TEACHING EXPERIENCE",
}
_TEMPLATE_SECTION_ALIASES = {
    "基本信息": {"个人信息", "基础信息", "联系方式", "personal information", "contact", "contact information"},
    "个人总结": {"个人简介", "自我评价", "职业总结", "profile", "professional summary", "summary", "about me"},
    "教育背景": {"教育经历", "学历背景", "education", "education background"},
    "专业技能": {"技能", "技能清单", "个人技能", "skills", "professional skills", "core skills"},
    "工作/实习经历": {
        "工作经历", "实习经历", "职业经历", "任职经历", "work experience",
        "employment history", "professional experience", "internship experience",
    },
    "项目经历": {"项目经验", "个人项目", "课程项目", "project experience", "projects", "selected projects"},
    "科研经历": {"研究经历", "实验室经历", "research experience", "research"},
    "校园与志愿经历": {
        "校园经历", "社团经历", "志愿经历", "社会实践", "campus experience",
        "volunteer experience", "activities", "leadership & activities",
    },
    "论文与专利成果": {
        "论文成果", "论文", "专利成果", "学术成果", "publications", "patents",
        "publications & patents",
    },
    "荣誉与奖项": {
        "荣誉奖项", "奖项", "证书", "证书与资质", "资格证书",
        "honors", "awards", "honors & awards", "certifications", "certificates",
    },
    "培训与进修": {"培训经历", "进修经历", "training", "professional development"},
    "教学经历": {"授课经历", "teaching experience", "teaching"},
}
_TEMPLATE_SECTION_TAGS = {
    "basic": "基本信息",
    "basic_info": "基本信息",
    "contact": "基本信息",
    "meta": "基本信息",
    "summary": "个人总结",
    "profile": "个人总结",
    "education": "教育背景",
    "experience": "工作/实习经历",
    "work": "工作/实习经历",
    "work_experience": "工作/实习经历",
    "projects": "项目经历",
    "project": "项目经历",
    "research": "科研经历",
    "activities": "校园与志愿经历",
    "campus": "校园与志愿经历",
    "skills": "专业技能",
    "publications": "论文与专利成果",
    "honors": "荣誉与奖项",
    "awards": "荣誉与奖项",
    "training": "培训与进修",
    "teaching": "教学经历",
}
_TEMPLATE_SECTION_TAG_RE = re.compile(
    r"^\s*\{\{\s*section\.([a-z][a-z0-9_]*)\s*\}\}\s*$",
    re.IGNORECASE,
)
_TEMPLATE_SCALAR_TAG_RE = re.compile(
    r"(?:\{\{\s*(?:name|phone|email|age|gender|work_experience|education_level|target_role|summary)\s*\}\}"
    r"|\$\{(?:name|phone|email|age|gender|work_experience|education_level|target_role|summary)\}"
    r"|\[\[(?:name|phone|email|age|gender|work_experience|education_level|target_role|summary)\]\])",
    re.IGNORECASE,
)
_INTERNAL_ADDITIONAL_SECTION = re.compile(
    r"^(?:待整理(?:的)?原始(?:信息|经历)|教育经历补充)$"
)


def _is_empty_profile_framework(resume_data: Any) -> bool:
    framework = resume_data.get("framework") if isinstance(resume_data, dict) else None
    return isinstance(framework, dict) and framework.get("mode") == "empty_profile"


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^\w\u4e00-\u9fff\-]+", "_", value.strip())
    return cleaned.strip("_") or "resume"


def _normalize_template(template: str) -> str:
    tpl = (template or "classic").strip()
    if not tpl:
        return "classic"
    # Accept both absolute paths and relative filenames.  The old separator
    # check silently downgraded a valid ``template.docx`` in the CWD.
    if Path(tpl).is_file():
        return tpl
    if "/" in tpl or "\\" in tpl:
        return "classic"
    if tpl not in SUPPORTED_TEMPLATES:
        return "classic"
    return tpl


def _contains_cjk(text: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", text or ""))


def _contains_east_asian_text(text: str) -> bool:
    """Return whether text needs the portable CJK font path."""

    return bool(
        re.search(
            r"[\u2e80-\u2fff\u3000-\u303f\u3040-\u30ff\u31f0-\u31ff"
            r"\u3400-\u4dbf\u4e00-\u9fff\uac00-\ud7af\uf900-\ufaff"
            r"\U00020000-\U0002ffff]",
            text or "",
        )
    )


def _font_family_names(font: Any) -> list[str]:
    names: list[str] = []
    name_table = font.get("name")
    if name_table is None:
        return names
    for record in name_table.names:
        if record.nameID != 1:
            continue
        try:
            value = record.toUnicode().strip()
        except Exception:
            continue
        if value and value not in names:
            names.append(value)
    return names


def _docx_cjk_font_paths() -> list[Path]:
    configured = os.getenv("RESUME_DOC_CJK_FONT_PATH", "").strip()
    values = ([configured] if configured else []) + list(_DOCX_CJK_FONT_CANDIDATES)
    paths: list[Path] = []
    for value in values:
        path = Path(value)
        if path.is_file() and path not in paths:
            paths.append(path)
    return paths


def _docx_max_embedded_font_bytes() -> int:
    raw = os.getenv(
        "RESUME_DOC_MAX_EMBEDDED_FONT_BYTES",
        str(DEFAULT_DOC_MAX_EMBEDDED_FONT_BYTES),
    ).strip()
    try:
        value = int(raw)
    except ValueError:
        return DEFAULT_DOC_MAX_EMBEDDED_FONT_BYTES
    return value if value > 0 else DEFAULT_DOC_MAX_EMBEDDED_FONT_BYTES


def _build_cjk_font_subset(characters: set[str]) -> tuple[str, bytes]:
    """Build a license-checked, used-glyph CJK font subset for one DOCX."""

    if font_subset is None or TTCollection is None or TTFont is None:
        raise RuntimeError("fonttools is unavailable")
    if not characters:
        raise ValueError("no characters were supplied for the CJK font subset")

    desired_family = DEFAULT_DOC_EAST_ASIA_FONT.casefold()
    last_error: Optional[Exception] = None
    for path in _docx_cjk_font_paths():
        collection = None
        standalone_font = None
        try:
            if path.suffix.lower() in {".ttc", ".otc"}:
                collection = TTCollection(str(path), lazy=False)
                faces = list(collection.fonts)
            else:
                standalone_font = TTFont(str(path), lazy=False)
                faces = [standalone_font]

            selected = next(
                (
                    face
                    for face in faces
                    if desired_family in {name.casefold() for name in _font_family_names(face)}
                ),
                None,
            )
            if selected is None:
                selected = next(
                    (
                        face
                        for face in faces
                        if any("cjk sc" in name.casefold() for name in _font_family_names(face))
                    ),
                    faces[0] if faces else None,
                )
            if selected is None:
                raise RuntimeError(f"no font face found in {path}")

            family_names = _font_family_names(selected)
            family = next(
                (name for name in family_names if name.casefold() == desired_family),
                family_names[0] if family_names else DEFAULT_DOC_EAST_ASIA_FONT,
            )
            os2 = selected.get("OS/2")
            fs_type = int(getattr(os2, "fsType", 0) or 0)
            if fs_type & 0x0002:
                raise RuntimeError(f"font embedding is restricted by OS/2 fsType: {path}")

            options = font_subset.Options()
            options.name_IDs = [0, 1, 2, 3, 4, 5, 6]
            options.name_legacy = True
            options.name_languages = [0x409, 0x804]
            options.notdef_glyph = True
            options.notdef_outline = True
            options.recalc_timestamp = False
            subsetter = font_subset.Subsetter(options=options)
            subsetter.populate(text="".join(sorted(characters)))
            subsetter.subset(selected)

            target = io.BytesIO()
            selected.save(target)
            payload = target.getvalue()
            maximum = _docx_max_embedded_font_bytes()
            if len(payload) > maximum:
                raise RuntimeError(
                    f"CJK font subset is {len(payload)} bytes, above configured {maximum}-byte cap"
                )
            return family, payload
        except Exception as exc:
            last_error = exc
        finally:
            if collection is not None:
                collection.close()
            elif standalone_font is not None:
                standalone_font.close()

    if last_error is not None:
        raise RuntimeError(f"unable to create CJK font subset: {last_error}") from last_error
    raise FileNotFoundError("no embeddable CJK font file was found")


def _obfuscate_ooxml_font(font_data: bytes, font_key: uuid.UUID) -> bytes:
    """Apply the reversible OOXML font obfuscation to the first 32 bytes."""

    result = bytearray(font_data)
    key = font_key.bytes[::-1]
    for index in range(min(32, len(result))):
        result[index] ^= key[index % 16]
    return bytes(result)


def _serialize_ooxml(root: Any) -> bytes:
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _parse_ooxml(payload: bytes) -> Any:
    parser = etree.XMLParser(resolve_entities=False, no_network=True, huge_tree=False)
    return etree.fromstring(payload, parser=parser)


def _ensure_run_east_asia_font(run: Any, font_family: str) -> None:
    r_pr = run.find(f"{{{_OOXML_WORD_NS}}}rPr")
    if r_pr is None:
        r_pr = etree.Element(f"{{{_OOXML_WORD_NS}}}rPr")
        run.insert(0, r_pr)
    r_fonts = r_pr.find(f"{{{_OOXML_WORD_NS}}}rFonts")
    if r_fonts is None:
        r_fonts = etree.Element(f"{{{_OOXML_WORD_NS}}}rFonts")
        r_pr.insert(0, r_fonts)
    # A user template may intentionally declare different Latin/CJK faces.
    # Preserve every explicit mapping.  Generated runs, however, frequently
    # contain no direct ascii/hAnsi declaration; mapping only eastAsia makes
    # LibreOffice ignore the embedded subset for mixed Chinese/Latin runs and
    # render square glyphs on hosts without CJK fonts.  Fill only the missing
    # direct mappings with the selected East Asia family.
    east_asia_key = f"{{{_OOXML_WORD_NS}}}eastAsia"
    selected_family = r_fonts.get(east_asia_key) or font_family
    if not r_fonts.get(east_asia_key):
        r_fonts.set(east_asia_key, selected_family)
        r_fonts.attrib.pop(f"{{{_OOXML_WORD_NS}}}eastAsiaTheme", None)
    for attribute in ("ascii", "hAnsi"):
        key = f"{{{_OOXML_WORD_NS}}}{attribute}"
        if r_fonts.get(key):
            continue
        r_fonts.set(key, selected_family)
        r_fonts.attrib.pop(f"{{{_OOXML_WORD_NS}}}{attribute}Theme", None)


def _patch_docx_cjk_xml(
    members: dict[str, bytes],
    font_family: str,
) -> tuple[bool, set[str]]:
    """Set direct CJK run fonts and replace non-portable Symbol bullets."""

    visible_parts = {
        name
        for name in members
        if name == "word/document.xml"
        or re.fullmatch(
            r"word/(?:header\d+|footer\d+|footnotes|endnotes|comments)\.xml",
            name,
        )
    }
    parsed: dict[str, Any] = {}
    characters: set[str] = set()
    has_east_asian_text = False
    namespaces = {"w": _OOXML_WORD_NS}

    for name in visible_parts:
        root = _parse_ooxml(members[name])
        parsed[name] = root
        for node in root.xpath("//w:t | //w:instrText", namespaces=namespaces):
            characters.update(node.text or "")
        for run in root.xpath("//w:r", namespaces=namespaces):
            text = "".join(
                node.text or ""
                for node in run.xpath(".//w:t | .//w:instrText", namespaces=namespaces)
            )
            if _contains_east_asian_text(text):
                has_east_asian_text = True
                _ensure_run_east_asia_font(run, font_family)

    if not has_east_asian_text:
        return False, characters

    numbering_name = "word/numbering.xml"
    if numbering_name in members:
        numbering = _parse_ooxml(members[numbering_name])
        parsed[numbering_name] = numbering
        for level in numbering.xpath(
            "//w:lvl[w:numFmt[@w:val='bullet']]",
            namespaces=namespaces,
        ):
            level_text = level.find(f"{{{_OOXML_WORD_NS}}}lvlText")
            if level_text is None:
                continue
            value = level_text.get(f"{{{_OOXML_WORD_NS}}}val", "")
            if any(0xE000 <= ord(char) <= 0xF8FF for char in value):
                value = "\u2022"
                level_text.set(f"{{{_OOXML_WORD_NS}}}val", value)
            characters.update(value.replace("%1", ""))
            r_pr = level.find(f"{{{_OOXML_WORD_NS}}}rPr")
            if r_pr is None:
                r_pr = etree.SubElement(level, f"{{{_OOXML_WORD_NS}}}rPr")
            r_fonts = r_pr.find(f"{{{_OOXML_WORD_NS}}}rFonts")
            if r_fonts is None:
                r_fonts = etree.SubElement(r_pr, f"{{{_OOXML_WORD_NS}}}rFonts")
            for attribute in ("ascii", "hAnsi", "eastAsia"):
                r_fonts.set(f"{{{_OOXML_WORD_NS}}}{attribute}", font_family)
            for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme"):
                r_fonts.attrib.pop(f"{{{_OOXML_WORD_NS}}}{attribute}", None)

    for name, root in parsed.items():
        members[name] = _serialize_ooxml(root)
    return True, characters


def _add_embedded_cjk_font(
    members: dict[str, bytes],
    font_family: str,
    font_data: bytes,
) -> str:
    required = {"word/fontTable.xml", "[Content_Types].xml", "word/settings.xml"}
    missing = sorted(required.difference(members))
    if missing:
        raise RuntimeError(f"DOCX is missing required OOXML parts: {', '.join(missing)}")

    font_key = uuid.uuid5(uuid.NAMESPACE_URL, hashlib.sha256(font_data).hexdigest())
    font_part_name = f"word/fonts/{font_key.hex}.odttf"
    relationship_id = "rIdPortableCjkFont"
    namespaces = {"w": _OOXML_WORD_NS}

    font_table = _parse_ooxml(members["word/fontTable.xml"])
    font_entry = next(
        (
            item
            for item in font_table.xpath("./w:font", namespaces=namespaces)
            if item.get(f"{{{_OOXML_WORD_NS}}}name", "").casefold()
            == font_family.casefold()
        ),
        None,
    )
    if font_entry is None:
        font_entry = etree.SubElement(font_table, f"{{{_OOXML_WORD_NS}}}font")
        font_entry.set(f"{{{_OOXML_WORD_NS}}}name", font_family)
    for old_embed in font_entry.findall(f"{{{_OOXML_WORD_NS}}}embedRegular"):
        font_entry.remove(old_embed)
    embed = etree.SubElement(font_entry, f"{{{_OOXML_WORD_NS}}}embedRegular")
    embed.set(f"{{{_OOXML_REL_NS}}}id", relationship_id)
    embed.set(f"{{{_OOXML_WORD_NS}}}fontKey", "{" + str(font_key).upper() + "}")
    members["word/fontTable.xml"] = _serialize_ooxml(font_table)

    relationships_name = "word/_rels/fontTable.xml.rels"
    if relationships_name in members:
        relationships = _parse_ooxml(members[relationships_name])
    else:
        relationships = etree.Element(
            f"{{{_OOXML_PACKAGE_REL_NS}}}Relationships",
            nsmap={None: _OOXML_PACKAGE_REL_NS},
        )
    existing_ids = {
        item.get("Id", "")
        for item in relationships.findall(f"{{{_OOXML_PACKAGE_REL_NS}}}Relationship")
    }
    if relationship_id in existing_ids:
        suffix = 2
        while f"{relationship_id}{suffix}" in existing_ids:
            suffix += 1
        relationship_id = f"{relationship_id}{suffix}"
        embed.set(f"{{{_OOXML_REL_NS}}}id", relationship_id)
        members["word/fontTable.xml"] = _serialize_ooxml(font_table)
    relationship = etree.SubElement(
        relationships,
        f"{{{_OOXML_PACKAGE_REL_NS}}}Relationship",
    )
    relationship.set("Id", relationship_id)
    relationship.set("Type", _OOXML_FONT_REL_TYPE)
    relationship.set("Target", f"fonts/{font_key.hex}.odttf")
    members[relationships_name] = _serialize_ooxml(relationships)

    content_types = _parse_ooxml(members["[Content_Types].xml"])
    defaults = content_types.findall(f"{{{_OOXML_CONTENT_TYPES_NS}}}Default")
    font_type = next(
        (item for item in defaults if item.get("Extension", "").casefold() == "odttf"),
        None,
    )
    if font_type is None:
        font_type = etree.SubElement(
            content_types,
            f"{{{_OOXML_CONTENT_TYPES_NS}}}Default",
        )
        font_type.set("Extension", "odttf")
    font_type.set("ContentType", _OOXML_OBFUSCATED_FONT_CONTENT_TYPE)
    members["[Content_Types].xml"] = _serialize_ooxml(content_types)

    settings = _parse_ooxml(members["word/settings.xml"])
    embed_setting = settings.find(f"{{{_OOXML_WORD_NS}}}embedTrueTypeFonts")
    if embed_setting is None:
        embed_setting = etree.SubElement(settings, f"{{{_OOXML_WORD_NS}}}embedTrueTypeFonts")
    embed_setting.set(f"{{{_OOXML_WORD_NS}}}val", "true")
    members["word/settings.xml"] = _serialize_ooxml(settings)
    members[font_part_name] = _obfuscate_ooxml_font(font_data, font_key)
    return font_part_name


def _rewrite_docx_package(
    output_path: Path,
    original_entries: list[zipfile.ZipInfo],
    members: dict[str, bytes],
) -> None:
    temporary = output_path.with_name(f".{output_path.name}.{uuid.uuid4().hex}.tmp")
    try:
        original_names = {entry.filename for entry in original_entries}
        with zipfile.ZipFile(temporary, "w") as archive:
            for entry in original_entries:
                archive.writestr(entry, members[entry.filename])
            for name in sorted(set(members).difference(original_names)):
                archive.writestr(name, members[name], compress_type=zipfile.ZIP_DEFLATED)
        os.replace(temporary, output_path)
    finally:
        if temporary.exists():
            temporary.unlink()


def _make_docx_cjk_portable(output_path: Path) -> bool:
    """Finalize CJK font declarations and embed a used-glyph portable fallback."""

    try:
        with zipfile.ZipFile(output_path, "r") as archive:
            original_entries = archive.infolist()
            members = {entry.filename: archive.read(entry.filename) for entry in original_entries}

        font_family = DEFAULT_DOC_EAST_ASIA_FONT
        has_cjk, characters = _patch_docx_cjk_xml(members, font_family)
        if not has_cjk:
            return False

        embedded = False
        font_bytes = 0
        if DEFAULT_DOC_EMBED_CJK_FONT:
            try:
                font_family, font_data = _build_cjk_font_subset(characters)
                # Reapply direct declarations if an operator-supplied font file
                # exposes a family name different from the requested default.
                _patch_docx_cjk_xml(members, font_family)
                _add_embedded_cjk_font(members, font_family, font_data)
                embedded = True
                font_bytes = len(font_data)
            except Exception as exc:
                logger.warning(
                    "CJK font embedding unavailable; keeping explicit East Asia font mapping | "
                    "path=%s error=%s",
                    output_path,
                    exc,
                )

        _rewrite_docx_package(output_path, original_entries, members)
        logger.info(
            "DOCX CJK portability finalized | path=%s family=%s embedded=%s font_bytes=%d",
            output_path,
            font_family,
            embedded,
            font_bytes,
        )
        return embedded
    except Exception as exc:
        logger.warning("DOCX CJK portability finalization failed | path=%s error=%s", output_path, exc)
        return False


def _rgb01(color: tuple[int, int, int]) -> tuple[float, float, float]:
    return (round(color[0] / 255, 4), round(color[1] / 255, 4), round(color[2] / 255, 4))


def _html_escape(value: Any) -> str:
    return html_lib.escape(str(value or ""), quote=True)


def _has_publication_content(pub: Any) -> bool:
    if isinstance(pub, dict):
        return any(str(pub.get(key, "")).strip() for key in ("venue", "title", "authors"))
    if isinstance(pub, str):
        return bool(pub.strip())
    return False


def _publication_line(pub: Any) -> str:
    if isinstance(pub, dict):
        venue = str(pub.get("venue", "")).strip()
        title = str(pub.get("title", "")).strip()
        authors = str(pub.get("authors", "")).strip()
        return " - ".join(part for part in (venue, title, authors) if part)
    if isinstance(pub, str):
        return pub.strip()
    return ""


def _normalize_text_items(value: Any) -> list[str]:
    if isinstance(value, str):
        text = value.strip()
        return [text] if text else []
    if not isinstance(value, list):
        return []
    result: list[str] = []
    for item in value:
        text = str(item).strip()
        if text:
            result.append(text)
    return result


def _collect_experience_bullets(exp: dict[str, Any]) -> list[str]:
    merged: list[str] = []
    for key in (
        "function_description",
        "result_description",
        "bullets",
        "responsibilities",
        "achievements",
        "highlights",
        "details",
        "description",
    ):
        merged.extend(_normalize_text_items(exp.get(key)))
    # Keep order while deduping.
    seen: set[str] = set()
    deduped: list[str] = []
    for item in merged:
        if item in seen:
            continue
        seen.add(item)
        deduped.append(item)
    return deduped


def _collect_text_entries(payload: Any, keys: tuple[str, ...]) -> list[str]:
    if not isinstance(payload, dict):
        return []
    merged: list[str] = []
    for key in keys:
        merged.extend(_normalize_text_items(payload.get(key)))
    seen: set[str] = set()
    deduped: list[str] = []
    for item in merged:
        if item in seen:
            continue
        seen.add(item)
        deduped.append(item)
    return deduped


def _collect_project_bullets(project: dict[str, Any]) -> list[str]:
    return _collect_text_entries(
        project,
        ("description", "function_description", "result_description", "bullets", "achievements", "highlights"),
    )


def _compress_resume_data_for_docx(resume_data: dict[str, Any]) -> dict[str, Any]:
    """Normalize render data without silently deleting verified content.

    Page count is a layout signal, not a content budget.  Any optional
    compaction happens only after inspecting a real rendered document.
    """
    data = copy.deepcopy(resume_data) if isinstance(resume_data, dict) else {}
    if isinstance(data.get("summary"), str):
        data["summary"] = data["summary"].strip()

    def _clean_texts(values: Any) -> list[str]:
        items = _normalize_text_items(values)
        result: list[str] = []
        for item in items:
            text = re.sub(r"\s+", " ", item).strip()
            if text and text not in result:
                result.append(text)
        return result

    for section_key in ("experience", "research", "campus_experience"):
        for exp in data.get(section_key, []) if isinstance(data.get(section_key), list) else []:
            if not isinstance(exp, dict):
                continue
            merged = _collect_experience_bullets(exp)
            exp["bullets"] = _clean_texts(merged)
            for proj in exp.get("projects", []) if isinstance(exp.get("projects"), list) else []:
                if isinstance(proj, dict):
                    proj["bullets"] = _clean_texts(_collect_project_bullets(proj))

    for proj in data.get("projects", []) if isinstance(data.get("projects"), list) else []:
        if isinstance(proj, dict):
            proj["bullets"] = _clean_texts(_collect_project_bullets(proj))

    skills = data.get("skills")
    if isinstance(skills, dict):
        for key, values in list(skills.items()):
            if isinstance(values, list):
                skills[key] = _clean_texts(values)

    for key in ("publications", "honors", "awards", "certifications", "personal_skills"):
        value = data.get(key)
        if isinstance(value, list):
            # Publication objects must keep their typed structure.
            if key == "publications":
                data[key] = [item for item in value if str(item).strip()]
            else:
                data[key] = _clean_texts(value)
    return data


def _tighten_resume_data_for_layout(resume_data: dict[str, Any]) -> dict[str, Any]:
    """Return retry data without deleting verified resume content.

    Multi-page resumes are valid. Layout retries may tighten typography and
    pagination rules, but must never shorten summaries, bullets, skills, or
    sections behind the user's back.
    """

    return copy.deepcopy(resume_data)


def _collect_profile_items(meta: Any) -> list[tuple[str, str]]:
    if not isinstance(meta, dict):
        return []
    mapping = [
        ("年龄", "age"),
        ("性别", "gender"),
        ("工作经验", "work_experience"),
        ("学历", "education_level"),
        ("政治面貌", "political_status"),
        ("期望城市", "expected_city"),
        ("目标岗位", "target_role"),
        ("求职意向", "job_intention"),
    ]
    rows: list[tuple[str, str]] = []
    for label, key in mapping:
        value = str(meta.get(key, "")).strip()
        if value:
            rows.append((label, value))
    return rows


def _resolve_avatar_path(meta: Any) -> Optional[Path]:
    if not isinstance(meta, dict):
        return None
    for key in ("avatar_path", "avatar", "photo_path", "portrait_path"):
        raw = str(meta.get(key, "") or "").strip()
        if not raw:
            continue
        path = Path(raw)
        if path.exists() and path.is_file():
            return path
    return None


def _avatar_image_uri(meta: Any) -> str:
    avatar_path = _resolve_avatar_path(meta)
    if avatar_path is None:
        return ""
    try:
        return avatar_path.resolve().as_uri()
    except Exception:
        return ""


def _append_docx_avatar(doc: DocxDocument, meta: Any, template: str) -> None:
    avatar_path = _resolve_avatar_path(meta)
    if avatar_path is None:
        return
    width_map = {
        "minimal": Inches(0.66),
        "classic": Inches(0.76),
        "modern": Inches(0.72),
    }
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run()
    try:
        run.add_picture(str(avatar_path), width=width_map.get(template, Inches(1.0)))
        para.space_after = Pt(4)
    except Exception as exc:
        logger.warning("Failed to insert avatar into DOCX | path=%s error=%s", avatar_path, exc)


def _append_docx_header_block(
    doc: DocxDocument,
    *,
    meta: Any,
    template: str,
    name: str,
    summary: str = "",
) -> None:
    avatar_path = _resolve_avatar_path(meta)
    has_avatar = avatar_path is not None

    # --- Avatar (centered above name when present) ---
    if avatar_path is not None:
        try:
            p_av = doc.add_paragraph()
            p_av.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_av = p_av.add_run()
            run_av.add_picture(str(avatar_path), width=Inches(0.8))
            p_av.space_after = Pt(4)
        except Exception as exc:
            logger.warning("Failed to insert header avatar into DOCX | path=%s error=%s", avatar_path, exc)

    # --- Name ---
    name_align = WD_ALIGN_PARAGRAPH.LEFT
    name_size = Pt(22)
    name_color = RGBColor(0x0F, 0x34, 0x60)

    if template == "classic":
        name_align = WD_ALIGN_PARAGRAPH.CENTER
        name_size = Pt(22)
        name_color = RGBColor(0x0F, 0x34, 0x60)
    elif template == "modern":
        name_size = Pt(22)
        name_color = RGBColor(0x00, 0x5A, 0x64)

    if str(name or "").strip():
        p_name = doc.add_paragraph()
        p_name.alignment = name_align
        run = p_name.add_run(name)
        run.bold = True
        run.font.size = name_size
        run.font.name = DEFAULT_DOC_FONT
        run.font.color.rgb = name_color
        p_name.space_after = Pt(2)

    # --- Contact line (horizontal) ---
    contacts = [
        v for k in ("email", "phone", "wechat", "github", "linkedin", "website")
        if isinstance(meta, dict) and (v := meta.get(k))
    ]
    if contacts:
        contact_sep = "  |  "
        p_contact = doc.add_paragraph(contact_sep.join(contacts))
        p_contact.alignment = name_align
        p_contact.paragraph_format.space_after = Pt(2)
        for item in p_contact.runs:
            item.font.size = Pt(9.5)
            item.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Profile tags (age/gender/edu/exp in one compact line) ---
    if isinstance(meta, dict):
        profile_bits = []
        for label, key in [
            ("性别", "gender"), ("年龄", "age"),
            ("学历", "education_level"), ("工作经验", "work_experience"),
        ]:
            val = str(meta.get(key, "")).strip()
            if val:
                profile_bits.append(f"{label}: {val}")
        if profile_bits:
            p_profile = doc.add_paragraph("  |  ".join(profile_bits))
            p_profile.alignment = name_align
            p_profile.paragraph_format.space_after = Pt(4)
            for item in p_profile.runs:
                item.font.size = Pt(9.5)
                item.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # --- Target role / intention (if set) ---
    if isinstance(meta, dict):
        target = str(meta.get("target_role", "") or meta.get("job_intention", "")).strip()
        if target:
            p_target = doc.add_paragraph(f"求职意向: {target}")
            p_target.alignment = name_align
            p_target.paragraph_format.space_after = Pt(6)
            for item in p_target.runs:
                item.font.size = Pt(10)
                item.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # --- Summary (if present) ---
    if isinstance(summary, str) and summary.strip():
        p_summary = doc.add_paragraph(summary.strip())
        p_summary.alignment = name_align
        p_summary.paragraph_format.space_after = Pt(8 if template != "classic" else 6)
        for item in p_summary.runs:
            item.font.size = Pt(10)
            item.font.color.rgb = RGBColor(0x44, 0x44, 0x44) if template == "classic" else RGBColor(0x33, 0x33, 0x33)
            if template == "classic":
                item.italic = True


def _collect_projects_from_experience(experience: Any) -> list[dict[str, Any]]:
    projects: list[dict[str, Any]] = []
    if not isinstance(experience, list):
        return projects

    for exp in experience:
        if not isinstance(exp, dict):
            continue

        company = str(exp.get("company", "")).strip()
        role = str(exp.get("role", "")).strip()
        period = str(exp.get("period", "")).strip()
        affiliation = " | ".join(part for part in [company, role, period] if part)

        exp_projects = exp.get("projects", [])
        if not isinstance(exp_projects, list):
            continue

        for proj in exp_projects:
            if not isinstance(proj, dict):
                continue
            name = str(proj.get("name", "项目")).strip() or "项目"
            bullets = _normalize_text_items(proj.get("bullets"))
            tech_stack = _normalize_text_items(proj.get("tech_stack"))
            for field_name in ("description", "function_description", "result_description"):
                value = str(proj.get(field_name, "")).strip()
                if value and value not in bullets:
                    bullets = [value] + bullets
            projects.append(
                {
                    "name": name,
                    "affiliation": affiliation,
                    "bullets": bullets,
                    "tech_stack": tech_stack,
                }
            )

    return projects


def _collect_projects_from_top_level(resume_data: Any) -> list[dict[str, Any]]:
    if not isinstance(resume_data, dict):
        return []

    projects: list[dict[str, Any]] = []
    candidate_keys = ("projects", "project_experience", "project_experiences")
    for key in candidate_keys:
        records = resume_data.get(key, [])
        if not isinstance(records, list):
            continue
        for proj in records:
            if not isinstance(proj, dict):
                continue
            name = str(proj.get("name", "项目")).strip() or "项目"
            period = str(proj.get("period", "")).strip()
            company = str(proj.get("company", "")).strip()
            role = str(proj.get("role", "")).strip()
            affiliation = " | ".join(part for part in (company, role, period) if part)

            bullets = _normalize_text_items(proj.get("bullets"))
            for field_name in ("description", "function_description", "result_description"):
                value = str(proj.get(field_name, "")).strip()
                if value and value not in bullets:
                    bullets = [value] + bullets
            tech_stack = _normalize_text_items(proj.get("tech_stack"))

            projects.append(
                {
                    "name": name,
                    "affiliation": affiliation,
                    "bullets": bullets,
                    "tech_stack": tech_stack,
                }
            )
    return projects


def _project_signature(project: dict[str, Any]) -> tuple[str, str, str]:
    name = _normalize_project_name(project.get("name", ""))
    affiliation = str(project.get("affiliation", "")).strip().lower()
    bullets = project.get("bullets", [])
    if isinstance(bullets, list):
        parts = [str(item).strip() for item in bullets[:2] if str(item).strip()]
    else:
        parts = []
    bullet_sig = " ".join(parts).lower()
    bullet_sig = re.sub(
        r"((19|20)\d{2}[./-]\d{1,2}\s*(?:[-–—至到~]\s*((19|20)\d{2}[./-]\d{1,2}|至今|present))?)",
        " ",
        bullet_sig,
        flags=re.IGNORECASE,
    )
    bullet_sig = re.sub(r"[-–—|:：()\[\]{}【】]+", " ", bullet_sig)
    bullet_sig = re.sub(r"\s+", " ", bullet_sig).strip()
    return name, affiliation, bullet_sig


def _is_duplicate_project(left: dict[str, Any], right: dict[str, Any]) -> bool:
    left_name, left_aff, left_sig = _project_signature(left)
    right_name, right_aff, right_sig = _project_signature(right)

    if left_name and right_name and left_name == right_name:
        if not left_sig or not right_sig or left_sig == right_sig:
            return True
        if left_sig and right_sig and (left_sig in right_sig or right_sig in left_sig):
            return True
        if left_aff and right_aff and left_aff == right_aff:
            return True

    if left_sig and right_sig and left_sig == right_sig:
        if not left_name or not right_name:
            return True
        if left_name in right_name or right_name in left_name:
            return True

    return False


def _collect_all_projects(resume_data: Any, experience: Any) -> list[dict[str, Any]]:
    merged = _collect_projects_from_experience(experience) + _collect_projects_from_top_level(resume_data)
    deduped: list[dict[str, Any]] = []
    for proj in merged:
        if not isinstance(proj, dict):
            continue
        if any(_is_duplicate_project(proj, existing) for existing in deduped):
            continue
        deduped.append(proj)
    return deduped


def _collect_all_experience(resume_data: Any) -> list[dict[str, Any]]:
    if not isinstance(resume_data, dict):
        return []

    merged: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str]] = set()
    candidate_keys = (
        "experience",
        "work_experience",
        "work_experiences",
        "internships",
        "internship_experience",
        "internship",
    )

    for key in candidate_keys:
        records = resume_data.get(key, [])
        if not isinstance(records, list):
            continue
        for record in records:
            if not isinstance(record, dict):
                continue
            company = str(record.get("company", "")).strip()
            role = str(record.get("role", "")).strip()
            period = str(record.get("period", "")).strip()
            dedupe_key = (company, role, period)
            if company or role or period:
                if dedupe_key in seen:
                    continue
                seen.add(dedupe_key)
            else:
                # All empty fields — still include, but no dedup key
                pass
            merged.append(record)

    return merged


def _build_resume_html(resume_data: dict[str, Any], template: str = "classic") -> str:
    tpl = _normalize_template(template)
    theme = {
        "classic": {
            "accent": "#0f3460",
            "title": "#10243f",
            "text": "#222222",
            "muted": "#5f6368",
            "divider": "#d8dde6",
            "bg": "#ffffff",
        },
        "modern": {
            "accent": "#005a64",
            "title": "#08373d",
            "text": "#1f2a2e",
            "muted": "#58666b",
            "divider": "#d1dde0",
            "bg": "#ffffff",
        },
        "minimal": {
            "accent": "#333333",
            "title": "#202020",
            "text": "#2a2a2a",
            "muted": "#616161",
            "divider": "#dddddd",
            "bg": "#ffffff",
        },
    }[tpl]

    meta = resume_data.get("meta", {}) if isinstance(resume_data, dict) else {}
    summary = resume_data.get("summary", "") if isinstance(resume_data, dict) else ""
    experience = _collect_all_experience(resume_data)
    project_items = _collect_all_projects(resume_data, experience)
    education = resume_data.get("education", []) if isinstance(resume_data, dict) else []
    skills = resume_data.get("skills", {}) if isinstance(resume_data, dict) else {}
    publications = resume_data.get("publications", []) if isinstance(resume_data, dict) else []

    name = _html_escape(meta.get("name", "Candidate")) if isinstance(meta, dict) else "Candidate"
    contacts: list[str] = []
    if isinstance(meta, dict):
        for key in ("email", "phone", "wechat", "github", "linkedin", "website"):
            value = str(meta.get(key) or "").strip()
            if value:
                contacts.append(_html_escape(value))
    contacts_html = " | ".join(contacts)
    avatar_uri = _avatar_image_uri(meta)
    profile_items = _collect_profile_items(meta)

    summary_html = ""
    if isinstance(summary, str) and summary.strip():
        summary_html = f'<p class="summary">{_html_escape(summary.strip())}</p>'

    exp_blocks: list[str] = []
    if isinstance(experience, list):
        for exp in experience:
            if not isinstance(exp, dict):
                continue
            company = _html_escape(exp.get("company", ""))
            role = _html_escape(exp.get("role", ""))
            period = _html_escape(exp.get("period", ""))
            head_parts = [part for part in [company, role, period] if part]
            if not head_parts:
                continue
            exp_bullets = _collect_experience_bullets(exp)
            exp_bullets_html = ""
            if exp_bullets:
                exp_bullets_html = (
                    '<ul class="bullets">'
                    + "".join(f"<li>{_html_escape(item)}</li>" for item in exp_bullets)
                    + "</ul>"
                )
            exp_blocks.append(
                f"""
                <article class="entry">
                  <h3>{" | ".join(head_parts)}</h3>
                  {exp_bullets_html}
                </article>
                """
            )

    project_blocks: list[str] = []
    for proj in project_items:
        proj_name = _html_escape(proj.get("name", "Project"))
        affiliation = _html_escape(proj.get("affiliation", ""))
        bullets_html = ""
        bullets = proj.get("bullets", [])
        if isinstance(bullets, list) and bullets:
            bullet_items = "".join(f"<li>{_html_escape(item)}</li>" for item in bullets if str(item).strip())
            if bullet_items:
                bullets_html = f'<ul class="bullets">{bullet_items}</ul>'
        tech_html = ""
        tech = proj.get("tech_stack", [])
        if isinstance(tech, list) and tech:
            clean_tech = [_html_escape(item) for item in tech if str(item).strip()]
            if clean_tech:
                tech_html = f'<p class="tech"><span>Tech</span>{", ".join(clean_tech)}</p>'

        project_blocks.append(
            f"""
            <article class="entry">
              <h4>{proj_name}</h4>
              {'<p class="project-meta">' + affiliation + '</p>' if affiliation else ''}
              {bullets_html}
              {tech_html}
            </article>
            """
        )

    edu_blocks: list[str] = []
    if isinstance(education, list):
        for edu in education:
            if not isinstance(edu, dict):
                continue
            school = _html_escape(edu.get("school", ""))
            degree = _html_escape(edu.get("degree", ""))
            major = _html_escape(edu.get("major", ""))
            period = _html_escape(edu.get("period", ""))
            head = " | ".join(part for part in [school, f"{degree} {major}".strip(), period] if part)
            if not head:
                continue
            highlights_html = ""
            highlights = edu.get("highlights", [])
            if isinstance(highlights, list) and highlights:
                rows = "".join(f"<li>{_html_escape(item)}</li>" for item in highlights if str(item).strip())
                if rows:
                    highlights_html = f'<ul class="bullets">{rows}</ul>'
            edu_blocks.append(f'<article class="entry"><h3>{head}</h3>{highlights_html}</article>')

    skill_blocks: list[str] = []
    if isinstance(skills, dict):
        labels = {
            "languages": "编程语言",
            "frameworks": "框架工具",
            "tools": "开发工具",
            "domains": "业务领域",
            "methodologies": "方法与流程",
            "certifications": "证书与资质",
            "natural_languages": "语言能力",
            "others": "其他专业技能",
        }
        for key, label in labels.items():
            items = skills.get(key, [])
            if isinstance(items, list) and items:
                tags = "".join(
                    f'<span class="tag">{_html_escape(item)}</span>' for item in items if str(item).strip()
                )
                if tags:
                    skill_blocks.append(
                        f"""
                        <div class="skill-row">
                          <div class="skill-label">{label}</div>
                          <div class="skill-tags">{tags}</div>
                        </div>
                        """
                    )

    pub_blocks: list[str] = []
    if isinstance(publications, list):
        for pub in publications:
            if not _has_publication_content(pub):
                continue
            line = _publication_line(pub)
            if line:
                pub_blocks.append(f"<li>{_html_escape(line)}</li>")

    honor_items = _collect_text_entries(
        resume_data if isinstance(resume_data, dict) else {},
        ("honors", "awards", "certifications"),
    )
    personal_skill_items = _collect_text_entries(
        resume_data if isinstance(resume_data, dict) else {},
        ("personal_skills",),
    )

    section_summary = ""
    if summary_html:
        section_summary = f'<section><h2>个人简介</h2>{summary_html}</section>'
    section_profile = ""
    if profile_items:
        profile_rows = "".join(
            f'<li><span class="meta-key">{_html_escape(k)}：</span>{_html_escape(v)}</li>'
            for k, v in profile_items
        )
        section_profile = f'<section><h2>求职信息</h2><ul class="bullets">{profile_rows}</ul></section>'
    empty_experience = '<p class="empty">暂无经历信息</p>'
    empty_projects = '<p class="empty">暂无项目信息</p>'
    empty_education = '<p class="empty">暂无教育信息</p>'
    section_exp = f'<section><h2>工作/实习经历</h2>{"".join(exp_blocks) or empty_experience}</section>'
    section_projects = f'<section><h2>项目经历</h2>{"".join(project_blocks) or empty_projects}</section>'
    section_edu = f'<section><h2>教育经历</h2>{"".join(edu_blocks) or empty_education}</section>'
    section_skills = ""
    if skill_blocks:
        section_skills = f'<section><h2>专业技能</h2>{"".join(skill_blocks)}</section>'
    section_pubs = ""
    if pub_blocks:
        section_pubs = f'<section><h2>论文成果</h2><ul class="bullets">{"".join(pub_blocks)}</ul></section>'
    section_honors = ""
    if honor_items:
        rows = "".join(f"<li>{_html_escape(item)}</li>" for item in honor_items)
        section_honors = f'<section><h2>荣誉与奖项</h2><ul class="bullets">{rows}</ul></section>'
    section_personal_skills = ""
    if personal_skill_items:
        rows = "".join(f"<li>{_html_escape(item)}</li>" for item in personal_skill_items)
        section_personal_skills = f'<section><h2>个人技能</h2><ul class="bullets">{rows}</ul></section>'
    section_additional = "".join(
        f'<section><h2>{_html_escape(title)}</h2><ul class="bullets">'
        + "".join(f"<li>{_html_escape(item)}</li>" for item in items)
        + "</ul></section>"
        for title, items in _collect_additional_sections(resume_data)
    )

    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <style>
    @page {{
      size: A4;
      margin: {PDF_MARGIN_MM};
    }}
    html {{
      background: #ffffff !important;
    }}
    * {{
      box-sizing: border-box;
    }}
    body {{
      margin: 0;
      background: #ffffff !important;
      font-family: "Noto Sans CJK SC", "Noto Sans CJK TC", "Noto Sans CJK JP", "WenQuanYi Zen Hei", "Microsoft YaHei", "PingFang SC", "Segoe UI", "Helvetica Neue", Arial, sans-serif;
      color: {theme["text"]};
      font-size: 12px;
      line-height: 1.55;
      -webkit-print-color-adjust: exact;
      print-color-adjust: exact;
    }}
    .resume-print {{
      width: 100%;
      max-width: none;
      margin: 0;
      background: #ffffff !important;
      padding: 0;
    }}
    .header {{
      display: flex;
      align-items: flex-start;
      justify-content: space-between;
      gap: 12px;
      padding-bottom: 12px;
      border-bottom: 2px solid {theme["divider"]};
      margin-bottom: 12px;
    }}
    .header-main {{
      flex: 1 1 auto;
      min-width: 0;
    }}
    .name {{
      margin: 0;
      color: {theme["title"]};
      font-size: 30px;
      line-height: 1.2;
      letter-spacing: 0.2px;
      font-weight: 700;
    }}
    .contacts {{
      margin-top: 8px;
      color: {theme["muted"]};
      font-size: 12px;
      overflow-wrap: anywhere;
      word-wrap: break-word;
    }}
    .avatar-wrap {{
      flex: 0 0 auto;
      width: 86px;
      text-align: right;
    }}
    .avatar {{
      width: 82px;
      height: 106px;
      object-fit: cover;
      border-radius: 3px;
      border: 1px solid {theme["divider"]};
      background: #fff;
    }}
    section {{
      margin-top: 10px;
      page-break-inside: auto;
      break-inside: auto;
    }}
    h2 {{
      margin: 0 0 8px;
      color: {theme["accent"]};
      font-size: 15px;
      border-bottom: 1px solid {theme["divider"]};
      padding-bottom: 4px;
      text-transform: uppercase;
      letter-spacing: 0.4px;
    }}
    h3 {{
      margin: 0 0 2px;
      color: {theme["title"]};
      font-size: 13px;
      font-weight: 600;
    }}
    h4 {{
      margin: 6px 0 2px;
      color: {theme["text"]};
      font-size: 12.5px;
      font-weight: 600;
    }}
    .summary {{
      margin: 0;
      color: {theme["muted"]};
      font-size: 12px;
      white-space: pre-wrap;
    }}
    .entry {{
      margin-bottom: 7px;
      page-break-inside: auto;
      break-inside: auto;
    }}
    .bullets {{
      margin: 3px 0 0 16px;
      padding: 0;
    }}
    .bullets li {{
      margin: 2px 0;
      white-space: pre-wrap;
    }}
    .tech {{
      margin: 4px 0 0;
      color: {theme["muted"]};
      font-size: 11px;
    }}
    .project-meta {{
      margin: 0 0 2px;
      color: {theme["muted"]};
      font-size: 11px;
    }}
    .project-list {{
      margin-top: 3px;
    }}
    .project-item {{
      margin-top: 4px;
      padding-left: 10px;
      border-left: 2px solid {theme["divider"]};
      break-inside: auto;
      page-break-inside: auto;
    }}
    .tech span {{
      color: {theme["accent"]};
      font-weight: 600;
      margin-right: 6px;
    }}
    .skill-row {{
      display: block;
      margin-bottom: 5px;
    }}
    .skill-label {{
      display: inline-block;
      width: auto;
      color: {theme["accent"]};
      font-weight: 600;
      font-size: 11.5px;
      margin-right: 6px;
    }}
    .skill-tags {{
      display: inline;
    }}
    .tag {{
      background: transparent;
      border: 0;
      border-radius: 0;
      padding: 0;
      font-size: 11px;
      margin-right: 4px;
    }}
    .empty {{
      margin: 0;
      color: {theme["muted"]};
      font-style: italic;
    }}
    .meta-key {{
      color: {theme["accent"]};
      font-weight: 600;
    }}
  </style>
</head>
<body>
  <main class="resume-print template-{tpl}">
    <header class="header">
      <div class="header-main">
        <h1 class="name">{name}</h1>
        {'<div class="contacts">' + contacts_html + '</div>' if contacts_html else ''}
      </div>
      {'<div class="avatar-wrap"><img class="avatar" src="' + _html_escape(avatar_uri) + '" alt="avatar" /></div>' if avatar_uri else ''}
    </header>
    {section_summary}
    {section_profile}
    {section_exp}
    {section_projects}
    {section_edu}
    {section_skills}
    {section_pubs}
    {section_honors}
    {section_personal_skills}
    {section_additional}
  </main>
</body>
</html>"""


def _render_pdf_via_weasyprint(resume_data: dict[str, Any], output_path: Path, template: str = "classic") -> bool:
    if WeasyHTML is None:
        return False

    html_content = _build_resume_html(resume_data, template=template)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        WeasyHTML(string=html_content, base_url=str(output_path.parent)).write_pdf(str(output_path))
        return output_path.exists() and output_path.stat().st_size > 0
    except Exception as exc:
        import logging
        logging.getLogger("resume_renderer").error("WeasyPrint PDF render failed: %s", exc, exc_info=True)
        return False


def _pdf_has_middle_blank_pages(pdf_path: Path) -> bool:
    """Detect blank pages introduced by unstable pagination in HTML->PDF engines."""
    if fitz is None or not pdf_path.exists():
        return False
    try:
        doc = fitz.open(str(pdf_path))
    except Exception:
        return False

    try:
        if doc.page_count < 3:
            return False
        for page_index in range(1, doc.page_count - 1):
            page = doc.load_page(page_index)
            text = (page.get_text("text") or "").strip()
            drawings = page.get_drawings() or []
            images = page.get_images(full=True) or []
            if len(text) < 12 and not drawings and not images:
                return True
        return False
    finally:
        doc.close()


def analyze_pdf_layout(pdf_path: Path) -> dict[str, Any]:
    """Inspect real rendered PDF pages for pagination problems.

    This is intentionally deterministic: no vision model and no quality score.
    It reports concrete pass/fail conditions that can trigger layout tightening.
    """

    report: dict[str, Any] = {
        "available": False,
        "page_count": 0,
        "page_char_counts": [],
        "blank_pages": [],
        "orphan_headings": [],
        "sparse_last_page": False,
        "first_page_has_core_section": False,
        "issues": [],
    }
    if fitz is None or not pdf_path.exists():
        report["issues"].append("pdf_layout_inspection_unavailable")
        return report
    try:
        doc = fitz.open(str(pdf_path))
    except Exception:
        report["issues"].append("pdf_open_failed")
        return report

    try:
        report["available"] = True
        report["page_count"] = doc.page_count
        page_texts: list[str] = []
        for index in range(doc.page_count):
            text = (doc.load_page(index).get_text("text") or "").strip()
            page_texts.append(text)
            count = len(re.sub(r"\s+", "", text))
            report["page_char_counts"].append(count)
            if count < 12:
                report["blank_pages"].append(index + 1)

        if page_texts:
            report["first_page_has_core_section"] = any(
                title in page_texts[0]
                for title in ("工作经历", "工作/实习经历", "实习经历", "科研经历", "项目经历")
            )
        for index, text in enumerate(page_texts[:-1]):
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            if lines and lines[-1] in _RESUME_SECTION_TITLES:
                report["orphan_headings"].append({"page": index + 1, "heading": lines[-1]})

        if len(report["page_char_counts"]) > 1:
            previous = max(report["page_char_counts"][:-1] or [1])
            last = report["page_char_counts"][-1]
            report["sparse_last_page"] = last < 140 or last < previous * 0.22

        if report["blank_pages"]:
            report["issues"].append("blank_page")
        if report["orphan_headings"]:
            report["issues"].append("orphan_heading")
        if report["sparse_last_page"]:
            report["issues"].append("sparse_last_page")
        if report["page_count"] > 1 and not report["first_page_has_core_section"]:
            report["issues"].append("core_experience_not_on_first_page")
        return report
    finally:
        doc.close()


def _preview_visual_layout(resume_data: dict[str, Any], template: str) -> dict[str, Any]:
    if os.getenv("ENABLE_RESUME_VISUAL_QA", "1").strip().lower() in {"0", "false", "no"}:
        return {"available": False, "issues": ["visual_qa_disabled"]}
    if WeasyHTML is None or fitz is None:
        return {"available": False, "issues": ["pdf_layout_inspection_unavailable"]}
    try:
        with tempfile.TemporaryDirectory() as directory:
            preview_path = Path(directory) / "layout-preview.pdf"
            if not _render_pdf_via_weasyprint(resume_data, preview_path, template=template):
                return {"available": False, "issues": ["preview_render_failed"]}
            return analyze_pdf_layout(preview_path)
    except Exception as exc:
        logger.warning("Visual layout preview failed: %s", exc)
        return {"available": False, "issues": ["preview_render_failed"]}


def inspect_docx_layout(docx_path: Path) -> dict[str, Any]:
    """Convert the actual DOCX with LibreOffice and inspect its real pages."""

    if os.getenv("ENABLE_RESUME_VISUAL_QA", "1").strip().lower() in {"0", "false", "no"}:
        return {"available": False, "renderer": "docx", "issues": ["visual_qa_disabled"]}
    if not shutil.which("libreoffice"):
        return {"available": False, "renderer": "docx", "issues": ["libreoffice_unavailable"]}
    try:
        with tempfile.TemporaryDirectory() as directory:
            pdf_path = Path(directory) / "docx-layout.pdf"
            if not _convert_docx_to_pdf(docx_path, pdf_path):
                return {"available": False, "renderer": "docx", "issues": ["docx_conversion_failed"]}
            report = analyze_pdf_layout(pdf_path)
            report["renderer"] = "libreoffice"
            return report
    except Exception as exc:
        logger.warning("Actual DOCX layout inspection failed: %s", exc)
        return {"available": False, "renderer": "docx", "issues": ["docx_conversion_failed"]}


def _layout_needs_tightening(report: dict[str, Any]) -> bool:
    return bool(report.get("available")) and any(
        issue in report.get("issues", [])
        for issue in (
            "core_experience_not_on_first_page",
            "blank_page", "sparse_last_page",
        )
    )


def _layout_retry_data(
    resume_data: dict[str, Any],
    report: dict[str, Any],
) -> dict[str, Any]:
    """Choose whether a retry may compact content or typography only."""

    issues = set(report.get("issues", []))
    # All layout problems are typography/pagination problems. Preserve every
    # verified item and let the retry use slightly tighter typography instead.
    return copy.deepcopy(resume_data)


def _apply_docx_compact_typography(doc: "DocxDocument") -> None:
    """Fit a nearly empty tail page without deleting resume content."""

    for section in doc.sections:
        section.top_margin = Inches(0.48)
        section.bottom_margin = Inches(0.48)
        section.left_margin = Inches(0.66)
        section.right_margin = Inches(0.66)
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0.5)
        paragraph.paragraph_format.line_spacing = 1.08
        for run in paragraph.runs:
            if run.font.size is not None and run.font.size.pt >= 9.5:
                run.font.size = Pt(max(9.0, run.font.size.pt - 0.5))


def _fitz_font_for_text(text: str) -> str:
    if _contains_cjk(text):
        # CJK fallback font name used by PyMuPDF on many platforms.
        return "china-s"
    return "helv"


def _measure_fitz_text(text: str, fontname: str, fontsize: float) -> float:
    if not text:
        return 0.0
    try:
        return float(fitz.get_text_length(text, fontname=fontname, fontsize=fontsize))
    except Exception:
        # Approximate width fallback: CJK ~1.0em, Latin ~0.55em
        width = 0.0
        for ch in text:
            width += fontsize * (1.0 if _contains_cjk(ch) else 0.55)
        return width


def _wrap_fitz_text(text: str, max_width: float, fontname: str, fontsize: float) -> list[str]:
    content = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    wrapped: list[str] = []
    for raw_line in content.split("\n"):
        line = raw_line.strip()
        if not line:
            wrapped.append("")
            continue
        buf = ""
        for ch in line:
            trial = f"{buf}{ch}"
            if buf and _measure_fitz_text(trial, fontname, fontsize) > max_width:
                wrapped.append(buf.rstrip())
                buf = ch.lstrip() if ch == " " else ch
            else:
                buf = trial
        if buf:
            wrapped.append(buf.rstrip())
    return wrapped or [""]


def _add_section_heading(doc: DocxDocument, title: str, template: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    run = p.add_run(title)

    if template == "minimal":
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = DEFAULT_DOC_FONT
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        p.space_before = Pt(10)
        p.space_after = Pt(4)
        return
    if template == "modern":
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = DEFAULT_DOC_FONT
        run.font.color.rgb = RGBColor(0x00, 0x5A, 0x64)
        p.space_before = Pt(12)
        p.space_after = Pt(5)
        return

    run.bold = True
    run.font.size = Pt(13)
    run.font.name = DEFAULT_DOC_FONT
    run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
    p.space_before = Pt(12)
    p.space_after = Pt(4)


def _apply_docx_pagination_guards(doc: DocxDocument) -> None:
    """Prevent section/record headings and bullets from being split poorly."""

    paragraphs = list(doc.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        style_name = str(getattr(paragraph.style, "name", "") or "")
        if text in _RESUME_SECTION_TITLES:
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        if "List Bullet" in style_name:
            paragraph.paragraph_format.keep_together = True
        if index + 1 < len(paragraphs):
            next_style = str(getattr(paragraphs[index + 1].style, "name", "") or "")
            # A non-bullet immediately followed by a bullet is normally a
            # company/project heading and should stay with its first detail.
            if text and "List Bullet" not in style_name and "List Bullet" in next_style:
                paragraph.paragraph_format.keep_with_next = True


def _append_docx_profile_section(doc: DocxDocument, meta: Any, template: str) -> None:
    items = _collect_profile_items(meta)
    if not items:
        return
    _add_section_heading(doc, "求职信息", template)
    for key, value in items:
        p = doc.add_paragraph(f"{key}: {value}")
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _append_docx_text_section(doc: DocxDocument, title: str, items: list[str], template: str) -> None:
    clean = [str(x).strip() for x in items if str(x).strip()]
    if not clean:
        return
    _add_section_heading(doc, title, template)
    for item in clean:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def _collect_additional_sections(resume_data: dict[str, Any]) -> list[tuple[str, list[str]]]:
    """Collect extensible, evidence-backed sections for every renderer."""

    sections: list[tuple[str, list[str]]] = []
    for key, title in (("training", "培训与进修"), ("teaching", "教学经历")):
        values = resume_data.get(key, []) if isinstance(resume_data, dict) else []
        if isinstance(values, list):
            clean = [str(value).strip() for value in values if str(value).strip()]
            if clean:
                sections.append((title, clean))
    additional = resume_data.get("additional_sections", {}) if isinstance(resume_data, dict) else {}
    if isinstance(additional, dict):
        for title, values in additional.items():
            normalized_title = str(title).strip()
            if (
                not normalized_title
                or _INTERNAL_ADDITIONAL_SECTION.fullmatch(normalized_title)
                or not isinstance(values, list)
            ):
                continue
            clean = [str(value).strip() for value in values if str(value).strip()]
            if clean:
                sections.append((normalized_title, clean))
    return sections


def _clear_docx_body_for_style_template(doc: DocxDocument) -> None:
    """Remove example body content while retaining the template package.

    Headers, footers, page geometry, theme, styles, numbering and section
    properties live outside (or at the end of) the body content.  Rendering
    into this sanitized document therefore follows a static DOCX style guide
    substantially more closely than discarding it and creating a new file.
    """

    body = doc._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _capture_unanchored_table_shell(doc: DocxDocument) -> Any:
    """Keep one bounded table shell while removing all sample-person content."""

    candidates = [
        table for table in doc.tables
        if table.rows and 1 <= len(table.columns) <= 4
    ]
    if not candidates:
        return None
    table = candidates[0]
    shell = copy.deepcopy(table._tbl)
    rows = shell.findall(qn("w:tr"))
    for extra in rows[1:]:
        shell.remove(extra)
    for node in shell.iter(qn("w:t")):
        node.text = ""
    # A portrait in an example template belongs to the sample person.  Keep the
    # cell geometry and drawing-compatible package, but never copy that image.
    for tag in ("w:drawing", "w:pict", "w:object"):
        for node in list(shell.iter(qn(tag))):
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)
    return shell


def _set_paragraph_text_preserving_style(paragraph: Any, value: str, *, bold: bool = False) -> None:
    text = str(value or "")
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run(text)
    if bold:
        run.bold = True


def _install_table_header_shell(
    doc: DocxDocument,
    shell: Any,
    resume_data: dict[str, Any],
) -> bool:
    """Reuse an unanchored DOCX table as a sanitized candidate header."""

    if shell is None:
        return False
    body = doc._element.body
    section_properties = body.find(qn("w:sectPr"))
    if section_properties is None:
        body.append(shell)
    else:
        section_properties.addprevious(shell)
    table = doc.tables[-1]
    unique_cells: list[Any] = []
    seen: set[int] = set()
    for cell in table.rows[0].cells:
        marker = id(cell._tc)
        if marker in seen:
            continue
        seen.add(marker)
        unique_cells.append(cell)
    if not unique_cells:
        return False

    meta = resume_data.get("meta", {}) if isinstance(resume_data, dict) else {}
    if not isinstance(meta, dict):
        meta = {}
    name = str(meta.get("name", "") or "").strip()
    contacts = [
        str(meta.get(key, "") or "").strip()
        for key in ("phone", "email", "wechat", "github", "linkedin", "website")
        if str(meta.get(key, "") or "").strip()
    ]
    profile = [
        f"{label}: {value}"
        for label, key in (
            ("目标岗位", "target_role"), ("工作经验", "work_experience"),
            ("学历", "education_level"), ("期望城市", "expected_city"),
        )
        if (value := str(meta.get(key, "") or "").strip())
    ]
    values = [name]
    if len(unique_cells) == 1:
        values = ["\n".join(item for item in (name, " | ".join(contacts), " | ".join(profile)) if item)]
    elif len(unique_cells) == 2:
        values = [name, "\n".join(item for item in (" | ".join(contacts), " | ".join(profile)) if item)]
    else:
        values = [name, " | ".join(profile), " | ".join(contacts)]
    values.extend([""] * (len(unique_cells) - len(values)))
    for index, cell in enumerate(unique_cells):
        paragraphs = cell.paragraphs or [cell.add_paragraph()]
        _set_paragraph_text_preserving_style(
            paragraphs[0], values[index], bold=(index == 0 and bool(name)),
        )
        for paragraph in paragraphs[1:]:
            _set_paragraph_text_preserving_style(paragraph, "")
    return True


def _resume_data_without_header_fields(resume_data: dict[str, Any]) -> dict[str, Any]:
    data = copy.deepcopy(resume_data)
    meta = data.get("meta")
    if isinstance(meta, dict):
        for key in (
            "name", "phone", "email", "wechat", "github", "linkedin", "website",
            "target_role", "work_experience", "education_level", "expected_city",
            "age", "gender", "political_status", "job_intention",
        ):
            meta[key] = ""
    return data


def _all_template_paragraphs(doc: DocxDocument) -> list[Any]:
    paragraphs: list[Any] = list(doc.paragraphs)
    seen = {paragraph._element for paragraph in paragraphs}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph._element in seen:
                        continue
                    seen.add(paragraph._element)
                    paragraphs.append(paragraph)
    return paragraphs


def _all_template_story_paragraphs(doc: DocxDocument) -> list[Any]:
    """Return body, header and footer paragraphs once per XML node."""

    paragraphs = _all_template_paragraphs(doc)
    seen = {paragraph._element for paragraph in paragraphs}
    # Accessing ``section.header.paragraphs`` creates a blank part when the
    # source package has no header. Walk only relationships that already exist
    # so reading a template is side-effect free.
    for relationship in doc.part.rels.values():
        if relationship.reltype not in {RT.HEADER, RT.FOOTER}:
            continue
        part = relationship.target_part
        for element in part.element.iter(qn("w:p")):
            if element in seen:
                continue
            seen.add(element)
            paragraphs.append(Paragraph(element, doc))
    return paragraphs


def _template_mode(doc: DocxDocument) -> str:
    body_paragraphs = _all_template_paragraphs(doc)
    story_paragraphs = _all_template_story_paragraphs(doc)
    if any(
        _TEMPLATE_SECTION_TAG_RE.fullmatch(paragraph.text or "")
        for paragraph in body_paragraphs
    ) or any(
        _TEMPLATE_SCALAR_TAG_RE.search(paragraph.text or "")
        for paragraph in story_paragraphs
    ):
        return "tagged"
    if any(
        any(
            _template_section_matches(paragraph.text.strip(), title)
            for title in _TEMPLATE_SECTION_ALIASES
        )
        for paragraph in body_paragraphs
    ):
        return "anchored"
    return "style_only"


def detect_docx_template_mode(template: str | Path) -> str:
    """Classify a user DOCX without mutating it."""

    try:
        path = Path(template)
        if path.is_file() and path.suffix.lower() == ".docx":
            return _template_mode(DocxDocument(str(path)))
    except Exception as exc:
        logger.warning("DOCX template mode detection failed | template=%s error=%s", template, exc)
    return "style_only"


def _xml_element_text(element: Any) -> str:
    return "".join(str(node.text or "") for node in element.iter(qn("w:t"))).strip()


def _set_xml_paragraph_text(
    paragraph_element: Any,
    value: str,
    *,
    bold: bool = False,
) -> None:
    """Replace a tagged prototype's text while retaining paragraph/run style."""

    for existing_text in paragraph_element.iter(qn("w:t")):
        existing_text.text = ""
    runs = list(paragraph_element.findall(qn("w:r")))
    if runs:
        first_run = runs[0]
    else:
        first_run = OxmlElement("w:r")
        paragraph_element.append(first_run)
        runs = [first_run]
    for run in runs:
        for child in list(run):
            if child.tag != qn("w:rPr"):
                run.remove(child)
    run_properties = first_run.find(qn("w:rPr"))
    if bold:
        if run_properties is None:
            run_properties = OxmlElement("w:rPr")
            first_run.insert(0, run_properties)
        if run_properties.find(qn("w:b")) is None:
            run_properties.append(OxmlElement("w:b"))
    text_element = OxmlElement("w:t")
    text = str(value or "")
    if text.startswith(" ") or text.endswith(" "):
        text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_element.text = text
    first_run.append(text_element)


def _section_item_records(items: list[Any]) -> list[list[tuple[str, bool]]]:
    records: list[list[tuple[str, bool]]] = []
    for item in items:
        if isinstance(item, tuple) and len(item) == 2:
            heading, values = item
            lines = [(str(heading).strip(), True)] if str(heading).strip() else []
            if isinstance(values, list):
                lines.extend(
                    (str(value).strip(), False)
                    for value in values
                    if str(value).strip()
                )
            if lines:
                records.append(lines)
        elif str(item).strip():
            records.append([(str(item).strip(), False)])
    return records


def _paragraph_section_tag(paragraph_element: Any) -> Optional[str]:
    match = _TEMPLATE_SECTION_TAG_RE.fullmatch(_xml_element_text(paragraph_element))
    return match.group(1).casefold() if match else None


def _fill_tagged_paragraph_prototype(
    prototype: Any,
    lines: list[tuple[str, bool]],
) -> list[Any]:
    output: list[Any] = []
    for text, bold in lines:
        paragraph = copy.deepcopy(prototype)
        _set_xml_paragraph_text(paragraph, text, bold=bold)
        output.append(paragraph)
    return output


def _apply_tagged_section_anchors(
    doc: DocxDocument,
    sections: list[tuple[str, list[Any]]],
) -> set[str]:
    """Expand explicit section anchors in place without changing the shell."""

    section_values = {title: values for title, values in sections}
    consumed: set[str] = set()
    processed_rows: set[int] = set()
    for paragraph in list(_all_template_paragraphs(doc)):
        key = _paragraph_section_tag(paragraph._element)
        if key is None:
            continue
        title = _TEMPLATE_SECTION_TAGS.get(key)
        if title is None:
            logger.warning("Unsupported DOCX section anchor removed | anchor=section.%s", key)
            parent = paragraph._element.getparent()
            if parent is not None:
                parent.remove(paragraph._element)
            continue
        consumed.add(title)
        records = _section_item_records(section_values.get(title, []))
        parent = paragraph._element.getparent()
        row = parent.getparent() if parent is not None and parent.tag == qn("w:tc") else None
        if row is not None and row.tag == qn("w:tr"):
            marker = id(row)
            if marker in processed_rows:
                continue
            processed_rows.add(marker)
            row_parent = row.getparent()
            if row_parent is None:
                continue
            for record in records:
                cloned_row = copy.deepcopy(row)
                tagged_paragraphs = [
                    candidate for candidate in cloned_row.iter(qn("w:p"))
                    if _paragraph_section_tag(candidate) == key
                ]
                tagged_paragraph = tagged_paragraphs[0] if tagged_paragraphs else None
                if tagged_paragraph is None:
                    continue
                # A repeating row is a formatting prototype, never a source of
                # candidate facts. Clear sample text/portraits from every cell
                # while retaining row, cell, paragraph and run properties.
                for candidate in cloned_row.iter(qn("w:p")):
                    if candidate is not tagged_paragraph:
                        _set_xml_paragraph_text(candidate, "")
                for tag in ("w:drawing", "w:pict", "w:object"):
                    for node in list(cloned_row.iter(qn(tag))):
                        node_parent = node.getparent()
                        if node_parent is not None:
                            node_parent.remove(node)
                inserted = _fill_tagged_paragraph_prototype(tagged_paragraph, record)
                for replacement in inserted:
                    tagged_paragraph.addprevious(replacement)
                tagged_paragraph.getparent().remove(tagged_paragraph)
                row.addprevious(cloned_row)
            row_parent.remove(row)
            continue

        if parent is None:
            continue
        flattened = [line for record in records for line in record]
        for replacement in _fill_tagged_paragraph_prototype(paragraph._element, flattened):
            paragraph._element.addprevious(replacement)
        parent.remove(paragraph._element)
    return consumed


def _paragraphs_before_first_template_section(
    doc: DocxDocument,
    section_headings: list[str],
) -> list[Any]:
    wrappers = {
        paragraph._element: paragraph
        for paragraph in _all_template_paragraphs(doc)
    }
    recognized_headings = list(dict.fromkeys([
        *section_headings,
        *_TEMPLATE_SECTION_ALIASES.keys(),
    ]))
    result: list[Any] = []
    for child in doc._element.body:
        elements = [child] if child.tag == qn("w:p") else list(child.iter(qn("w:p")))
        for element in elements:
            text = _xml_element_text(element)
            if _TEMPLATE_SECTION_TAG_RE.fullmatch(text) or any(
                _template_section_matches(text, heading) for heading in recognized_headings
            ):
                return result
            wrapper = wrappers.get(element)
            if wrapper is not None:
                result.append(wrapper)
    return result


def _sanitize_template_identity_block(
    doc: DocxDocument,
    resume_data: dict[str, Any],
    section_headings: list[str],
) -> None:
    """Remove sample-person text and portraits before the first section."""

    paragraphs = _paragraphs_before_first_template_section(doc, section_headings)
    for paragraph in paragraphs:
        for tag in ("w:drawing", "w:pict", "w:object"):
            for node in list(paragraph._element.iter(qn(tag))):
                parent = node.getparent()
                if parent is not None:
                    parent.remove(node)

    decorative = {"个人简历", "求职简历", "resume", "cv", "curriculumvitae"}
    candidates = [
        paragraph for paragraph in paragraphs
        if paragraph.text.strip()
        and re.sub(r"\W+", "", paragraph.text).casefold() not in decorative
    ]
    meta = resume_data.get("meta", {}) if isinstance(resume_data, dict) else {}
    if not isinstance(meta, dict):
        meta = {}
    name = str(meta.get("name", "") or "").strip()
    contact = " | ".join(
        str(meta.get(key, "") or "").strip()
        for key in ("phone", "email", "wechat", "github", "linkedin", "website")
        if str(meta.get(key, "") or "").strip()
    )
    target = str(meta.get("target_role", "") or "").strip()
    values = [value for value in (name, contact, target) if value]
    for index, paragraph in enumerate(candidates):
        _set_paragraph_text_preserving_style(
            paragraph,
            values[index] if index < len(values) else "",
            bold=(index == 0 and bool(name)),
        )


def _structured_template_allowed_fragments(
    resume_data: dict[str, Any],
    sections: list[tuple[str, list[Any]]],
) -> set[str]:
    fragments: set[str] = set()
    meta = resume_data.get("meta", {}) if isinstance(resume_data, dict) else {}
    if isinstance(meta, dict):
        fragments.update(
            str(value).strip() for value in meta.values()
            if isinstance(value, (str, int, float)) and str(value).strip()
        )
    summary = str(resume_data.get("summary", "") or "").strip()
    if summary:
        fragments.add(summary)
    for heading, items in sections:
        fragments.add(str(heading).strip())
        for record in _section_item_records(items):
            fragments.update(text for text, _bold in record if text)
    return fragments


def _sanitize_structured_template_residuals(
    doc: DocxDocument,
    resume_data: dict[str, Any],
    sections: list[tuple[str, list[Any]]],
) -> None:
    """Remove example-person body content left outside recognized anchors."""

    allowed = _structured_template_allowed_fragments(resume_data, sections)
    decorative = {"个人简历", "求职简历", "resume", "cv", "curriculumvitae"}
    safe_label = re.compile(
        r"^(?:姓名|电话|手机|邮箱|微信|联系方式|求职意向|目标岗位|时间|公司|岗位|职责|项目)\s*[:：]?$",
        re.IGNORECASE,
    )
    for paragraph in list(_all_template_paragraphs(doc)):
        element = paragraph._element
        text = paragraph.text.strip()
        compact = re.sub(r"\W+", "", text).casefold()
        is_heading = any(
            _template_section_matches(text, heading)
            for heading in _TEMPLATE_SECTION_ALIASES
        )
        is_supported = any(
            fragment and (fragment in text or text in fragment)
            for fragment in allowed
        )
        keep = bool(
            not text
            or compact in decorative
            or safe_label.fullmatch(text)
            or is_heading
            or is_supported
        )

        # Body drawings are commonly the sample candidate's portrait. Header
        # and footer relationships are deliberately outside this traversal.
        for tag in ("w:drawing", "w:pict", "w:object"):
            for node in list(element.iter(qn(tag))):
                node_parent = node.getparent()
                if node_parent is not None:
                    node_parent.remove(node)
        if not text:
            # Empty sample spacer paragraphs may still carry hard page breaks.
            # Keep ordinary spacing but remove constructs that can create an
            # otherwise blank page in the middle of the generated resume.
            for node in list(element.iter(qn("w:br"))):
                node_parent = node.getparent()
                if node_parent is not None:
                    node_parent.remove(node)
            paragraph_properties = element.find(qn("w:pPr"))
            if paragraph_properties is not None:
                page_break = paragraph_properties.find(qn("w:pageBreakBefore"))
                if page_break is not None:
                    paragraph_properties.remove(page_break)
        if keep:
            continue
        parent = element.getparent()
        if parent is not None and parent.tag == qn("w:body"):
            parent.remove(element)
        else:
            _set_xml_paragraph_text(element, "")


def _is_known_template_heading(text: str) -> bool:
    return any(
        _template_section_matches(text, heading)
        for heading in _TEMPLATE_SECTION_ALIASES
    )


def _remove_empty_structured_headings(doc: DocxDocument) -> None:
    """Drop headings whose sample body was cleared and received no facts."""

    body = doc._element.body
    children = list(body)

    def paragraph_elements(element: Any) -> list[Any]:
        return [element] if element.tag == qn("w:p") else list(element.iter(qn("w:p")))

    def has_heading(element: Any) -> bool:
        return any(
            _is_known_template_heading(_xml_element_text(paragraph))
            for paragraph in paragraph_elements(element)
        )

    def has_non_heading_text(element: Any) -> bool:
        return any(
            bool(text := _xml_element_text(paragraph))
            and not _is_known_template_heading(text)
            for paragraph in paragraph_elements(element)
        )

    for index, child in enumerate(children):
        if child.tag != qn("w:p") or not has_heading(child):
            continue
        populated = False
        for sibling in children[index + 1:]:
            if sibling.tag == qn("w:sectPr") or has_heading(sibling):
                break
            if has_non_heading_text(sibling):
                populated = True
                break
        if not populated and child.getparent() is body:
            body.remove(child)

    # A section heading may live in one cell while its generated facts occupy
    # a companion cell or following rows. Empty table headings are cleared but
    # the row/grid shell remains intact.
    for table in doc.tables:
        rows = list(table.rows)
        for index, row in enumerate(rows):
            headings = [
                paragraph
                for cell in row.cells
                for paragraph in cell.paragraphs
                if _is_known_template_heading(paragraph.text.strip())
            ]
            if not headings:
                continue
            current_content = any(
                paragraph.text.strip()
                and paragraph not in headings
                and not _is_known_template_heading(paragraph.text.strip())
                for cell in row.cells
                for paragraph in cell.paragraphs
            )
            populated = current_content
            for following in rows[index + 1:]:
                if any(
                    _is_known_template_heading(paragraph.text.strip())
                    for cell in following.cells
                    for paragraph in cell.paragraphs
                ):
                    break
                if any(
                    paragraph.text.strip()
                    for cell in following.cells
                    for paragraph in cell.paragraphs
                ):
                    populated = True
                    break
            if not populated:
                for heading in headings:
                    _set_xml_paragraph_text(heading._element, "")


_SECTION_GEOMETRY_TAGS = (
    "w:pgSz",
    "w:pgMar",
    "w:cols",
    "w:docGrid",
    "w:mirrorMargins",
    "w:gutterAtTop",
    "w:paperSrc",
)


def _capture_section_geometry(doc: DocxDocument) -> list[dict[str, Any]]:
    snapshots: list[dict[str, Any]] = []
    for section in doc.sections:
        snapshots.append({
            tag: copy.deepcopy(node)
            for tag in _SECTION_GEOMETRY_TAGS
            if (node := section._sectPr.find(qn(tag))) is not None
        })
    return snapshots


def _restore_section_geometry(
    doc: DocxDocument,
    snapshots: list[dict[str, Any]],
) -> None:
    for section, snapshot in zip(doc.sections, snapshots):
        section_properties = section._sectPr
        for tag in _SECTION_GEOMETRY_TAGS:
            current = section_properties.find(qn(tag))
            if current is not None:
                section_properties.remove(current)
            saved = snapshot.get(tag)
            if saved is not None:
                section_properties.append(copy.deepcopy(saved))


def _template_supports_structured_injection(doc: DocxDocument) -> bool:
    """Return whether a DOCX exposes anchors that can be filled safely.

    A table by itself is not an anchor.  Treating every table-based document
    as injectable retained example names and stale rows, while appending the
    generated resume underneath.  Require an explicit placeholder or a known
    section heading in either normal paragraphs or table cells.
    """

    return _template_mode(doc) != "style_only"


def _initialize_docx_defaults(doc: DocxDocument, *, reset_geometry: bool) -> None:
    normal_style = doc.styles["Normal"]
    if not normal_style.font.name:
        normal_style.font.name = DEFAULT_DOC_FONT
    if normal_style.font.size is None:
        normal_style.font.size = Pt(10.5)
    if normal_style.paragraph_format.space_after is None:
        normal_style.paragraph_format.space_after = Pt(2)
    if normal_style.paragraph_format.line_spacing is None:
        normal_style.paragraph_format.line_spacing = 1.2

    if reset_geometry:
        section = doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def _render_docx_with_preset(
    doc: DocxDocument,
    resume_data: dict[str, Any],
    preset: str,
    framework_mode: bool,
) -> None:
    selected = preset if preset in SUPPORTED_TEMPLATES else "classic"
    if framework_mode:
        _render_docx_empty_profile_framework(doc, resume_data, selected)
    elif selected == "minimal":
        _render_docx_minimal(doc, resume_data)
    elif selected == "modern":
        _render_docx_modern(doc, resume_data)
    else:
        _render_docx_classic(doc, resume_data)


def _render_docx_empty_profile_framework(
    doc: DocxDocument,
    resume_data: dict[str, Any],
    template: str,
) -> None:
    """Render a useful skeleton while keeping placeholders outside fact fields."""

    framework = resume_data.get("framework", {})
    target_role = str(
        framework.get("target_role")
        or (resume_data.get("meta", {}) or {}).get("target_role", "")
    ).strip()
    _append_docx_header_block(
        doc,
        meta={"target_role": target_role},
        template=template,
        name="个人简历框架",
        summary=str(framework.get("notice", "以下内容均为待填写结构，不代表候选人已有事实。")),
    )
    sections = framework.get("sections", [])
    if not isinstance(sections, list):
        return
    detail_prompts = {
        "summary": "建议用2–4个完整句子概括真实背景、核心优势和目标方向。",
        "experience": "建议填写3–5条：个人职责、采取的方法、交付物和可核验结果。",
        "projects": "建议填写2–4条：项目背景、个人动作、使用方法和真实成果。",
        "skills": "仅填写实际掌握的工具、技术、行业方法、证书和语言能力。",
    }
    for section in sections:
        if not isinstance(section, dict):
            continue
        title = str(section.get("title", "")).strip()
        fields = [str(item).strip() for item in section.get("fields", []) if str(item).strip()]
        if not title:
            continue
        _add_section_heading(doc, title, template)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("[待填写] " + " ｜ ".join(fields))
        run.italic = True
        run.font.size = Pt(9.8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        key = str(section.get("key", ""))
        if key in detail_prompts:
            hint = doc.add_paragraph(style="List Bullet")
            hint_run = hint.add_run(detail_prompts[key])
            hint_run.font.size = Pt(9.5)
            hint_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _append_docx_record_section(
    doc: DocxDocument,
    resume_data: dict[str, Any],
    key: str,
    title: str,
    template: str,
) -> None:
    records = resume_data.get(key, []) if isinstance(resume_data, dict) else []
    if not isinstance(records, list) or not records:
        return
    _add_section_heading(doc, title, template)
    for record in records:
        if not isinstance(record, dict):
            continue
        company = str(record.get("company") or record.get("organization") or "").strip()
        role = str(record.get("role") or record.get("topic") or "").strip()
        period = str(record.get("period") or "").strip()
        heading = " | ".join(item for item in (company, role, period) if item)
        if heading:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(heading)
            run.bold = True
            run.font.size = Pt(10.5)
        for bullet in _collect_experience_bullets(record):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(str(bullet))


def _render_docx_minimal(doc: DocxDocument, resume_data: dict[str, Any]) -> None:
    meta = resume_data.get("meta", {}) if isinstance(resume_data, dict) else {}
    name = meta.get("name", "候选人") if isinstance(meta, dict) else "候选人"
    summary = resume_data.get("summary")
    _append_docx_header_block(
        doc,
        meta=meta,
        template="minimal",
        name=str(name),
        summary=summary if isinstance(summary, str) else "",
    )
    experience = _collect_all_experience(resume_data)
    project_items = _collect_all_projects(resume_data, experience)
    if isinstance(experience, list) and experience:
        _add_section_heading(doc, "工作/实习经历", "minimal")
        for exp in experience:
            if not isinstance(exp, dict):
                continue
            head = " | ".join([x for x in [exp.get("company", ""), exp.get("role", ""), exp.get("period", "")] if x])
            if head:
                doc.add_paragraph(head)
            for bullet in _collect_experience_bullets(exp):
                pb = doc.add_paragraph(style="List Bullet")
                pb.add_run(str(bullet))

    _append_docx_record_section(doc, resume_data, "research", "科研经历", "minimal")
    _append_docx_record_section(doc, resume_data, "campus_experience", "校园与志愿经历", "minimal")

    if project_items:
        _add_section_heading(doc, "项目经历", "minimal")
        for proj in project_items:
            doc.add_paragraph(f"- {proj.get('name', '项目')}")
            affiliation = str(proj.get("affiliation", "")).strip()
            if affiliation:
                p_aff = doc.add_paragraph(affiliation)
                for run in p_aff.runs:
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            bullets = proj.get("bullets", [])
            if isinstance(bullets, list):
                for bullet in bullets:
                    pb = doc.add_paragraph(style="List Bullet")
                    pb.add_run(str(bullet))
            tech = proj.get("tech_stack", [])
            if isinstance(tech, list) and tech:
                doc.add_paragraph("Tech: " + " · ".join(str(x) for x in tech))

    education = resume_data.get("education", []) if isinstance(resume_data, dict) else []
    if isinstance(education, list) and education:
        _add_section_heading(doc, "教育经历", "minimal")
        for edu in education:
            if not isinstance(edu, dict):
                continue
            line = " | ".join(
                str(value).strip()
                for value in (
                    edu.get("school", ""),
                    " ".join(str(value).strip() for value in (edu.get("degree", ""), edu.get("major", "")) if str(value).strip()),
                    edu.get("period", ""),
                )
                if str(value).strip()
            )
            if line:
                doc.add_paragraph(line)

    skills = resume_data.get("skills", {}) if isinstance(resume_data, dict) else {}
    if isinstance(skills, dict) and any(skills.values()):
        _add_section_heading(doc, "专业技能", "minimal")
        for key, label in {
            "languages": "编程语言", "frameworks": "框架工具", "tools": "工具",
            "domains": "专业领域", "methodologies": "方法与流程",
            "certifications": "证书与资质", "natural_languages": "语言能力",
            "others": "其他专业技能",
        }.items():
            values = skills.get(key, [])
            if isinstance(values, list) and values:
                doc.add_paragraph(f"{label}：" + " · ".join(str(value) for value in values))

    publications = resume_data.get("publications", []) if isinstance(resume_data, dict) else []
    if isinstance(publications, list):
        _append_docx_text_section(
            doc,
            "论文与专利成果",
            [_publication_line(item) for item in publications if _has_publication_content(item)],
            "minimal",
        )

    _append_docx_text_section(
        doc,
        "荣誉与奖项",
        _collect_text_entries(resume_data, ("honors", "awards", "certifications")),
        "minimal",
    )
    _append_docx_text_section(
        doc,
        "个人技能",
        _collect_text_entries(resume_data, ("personal_skills",)),
        "minimal",
    )
    for title, items in _collect_additional_sections(resume_data):
        _append_docx_text_section(doc, title, items, "minimal")


def _render_docx_classic(doc: DocxDocument, resume_data: dict[str, Any]) -> None:
    meta = resume_data.get("meta", {}) if isinstance(resume_data, dict) else {}
    experience = _collect_all_experience(resume_data)
    project_items = _collect_all_projects(resume_data, experience)
    education = resume_data.get("education", []) if isinstance(resume_data, dict) else []
    skills = resume_data.get("skills", {}) if isinstance(resume_data, dict) else {}
    summary = resume_data.get("summary", "") if isinstance(resume_data, dict) else ""
    publications = resume_data.get("publications", []) if isinstance(resume_data, dict) else []

    name = meta.get("name", "候选人") if isinstance(meta, dict) else "候选人"
    _append_docx_header_block(
        doc,
        meta=meta,
        template="classic",
        name=str(name),
        summary=summary if isinstance(summary, str) else "",
    )
    if isinstance(experience, list) and experience:
        _add_section_heading(doc, "工作/实习经历", "classic")
        for exp in experience:
            if not isinstance(exp, dict):
                continue

            p_company = doc.add_paragraph()
            run = p_company.add_run(exp.get("company", ""))
            run.bold = True
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)

            run = p_company.add_run(f"  |  {exp.get('role', '')}")
            run.font.size = Pt(10.5)

            run = p_company.add_run(f"    {exp.get('period', '')}")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.italic = True
            p_company.space_before = Pt(6)
            p_company.space_after = Pt(2)

            exp_bullets = _collect_experience_bullets(exp)
            for bullet in exp_bullets:
                p_bullet = doc.add_paragraph(style="List Bullet")
                run = p_bullet.add_run(str(bullet))
                run.font.size = Pt(10)
                p_bullet.space_after = Pt(1)

    _append_docx_record_section(doc, resume_data, "research", "科研经历", "classic")
    _append_docx_record_section(doc, resume_data, "campus_experience", "校园与志愿经历", "classic")

    if project_items:
        _add_section_heading(doc, "项目经历", "classic")
        for proj in project_items:
            p_proj = doc.add_paragraph()
            run = p_proj.add_run(f"▸ {proj.get('name', '')}")
            run.bold = True
            run.font.size = Pt(10.5)
            p_proj.space_before = Pt(4)
            p_proj.space_after = Pt(1)

            affiliation = str(proj.get("affiliation", "")).strip()
            if affiliation:
                p_aff = doc.add_paragraph(affiliation)
                for aff_run in p_aff.runs:
                    aff_run.font.size = Pt(9.2)
                    aff_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                    aff_run.italic = True
                p_aff.space_after = Pt(1)

            bullets = proj.get("bullets", [])
            if isinstance(bullets, list):
                for bullet in bullets:
                    p_bullet = doc.add_paragraph(style="List Bullet")
                    run = p_bullet.add_run(str(bullet))
                    run.font.size = Pt(10)
                    p_bullet.space_after = Pt(1)

            tech = proj.get("tech_stack", [])
            if isinstance(tech, list) and tech:
                p_tech = doc.add_paragraph()
                run = p_tech.add_run("Tech: ")
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                run = p_tech.add_run("  ·  ".join(str(x) for x in tech))
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                run.italic = True
                p_tech.space_after = Pt(4)

    if isinstance(education, list) and education:
        _add_section_heading(doc, "教育经历", "classic")
        for edu in education:
            if not isinstance(edu, dict):
                continue

            p_edu = doc.add_paragraph()
            run = p_edu.add_run(edu.get("school", ""))
            run.bold = True
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)

            degree_major = f"  |  {edu.get('degree', '')} · {edu.get('major', '')}"
            run = p_edu.add_run(degree_major)
            run.font.size = Pt(10.5)

            run = p_edu.add_run(f"    {edu.get('period', '')}")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.italic = True
            p_edu.space_after = Pt(2)

            highlights = edu.get("highlights", [])
            if isinstance(highlights, list) and highlights:
                p_hl = doc.add_paragraph()
                run = p_hl.add_run("  ·  ".join(str(x) for x in highlights))
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                p_hl.paragraph_format.left_indent = Inches(0.25)
                p_hl.space_after = Pt(4)

    if isinstance(skills, dict) and any(skills.values()):
        _add_section_heading(doc, "专业技能", "classic")
        label_map = {
            "languages": "编程语言",
            "frameworks": "框架工具",
            "tools": "基础设施",
            "domains": "专业领域",
            "methodologies": "方法与流程",
            "certifications": "证书与资质",
            "natural_languages": "语言能力",
            "others": "其他专业技能",
        }

        for key, label in label_map.items():
            items = skills.get(key, [])
            if isinstance(items, list) and items:
                p_skill = doc.add_paragraph()
                run = p_skill.add_run(f"{label}：")
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
                run = p_skill.add_run("  ·  ".join(str(x) for x in items))
                run.font.size = Pt(10)
                p_skill.space_after = Pt(1)
                p_skill.paragraph_format.left_indent = Inches(0.15)

    valid_publications = [pub for pub in publications if _has_publication_content(pub)] if isinstance(publications, list) else []
    if valid_publications:
        _add_section_heading(doc, "论文成果", "classic")
        for pub in valid_publications:
            if not isinstance(pub, dict):
                p_pub = doc.add_paragraph(_publication_line(pub))
                p_pub.space_after = Pt(2)
                p_pub.paragraph_format.left_indent = Inches(0.15)
                continue
            p_pub = doc.add_paragraph()
            venue = pub.get("venue", "")
            title = pub.get("title", "")
            run = p_pub.add_run(f"[{venue}]  ")
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
            run = p_pub.add_run(str(title))
            run.font.size = Pt(9.5)
            run.italic = True
            p_pub.space_after = Pt(2)
            p_pub.paragraph_format.left_indent = Inches(0.15)
    _append_docx_text_section(
        doc,
        "荣誉与奖项",
        _collect_text_entries(resume_data, ("honors", "awards", "certifications")),
        "classic",
    )
    _append_docx_text_section(
        doc,
        "个人技能",
        _collect_text_entries(resume_data, ("personal_skills",)),
        "classic",
    )
    for title, items in _collect_additional_sections(resume_data):
        _append_docx_text_section(doc, title, items, "classic")


def _render_docx_modern(doc: DocxDocument, resume_data: dict[str, Any]) -> None:
    meta = resume_data.get("meta", {}) if isinstance(resume_data, dict) else {}
    experience = _collect_all_experience(resume_data)
    project_items = _collect_all_projects(resume_data, experience)
    education = resume_data.get("education", []) if isinstance(resume_data, dict) else []
    skills = resume_data.get("skills", {}) if isinstance(resume_data, dict) else {}
    summary = resume_data.get("summary", "") if isinstance(resume_data, dict) else ""
    publications = resume_data.get("publications", []) if isinstance(resume_data, dict) else []

    name = meta.get("name", "候选人") if isinstance(meta, dict) else "候选人"
    _append_docx_header_block(
        doc,
        meta=meta,
        template="modern",
        name=str(name),
        summary=summary if isinstance(summary, str) else "",
    )
    if isinstance(experience, list) and experience:
        _add_section_heading(doc, "工作/实习经历", "modern")
        for exp in experience:
            if not isinstance(exp, dict):
                continue

            p_head = doc.add_paragraph()
            run = p_head.add_run(exp.get("company", ""))
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x00, 0x5A, 0x64)
            role = exp.get("role", "")
            period = exp.get("period", "")
            if role:
                rr = p_head.add_run(f"  |  {role}")
                rr.font.size = Pt(10.5)
            if period:
                pr = p_head.add_run(f"  ·  {period}")
                pr.font.size = Pt(9.5)
                pr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
            p_head.space_before = Pt(4)
            p_head.space_after = Pt(2)

            exp_bullets = _collect_experience_bullets(exp)
            for bullet in exp_bullets:
                p_bullet = doc.add_paragraph(f"• {bullet}")
                p_bullet.paragraph_format.left_indent = Inches(0.2)
                p_bullet.space_after = Pt(1)
                for br in p_bullet.runs:
                    br.font.size = Pt(10)

    _append_docx_record_section(doc, resume_data, "research", "科研经历", "modern")
    _append_docx_record_section(doc, resume_data, "campus_experience", "校园与志愿经历", "modern")

    if project_items:
        _add_section_heading(doc, "项目经历", "modern")
        for proj in project_items:
            p_proj = doc.add_paragraph()
            r = p_proj.add_run(f"- {proj.get('name', '')}")
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
            p_proj.space_before = Pt(2)
            p_proj.space_after = Pt(1)

            affiliation = str(proj.get("affiliation", "")).strip()
            if affiliation:
                p_aff = doc.add_paragraph(affiliation)
                p_aff.paragraph_format.left_indent = Inches(0.2)
                p_aff.space_after = Pt(1)
                for ar in p_aff.runs:
                    ar.font.size = Pt(9)
                    ar.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            bullets = proj.get("bullets", [])
            if isinstance(bullets, list):
                for bullet in bullets:
                    p_bullet = doc.add_paragraph(f"• {bullet}")
                    p_bullet.paragraph_format.left_indent = Inches(0.2)
                    p_bullet.space_after = Pt(1)
                    for br in p_bullet.runs:
                        br.font.size = Pt(10)

            tech = proj.get("tech_stack", [])
            if isinstance(tech, list) and tech:
                p_tech = doc.add_paragraph("TECH  " + " / ".join(str(x) for x in tech))
                p_tech.paragraph_format.left_indent = Inches(0.2)
                p_tech.space_after = Pt(3)
                for tr in p_tech.runs:
                    tr.font.size = Pt(9)
                    tr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if isinstance(education, list) and education:
        _add_section_heading(doc, "教育经历", "modern")
        for edu in education:
            if not isinstance(edu, dict):
                continue
            p_edu = doc.add_paragraph()
            run = p_edu.add_run(edu.get("school", ""))
            run.bold = True
            run.font.size = Pt(10.8)
            run.font.color.rgb = RGBColor(0x00, 0x5A, 0x64)
            dm = f"{edu.get('degree', '')} {edu.get('major', '')}".strip()
            if dm:
                rr = p_edu.add_run(f"  |  {dm}")
                rr.font.size = Pt(10)
            if edu.get("period", ""):
                pr = p_edu.add_run(f"  ·  {edu.get('period', '')}")
                pr.font.size = Pt(9.5)
                pr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
            p_edu.space_after = Pt(2)

            highlights = edu.get("highlights", [])
            if isinstance(highlights, list) and highlights:
                p_hl = doc.add_paragraph(" / ".join(str(x) for x in highlights))
                p_hl.paragraph_format.left_indent = Inches(0.2)
                p_hl.space_after = Pt(3)
                for hr in p_hl.runs:
                    hr.font.size = Pt(9.5)
                    hr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if isinstance(skills, dict) and any(skills.values()):
        _add_section_heading(doc, "专业技能", "modern")
        for key, label in {
            "languages": "编程语言",
            "frameworks": "框架工具",
            "tools": "开发工具",
            "domains": "业务领域",
            "methodologies": "方法与流程",
            "certifications": "证书与资质",
            "natural_languages": "语言能力",
            "others": "其他专业技能",
        }.items():
            items = skills.get(key, [])
            if isinstance(items, list) and items:
                p_skill = doc.add_paragraph()
                run = p_skill.add_run(f"{label}: ")
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x00, 0x5A, 0x64)
                rr = p_skill.add_run(" · ".join(str(x) for x in items))
                rr.font.size = Pt(9.8)
                p_skill.space_after = Pt(1)

    valid_publications = [pub for pub in publications if _has_publication_content(pub)] if isinstance(publications, list) else []
    if valid_publications:
        _add_section_heading(doc, "论文成果", "modern")
        for pub in valid_publications:
            if not isinstance(pub, dict):
                p_pub = doc.add_paragraph(_publication_line(pub))
                p_pub.space_after = Pt(2)
                continue
            p_pub = doc.add_paragraph()
            venue = pub.get("venue", "")
            title = pub.get("title", "")
            run = p_pub.add_run(f"{venue}  ")
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x00, 0x5A, 0x64)
            rr = p_pub.add_run(str(title))
            rr.italic = True
            rr.font.size = Pt(9.5)
            p_pub.space_after = Pt(2)
    _append_docx_text_section(
        doc,
        "荣誉与奖项",
        _collect_text_entries(resume_data, ("honors", "awards", "certifications")),
        "modern",
    )
    _append_docx_text_section(
        doc,
        "个人技能",
        _collect_text_entries(resume_data, ("personal_skills",)),
        "modern",
    )
    for title, items in _collect_additional_sections(resume_data):
        _append_docx_text_section(doc, title, items, "modern")


def render_docx(
    resume_data: dict[str, Any],
    output_path: Path,
    template: str = "classic",
    _layout_retry: bool = False,
) -> None:
    tpl = _normalize_template(template)
    resume_data = _compress_resume_data_for_docx(resume_data)
    framework_mode = _is_empty_profile_framework(resume_data)
    profile = extract_template_style_profile(tpl)
    template_path = Path(tpl) if Path(tpl).is_file() else None
    source_docx: Optional[DocxDocument] = None
    table_header_shell = None
    section_geometry: list[dict[str, Any]] = []

    if template_path is not None and template_path.suffix.lower() == ".docx":
        source_docx = DocxDocument(str(template_path))
        section_geometry = _capture_section_geometry(source_docx)
        template_mode = _template_mode(source_docx)
        if not framework_mode and template_mode != "style_only":
            _apply_resume_data_to_template(
                source_docx,
                resume_data,
                template_mode=template_mode,
            )
            # The anchored/tagged document is already the authoritative shell.
            # Do not re-derive and rewrite its margins, theme or heading styles.
            if _layout_retry:
                _apply_docx_compact_typography(source_docx)
            _apply_docx_pagination_guards(source_docx)
            source_docx.save(str(output_path))
            _make_docx_cjk_portable(output_path)
            actual_report = inspect_docx_layout(output_path)
            logger.info(
                "DOCX Visual QA: renderer=%s pages=%s issues=%s template_source=%s",
                actual_report.get("renderer"),
                actual_report.get("page_count"),
                actual_report.get("issues"),
                profile.source_kind,
            )
            logger.info("Rendered user DOCX template | mode=%s path=%s", template_mode, template_path)
            if _layout_needs_tightening(actual_report) and not _layout_retry:
                render_docx(
                    _layout_retry_data(resume_data, actual_report),
                    output_path,
                    template=template,
                    _layout_retry=True,
                )
            return
        if not framework_mode:
            table_header_shell = _capture_unanchored_table_shell(source_docx)

    # A static DOCX remains the package source so headers, footers, theme,
    # numbering and page settings survive.  PDF/image templates contribute a
    # bounded visual profile to a new editable DOCX.
    if source_docx is not None:
        doc = source_docx
        _clear_docx_body_for_style_template(doc)
        _initialize_docx_defaults(doc, reset_geometry=False)
    else:
        doc = DocxDocument()
        _initialize_docx_defaults(doc, reset_geometry=True)

    render_data = resume_data
    if table_header_shell is not None and _install_table_header_shell(
        doc, table_header_shell, resume_data,
    ):
        render_data = _resume_data_without_header_fields(resume_data)
    _render_docx_with_preset(doc, render_data, profile.preset, framework_mode)
    apply_template_style_profile(doc, resume_data, profile)
    if source_docx is not None:
        _restore_section_geometry(doc, section_geometry)

    _apply_docx_pagination_guards(doc)
    if _layout_retry:
        _apply_docx_compact_typography(doc)
    doc.save(str(output_path))
    _make_docx_cjk_portable(output_path)

    actual_report = inspect_docx_layout(output_path)
    if actual_report.get("available"):
        logger.info(
            "DOCX Visual QA: renderer=%s pages=%s chars=%s issues=%s template_source=%s",
            actual_report.get("renderer"), actual_report.get("page_count"),
            actual_report.get("page_char_counts"), actual_report.get("issues"), profile.source_kind,
        )
        if _layout_needs_tightening(actual_report) and not _layout_retry:
            render_docx(
                _layout_retry_data(resume_data, actual_report), output_path,
                template=template, _layout_retry=True,
            )
            return

    # Estimate page count for logging only.  Multi-page resumes are allowed;
    # actual visual QA handles blank pages and pagination defects.
    doc_reloaded = DocxDocument(str(output_path))
    estimated = _estimate_docx_pages(doc_reloaded)
    if estimated > 3.0:
        logger.warning("Generated DOCX estimated at %.1f pages (may affect visual score) | path=%s", estimated, output_path)


def _apply_resume_data_to_template(
    doc: "DocxDocument",
    resume_data: dict[str, Any],
    *,
    template_mode: Optional[str] = None,
) -> None:
    """Apply resume data to a user-provided template by replacing placeholder runs."""
    # Collect resume sections as a flat list of (heading, items)
    sections = _build_renderable_sections(resume_data)
    section_headings = [heading for heading, _items in sections]
    _replace_template_placeholders(doc, resume_data)
    _sanitize_template_identity_block(doc, resume_data, section_headings)
    tagged_sections = (
        _apply_tagged_section_anchors(doc, sections)
        if (template_mode or _template_mode(doc)) == "tagged"
        else set()
    )

    def all_paragraphs() -> list[Any]:
        paragraphs: list[Any] = list(doc.paragraphs)
        seen = {id(paragraph._element) for paragraph in paragraphs}
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if id(paragraph._element) not in seen:
                            seen.add(id(paragraph._element))
                            paragraphs.append(paragraph)
        return paragraphs

    def element_text(element: Any) -> str:
        return "".join(
            str(node.text or "") for node in element.iter(qn("w:t"))
        ).strip()

    def is_section_heading(text: str) -> bool:
        return any(_template_section_matches(text, heading) for heading in section_headings)

    def insert_after(
        anchor: Any,
        text: str,
        *,
        template_element: Any = None,
        bold: bool = False,
    ) -> Any:
        paragraph = OxmlElement("w:p")
        if template_element is not None:
            paragraph_properties = template_element.find(qn("w:pPr"))
            if paragraph_properties is not None:
                paragraph.append(copy.deepcopy(paragraph_properties))
        run = OxmlElement("w:r")
        if template_element is not None:
            prototype_run = template_element.find(qn("w:r"))
            prototype_run_properties = (
                prototype_run.find(qn("w:rPr"))
                if prototype_run is not None else None
            )
            if prototype_run_properties is not None:
                run.append(copy.deepcopy(prototype_run_properties))
        if bold:
            run_properties = run.find(qn("w:rPr"))
            if run_properties is None:
                run_properties = OxmlElement("w:rPr")
                run.insert(0, run_properties)
            if run_properties.find(qn("w:b")) is None:
                run_properties.append(OxmlElement("w:b"))
        text_element = OxmlElement("w:t")
        if text.startswith(" ") or text.endswith(" "):
            text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        text_element.text = text
        run.append(text_element)
        paragraph.append(run)
        anchor.addnext(paragraph)
        return paragraph

    def insert_before(
        anchor: Any,
        text: str,
        *,
        template_element: Any = None,
        bold: bool = False,
    ) -> Any:
        placeholder = OxmlElement("w:p")
        anchor.addprevious(placeholder)
        paragraph = insert_after(
            placeholder,
            text,
            template_element=template_element,
            bold=bold,
        )
        placeholder.getparent().remove(placeholder)
        return paragraph

    def insert_items_after(
        anchor: Any,
        values: list,
        *,
        template_element: Any = None,
    ) -> Any:
        for item in values:
            if isinstance(item, tuple) and len(item) == 2:
                sub_heading, sub_items = item
                anchor = insert_after(
                    anchor,
                    str(sub_heading),
                    template_element=template_element,
                    bold=True,
                )
                for sub_item in sub_items:
                    anchor = insert_after(
                        anchor,
                        str(sub_item),
                        template_element=template_element,
                    )
            else:
                anchor = insert_after(
                    anchor,
                    str(item),
                    template_element=template_element,
                )
        return anchor

    original_paragraphs = all_paragraphs()
    meta = resume_data.get("meta", {}) if isinstance(resume_data, dict) else {}
    if not isinstance(meta, dict):
        meta = {}
    heading_prototype = next(
        (
            paragraph._element
            for paragraph in original_paragraphs
            if any(
                _template_section_matches(paragraph.text.strip(), heading)
                for heading in section_headings
            )
        ),
        None,
    )
    body_prototype = next(
        (
            paragraph._element
            for paragraph in original_paragraphs
            if paragraph.text.strip()
            and not any(
                _template_section_matches(paragraph.text.strip(), heading)
                for heading in section_headings
            )
            and not re.search(r"(?:\{\{|\$\{|\[\[)", paragraph.text)
        ),
        None,
    )
    body = doc._element.body
    def contains_section_anchor(element: Any) -> bool:
        paragraphs = (
            [element] if element.tag == qn("w:p")
            else list(element.iter(qn("w:p")))
        )
        return any(
            _template_section_matches(element_text(paragraph), heading)
            for paragraph in paragraphs
            for heading in section_headings
        )

    first_section_anchor = next(
        (
            child for child in body
            if contains_section_anchor(child)
        ),
        None,
    )

    for section_heading, items in sections:
        if section_heading in tagged_sections:
            continue
        matched = False
        for para in all_paragraphs():
            full_text = para.text.strip()
            if _template_section_matches(full_text, section_heading):
                # Preserve the heading's exact position/style (including table
                # cells), clear only stale body paragraphs, then inject the new
                # section immediately below it using the template body style.
                anchor = para._element
                parent = anchor.getparent()
                if parent is not None and parent.tag == qn("w:tc"):
                    row = parent.getparent()
                    row_cells = list(row.findall(qn("w:tc"))) if row is not None else []
                    current_index = row_cells.index(parent) if parent in row_cells else -1
                    companion_cells = [
                        cell for index, cell in enumerate(row_cells)
                        if index != current_index
                        and not is_section_heading(element_text(cell))
                    ]
                    companion_cells.sort(
                        key=lambda cell: row_cells.index(cell) < current_index,
                    )
                    if companion_cells:
                        target_cell = companion_cells[0]
                        target_paragraphs = list(target_cell.findall(qn("w:p")))
                        content_template = next(
                            (element for element in target_paragraphs if element_text(element)),
                            target_paragraphs[0] if target_paragraphs else None,
                        )
                        for child in list(target_cell):
                            if child.tag != qn("w:tcPr"):
                                target_cell.remove(child)
                        placeholder = OxmlElement("w:p")
                        target_cell.append(placeholder)
                        insert_items_after(
                            placeholder,
                            items,
                            template_element=content_template,
                        )
                        target_cell.remove(placeholder)
                        matched = True
                        break
                sibling = anchor.getnext()
                content_template = None
                stale: list[Any] = []
                while sibling is not None:
                    next_sibling = sibling.getnext()
                    if sibling.tag != qn("w:p"):
                        break
                    if is_section_heading(element_text(sibling)):
                        break
                    if content_template is None and element_text(sibling):
                        content_template = sibling
                    stale.append(sibling)
                    sibling = next_sibling
                for element in stale:
                    element.getparent().remove(element)

                insert_items_after(
                    anchor,
                    items,
                    template_element=content_template,
                )
                matched = True
                break
        if matched:
            continue
        current_document_text = "\n".join(
            paragraph.text for paragraph in all_paragraphs()
        )
        if section_heading == "基本信息":
            missing_basic_items: list[Any] = []
            for item in items:
                text = str(item).strip()
                # A rendered item such as ``目标岗位: 产品经理`` is already
                # present when its factual value appears in a custom identity
                # block, even if the template uses a different label.
                factual_value = re.split(r"\s*[:：]\s*", text, maxsplit=1)[-1]
                if factual_value and factual_value not in current_document_text:
                    missing_basic_items.append(item)
            if not missing_basic_items:
                continue
            items = missing_basic_items
        if section_heading == "个人总结":
            summary_value = str(resume_data.get("summary", "") or "").strip()
            if summary_value and summary_value in current_document_text:
                continue
        if section_heading in {"基本信息", "个人总结"} and first_section_anchor is not None:
            insert_before(
                first_section_anchor,
                section_heading,
                template_element=heading_prototype,
                bold=True,
            )
            for item in items:
                # Each new paragraph is inserted immediately before the anchor,
                # after prior inserted siblings, so source order is preserved.
                if isinstance(item, tuple) and len(item) == 2:
                    sub_heading, sub_items = item
                    insert_before(
                        first_section_anchor,
                        str(sub_heading),
                        template_element=body_prototype,
                        bold=True,
                    )
                    for sub_item in sub_items:
                        insert_before(
                            first_section_anchor,
                            str(sub_item),
                            template_element=body_prototype,
                        )
                else:
                    insert_before(
                        first_section_anchor,
                        str(item),
                        template_element=body_prototype,
                    )
            continue
        # A user template rarely contains every profession-specific heading.
        # Append every non-empty unmatched section using the template's first
        # heading/body prototypes so font, color, border and indentation are
        # not silently replaced with hard-coded defaults.
        body = doc._element.body
        anchor = body[-2] if len(body) >= 2 and body[-1].tag == qn("w:sectPr") else body[-1]
        anchor = insert_after(
            anchor,
            section_heading,
            template_element=heading_prototype,
            bold=True,
        )
        insert_items_after(
            anchor,
            items,
            template_element=body_prototype,
        )
    _sanitize_structured_template_residuals(doc, resume_data, sections)
    _remove_empty_structured_headings(doc)


def _template_section_matches(text: str, section_heading: str) -> bool:
    def normalize(value: str) -> str:
        compact = re.sub(
            r"[\s：:|｜/\\【】\[\]()（）·•—_-]+", "", str(value or "").casefold(),
        )
        return re.sub(r"^(?:\d{1,2}|[一二三四五六七八九十])[.、)]*", "", compact)

    normalized = normalize(text)
    target = normalize(section_heading)
    aliases = _TEMPLATE_SECTION_ALIASES.get(section_heading, set())
    candidates = {target, *(normalize(value) for value in aliases)}
    return any(candidate and normalized == candidate for candidate in candidates)


def _replace_template_placeholders(doc: "DocxDocument", resume_data: dict[str, Any]) -> None:
    meta = resume_data.get("meta", {}) if isinstance(resume_data, dict) else {}
    if not isinstance(meta, dict):
        meta = {}
    values = {
        "name": meta.get("name", ""),
        "phone": meta.get("phone", ""),
        "email": meta.get("email", ""),
        "age": meta.get("age", ""),
        "gender": meta.get("gender", ""),
        "work_experience": meta.get("work_experience", ""),
        "education_level": meta.get("education_level", ""),
        "target_role": meta.get("target_role", ""),
        "summary": resume_data.get("summary", "") if isinstance(resume_data, dict) else "",
    }

    def _replace_text(text: str) -> str:
        result = text
        for key, value in values.items():
            for token in (f"{{{{{key}}}}}", f"${{{key}}}", f"[[{key}]]"):
                result = result.replace(token, str(value or ""))
        return result

    def _replace_para(para: Any) -> None:
        original = para.text
        replaced = _replace_text(original)
        if replaced == original:
            return
        _set_xml_paragraph_text(para._element, replaced)

    for para in _all_template_story_paragraphs(doc):
        _replace_para(para)


def _build_renderable_sections(resume_data: dict[str, Any]) -> list[tuple[str, list]]:
    """Build sections list for template rendering."""
    meta = resume_data.get("meta", {})
    if not isinstance(meta, dict):
        meta = {}
    sections = []

    # Basic info section
    basic_items = []
    for key, label in [("name", "姓名"), ("phone", "电话"), ("email", "邮箱"),
                        ("wechat", "微信"), ("github", "GitHub"), ("linkedin", "LinkedIn"),
                        ("political_status", "政治面貌")]:
        val = str(meta.get(key, "")).strip()
        if val:
            basic_items.append(val)
    if meta.get("expected_city") or meta.get("target_role"):
        info_parts = []
        if meta.get("expected_city"):
            info_parts.append(f"期望城市: {meta['expected_city']}")
        if meta.get("target_role"):
            info_parts.append(f"目标岗位: {meta['target_role']}")
        if info_parts:
            basic_items.append(" | ".join(info_parts))
    if basic_items:
        sections.append(("基本信息", basic_items))

    # Summary
    summary = str(resume_data.get("summary", "")).strip()
    if summary:
        sections.append(("个人总结", [summary]))

    # Education
    education = resume_data.get("education", [])
    if isinstance(education, list) and education:
        edu_items = []
        for edu in education:
            if not isinstance(edu, dict):
                continue
            parts = []
            if edu.get("school"):
                parts.append(f"{edu['school']}")
            if edu.get("degree") and edu.get("major"):
                parts.append(f"{edu['degree']} {edu['major']}")
            elif edu.get("degree"):
                parts.append(edu["degree"])
            elif edu.get("major"):
                parts.append(edu["major"])
            if edu.get("period"):
                parts.append(edu["period"])
            if parts:
                edu_items.append(" | ".join(parts))
            highlights = edu.get("highlights", [])
            if isinstance(highlights, list):
                edu_items.extend([h for h in highlights if h])
        sections.append(("教育背景", edu_items))

    # Experience
    experience = resume_data.get("experience", [])
    if isinstance(experience, list) and experience:
        experience_items = []
        for exp in experience:
            if not isinstance(exp, dict):
                continue
            exp_heading = f"{exp.get('company', '')} - {exp.get('role', '')}"
            exp_parts = []
            if exp.get("period"):
                exp_parts.append(exp["period"])
            if exp.get("team"):
                exp_parts.append(exp["team"])
            if exp_parts:
                exp_heading += f" ({' | '.join(exp_parts)})"
            items = []
            bullets = exp.get("bullets", [])
            if isinstance(bullets, list):
                items.extend([b for b in bullets if b])
            resp = exp.get("responsibilities", [])
            if isinstance(resp, list):
                items.extend([r for r in resp if r])
            ach = exp.get("achievements", [])
            if isinstance(ach, list):
                items.extend([a for a in ach if a])
            experience_items.append((exp_heading, items))
        if experience_items:
            sections.append(("工作/实习经历", experience_items))

    for key, title in (("research", "科研经历"), ("campus_experience", "校园与志愿经历")):
        records = resume_data.get(key, [])
        if not isinstance(records, list) or not records:
            continue
        record_items = []
        for record in records:
            if not isinstance(record, dict):
                continue
            heading = " | ".join(
                str(value).strip()
                for value in (
                    record.get("company") or record.get("organization"),
                    record.get("role") or record.get("topic"),
                    record.get("period"),
                )
                if str(value or "").strip()
            ) or title
            record_items.append((heading, _collect_experience_bullets(record)))
        if record_items:
            sections.append((title, record_items))

    # Projects
    projects = resume_data.get("projects", [])
    if isinstance(projects, list) and projects:
        project_items = []
        for proj in projects:
            if not isinstance(proj, dict):
                continue
            heading = proj.get("name", "")
            if not heading:
                continue
            parts = []
            if proj.get("period"):
                parts.append(proj["period"])
            if proj.get("description"):
                parts.append(proj["description"])
            if parts:
                heading += f" ({' | '.join(parts)})"
            items = []
            bullets = proj.get("bullets", [])
            if isinstance(bullets, list):
                items.extend([b for b in bullets if b])
            tech = proj.get("tech_stack", [])
            if isinstance(tech, list) and tech:
                items.append("技术栈: " + " · ".join(str(t) for t in tech))
            project_items.append((heading, items))
        if project_items:
            sections.append(("项目经历", project_items))

    # Skills
    skills = resume_data.get("skills", {})
    if isinstance(skills, dict) and any(skills.values()):
        skill_items = []
        for key, label in [("languages", "编程语言"), ("frameworks", "框架"),
                           ("tools", "工具"), ("domains", "领域"),
                           ("methodologies", "方法与流程"),
                           ("certifications", "证书与资质"),
                           ("natural_languages", "语言能力"),
                           ("others", "其他专业技能")]:
            items = skills.get(key, [])
            if isinstance(items, list) and items:
                skill_items.append(f"{label}: " + " · ".join(str(x) for x in items))
        if skill_items:
            sections.append(("专业技能", skill_items))

    publications = resume_data.get("publications", [])
    if isinstance(publications, list):
        publication_items = [
            _publication_line(item) for item in publications if _has_publication_content(item)
        ]
        if publication_items:
            sections.append(("论文与专利成果", publication_items))

    # Honors
    honors = []
    for key in ("honors", "awards", "certifications"):
        items = resume_data.get(key, [])
        if isinstance(items, list):
            honors.extend([h for h in items if h])
    if honors:
        sections.append(("荣誉与奖项", honors))

    # Personal skills
    personal_skills = resume_data.get("personal_skills", [])
    if isinstance(personal_skills, list):
        personal_skills = [s for s in personal_skills if s]
    if personal_skills:
        sections.append(("个人技能", personal_skills))

    sections.extend(_collect_additional_sections(resume_data))

    return sections


def _estimate_docx_pages(doc: "DocxDocument") -> float:
    """Estimate the number of pages a DOCX document would have."""
    # Count total characters and paragraphs
    total_chars = 0
    total_paragraphs = len(doc.paragraphs)
    for para in doc.paragraphs:
        total_chars += len(para.text)
        # Tables contribute too
        if para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl'):
            total_chars += 500  # Approximate for tables

    # Rough estimate: ~3000 chars per page for 10.5pt font, 1.2 line spacing
    # With CJK characters, ~2500 chars per page
    has_cjk = bool(re.search(r"[\u4e00-\u9fff]", doc.element.xml))
    chars_per_page = 2500 if has_cjk else 3500
    return max(1.0, total_chars / chars_per_page)


def _truncate_to_pages(doc: "DocxDocument", max_pages: float = 3.0) -> None:
    """Truncate paragraphs to fit within max_pages."""
    # Estimate current pages
    estimated = _estimate_docx_pages(doc)
    if estimated <= max_pages:
        return

    # Calculate how many paragraphs to keep
    total_paragraphs = len(doc.paragraphs)
    keep_ratio = max_pages / estimated
    keep_count = max(1, int(total_paragraphs * keep_ratio))

    # Remove paragraphs from the end
    for _ in range(total_paragraphs - keep_count):
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)

    # Add a note
    note_para = doc.add_paragraph()
    run = note_para.add_run("（注：为控制在3页以内，部分内容已简化。详见完整版简历。）")
    run.font.size = Pt(9)
    run.italic = True


def _convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    libreoffice = shutil.which("libreoffice")
    if not libreoffice:
        return False

    with tempfile.TemporaryDirectory() as temp_out:
        # Every request gets an isolated LibreOffice profile.  Without this,
        # concurrent headless conversions can contend for the default profile
        # lock and silently fail to produce a PDF.
        profile_dir = Path(temp_out) / "lo-profile"
        profile_dir.mkdir(parents=True, exist_ok=True)
        result = subprocess.run(
            [
                libreoffice,
                "--headless",
                f"-env:UserInstallation={profile_dir.as_uri()}",
                "--convert-to", "pdf",
                "--outdir", temp_out,
                str(docx_path),
            ],
            capture_output=True,
            text=True,
            timeout=90,
        )
        if result.returncode != 0:
            return False

        generated_pdf = Path(temp_out) / f"{docx_path.stem}.pdf"
        if not generated_pdf.exists():
            return False

        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(generated_pdf, pdf_path)
        return True


def _pdf_insert_wrapped_lines(
    *,
    pdf: Any,
    page: Any,
    text: str,
    x: float,
    y: float,
    max_width: float,
    fontsize: float,
    color: tuple[int, int, int],
    line_spacing: float = 1.35,
) -> tuple[Any, float]:
    lines = _wrap_fitz_text(text, max_width=max_width, fontname=_fitz_font_for_text(text), fontsize=fontsize)
    step = max(11.0, fontsize * line_spacing)

    for line in lines:
        if y + step > 800:
            page = pdf.new_page(width=595, height=842)
            y = 42

        draw_text = line if line else " "
        fontname = _fitz_font_for_text(draw_text)
        draw_kwargs = {
            "fontsize": fontsize,
            "color": _rgb01(color),
        }
        try:
            page.insert_text((x, y), draw_text, fontname=fontname, **draw_kwargs)
        except Exception:
            page.insert_text((x, y), draw_text, **draw_kwargs)
        y += step

    return page, y


def _fallback_render_pdf(resume_data: dict[str, Any], output_path: Path, template: str = "classic") -> None:
    if fitz is None:
        raise RuntimeError("PDF export needs LibreOffice or pymupdf")

    tpl = _normalize_template(template)
    accent = PDF_ACCENT_COLORS.get(tpl, PDF_ACCENT_COLORS["classic"])

    left = 42.0
    right = 42.0
    content_width = 595 - left - right

    lines: list[tuple[str, str]] = []
    meta = resume_data.get("meta", {}) if isinstance(resume_data, dict) else {}
    name = meta.get("name", "候选人") if isinstance(meta, dict) else "候选人"
    lines.append(("title", str(name)))

    contacts = [v for k in ("email", "phone", "github", "linkedin", "website") if isinstance(meta, dict) and (v := meta.get(k))]
    if contacts:
        lines.append(("meta", "  |  ".join(str(v) for v in contacts)))

    summary = resume_data.get("summary")
    if isinstance(summary, str) and summary.strip():
        lines.append(("section", "个人简介"))
        lines.append(("body", summary.strip()))

    profile_items = _collect_profile_items(meta)
    if profile_items:
        lines.append(("section", "求职信息"))
        for key, value in profile_items:
            lines.append(("body", f"{key}: {value}"))

    experience = _collect_all_experience(resume_data)
    project_items = _collect_all_projects(resume_data, experience)
    if isinstance(experience, list):
        lines.append(("section", "工作/实习经历"))
        for exp in experience:
            if not isinstance(exp, dict):
                continue
            head = " | ".join([x for x in [exp.get("company", ""), exp.get("role", ""), exp.get("period", "")] if x])
            if head:
                lines.append(("item_head", head))
            for bullet in _collect_experience_bullets(exp):
                lines.append(("bullet", f"• {bullet}"))

    if project_items:
        lines.append(("section", "项目经历"))
        for proj in project_items:
            lines.append(("project", f"- {proj.get('name', '')}"))
            affiliation = str(proj.get("affiliation", "")).strip()
            if affiliation:
                lines.append(("body", affiliation))
            bullets = proj.get("bullets", [])
            if isinstance(bullets, list):
                for bullet in bullets:
                    lines.append(("bullet", f"• {bullet}"))
            tech = proj.get("tech_stack", [])
            if isinstance(tech, list) and tech:
                lines.append(("body", "Tech: " + " · ".join(str(x) for x in tech)))

    education = resume_data.get("education", []) if isinstance(resume_data, dict) else []
    if isinstance(education, list) and education:
        lines.append(("section", "教育经历"))
        for edu in education:
            if not isinstance(edu, dict):
                continue
            edu_head = " | ".join([x for x in [edu.get("school", ""), f"{edu.get('degree', '')} {edu.get('major', '')}".strip(), edu.get("period", "")] if x])
            if edu_head:
                lines.append(("item_head", edu_head))
            highlights = edu.get("highlights", [])
            if isinstance(highlights, list):
                for item in highlights:
                    text = str(item).strip()
                    if text:
                        lines.append(("bullet", f"• {text}"))

    skills = resume_data.get("skills", {}) if isinstance(resume_data, dict) else {}
    if isinstance(skills, dict) and any(skills.values()):
        lines.append(("section", "专业技能"))
        for key, label in {
            "languages": "编程语言",
            "frameworks": "框架工具",
            "tools": "开发工具",
            "domains": "业务领域",
            "methodologies": "方法与流程",
            "certifications": "证书与资质",
            "natural_languages": "语言能力",
            "others": "其他专业技能",
        }.items():
            items = skills.get(key, [])
            if isinstance(items, list) and items:
                lines.append(("body", f"{label}: " + " · ".join(str(x) for x in items)))

    publications = resume_data.get("publications", []) if isinstance(resume_data, dict) else []
    if isinstance(publications, list):
        publication_items = [
            _publication_line(item) for item in publications if _has_publication_content(item)
        ]
        if publication_items:
            lines.append(("section", "论文与专利成果"))
            lines.extend(("bullet", f"• {item}") for item in publication_items)

    honor_items = _collect_text_entries(resume_data, ("honors", "awards", "certifications"))
    if honor_items:
        lines.append(("section", "荣誉与奖项"))
        for item in honor_items:
            lines.append(("bullet", f"• {item}"))

    personal_skill_items = _collect_text_entries(resume_data, ("personal_skills",))
    if personal_skill_items:
        lines.append(("section", "个人技能"))
        for item in personal_skill_items:
            lines.append(("bullet", f"• {item}"))

    for title, items in _collect_additional_sections(resume_data):
        lines.append(("section", title))
        lines.extend(("bullet", f"• {item}") for item in items)

    pdf = fitz.open()
    page = pdf.new_page(width=595, height=842)

    avatar_reserved = 0.0
    avatar_path = _resolve_avatar_path(meta)
    if avatar_path is not None:
        try:
            avatar_w = 66.0
            avatar_h = 86.0
            x1 = 595 - right
            x0 = x1 - avatar_w
            y0 = 42.0
            y1 = y0 + avatar_h
            page.insert_image(fitz.Rect(x0, y0, x1, y1), filename=str(avatar_path), keep_proportion=True)
            avatar_reserved = avatar_w + 12.0
        except Exception as exc:
            logger.warning("Fallback PDF failed to draw avatar | path=%s error=%s", avatar_path, exc)

    content_width = max(220.0, 595 - left - right - avatar_reserved)

    y = 42.0
    for kind, text in lines:
        if kind == "title":
            page, y = _pdf_insert_wrapped_lines(
                pdf=pdf,
                page=page,
                text=text,
                x=left,
                y=y,
                max_width=content_width,
                fontsize=20,
                color=accent,
                line_spacing=1.2,
            )
            y += 4
            continue
        if kind == "meta":
            page, y = _pdf_insert_wrapped_lines(
                pdf=pdf,
                page=page,
                text=text,
                x=left,
                y=y,
                max_width=content_width,
                fontsize=9.5,
                color=PDF_MUTED_COLOR,
                line_spacing=1.25,
            )
            y += 4
            continue
        if kind == "section":
            page, y = _pdf_insert_wrapped_lines(
                pdf=pdf,
                page=page,
                text=text,
                x=left,
                y=y,
                max_width=content_width,
                fontsize=11.5,
                color=accent,
                line_spacing=1.2,
            )
            if y + 3 > 800:
                page = pdf.new_page(width=595, height=842)
                y = 42
            page.draw_line((left, y), (left + content_width, y), color=_rgb01(accent), width=0.8)
            y += 7
            continue
        if kind == "item_head":
            page, y = _pdf_insert_wrapped_lines(
                pdf=pdf,
                page=page,
                text=text,
                x=left,
                y=y,
                max_width=content_width,
                fontsize=10.2,
                color=PDF_BODY_COLOR,
                line_spacing=1.25,
            )
            y += 1
            continue
        if kind == "project":
            page, y = _pdf_insert_wrapped_lines(
                pdf=pdf,
                page=page,
                text=text,
                x=left + 6,
                y=y,
                max_width=content_width - 6,
                fontsize=10.0,
                color=PDF_BODY_COLOR,
                line_spacing=1.25,
            )
            continue
        if kind == "bullet":
            page, y = _pdf_insert_wrapped_lines(
                pdf=pdf,
                page=page,
                text=text,
                x=left + 18,
                y=y,
                max_width=content_width - 18,
                fontsize=9.7,
                color=PDF_BODY_COLOR,
                line_spacing=1.3,
            )
            continue
        page, y = _pdf_insert_wrapped_lines(
            pdf=pdf,
            page=page,
            text=text,
            x=left,
            y=y,
            max_width=content_width,
            fontsize=9.8,
            color=PDF_BODY_COLOR,
        )

    pdf.save(str(output_path))
    pdf.close()


def render_pdf(
    resume_data: dict[str, Any],
    output_path: Path,
    template: str = "classic",
    docx_source: Optional[Path] = None,
) -> None:
    tpl = _normalize_template(template)
    _ = docx_source
    if _render_pdf_via_weasyprint(resume_data, output_path, template=tpl):
        if _pdf_has_middle_blank_pages(output_path):
            _fallback_render_pdf(resume_data, output_path, template=tpl)
        return

    if fitz is not None:
        _fallback_render_pdf(resume_data, output_path, template=tpl)
        return

    raise RuntimeError(
        "WeasyPrint HTML-to-PDF failed and pymupdf is unavailable. "
        "Please install Linux system libs: cairo/pango/gdk-pixbuf or install pymupdf."
    )


def export_resume_files(
    resume_data: dict[str, Any],
    output_dir: Path,
    output_format: str = "both",
    template: str = "classic",
    file_prefix: Optional[str] = None,
) -> dict[str, Optional[str]]:
    output_dir.mkdir(parents=True, exist_ok=True)

    fmt = (output_format or "both").lower()
    if fmt not in {"docx", "pdf", "both"}:
        raise ValueError("output_format must be docx/pdf/both")

    meta = resume_data.get("meta", {}) if isinstance(resume_data, dict) else {}
    name = str(meta.get("name", "") or "").strip() if isinstance(meta, dict) else ""
    target_role = str(meta.get("target_role", "") or "").strip() if isinstance(meta, dict) else ""
    identity = name or target_role or "候选人"
    unique = uuid.uuid4().hex[:8]
    prefix = _safe_filename(file_prefix or f"简历_{identity}_优化版_{unique}")

    docx_path: Optional[Path] = None
    pdf_path: Optional[Path] = None

    tpl = _normalize_template(template)

    if fmt in {"docx", "both"}:
        docx_path = output_dir / f"{prefix}.docx"
        render_docx(resume_data, docx_path, template=tpl)

    if fmt in {"pdf", "both"}:
        pdf_path = output_dir / f"{prefix}.pdf"

        temp_docx_for_pdf: Optional[Path] = None
        source_docx = docx_path
        if source_docx is None:
            temp_docx_for_pdf = output_dir / f"{prefix}.__pdf_tmp__.docx"
            render_docx(resume_data, temp_docx_for_pdf, template=tpl)
            source_docx = temp_docx_for_pdf

        try:
            render_pdf(resume_data, pdf_path, template=tpl, docx_source=source_docx)
        finally:
            if temp_docx_for_pdf and temp_docx_for_pdf.exists():
                temp_docx_for_pdf.unlink(missing_ok=True)

    return {
        "docx": str(docx_path) if docx_path else None,
        "pdf": str(pdf_path) if pdf_path else None,
    }
