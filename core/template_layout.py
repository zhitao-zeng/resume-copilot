"""ATS-safe style extraction for user-provided resume templates.

The renderer intentionally supports a bounded set of layout primitives rather
than using a screenshot as a page background.  The resulting DOCX therefore
remains editable and readable by ordinary ATS parsers while still reflecting
the supplied template's most visible design choices.
"""

from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
import logging
from pathlib import Path
import re
from typing import Any, Iterable, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    import fitz
except ImportError:  # pragma: no cover - optional deployment dependency
    fitz = None


logger = logging.getLogger(__name__)

_BUILTIN_ACCENTS = {
    "classic": "0F3460",
    "modern": "005A64",
    "minimal": "333333",
}
_SECTION_HEADINGS = {
    "个人简介", "个人总结", "基本信息", "求职信息", "工作经历", "工作/实习经历", "实习经历",
    "科研经历", "校园与志愿经历", "项目经历", "教育经历", "教育背景", "专业技能",
    "论文成果", "论文与专利成果", "荣誉与奖项", "个人技能", "培训与进修", "教学经历",
}


@dataclass(frozen=True)
class TemplateStyleProfile:
    preset: str = "classic"
    accent_hex: str = "0F3460"
    name_alignment: Optional[str] = None
    heading_decoration: str = "plain"
    compact: bool = False
    body_font: Optional[str] = None
    margins_in: Optional[tuple[float, float, float, float]] = None
    source_kind: str = "builtin"
    reasons: tuple[str, ...] = ()


def _hex_color(value: Any) -> Optional[str]:
    if value is None:
        return None
    text = str(value).strip().lstrip("#").upper()
    if re.fullmatch(r"[0-9A-F]{6}", text):
        return text
    return None


def _rgb_tuple(value: str) -> tuple[int, int, int]:
    color = _hex_color(value) or "0F3460"
    return tuple(int(color[index:index + 2], 16) for index in (0, 2, 4))


def _is_chromatic(color: str) -> bool:
    red, green, blue = _rgb_tuple(color)
    return max(red, green, blue) - min(red, green, blue) >= 24 and max(red, green, blue) < 245


def _paragraphs_including_tables(doc: Any) -> Iterable[Any]:
    yield from doc.paragraphs
    seen = {id(paragraph._element) for paragraph in doc.paragraphs}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if id(paragraph._element) not in seen:
                        seen.add(id(paragraph._element))
                        yield paragraph


def _preset_from_text(text: str, fallback: str = "classic") -> tuple[str, list[str]]:
    normalized = str(text or "").casefold()
    reasons: list[str] = []
    if any(token in normalized for token in ("极简", "minimal", "黑白", "无装饰")):
        reasons.append("minimal style hint")
        return "minimal", reasons
    if any(token in normalized for token in ("现代", "modern", "浅灰分隔线", "姓名居中")):
        reasons.append("modern style hint")
        return "modern", reasons
    if any(token in normalized for token in ("经典", "classic", "深蓝", "商务")):
        reasons.append("classic style hint")
        return "classic", reasons
    return fallback, reasons


def _docx_profile(path: Path) -> TemplateStyleProfile:
    doc = Document(str(path))
    paragraphs = list(_paragraphs_including_tables(doc))
    all_text = "\n".join(paragraph.text for paragraph in paragraphs)
    preset, reasons = _preset_from_text(f"{path.stem}\n{all_text}")

    colors: Counter[str] = Counter()
    fonts: Counter[str] = Counter()
    explicit_bar = False
    explicit_divider = False
    for paragraph in paragraphs:
        p_pr = paragraph._element.find(qn("w:pPr"))
        if p_pr is not None:
            shading = p_pr.find(qn("w:shd"))
            if shading is not None and shading.get(qn("w:fill"), "").upper() not in {"", "AUTO", "FFFFFF"}:
                explicit_bar = True
            borders = p_pr.find(qn("w:pBdr"))
            if borders is not None and len(borders):
                explicit_divider = True
        for run in paragraph.runs:
            color = _hex_color(getattr(getattr(run.font, "color", None), "rgb", None))
            if color and _is_chromatic(color):
                weight = 3 if run.bold else 1
                colors[color] += weight
            if run.font.name:
                fonts[str(run.font.name)] += max(1, len(run.text.strip()))

    normal_font = None
    try:
        normal_font = doc.styles["Normal"].font.name
    except Exception:
        pass
    body_font = fonts.most_common(1)[0][0] if fonts else normal_font
    accent = colors.most_common(1)[0][0] if colors else _BUILTIN_ACCENTS[preset]

    hint_text = all_text.casefold()
    if explicit_bar or any(token in hint_text for token in ("色条", "色块", "标题栏")):
        decoration = "bar"
        reasons.append("colored heading bar")
    elif explicit_divider or any(token in hint_text for token in ("分隔线", "下划线", "divider")):
        decoration = "divider"
        reasons.append("heading divider")
    else:
        decoration = "plain"

    name_alignment: Optional[str] = None
    if any(token in hint_text for token in ("姓名居中", "标题居中", "centered name")):
        name_alignment = "center"
        reasons.append("centered name hint")
    else:
        first_nonempty = next((paragraph for paragraph in paragraphs if paragraph.text.strip()), None)
        if first_nonempty is not None:
            if first_nonempty.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                name_alignment = "center"
                reasons.append("centered first block")
            elif first_nonempty.alignment in {WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT}:
                name_alignment = "right" if first_nonempty.alignment == WD_ALIGN_PARAGRAPH.RIGHT else "left"

    compact = any(token in hint_text for token in ("紧凑", "compact", "窄间距"))
    if compact:
        reasons.append("compact spacing hint")

    margins = None
    if doc.sections:
        section = doc.sections[0]
        try:
            margins = (
                round(section.top_margin.inches, 3),
                round(section.bottom_margin.inches, 3),
                round(section.left_margin.inches, 3),
                round(section.right_margin.inches, 3),
            )
        except Exception:
            margins = None

    return TemplateStyleProfile(
        preset=preset,
        accent_hex=accent,
        name_alignment=name_alignment,
        heading_decoration=decoration,
        compact=compact,
        body_font=body_font,
        margins_in=margins,
        source_kind="docx",
        reasons=tuple(reasons),
    )


def _visual_profile(path: Path) -> TemplateStyleProfile:
    preset, reasons = _preset_from_text(path.stem)
    source_kind = "pdf" if path.suffix.lower() == ".pdf" else "image"
    accent = _BUILTIN_ACCENTS[preset]
    decoration = "plain"
    name_alignment: Optional[str] = None
    compact = any(token in path.stem.casefold() for token in ("compact", "紧凑", "minimal", "极简"))
    margins = None

    if fitz is None:
        return TemplateStyleProfile(
            preset=preset, accent_hex=accent, compact=compact,
            source_kind=source_kind, reasons=tuple([*reasons, "visual parser unavailable"]),
        )

    try:
        document = fitz.open(str(path))
        if not document.page_count:
            raise ValueError("template has no pages")
        page = document[0]
        native_text = page.get_text("text") if source_kind == "pdf" else ""
        text_preset, text_reasons = _preset_from_text(f"{path.stem}\n{native_text}", fallback=preset)
        preset = text_preset
        reasons.extend(text_reasons)
        # Bound rasterization independently of uploaded page dimensions.  The
        # profile needs layout/color evidence, not print-resolution pixels.
        longest_side = max(float(page.rect.width), float(page.rect.height), 1.0)
        render_scale = min(1.0, 1600.0 / longest_side)
        pixmap = page.get_pixmap(matrix=fitz.Matrix(render_scale, render_scale), alpha=False)
        width, height, channels = pixmap.width, pixmap.height, pixmap.n
        samples = pixmap.samples
        color_counts: Counter[tuple[int, int, int]] = Counter()
        occupied_x: list[int] = []
        bar_rows = 0
        stride = max(1, min(width, height) // 700)
        top_limit = max(1, int(height * 0.28))

        for y in range(0, height, stride):
            tinted_in_row = 0
            occupied_in_row: list[int] = []
            for x in range(0, width, stride):
                offset = (y * width + x) * channels
                red, green, blue = samples[offset:offset + 3]
                maximum, minimum = max(red, green, blue), min(red, green, blue)
                if maximum < 248 and minimum < 244:
                    occupied_in_row.append(x)
                # Very light tinted bars (for example RGB 232/240/250) are a
                # common resume heading treatment and must not be mistaken for
                # a white background.
                if maximum - minimum >= 5 and maximum < 254:
                    tinted_in_row += 1
                if maximum - minimum >= 24 and 35 <= maximum < 235:
                    quantized = (red // 16 * 16, green // 16 * 16, blue // 16 * 16)
                    color_counts[quantized] += 1
            sampled_width = (width + stride - 1) // stride
            if sampled_width and tinted_in_row / sampled_width >= 0.32:
                bar_rows += 1
            elif y < top_limit and occupied_in_row:
                occupied_x.extend(occupied_in_row)

        if color_counts:
            selected = max(
                color_counts,
                key=lambda color: color_counts[color]
                * (1.0 + (max(color) - min(color)) / 255.0)
                * (1.0 + (255 - sum(color) / 3) / 255.0),
            )
            accent = "".join(f"{value:02X}" for value in selected)
            reasons.append("dominant visual accent")
        else:
            accent = _BUILTIN_ACCENTS[preset]

        if bar_rows >= 2:
            decoration = "bar"
            reasons.append("wide colored row")
        elif preset == "modern":
            decoration = "divider"

        if occupied_x:
            center = sum(occupied_x) / len(occupied_x)
            if abs(center - width / 2) <= width * 0.10:
                name_alignment = "center"
            elif center >= width * 0.67:
                name_alignment = "right"
            else:
                name_alignment = "left"

            left = min(occupied_x) / max(1, width)
            right = (width - max(occupied_x)) / max(1, width)
            left_in = min(0.90, max(0.50, left * 8.27))
            right_in = min(0.90, max(0.50, right * 8.27))
            margins = (0.60, 0.60, round(left_in, 3), round(right_in, 3))
    except Exception as exc:
        logger.warning("Template visual profile extraction failed | path=%s error=%s", path, exc)
        reasons.append("visual parsing failed")

    return TemplateStyleProfile(
        preset=preset,
        accent_hex=accent,
        name_alignment=name_alignment,
        heading_decoration=decoration,
        compact=compact,
        margins_in=margins,
        source_kind=source_kind,
        reasons=tuple(dict.fromkeys(reasons)),
    )


def extract_template_style_profile(template: str) -> TemplateStyleProfile:
    value = str(template or "classic").strip() or "classic"
    if value in _BUILTIN_ACCENTS:
        return TemplateStyleProfile(
            preset=value,
            accent_hex=_BUILTIN_ACCENTS[value],
            name_alignment="center" if value == "classic" else None,
            source_kind="builtin",
        )

    path = Path(value)
    if not path.is_file():
        return TemplateStyleProfile(reasons=("template path unavailable",))
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _docx_profile(path)
    if suffix == ".pdf" or suffix in {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".gif", ".tif", ".tiff"}:
        return _visual_profile(path)
    return TemplateStyleProfile(reasons=("unsupported template format",))


def _get_or_add_child(parent: Any, tag: str) -> Any:
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _set_heading_decoration(paragraph: Any, profile: TemplateStyleProfile) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    if profile.heading_decoration == "bar":
        shading = _get_or_add_child(p_pr, "w:shd")
        red, green, blue = _rgb_tuple(profile.accent_hex)
        light = tuple(round(channel * 0.12 + 255 * 0.88) for channel in (red, green, blue))
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "".join(f"{value:02X}" for value in light))
        paragraph.paragraph_format.left_indent = Inches(0.06)
        paragraph.paragraph_format.right_indent = Inches(0.04)
    elif profile.heading_decoration == "divider":
        borders = _get_or_add_child(p_pr, "w:pBdr")
        bottom = _get_or_add_child(borders, "w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), "D9DDE3")


def _normalized_heading(text: str) -> str:
    return re.sub(r"[\s：:]+", "", str(text or ""))


def apply_template_style_profile(
    doc: Any,
    resume_data: dict[str, Any],
    profile: TemplateStyleProfile,
) -> None:
    """Apply supported template primitives to an already-rendered DOCX."""

    if profile.margins_in:
        top, bottom, left, right = profile.margins_in
        for section in doc.sections:
            section.top_margin = Inches(max(0.35, min(1.25, top)))
            section.bottom_margin = Inches(max(0.35, min(1.25, bottom)))
            section.left_margin = Inches(max(0.45, min(1.35, left)))
            section.right_margin = Inches(max(0.45, min(1.35, right)))

    if profile.body_font:
        try:
            doc.styles["Normal"].font.name = profile.body_font
        except Exception:
            pass

    meta = resume_data.get("meta", {}) if isinstance(resume_data, dict) else {}
    name = str(meta.get("name", "") if isinstance(meta, dict) else "").strip()
    framework_names = {"个人简历框架", "RESUME FRAMEWORK"}
    heading_keys = {_normalized_heading(value) for value in _SECTION_HEADINGS}
    accent = RGBColor(*_rgb_tuple(profile.accent_hex))
    builtin_accents = set(_BUILTIN_ACCENTS.values())

    paragraphs = list(_paragraphs_including_tables(doc))
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        normalized = _normalized_heading(text)
        is_heading = normalized in heading_keys
        is_name = bool(text and (text == name or text in framework_names))

        if is_name and profile.name_alignment:
            paragraph.alignment = {
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "left": WD_ALIGN_PARAGRAPH.LEFT,
            }.get(profile.name_alignment, paragraph.alignment)

        if is_heading:
            _set_heading_decoration(paragraph, profile)
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = accent

        for run in paragraph.runs:
            if profile.body_font and not run.font.name:
                run.font.name = profile.body_font
            current = _hex_color(getattr(getattr(run.font, "color", None), "rgb", None))
            if current in builtin_accents:
                run.font.color.rgb = accent

        if profile.compact and not is_name:
            paragraph.paragraph_format.space_after = Pt(
                min(1.5, paragraph.paragraph_format.space_after.pt)
                if paragraph.paragraph_format.space_after is not None else 1.0
            )
            if not is_heading:
                paragraph.paragraph_format.line_spacing = 1.08

    logger.info(
        "Applied template profile | source=%s preset=%s accent=%s align=%s decoration=%s compact=%s reasons=%s",
        profile.source_kind,
        profile.preset,
        profile.accent_hex,
        profile.name_alignment,
        profile.heading_decoration,
        profile.compact,
        list(profile.reasons),
    )
