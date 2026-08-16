#!/usr/bin/env python3
"""Rendered-document gate (R24 Phase 7) contract tests."""

from __future__ import annotations

import json
import sys
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parents[1]
HOLDOUT_ROOT = REPO_ROOT / "validation_sets" / "public_resume_holdout"
for path in (str(REPO_ROOT), str(REPO_ROOT / "core"), str(HOLDOUT_ROOT)):
    if path not in sys.path:
        sys.path.insert(0, path)

from rendered_audit import (  # noqa: E402
    audit_rendered_docx,
    template_fidelity_score01,
    visual_layout_score01,
)
from resume_renderer import export_resume_files  # noqa: E402


def _resume_data():
    return {
        "meta": {"name": "李四", "phone": "13800001111"},
        "summary": "三年后端经验。",
        "experience": [{
            "organization": "甲公司",
            "role": "后端工程师",
            "period": "2021-2024",
            "bullets": ["负责订单系统重构，接口耗时降低 30%"],
        }],
        "education": [{"school": "乙大学", "degree": "软件工程学士", "period": "2016-2020"}],
        "skills": {"技术": ["Python", "SQL"], "languages": ["Go"]},
    }


def _render(tmp_path, resume_data, template=None):
    out = export_resume_files(
        resume_data=resume_data,
        output_dir=tmp_path,
        output_format="docx",
        template=template or "classic",
    )
    return out["docx"]


def _tagged_template(tmp_path):
    from docx import Document

    path = tmp_path / "tagged.docx"
    doc = Document()
    doc.add_paragraph("{{name}}")
    doc.add_paragraph("{{section.experience}}")
    doc.add_paragraph("{{section.education}}")
    doc.add_paragraph("{{section.skills}}")
    doc.save(path)
    return path


def test_default_render_passes_gate_with_full_fact_retention(tmp_path):
    resume_data = _resume_data()
    docx = _render(tmp_path, resume_data)
    report = audit_rendered_docx(docx, resume_data=resume_data)
    assert report["pass"] is True
    assert report["label_remnants"] == []
    assert report["separator_artifacts"] == []
    assert report["fact_retention"] == 1.0
    assert report["editable"] is True
    assert report["cjk_present"] is True
    assert visual_layout_score01(report) == 1.0
    assert template_fidelity_score01(report) == 1.0


def test_tagged_template_preserves_shell_and_all_facts(tmp_path):
    resume_data = _resume_data()
    template = _tagged_template(tmp_path)
    docx = _render(tmp_path, resume_data, template=str(template))
    report = audit_rendered_docx(docx, resume_data=resume_data, template_path=str(template))
    assert report["template_mode"] == "tagged"
    assert report["leftover_tags"] == []
    assert report["pass"] is True
    # 组织名与自定义技能键不得在模板渲染中丢失
    assert report["fact_retention"] == 1.0
    assert report["fact_missing"] == []


def test_gate_flags_label_remnants_and_separator_artifacts(tmp_path):
    from docx import Document

    path = tmp_path / "dirty.docx"
    doc = Document()
    doc.add_paragraph("职位：")
    doc.add_paragraph("负责后端开发____")
    doc.save(path)
    report = audit_rendered_docx(path)
    assert report["pass"] is False
    assert "职位：" in report["label_remnants"]
    assert report["separator_artifacts"]
    assert visual_layout_score01(report) < 1.0


def test_gate_flags_missing_docx_and_unreadable(tmp_path):
    missing = audit_rendered_docx(tmp_path / "nope.docx")
    assert missing["pass"] is False if "pass" in missing else missing["pass_"] is False
    assert missing["error"] == "docx_missing"
    assert visual_layout_score01(missing) is None


def test_gate_flags_facts_lost_in_render(tmp_path):
    resume_data = _resume_data()
    docx = _render(tmp_path, resume_data)
    shrunk = dict(resume_data)
    shrunk["experience"] = [{
        **resume_data["experience"][0],
        "bullets": ["负责订单系统重构，接口耗时降低 30%", "设计了从不存在的缓存层"],
    }]
    report = audit_rendered_docx(docx, resume_data=shrunk)
    assert report["fact_retention"] < 1.0
    assert any("缓存层" in item for item in report["fact_missing"])
    assert report["pass"] is False
