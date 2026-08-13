from __future__ import annotations

import asyncio
import base64
import copy
from pathlib import Path
from unittest.mock import patch
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import resume_copilot_pipeline
from resume_generate_api import _MockUpload
from resume_renderer import detect_docx_template_mode, render_docx
from template_layout import extract_template_style_profile


FIXTURES = Path(__file__).parent.parent / "acceptance_testset" / "files" / "templates"


def _payload() -> dict:
    return {
        "meta": {
            "name": "李明",
            "phone": "13800000000",
            "email": "liming@example.com",
            "target_role": "高级产品经理",
        },
        "summary": "具备五年B端产品经验，负责需求分析、产品规划与跨团队交付。",
        "experience": [{
            "company": "星河科技有限公司",
            "role": "产品经理",
            "period": "2021.03-2025.06",
            "bullets": [
                "访谈业务用户并梳理核心流程，输出需求文档与产品方案。",
                "协同研发、测试推进版本上线，并基于反馈持续迭代。",
            ],
        }],
        "projects": [{
            "name": "企业数据平台",
            "bullets": ["负责指标体系与看板规划，支持运营复盘。"],
            "tech_stack": ["Axure", "SQL"],
        }],
        "education": [{
            "school": "江南大学",
            "degree": "本科",
            "major": "工业工程",
            "period": "2017.09-2021.06",
        }],
        "skills": {"tools": ["Axure", "SQL", "Excel"], "domains": ["B端产品", "数据分析"]},
    }


def _render(payload: dict, output: Path, template: Path) -> Document:
    with patch("resume_renderer.inspect_docx_layout", return_value={"available": False, "issues": []}):
        render_docx(payload, output, template=str(template))
    return Document(output)


def _document_xml(path: Path) -> bytes:
    with zipfile.ZipFile(path) as archive:
        return archive.read("word/document.xml")


def _paragraph_shading(paragraph) -> str:
    shading = paragraph._element.find("./w:pPr/w:shd", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
    return shading.get(qn("w:fill"), "") if shading is not None else ""


def _paragraph_bottom_border(paragraph) -> str:
    border = paragraph._element.find("./w:pPr/w:pBdr/w:bottom", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
    return border.get(qn("w:val"), "") if border is not None else ""


def _all_document_text(doc: Document) -> str:
    body = [paragraph.text for paragraph in doc.paragraphs]
    tables = [
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join([*body, *tables])


def _table_grid_widths(table) -> tuple[str, ...]:
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        return ()
    return tuple(
        str(column.get(qn("w:w"), ""))
        for column in grid.findall(qn("w:gridCol"))
    )


def test_compact_and_modern_docx_templates_produce_distinct_editable_layouts(tmp_path):
    compact_path = tmp_path / "compact.docx"
    modern_path = tmp_path / "modern.docx"
    compact = _render(_payload(), compact_path, FIXTURES / "template_compact.docx")
    modern = _render(_payload(), modern_path, FIXTURES / "template_modern.docx")

    assert _document_xml(compact_path) != _document_xml(modern_path)
    for expected in ("李明", "星河科技有限公司", "企业数据平台", "江南大学", "Axure"):
        assert expected in "\n".join(paragraph.text for paragraph in compact.paragraphs)
        assert expected in "\n".join(paragraph.text for paragraph in modern.paragraphs)

    modern_name = next(paragraph for paragraph in modern.paragraphs if paragraph.text == "李明")
    modern_heading = next(paragraph for paragraph in modern.paragraphs if paragraph.text == "工作/实习经历")
    assert modern_name.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert _paragraph_bottom_border(modern_heading) == "single"


def test_image_template_drives_accent_alignment_and_heading_bar(tmp_path):
    template = FIXTURES / "template_blue.png"
    profile = extract_template_style_profile(str(template))
    assert profile.source_kind == "image"
    assert profile.name_alignment == "left"
    assert profile.heading_decoration == "bar"
    assert profile.accent_hex != "0F3460"

    output = tmp_path / "blue.docx"
    rendered = _render(_payload(), output, template)
    name = next(paragraph for paragraph in rendered.paragraphs if paragraph.text == "李明")
    heading = next(paragraph for paragraph in rendered.paragraphs if paragraph.text == "工作/实习经历")
    assert name.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert _paragraph_shading(heading) not in {"", "FFFFFF", "auto"}


def test_pdf_template_is_retained_as_a_minimal_style_source(tmp_path):
    template = FIXTURES / "template_minimal.pdf"
    profile = extract_template_style_profile(str(template))
    assert profile.source_kind == "pdf"
    assert profile.preset == "minimal"
    assert profile.compact is True

    output = tmp_path / "minimal.docx"
    rendered = _render(_payload(), output, template)
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "李明" in text and "星河科技有限公司" in text and "企业数据平台" in text


def test_framework_mode_uses_uploaded_docx_style(tmp_path):
    payload = {
        "meta": {"target_role": "病理科技师"},
        "framework": {
            "mode": "empty_profile",
            "target_role": "病理科技师",
            "notice": "以下内容均为待填写结构，不代表候选人已有事实。",
            "sections": [
                {"title": "基本信息"},
                {"title": "个人总结"},
                {"title": "教育经历"},
                {"title": "工作/实习经历"},
                {"title": "项目经历"},
                {"title": "专业技能"},
            ],
        },
    }
    output = tmp_path / "framework.docx"
    rendered = _render(payload, output, FIXTURES / "template_modern.docx")
    name = next(paragraph for paragraph in rendered.paragraphs if paragraph.text == "个人简历框架")
    heading = next(paragraph for paragraph in rendered.paragraphs if paragraph.text == "基本信息")
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert name.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert _paragraph_bottom_border(heading) == "single"
    assert "现代简历模板" not in text and "版式要求" not in text


def test_injected_section_content_clones_run_level_formatting(tmp_path):
    template = tmp_path / "section-template.docx"
    source = Document()
    heading = source.add_paragraph("工作经历")
    heading.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0x78)
    sample = source.add_paragraph()
    sample_run = sample.add_run("示例工作内容")
    sample_run.font.name = "Arial"
    sample_run.font.size = Pt(8.5)
    sample_run.font.italic = True
    sample_run.font.color.rgb = RGBColor(0x77, 0x22, 0x22)
    source.save(template)

    output = tmp_path / "injected.docx"
    rendered = _render(_payload(), output, template)
    paragraph = next(
        paragraph for paragraph in rendered.paragraphs
        if paragraph.text == "访谈业务用户并梳理核心流程，输出需求文档与产品方案。"
    )
    run = paragraph.runs[0]
    assert run.font.name == "Arial"
    assert round(run.font.size.pt, 1) == 8.5
    assert run.font.italic is True
    assert str(run.font.color.rgb) == "772222"


def test_unanchored_table_template_does_not_leak_sample_content(tmp_path):
    template = tmp_path / "sample-table.docx"
    source = Document()
    table = source.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "示例姓名"
    table.cell(0, 1).text = "示例电话"
    source.save(template)

    output = tmp_path / "from-table.docx"
    rendered = _render(_payload(), output, template)
    table_text = "\n".join(
        cell.text for table in rendered.tables for row in table.rows for cell in row.cells
    )
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs) + "\n" + table_text
    assert "示例姓名" not in text and "示例电话" not in text
    assert "李明" in text and "星河科技有限公司" in text
    assert len(rendered.tables) == 1
    assert len(rendered.tables[0].columns) == 2
    assert "李明" in rendered.tables[0].cell(0, 0).text
    assert "13800000000" in rendered.tables[0].cell(0, 1).text


def test_english_section_anchors_keep_order_styles_and_remove_sample_resume(tmp_path):
    template = tmp_path / "english-sections.docx"
    source = Document()
    expected_headings = [
        "PERSONAL INFORMATION",
        "PROFESSIONAL SUMMARY",
        "WORK EXPERIENCE",
        "EDUCATION",
        "SKILLS",
    ]
    for heading_text in expected_headings:
        heading = source.add_paragraph(heading_text)
        heading.runs[0].font.color.rgb = RGBColor(0x21, 0x4E, 0x72)
        sample = source.add_paragraph(f"Sample content under {heading_text}")
        sample.runs[0].font.name = "Arial"
        sample.runs[0].font.size = Pt(9)
    source.save(template)

    output = tmp_path / "english-output.docx"
    rendered = _render(_payload(), output, template)
    paragraphs = [paragraph.text for paragraph in rendered.paragraphs]
    text = "\n".join(paragraphs)

    assert "Sample content" not in text
    for expected in ("李明", "星河科技有限公司", "江南大学", "Axure"):
        assert expected in text
    positions = [paragraphs.index(heading) for heading in expected_headings]
    assert positions == sorted(positions)
    injected = next(
        paragraph for paragraph in rendered.paragraphs
        if paragraph.text == "访谈业务用户并梳理核心流程，输出需求文档与产品方案。"
    )
    assert injected.runs[0].font.name == "Arial"


def test_table_section_anchor_injects_into_companion_cell_without_sample_leak(tmp_path):
    template = tmp_path / "table-section.docx"
    source = Document()
    table = source.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "工作经历"
    sample = table.cell(0, 1).paragraphs[0]
    sample.text = "示例公司｜示例岗位｜示例职责"
    sample.runs[0].font.name = "Arial"
    sample.runs[0].font.size = Pt(8.5)
    source.save(template)

    output = tmp_path / "table-section-output.docx"
    rendered = _render(_payload(), output, template)
    assert len(rendered.tables) == 1
    left = rendered.tables[0].cell(0, 0).text
    right = rendered.tables[0].cell(0, 1).text
    assert left == "工作经历"
    assert "示例" not in right
    assert "星河科技有限公司" in right
    assert "访谈业务用户" in right
    document_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "李明" in document_text
    assert "具备五年B端产品经验" in document_text
    assert document_text.index("李明") < document_text.index("具备五年B端产品经验")
    content_run = next(
        paragraph.runs[0]
        for paragraph in rendered.tables[0].cell(0, 1).paragraphs
        if "访谈业务用户" in paragraph.text
    )
    assert content_run.font.name == "Arial"


def test_template_upload_keeps_pdf_and_image_for_layout_analysis(tmp_path, monkeypatch):
    monkeypatch.setattr(resume_copilot_pipeline, "AVATAR_DIR", tmp_path)

    async def resolve(filename: str, payload: str):
        warnings = []
        upload = _MockUpload(None, filename, content=payload)
        path = await resume_copilot_pipeline._resolve_template_path(upload, warnings)
        return Path(path), warnings

    pdf_path, pdf_warnings = asyncio.run(resolve("style.pdf", "%PDF-style"))
    image_path, image_warnings = asyncio.run(resolve("style.png", "png-style"))
    assert pdf_path.suffix == ".pdf" and pdf_path.read_bytes() == b"%PDF-style"
    assert image_path.suffix == ".png" and image_path.read_bytes() == b"png-style"
    assert pdf_warnings == [] and image_warnings == []


def test_docx_template_modes_are_deterministic(tmp_path):
    tagged = tmp_path / "tagged.docx"
    tagged_doc = Document()
    tagged_doc.add_paragraph("{{name}}")
    tagged_doc.add_paragraph("{{section.experience}}")
    tagged_doc.save(tagged)

    anchored = tmp_path / "anchored.docx"
    anchored_doc = Document()
    anchored_doc.add_paragraph("WORK EXPERIENCE")
    anchored_doc.add_paragraph("Example work")
    anchored_doc.save(anchored)

    style_only = tmp_path / "style-only.docx"
    style_doc = Document()
    style_doc.add_paragraph("示例姓名和示例履历")
    style_doc.add_table(rows=1, cols=2)
    style_doc.save(style_only)

    assert detect_docx_template_mode(tagged) == "tagged"
    assert detect_docx_template_mode(anchored) == "anchored"
    assert detect_docx_template_mode(style_only) == "style_only"


def test_tagged_paragraph_anchors_preserve_shell_and_place_every_fact(tmp_path):
    template = tmp_path / "tagged-shell.docx"
    source = Document()
    section = source.sections[0]
    section.top_margin = Inches(0.43)
    section.bottom_margin = Inches(0.57)
    section.left_margin = Inches(0.71)
    section.right_margin = Inches(0.83)
    section.header.paragraphs[0].text = "CAREER PROFILE"
    section.footer.paragraphs[0].text = "CONFIDENTIAL"
    source.add_paragraph("{{name}}")
    source.add_paragraph("{{phone}} | {{email}}")
    source.add_paragraph("{{target_role}}")
    source.add_paragraph("PROFILE")
    source.add_paragraph("{{section.summary}}")
    source.add_paragraph("WORK EXPERIENCE")
    prototype = source.add_paragraph("{{section.experience}}")
    prototype.runs[0].font.name = "Arial"
    prototype.runs[0].font.size = Pt(9.5)
    prototype.runs[0].font.italic = True
    source.add_paragraph("PROJECTS")
    source.add_paragraph("{{section.projects}}")
    source.save(template)

    output = tmp_path / "tagged-shell-output.docx"
    rendered = _render(_payload(), output, template)
    text = _all_document_text(rendered)
    expected = (
        "李明", "13800000000", "liming@example.com", "高级产品经理",
        "具备五年B端产品经验", "星河科技有限公司", "产品经理",
        "2021.03-2025.06", "访谈业务用户并梳理核心流程",
        "协同研发、测试推进版本上线", "企业数据平台", "支持运营复盘",
        "江南大学", "工业工程", "Axure", "SQL",
    )
    assert all(value in text for value in expected)
    assert "{{" not in text and "示例" not in text
    assert text.index("WORK EXPERIENCE") < text.index("星河科技有限公司")
    assert text.index("PROJECTS") < text.index("企业数据平台")
    injected = next(
        paragraph for paragraph in rendered.paragraphs
        if "访谈业务用户并梳理核心流程" in paragraph.text
    )
    assert injected.runs[0].font.name == "Arial"
    assert round(injected.runs[0].font.size.pt, 1) == 9.5
    assert injected.runs[0].font.italic is True
    rendered_section = rendered.sections[0]
    assert rendered_section.top_margin == section.top_margin
    assert rendered_section.bottom_margin == section.bottom_margin
    assert rendered_section.left_margin == section.left_margin
    assert rendered_section.right_margin == section.right_margin
    assert rendered_section.header.paragraphs[0].text == "CAREER PROFILE"
    assert rendered_section.footer.paragraphs[0].text == "CONFIDENTIAL"


def test_tagged_table_row_repeats_records_without_sample_leak(tmp_path):
    template = tmp_path / "tagged-table.docx"
    source = Document()
    source.add_paragraph("WORK EXPERIENCE")
    table = source.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(1.2)
    table.columns[1].width = Inches(5.1)
    table.cell(0, 0).text = "示例公司与示例头像"
    anchor = table.cell(0, 1).paragraphs[0]
    anchor.text = "{{section.experience}}"
    anchor.runs[0].font.name = "Arial"
    grid_before = _table_grid_widths(table)
    source.save(template)

    payload = copy.deepcopy(_payload())
    payload["experience"].append({
        "company": "远帆医疗中心",
        "role": "运营专员",
        "period": "2019.07-2021.02",
        "bullets": ["整理患者服务流程并维护月度运营台账。"],
    })
    output = tmp_path / "tagged-table-output.docx"
    rendered = _render(payload, output, template)
    assert len(rendered.tables) == 1
    rendered_table = rendered.tables[0]
    assert len(rendered_table.rows) == 2
    assert len(rendered_table.columns) == 2
    assert _table_grid_widths(rendered_table) == grid_before
    table_text = "\n".join(cell.text for row in rendered_table.rows for cell in row.cells)
    assert "示例" not in table_text and "{{" not in table_text
    assert "星河科技有限公司" in table_text
    assert "远帆医疗中心" in table_text
    assert "整理患者服务流程并维护月度运营台账" in table_text
    assert all(row.cells[0].text.strip() == "" for row in rendered_table.rows)
    with zipfile.ZipFile(output) as archive:
        assert not any(
            name.startswith(("word/header", "word/footer"))
            for name in archive.namelist()
        )


def test_anchored_template_removes_sample_identity_portrait_and_absent_sections(tmp_path):
    template = tmp_path / "anchored-sample.docx"
    portrait = tmp_path / "sample.png"
    portrait.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
    ))
    source = Document()
    identity = source.add_paragraph("王样例")
    identity.add_run().add_picture(str(portrait), width=Inches(0.2))
    source.add_paragraph("13911112222 | sample-person@example.com")
    source.add_paragraph("WORK EXPERIENCE")
    source.add_paragraph("旧公司 - 虚构岗位")
    source.add_paragraph("负责旧模板中的示例经历，增长999%。")
    source.add_paragraph("AWARDS")
    source.add_paragraph("旧模板人物获得国家级示例奖项")
    source.save(template)

    output = tmp_path / "anchored-sample-output.docx"
    rendered = _render(_payload(), output, template)
    text = _all_document_text(rendered)
    assert "王样例" not in text
    assert "13911112222" not in text
    assert "sample-person@example.com" not in text
    assert "旧公司" not in text and "增长999%" not in text
    assert "国家级示例奖项" not in text
    assert "AWARDS" not in text
    assert "李明" in text and "高级产品经理" in text
    assert "星河科技有限公司" in text
    assert len(rendered.inline_shapes) == 0


def test_style_only_template_keeps_geometry_headers_footer_and_table_grid(tmp_path):
    template = tmp_path / "style-shell.docx"
    source = Document()
    section = source.sections[0]
    section.top_margin = Inches(0.27)
    section.bottom_margin = Inches(1.31)
    section.left_margin = Inches(0.41)
    section.right_margin = Inches(1.47)
    section.header.paragraphs[0].text = "简历模板页眉"
    section.footer.paragraphs[0].text = "简历模板页脚"
    table = source.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(1.4)
    table.columns[1].width = Inches(4.8)
    table.cell(0, 0).text = "示例姓名"
    table.cell(0, 1).text = "sample@example.com"
    table.cell(1, 0).text = "示例公司"
    table.cell(1, 1).text = "旧模板工作经历"
    grid_before = _table_grid_widths(table)
    geometry_before = (
        section.top_margin, section.bottom_margin,
        section.left_margin, section.right_margin,
    )
    source.save(template)

    output = tmp_path / "style-shell-output.docx"
    rendered = _render(_payload(), output, template)
    text = _all_document_text(rendered)
    assert "示例" not in text and "sample@example.com" not in text
    assert all(value in text for value in (
        "李明", "星河科技有限公司", "企业数据平台", "江南大学", "Axure",
    ))
    rendered_section = rendered.sections[0]
    assert (
        rendered_section.top_margin, rendered_section.bottom_margin,
        rendered_section.left_margin, rendered_section.right_margin,
    ) == geometry_before
    assert rendered_section.header.paragraphs[0].text == "简历模板页眉"
    assert rendered_section.footer.paragraphs[0].text == "简历模板页脚"
    assert len(rendered.tables) == 1
    assert len(rendered.tables[0].columns) == 2
    assert _table_grid_widths(rendered.tables[0]) == grid_before
    assert rendered.tables[0].style.name == "Table Grid"


def test_docx_rendering_never_drops_facts_to_enforce_a_page_budget(tmp_path):
    payload = copy.deepcopy(_payload())
    facts = [
        f"唯一事实{i:03d}：完成第{i:03d}项验证并保留对应交付记录。"
        for i in range(140)
    ]
    payload["experience"][0]["bullets"] = facts
    output = tmp_path / "multi-page.docx"
    with patch("resume_renderer.inspect_docx_layout", return_value={"available": False, "issues": []}):
        render_docx(payload, output, template="classic")
    rendered = Document(output)
    text = _all_document_text(rendered)
    assert all(fact in text for fact in facts)
