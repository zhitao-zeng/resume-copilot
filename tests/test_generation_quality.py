import asyncio
import json
from pathlib import Path
from unittest.mock import patch

from input_normalization import html_to_visible_text, split_url_and_text
from evidence_binding import bind_resume_evidence, enforce_resume_evidence, measure_source_coverage
from resume_classifier import classify_resume_request
from resume_classifier import normalize_target_role, reconcile_user_stage
from resume_optimizer import optimize_resume
from resume_renderer import (
    _collect_all_projects,
    _compress_resume_data_for_docx,
    _convert_docx_to_pdf,
    _layout_needs_tightening,
    _layout_retry_data,
    _tighten_resume_data_for_layout,
    analyze_pdf_layout,
    render_docx,
    render_pdf,
)
from resume_validator import check_required_fields
from resume_copilot_pipeline import build_reply_text, _build_targeted_suggestions
from resume_verifier import (
    _ground_fixed_fields,
    _merge_adjacent_duplicate_experiences,
    _reclassify_non_work,
)
from source_adapter import build_source_bundle, candidate_blocks
from v2_pipeline import (
    _bullet_rewrite_changes,
    _build_evidence_summary,
    _canonical_to_v1_format,
    _compact_canonical,
    _deterministic_verify_draft,
    _empty_profile_framework,
    _needs_optimizer,
    _rank_resume_content,
    _atomize_resume_bullets,
    run_v2_pipeline,
)
from v2_schemas import CanonicalResume, DraftResume


def test_mixed_jd_url_preserves_inline_text():
    value = "https://hr.xiaomi.com/job/view/799、负责智能硬件产品（IoT设备及配套App）的需求定义"
    url, text = split_url_and_text(value)
    assert url == "https://hr.xiaomi.com/job/view/799"
    assert text.startswith("负责智能硬件产品")


def test_html_extraction_removes_css_navigation_and_scripts():
    payload = b"""<html><style>.x{color:red}</style><nav>menu</nav>
    <script>alert(1)</script><h1>IoT PM</h1><p>Manage devices</p></html>"""
    text = html_to_visible_text(payload)
    assert "IoT PM" in text
    assert "Manage devices" in text
    assert "color:red" not in text
    assert "menu" not in text
    assert "alert" not in text


def test_mixed_jd_resolution_prioritizes_inline_text():
    from resume_copilot_pipeline import _resolve_jd_text

    async def fake_fetch(url, warnings):
        return "网页岗位：电视 OTT 产品经理"

    with patch("resume_copilot_pipeline._fetch_jd_url", side_effect=fake_fetch):
        perf, warnings = {}, []
        result = asyncio.run(_resolve_jd_text(
            target_jd="https://example.com/job/1、用户补充：IoT 智能硬件产品经理",
            jd_text=None,
            target_jd_url=None,
            jd_url=None,
            target_jd_file=None,
            perf=perf,
            warnings=warnings,
        ))
    assert result.index("IoT 智能硬件产品经理") < result.index("电视 OTT")
    assert "（优先）" in result


def test_query_facts_survive_watermark_cleaning_with_cv():
    from resume_copilot_pipeline import _clean_template_watermarks, _prune_empty_resume_values

    resume = {
        "meta": {"target_role": "算法工程师"},
        "skills": {"languages": ["Python"], "frameworks": ["PyTorch", ""]},
    }
    query = "我熟悉Python和PyTorch，希望从事算法工程师，帮我整理简历。"
    warnings = _clean_template_watermarks(resume, query, has_cv=True, cv_text="北京邮电大学硕士")
    _prune_empty_resume_values(resume)
    assert resume["skills"]["languages"] == ["Python"]
    assert resume["skills"]["frameworks"] == ["PyTorch"]
    assert not warnings


def test_fixed_fields_are_grounded_and_non_work_is_reclassified():
    parsed = {
        "meta": {"name": "", "work_experience": "1年", "target_role": "产品经理"},
        "education": [{"school": "北京邮电大学", "degree": "硕士", "major": "人工智能", "period": ""}],
        "experience": [
            {"organization": "学生会", "role": "宣传负责人", "period": "", "bullets": ["负责宣传工作"]},
            {"organization": "实验室", "role": "研究助理", "period": "", "bullets": ["参与模型训练"]},
        ],
        "research": [],
        "activities": [],
        "projects": [],
        "skills": {"items": [{"name": "PyTorch", "category": "framework"}, {"name": "", "category": "tool"}]},
        "awards": [],
    }
    evidence = "北京邮电大学人工智能专业硕士在读，在实验室参与模型训练；负责学生会宣传工作；熟悉PyTorch。"
    _ground_fixed_fields(parsed, evidence)
    _reclassify_non_work(parsed, evidence)
    assert parsed["meta"]["work_experience"] == ""
    assert parsed["education"][0]["school"] == "北京邮电大学"
    assert len(parsed["activities"]) == 1
    assert len(parsed["research"]) == 1
    assert parsed["experience"] == []
    assert parsed["skills"]["items"] == [{"name": "PyTorch", "category": "framework"}]


def test_equivalent_date_formats_remain_grounded_and_bindable():
    parsed = {
        "meta": {},
        "education": [{"school": "复旦大学", "degree": "本科", "major": "计算机科学", "period": "2016.09 - 2020.06"}],
        "experience": [], "research": [], "activities": [], "projects": [],
        "skills": {"items": []}, "awards": [],
    }
    evidence = "复旦大学｜计算机科学｜本科｜09-2016 - 06-2020"
    _ground_fixed_fields(parsed, evidence)
    assert parsed["education"][0]["period"] == "2016.09 - 2020.06"

    source = build_source_bundle(evidence, "", "")
    resume = CanonicalResume.model_validate(parsed)
    bindings = bind_resume_evidence(resume, source)
    assert any(binding.path == "education[0].period" for binding in bindings)


def test_optimizer_applies_safe_patch_but_rejects_ownership_upgrade():
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "超级公司",
            "role": "产品助理实习生",
            "bullets": ["协助产品经理进行竞品分析", "负责收集并整理用户反馈"],
        }],
    })
    response = json.dumps({
        "summary": "",
        "experience": [{
            "index": 0,
            "bullets": [
                "主导竞品分析并形成关键决策依据",
                "负责收集和整理用户反馈，形成需求记录",
            ],
        }],
    }, ensure_ascii=False)
    with patch("resume_optimizer.llm_enabled", return_value=True), patch(
        "resume_optimizer.call_llm_text", return_value=response
    ):
        result = optimize_resume(resume, "产品经理")
    assert result.experience[0].bullets[0] == "协助产品经理进行竞品分析"
    assert result.experience[0].bullets[1] == "负责收集和整理用户反馈，形成需求记录"


def test_compaction_removes_blank_skills_and_preserves_typed_sections():
    resume = CanonicalResume.model_validate({
        "research": [{"institution": "北京邮电大学", "topic": "计算机视觉", "bullets": ["参与OCR项目"]}],
        "activities": [{"organization": "学生会", "role": "", "bullets": ["负责宣传工作"]}],
        "skills": {"items": [
            {"name": "", "category": "language"},
            {"name": "Python", "category": "language"},
        ]},
    })
    rendered = _canonical_to_v1_format(_compact_canonical(resume))
    assert rendered["skills"]["languages"] == ["Python"]
    assert rendered["research"][0]["company"] == "北京邮电大学"
    assert rendered["campus_experience"][0]["company"] == "学生会"
    assert rendered["experience"] == []


def test_case39_rules_classify_ai_student():
    query = "马上要毕业了，我在北京邮电大学读人工智能专业硕士，想找算法工程师工作。"
    result = classify_resume_request(query=query, cv_text=query, jd_text="", has_cv=True, has_jd=False)
    assert result.industry == "ai_engineering"
    assert result.user_stage == "student"
    assert result.target_role == "算法工程师"


def test_query_role_wins_over_conflicting_fetched_jd_title():
    query = "我是做智能硬件产品的，给你一个IoT智能硬件产品经理的岗位JD。"
    result = classify_resume_request(
        query=query,
        cv_text="",
        jd_text="招聘岗位：天线工程师\n负责射频天线设计",
        has_cv=False,
        has_jd=True,
    )
    assert result.target_role == "IoT智能硬件产品经理"


def test_target_role_removes_non_role_related_suffix():
    assert normalize_target_role("算法工程师相关") == "算法工程师"
    assert normalize_target_role("计算机视觉相关方向") == "计算机视觉"


def test_internal_role_label_is_mapped_and_final_work_history_fixes_stage():
    assert normalize_target_role("product_pm") == "产品经理"
    resume = {
        "meta": {"work_experience": "4年经验"},
        "experience": [
            {"company": "第四范式", "role": "产品经理", "period": "07-2022 - 05-2026"},
            {"company": "星河科技", "role": "产品助理", "period": "07-2020 - 06-2022"},
        ],
        "education": [{"school": "复旦大学", "degree": "本科"}],
    }
    assert reconcile_user_stage("student", resume, "复旦大学本科，随后从事产品工作") == "experienced"


def test_final_stage_reconciliation_keeps_intern_as_student():
    resume = {
        "meta": {},
        "experience": [{
            "company": "某互联网公司", "role": "产品实习生",
            "period": "06-2025 - 09-2025",
        }],
        "education": [{"school": "某大学", "degree": "本科"}],
    }
    assert reconcile_user_stage("student", resume, "在校本科生，产品实习") == "student"


def test_reply_uses_readable_stage_and_no_false_missing_advice():
    reply = build_reply_text(
        scenario="scenario1", industry="product_research",
        user_stage="experienced", missing_fields=[], conflicts=[],
        ocr_warnings=[], direction="建议投递产品经理岗位", score_total=0,
    )
    assert "职场人士" in reply
    assert "experienced" not in reply
    assert "未检测到必填信息缺失" in reply
    assert "优先补齐联系方式" not in reply
    assert "生成方向总结：" in reply
    assert "缺失信息：" in reply
    assert "岗位匹配与建议：" in reply
    assert "时间或内容冲突：" in reply


def test_reply_lists_every_missing_field_and_gives_targeted_advice():
    missing = [
        {"field": f"field.{index}", "label": f"字段{index}", "reason": f"需要补充字段{index}"}
        for index in range(1, 8)
    ]
    resume = {
        "meta": {"target_role": "产品经理"},
        "experience": [{"company": "某公司", "bullets": ["负责用户调研"]}],
        "projects": [],
    }
    suggestions = _build_targeted_suggestions(
        "1、负责用户调研与竞品分析\n2、输出PRD并推动研发交付",
        resume,
        "产品经理",
    )
    reply = build_reply_text(
        scenario="scenario3", industry="product_research",
        user_stage="experienced", missing_fields=missing, conflicts=[],
        ocr_warnings=[], direction="建议投递产品经理岗位", score_total=0,
        targeted_suggestions=suggestions,
    )
    assert "缺失或待补充信息（7项）" in reply
    assert "字段7" in reply
    assert "针对岗位的建议" in reply
    assert "PRD" in reply

    empty_profile_advice = _build_targeted_suggestions(
        "负责IoT产品规划与全生命周期管理\n小米-北京-产品经理-职位详情",
        {"meta": {"target_role": "IoT产品经理"}, "framework": _empty_profile_framework("IoT产品经理")},
        "IoT产品经理",
    )
    assert "当前简历缺少对应证据" in empty_profile_advice[0]
    assert all("职位详情" not in item for item in empty_profile_advice)


def test_accepted_bullet_rewrite_is_exposed_in_changes_and_reply():
    before = CanonicalResume.model_validate({
        "experience": [{
            "organization": "某公司",
            "role": "产品经理",
            "bullets": ["参与公司核心产品的功能迭代，负责收集、整理用户反馈。"],
        }],
    })
    after = before.model_copy(deep=True)
    after.experience[0].bullets[0] = "参与核心产品功能迭代，收集并整理用户反馈。"

    changes = _bullet_rewrite_changes(before, after)
    assert len(changes) == 1
    assert changes[0].path == "experience[0].bullets[0]"
    assert changes[0].action == "replace"

    reply = build_reply_text(
        scenario="scenario1", industry="product_research",
        user_stage="experienced", missing_fields=[], conflicts=[],
        ocr_warnings=[], direction="建议投递产品经理岗位", score_total=0,
        changes=[change.model_dump() for change in changes],
    )
    assert "优化 1 条经历/项目表述" in reply


def test_sparse_query_does_not_claim_unprovided_experience():
    generated = CanonicalResume.model_validate({
        "meta": {"target_role": "IoT智能硬件产品经理"},
        "summary": "拥有智能硬件产品全生命周期管理和完整落地经验。",
        "skills": {"items": [{"name": "智能硬件", "category": "domain"}]},
    })
    with patch("v2_pipeline.compose_from_query", return_value=generated), patch(
        "v2_pipeline._needs_optimizer", return_value=False,
    ):
        result = run_v2_pipeline("", "我是做智能硬件产品的，帮我优化简历。", "IoT产品经理")
    assert "全生命周期管理" not in result.resume.summary
    assert "完整落地经验" not in result.resume.summary
    assert "智能硬件" in result.resume.summary
    assert result.resume.meta.target_role == "IoT智能硬件产品经理"
    assert result.resume.experience == []
    assert result.resume.projects == []
    assert "framework" not in result.resume_dict


def test_jd_derived_temporary_records_still_produce_structured_framework():
    generated = CanonicalResume.model_validate({
        "meta": {"target_role": "IoT智能硬件产品经理"},
        "experience": [{
            "organization": "目标公司",
            "role": "产品经理",
            "bullets": ["负责智能硬件产品规划和跨团队协同"],
        }],
        "projects": [{
            "name": "智能硬件项目",
            "bullets": ["开展用户研究并推动产品上市"],
        }],
    })
    with patch("v2_pipeline.compose_from_query", return_value=generated), patch(
        "v2_pipeline._needs_optimizer", return_value=False,
    ):
        result = run_v2_pipeline(
            "",
            "申请IoT智能硬件产品经理，请根据JD给我待填写框架，不要编造个人经历",
            "负责智能硬件产品规划、用户研究和跨团队协同",
        )
    assert result.resume.experience == []
    assert result.resume.projects == []
    assert result.resume_dict["framework"]["mode"] == "empty_profile"
    assert len(result.resume_dict["framework"]["sections"]) >= 6


def test_empty_profile_framework_renders_structured_docx(tmp_path):
    resume = {
        "meta": {"target_role": "IoT智能硬件产品经理"},
        "education": [], "experience": [], "projects": [], "skills": {},
        "summary": "", "framework": _empty_profile_framework("IoT智能硬件产品经理"),
    }
    output = tmp_path / "framework.docx"
    render_docx(resume, output, template="classic")
    from docx import Document
    text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert "个人简历框架" in text
    assert "以下内容均为待填写结构，不代表候选人已有事实" in text
    assert "基本信息" in text
    assert "教育经历" in text
    assert "工作/实习经历" in text
    assert "项目经历" in text
    assert "专业技能" in text


def test_generated_files_are_unique_and_have_role_fallback(tmp_path):
    from resume_renderer import export_resume_files

    resume = {
        "meta": {"name": "", "target_role": "算法工程师"},
        "summary": "人工智能专业硕士在读。",
        "education": [],
        "experience": [],
        "projects": [],
        "skills": {},
    }
    first = export_resume_files(resume, tmp_path, output_format="docx")
    second = export_resume_files(resume, tmp_path, output_format="docx")
    assert first["docx"] != second["docx"]
    assert "算法工程师" in first["docx"]


def test_compaction_keeps_education_schema_valid():
    resume = CanonicalResume.model_validate({
        "education": [{"school": "北京邮电大学", "degree": "硕士", "major": "人工智能"}],
    })
    compacted = _compact_canonical(resume)
    assert compacted.education[0].school == "北京邮电大学"


def test_unknown_industry_skills_are_preserved_in_semantic_buckets():
    resume = CanonicalResume.model_validate({
        "education": [{"school": "某建筑大学", "major": "建筑学"}],
        "skills": {"items": [
            {"name": "Revit", "category": "software"},
            {"name": "装配式建筑深化", "category": "craft"},
            {"name": "一级注册建筑师资格证", "category": "other"},
            {"name": "CET-6", "category": "language"},
        ]},
    })
    rendered = _canonical_to_v1_format(_compact_canonical(resume))
    assert rendered["skills"]["tools"] == ["Revit"]
    assert rendered["skills"]["others"] == ["装配式建筑深化"]
    assert rendered["skills"]["certifications"] == ["一级注册建筑师资格证"]
    assert rendered["skills"]["natural_languages"] == ["CET-6"]


def test_summary_is_deterministic_and_contains_no_subjective_filler():
    resume = CanonicalResume.model_validate({
        "meta": {"target_role": "病理科技师"},
        "education": [{"school": "某医科大学", "major": "医学检验", "degree": "本科"}],
        "experience": [{"organization": "某医院", "role": "检验科实习生"}],
        "projects": [{"name": "病理标本处理项目", "bullets": ["完成标本处理流程记录"]}],
        "skills": {"items": [
            {"name": "病理切片制备", "category": "other"},
            {"name": "临床标本处理", "category": "domain"},
        ]},
        "summary": "具备扎实基础、敏锐洞察力，致力于成为优秀人才。",
    })
    summary = _build_evidence_summary(resume)
    assert "某医院检验科实习生" in summary
    assert "病理切片制备" in summary
    assert "求职方向为病理科技师" in summary
    assert len(summary) <= 100
    assert not any(word in summary for word in ("扎实", "敏锐", "致力于", "优秀"))


def test_generic_target_role_does_not_require_known_role_dictionary():
    result = classify_resume_request(
        query="我想申请新能源电池工艺工程师岗位，请按这个方向优化简历。",
        cv_text="",
        jd_text="",
        has_cv=False,
        has_jd=False,
    )
    assert result.target_role == "新能源电池工艺工程师"
    assert result.industry == "新能源电池工艺工程师"


def test_date_range_is_not_rendered_as_work_seniority():
    resume = CanonicalResume.model_validate({
        "meta": {"work_experience": "2024.07-2024.12"},
        "experience": [{
            "organization": "某新能源公司",
            "role": "电池工艺实习生",
            "period": "2024.07-2024.12",
        }],
    })
    assert resume.meta.work_experience == ""


def test_summary_prioritizes_grounded_quantified_achievement():
    resume = CanonicalResume.model_validate({
        "education": [{"school": "华南理工大学", "major": "材料科学与工程", "degree": "本科"}],
        "experience": [{"organization": "华创新能源有限公司", "role": "电池工艺实习生"}],
        "projects": [{
            "name": "涂布工序参数优化项目",
            "bullets": [
                "清洗汇总12批次生产数据。",
                "依据实验记录将厚度波动率从4.8%降至3.1%，结果经试产批次复核。",
            ],
        }],
        "skills": {"items": [{"name": "Minitab", "category": "tool"}]},
    })
    ranked = _rank_resume_content(resume, "涂布工艺参数优化")
    compacted = _compact_canonical(ranked)
    assert "4.8%降至3.1%" in compacted.summary
    assert len(compacted.summary) <= 220
    assert compacted.summary.endswith("。")
    assert "4.8%降至3.1%" in compacted.projects[0].bullets[0]


def test_compound_grounded_bullet_preserves_action_result_chain():
    original = "负责企业数据平台需求调研、版本规划与跨团队推进，推动报表配置效率提升30%"
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "第四范式",
            "role": "产品经理",
            "bullets": [original],
        }],
    })

    atomized, provenance = _atomize_resume_bullets(resume)

    assert atomized.experience[0].bullets == [original]
    assert provenance == {}


def test_narrative_education_fragments_are_normalized_and_deduplicated():
    resume = CanonicalResume.model_validate({
        "education": [
            {"school": "北京邮电大学", "degree": "硕士", "major": "人工智能"},
            {
                "school": "我目前在北京邮电大学",
                "degree": "硕士",
                "major": "读人工智能专业",
                "period": "2026年",
            },
            {
                "school": "北京邮电大学",
                "degree": "硕士",
                "major": "",
                "period": "预计2026年毕业",
            },
            {
                "school": "马上要毕业了",
                "degree": "硕士",
                "major": "最近开始准备找算法工程师相关的工作",
            },
        ],
        "experience": [{"organization": "在实验室", "bullets": ["负责模型训练"]}],
    })

    compacted = _compact_canonical(resume)

    assert [item.model_dump() for item in compacted.education] == [{
        "school": "北京邮电大学",
        "degree": "硕士",
        "major": "人工智能",
        "period": "预计2026年毕业",
    }]
    assert compacted.experience[0].organization == "实验室"


def test_independent_project_with_bullets_has_no_false_missing_warnings():
    resume_data = {
        "meta": {"name": "李明", "phone": "13812345678", "email": "liming@example.cn", "education_level": "本科"},
        "summary": "华南理工大学材料科学与工程本科。",
        "education": [{"school": "华南理工大学", "degree": "本科", "major": "材料科学与工程", "period": "2021-2025"}],
        "experience": [],
        "projects": [{
            "name": "涂布工序参数优化项目",
            "period": "2024.09-2024.11",
            "company": "",
            "description": "",
            "bullets": ["参与设计3组工艺参数对照实验。"],
        }],
        "skills": {"certifications": ["质量管理体系内审员"]},
    }
    missing = check_required_fields(resume_data, user_stage="student")
    fields = {item.field for item in missing}
    assert "projects[0].company" not in fields
    assert "projects[0].description" not in fields
    assert "skills" not in fields


def test_short_resume_keeps_three_ranked_bullets_in_docx_payload():
    resume_data = {
        "summary": "简短摘要。",
        "experience": [{"company": "某公司", "role": "实习生", "bullets": ["成果一", "成果二", "成果三"]}],
        "projects": [],
    }
    compressed = _compress_resume_data_for_docx(resume_data)
    assert compressed["experience"][0]["bullets"] == ["成果一", "成果二", "成果三"]


def test_dense_resume_keeps_verified_content_and_allows_sparse_last_page():
    bullets = [f"已验证经历描述{i}" for i in range(1, 7)]
    project_bullets = [f"已验证项目描述{i}" for i in range(1, 7)]
    awards = [f"奖项{i}" for i in range(1, 7)]
    summary = "完整保留的个人总结。" * 12
    resume_data = {
        "summary": summary,
        "experience": [{"company": "某公司", "role": "工程师", "bullets": bullets}],
        "projects": [{"name": "某项目", "bullets": project_bullets}],
        "awards": awards,
    }

    normalized = _compress_resume_data_for_docx(resume_data)
    assert normalized["summary"] == summary
    assert normalized["experience"][0]["bullets"] == bullets
    assert normalized["projects"][0]["bullets"] == project_bullets
    assert normalized["awards"] == awards

    sparse_two_page_report = {
        "available": True,
        "issues": ["sparse_last_page"],
    }
    assert _layout_needs_tightening(sparse_two_page_report) is True
    assert _layout_needs_tightening({
        "available": True,
        "issues": ["more_than_three_pages"],
    }) is False
    assert _tighten_resume_data_for_layout(normalized) == normalized
    assert _layout_retry_data(normalized, {
        "available": True,
        "issues": ["core_experience_not_on_first_page"],
    }) == normalized

    collected_projects = _collect_all_projects(normalized, normalized["experience"])
    assert collected_projects[0]["bullets"] == project_bullets


def test_source_blocks_keep_deterministic_section_hints():
    bundle = build_source_bundle(
        "工作经历\n某公司｜工艺实习生｜2025\n记录工艺参数\n专业技能：Minitab、SPC",
        "",
        "",
    )
    blocks = bundle.blocks
    assert [block.block_id for block in blocks] == ["resume_0", "resume_1", "resume_2", "resume_3"]
    assert [block.section_hint for block in blocks] == ["experience", "experience", "experience", "skills"]


def test_source_record_boundaries_keep_company_header_with_following_period():
    bundle = build_source_bundle(
        "工作经历\n"
        "XXX有限公司 UG用户增长运营专员用户运营部\n"
        "2018年4月 - 至今\n"
        "- 根据数据分析制定用户增长运营方案。\n"
        "XXX有限公司 UG用户增长运营专员用户运营部\n"
        "2015年7月 - 2017年12月\n"
        "- 分析销售数据并优化运营策略。",
        "",
        "",
    )
    records = [
        (block.text, block.record_id)
        for block in bundle.blocks
        if block.section_hint == "experience" and block.record_id
    ]

    assert len({record_id for _, record_id in records}) == 2
    assert len({record_id for _, record_id in records[:3]}) == 1
    assert len({record_id for _, record_id in records[3:]}) == 1
    assert records[0][1] != records[3][1]


def test_multicolumn_ocr_dates_and_qualifications_stay_in_two_records():
    cv_text = (
        "教育背景\n"
        "2020-09至\n中国医科大学\n内科|硕士\n2023-06\nGPA: 4.5\n"
        "2016-09至\n中国医科大学\n内科|本科\n2020-06\nGPA: 4.4\n"
        "工作经历\n"
        "2021-06至\n大连市第三人民医院\n2022-02\n实习生\n"
        "2020-06至\n辽宁省人民医院\n2020-12\n内科实习生"
    )
    source = build_source_bundle(cv_text, "", "")
    education_ids = {
        block.record_id
        for block in source.blocks
        if block.section_hint == "education" and block.record_id
    }
    experience_ids = {
        block.record_id
        for block in source.blocks
        if block.section_hint == "experience" and block.record_id
    }

    assert len(education_ids) == 2
    assert len(experience_ids) == 2

    data = CanonicalResume.model_validate({
        "education": [
            {
                "school": "中国医科大学",
                "major": "内科",
                "degree": "硕士",
                "period": "2020-09至2023-06",
            },
            {
                "school": "中国医科大学",
                "major": "内科",
                "degree": "本科",
                "period": "2016-09至2020-06",
            },
        ],
        "experience": [
            {
                "organization": "大连市第三人民医院",
                "role": "实习生",
                "period": "2021-06至2022-02",
            },
            {
                "organization": "辽宁省人民医院",
                "role": "内科实习生",
                "period": "2020-06至2020-12",
            },
        ],
    }).model_dump()
    _ground_fixed_fields(data, cv_text)
    gated, _, removed = enforce_resume_evidence(
        CanonicalResume.model_validate(data),
        source,
    )

    assert len(gated.education) == 2
    assert [item.period for item in gated.education] == [
        "2020-09至2023-06",
        "2016-09至2020-06",
    ]
    assert [item.period for item in gated.experience] == [
        "2021-06至2022-02",
        "2020-06至2020-12",
    ]
    assert removed == []


def test_ocr_section_aliases_and_detached_numbered_duties_are_typed_safely():
    source = build_source_bundle(
        "职业技能\n分析心电图\n"
        "资格证书\n执业医师资格证\n"
        "在校经历\n社会实践\n"
        "奖学金\n国家奖学金\n"
        "1. 负责配合医生管理住院患者。\n"
        "2.学会使用多种设备开展诊疗。\n"
        "3. 在上级指导下分管病床并担任值班工作。",
        "",
        "",
    )
    by_text = {block.text: block for block in source.blocks}

    assert by_text["职业技能"].section_hint == "skills"
    assert by_text["资格证书"].section_hint == "certifications"
    assert by_text["在校经历"].section_hint == "activities"
    assert by_text["奖学金"].section_hint == "awards"
    assert by_text["1. 负责配合医生管理住院患者。"].section_hint is None
    assert by_text["2.学会使用多种设备开展诊疗。"].section_hint is None
    assert by_text["3. 在上级指导下分管病床并担任值班工作。"].section_hint is None


def test_anonymous_professional_duty_group_moves_out_of_activities():
    parsed = {
        "experience": [{
            "organization": "某医院",
            "role": "实习生",
            "period": "",
            "bullets": [],
        }],
        "activities": [
            {
                "organization": "",
                "role": "",
                "period": "",
                "bullets": [
                    "负责配合医生管理住院患者",
                    "配合上级完成查房和医疗文书书写",
                    "在上级指导下分管病床并担任值班工作",
                ],
            },
            {
                "organization": "",
                "role": "",
                "period": "",
                "bullets": ["获国家奖学金"],
            },
        ],
        "research": [],
    }

    _reclassify_non_work(parsed, "本科在读")

    assert len(parsed["experience"]) == 2
    assert parsed["experience"][1]["organization"] == ""
    assert len(parsed["experience"][1]["bullets"]) == 3
    assert parsed["activities"] == [{
        "organization": "",
        "role": "",
        "period": "",
        "bullets": ["获国家奖学金"],
    }]


def test_source_record_boundaries_split_undated_projects_and_wrapped_activity_heading():
    bundle = build_source_bundle(
        "项目经历\n"
        "校园二手交易平台产品负责人\n"
        "- 负责需求分析和产品设计。\n"
        "智能家居APP设计核心成员\n"
        "- 负责界面设计和用户测试。\n"
        "社团和\n"
        "组织经历\n"
        "学生会宣传部部长\n"
        "- 负责学生会活动宣传。\n"
        "- 通过公众号发布学生会活动信息。",
        "",
        "",
    )
    by_text = {block.text: block for block in bundle.blocks}

    assert by_text["校园二手交易平台产品负责人"].record_id != by_text[
        "智能家居APP设计核心成员"
    ].record_id
    assert by_text["组织经历"].section_hint == "activities"
    assert by_text["学生会宣传部部长"].section_hint == "activities"
    assert by_text["- 负责学生会活动宣传。"].record_id == by_text[
        "- 通过公众号发布学生会活动信息。"
    ].record_id


def test_duplicate_bullet_is_disambiguated_within_each_source_record():
    source = build_source_bundle(
        "项目经历\n"
        "智能医疗项目\n"
        "2025年\n"
        "- 收集医疗数据并支持模型训练。\n"
        "- 进行用户测试，收集用户反馈，并对系统进行改进。\n"
        "AI客服项目\n"
        "2025年\n"
        "- 收集客服数据并支持模型训练。\n"
        "- 进行用户测试，收集用户反馈，并对系统进行改进。",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "projects": [
            {
                "name": "智能医疗项目",
                "period": "2025年",
                "bullets": [
                    "收集医疗数据并支持模型训练。",
                    "进行用户测试，收集用户反馈，并对系统进行改进。",
                ],
            },
            {
                "name": "AI客服项目",
                "period": "2025年",
                "bullets": [
                    "收集客服数据并支持模型训练。",
                    "进行用户测试，收集用户反馈，并对系统进行改进。",
                ],
            },
        ],
    })

    gated, _, removed = enforce_resume_evidence(resume, source)

    assert len(gated.projects) == 2
    assert "projects[1]" not in removed


def test_identityless_activity_continuation_merges_with_same_source_record():
    source = build_source_bundle(
        "组织经历\n"
        "志愿者协会 活动策划\n"
        "- 参与社区志愿活动策划。\n"
        "- 负责活动现场组织协调。",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "activities": [
            {
                "organization": "志愿者协会",
                "role": "活动策划",
                "bullets": ["参与社区志愿活动策划。"],
            },
            {
                "bullets": ["负责活动现场组织协调。"],
            },
        ],
    })

    gated, _, removed = enforce_resume_evidence(resume, source)

    assert removed == []
    assert len(gated.activities) == 1
    assert gated.activities[0].bullets == [
        "参与社区志愿活动策划。",
        "负责活动现场组织协调。",
    ]


def test_activity_organization_and_department_role_are_coalesced():
    source = build_source_bundle(
        "组织经历\n"
        "计算机协会\n"
        "技术部部长技术部\n"
        "- 组织成员学习Java。\n"
        "- 举办技术讲座。",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "activities": [
            {"organization": "计算机协会"},
            {
                "organization": "技术部",
                "role": "技术部部长",
                "bullets": ["组织成员学习Java。", "举办技术讲座。"],
            },
        ],
    })

    gated, _, removed = enforce_resume_evidence(resume, source)

    assert removed == []
    assert len(gated.activities) == 1
    assert gated.activities[0].organization == "计算机协会"
    assert gated.activities[0].role == "技术部部长"
    assert gated.activities[0].bullets == ["组织成员学习Java。", "举办技术讲座。"]


def test_single_education_record_survives_multicolumn_ocr_reordering():
    source = build_source_bundle(
        "教育经历\n"
        "xx大学\n"
        "统计学本科\n"
        "其他\n"
        "负责撰写数据分析报告\n"
        "2018年6月-2019年10月\n"
        "商业数据分析师\n"
        "2013年9月-2017年6月",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "education": [{
            "school": "xx大学",
            "degree": "本科",
            "major": "统计学",
            "period": "2013年9月-2017年6月",
        }],
    })

    gated, _, removed = enforce_resume_evidence(resume, source)

    assert len(gated.education) == 1
    assert "education[0]" not in removed


def test_verifier_does_not_merge_same_role_across_distinct_periods():
    entries = [
        {
            "organization": "XXX有限公司",
            "role": "用户增长运营专员",
            "period": "2018年4月 - 至今",
            "bullets": ["负责用户增长。"],
        },
        {
            "organization": "XXX有限公司",
            "role": "用户增长运营专员",
            "period": "2015年7月 - 2017年12月",
            "bullets": ["负责运营策略。"],
        },
    ]

    merged = _merge_adjacent_duplicate_experiences(entries)

    assert merged == entries


def test_evidence_binding_traces_bullets_and_excludes_jd():
    source = build_source_bundle(
        "项目经历\n涂布优化项目\n根据实验记录将厚度波动率从4.8%降至3.1%。",
        "想申请工艺工程师",
        "要求主导整条产品线并降低50%成本",
    )
    resume = CanonicalResume.model_validate({
        "meta": {"target_role": "工艺工程师"},
        "projects": [{
            "name": "涂布优化项目",
            "bullets": ["依据实验记录将厚度波动率从4.8%降至3.1%。"],
        }],
    })
    bindings = bind_resume_evidence(resume, source)
    paths = {binding.path: binding for binding in bindings}
    assert paths["projects[0].name"].block_id == "resume_1"
    assert paths["projects[0].bullets[0]"].block_id == "resume_2"
    assert all(not binding.block_id.startswith("jd_") for binding in bindings)


def test_evidence_gate_removes_fabricated_and_jd_only_claims():
    source = build_source_bundle(
        "项目经历\n涂布优化项目\n记录12批次生产参数。\n技能：Minitab",
        "申请工艺工程师",
        "要求精通六西格玛并主导降本50%",
    )
    resume = CanonicalResume.model_validate({
        "meta": {"target_role": "工艺工程师"},
        "projects": [{
            "name": "涂布优化项目",
            "bullets": ["记录12批次生产参数。", "主导全公司改革并节省一亿元。"],
        }],
        "skills": {"items": [
            {"name": "Minitab", "category": "tool"},
            {"name": "六西格玛", "category": "methodology"},
        ]},
    })
    gated, bindings, removed = enforce_resume_evidence(resume, source)
    assert gated.projects[0].bullets == ["记录12批次生产参数。"]
    assert [item.name for item in gated.skills.items] == ["Minitab"]
    assert "projects[0].bullets[1]" in removed
    assert "skills.items[1].name" in removed
    assert all(not binding.block_id.startswith("jd_") for binding in bindings)


def test_clean_draft_uses_deterministic_verifier_and_optimizer_pass():
    source = build_source_bundle(
        "姓名：李明\n教育经历\n华南理工大学 材料科学与工程 本科\n"
        "项目经历\n涂布优化项目\n记录12批次生产参数。\n技能：Minitab",
        "申请工艺工程师",
        "要求工艺分析经验",
    )
    draft = DraftResume.model_validate({
        "meta": {"name": "李明", "target_role": "工艺工程师"},
        "education": [{"school": "华南理工大学", "major": "材料科学与工程", "degree": "本科"}],
        "projects": [{"name": "涂布优化项目", "bullets": ["记录12批次生产参数。"]}],
        "skills": {"items": [{"name": "Minitab", "category": "tool"}]},
    })
    result = _deterministic_verify_draft(source, draft)
    assert result is not None
    assert len(result.evidence_bindings) >= 6
    assert _needs_optimizer(result.resume) is True

    duplicated = result.resume.model_copy(deep=True)
    duplicated.projects[0].bullets = ["记录12批次生产参数。", "记录12批次生产参数。"]
    assert _needs_optimizer(duplicated) is True

    empty = result.resume.model_copy(deep=True)
    empty.projects[0].bullets = []
    assert _needs_optimizer(empty) is False


def test_optimizer_rejects_new_tool_name_even_when_other_text_overlaps():
    from resume_optimizer import _safe_rewrite

    original = "负责整理用户反馈并维护需求文档"
    assert _safe_rewrite(original, "负责整理用户反馈，维护需求文档") is True
    assert _safe_rewrite(original, "使用SQL整理用户反馈并维护需求文档") is False
    chinese_original = "负责用户调研和需求分析，输出产品文档。"
    assert _safe_rewrite(
        chinese_original,
        "负责用户调研和需求分析，使用六西格玛方法输出产品文档。",
    ) is False


def test_query_instructions_are_not_candidate_evidence():
    source = build_source_bundle(
        "甲公司｜产品经理",
        "不要写管理经验\n不要写我负责团队管理\n补充：我会SQL\n申请数据分析师",
        "负责数据分析",
    )
    facts = [block.text for block in candidate_blocks(source)]
    assert "甲公司｜产品经理" in facts
    assert "补充：我会SQL" in facts
    assert "不要写管理经验" not in facts
    assert "不要写我负责团队管理" not in facts
    assert "申请数据分析师" not in facts


def test_evidence_gate_rejects_cross_record_identity_splice():
    source = build_source_bundle(
        "工作经历\n"
        "甲公司｜产品经理｜2020.01-2022.01\n"
        "负责需求分析并输出PRD。\n"
        "乙公司｜销售经理｜2022.02-2024.01\n"
        "负责客户开拓与商务谈判。",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "销售经理",
            "period": "2020.01-2022.01",
            "bullets": ["负责需求分析并输出PRD。"],
        }],
    })
    gated, _, removed = enforce_resume_evidence(resume, source)
    assert gated.experience == []
    assert "experience[0]" in removed


def test_source_coverage_detects_composer_omissions():
    source = build_source_bundle(
        "姓名：李明\n教育经历\n华南理工大学 材料科学与工程 本科\n"
        "项目经历\n涂布优化项目\n记录12批次生产参数。\n"
        "分析异常批次原因。\n输出工艺改进建议。",
        "",
        "",
    )
    draft = DraftResume.model_validate({
        "meta": {"name": "李明"},
        "education": [{"school": "华南理工大学", "major": "材料科学与工程", "degree": "本科"}],
        "projects": [{"name": "涂布优化项目", "bullets": ["记录12批次生产参数。"]}],
    })
    resume = CanonicalResume.model_validate(draft.model_dump())
    bindings = bind_resume_evidence(resume, source)
    coverage, missing = measure_source_coverage(source, bindings)
    assert coverage < 0.75
    assert {"resume_6", "resume_7"}.issubset(set(missing))
    assert _deterministic_verify_draft(source, draft) is None


def test_cross_industry_sections_remain_typed_and_rendered(tmp_path):
    source = build_source_bundle(
        "论文成果\n医学影像分割研究，第一作者，2024\n"
        "专利成果\n医学图像处理方法发明专利\n"
        "证书与资质\n医师执业证书\n"
        "培训经历\n住院医师规范化培训\n"
        "教学经历\n承担本科生临床带教\n"
        "专业会员\n中华医学会会员",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "publications": ["医学影像分割研究，第一作者，2024"],
        "patents": ["医学图像处理方法发明专利"],
        "certifications": ["医师执业证书"],
        "training": ["住院医师规范化培训"],
        "teaching": ["承担本科生临床带教"],
        "additional_sections": {"专业会员": ["中华医学会会员"]},
    })
    gated, _, removed = enforce_resume_evidence(resume, source)
    assert removed == []
    rendered = _canonical_to_v1_format(gated)
    output = tmp_path / "cross-industry.docx"
    render_docx(rendered, output, template="classic")
    from docx import Document
    text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    for expected in (
        "医学影像分割研究", "医学图像处理方法发明专利", "医师执业证书",
        "住院医师规范化培训", "承担本科生临床带教", "中华医学会会员",
    ):
        assert expected in text


def test_publication_evidence_cannot_spawn_duplicate_research_record():
    source = build_source_bundle(
        "论文成果\n老年慢病管理研究，核心期刊，2022",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "research": [{
            "topic": "老年慢病管理研究",
            "period": "2022",
            "bullets": ["发表核心期刊论文"],
        }],
        "publications": ["老年慢病管理研究，核心期刊，2022"],
    })
    gated, _, removed = enforce_resume_evidence(resume, source)
    assert gated.research == []
    assert gated.publications == ["老年慢病管理研究，核心期刊，2022"]
    assert "research[0]" in removed


def test_optimizer_preserves_ownership_level_and_rejects_overcompression():
    from resume_optimizer import _safe_rewrite

    strong = "独立负责小型功能模块的需求分析、产品设计和上线推进"
    assert _safe_rewrite(strong, "参与小型功能模块的设计与上线") is False
    assert _safe_rewrite(strong, "独立负责小型功能模块的需求分析、设计及上线推进") is True

    detailed = "负责收集用户反馈、开展竞品分析并维护需求文档，协同研发跟进交付"
    assert _safe_rewrite(detailed, "负责产品工作") is False


def test_optimizer_provenance_preserves_reviewed_low_overlap_rewrite():
    from resume_optimizer import optimize_resume_with_provenance

    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "产品经理",
            "bullets": ["负责收集用户反馈并维护需求文档"],
        }],
    })
    response = json.dumps({
        "experience": [{
            "index": 0,
            "bullets": ["负责围绕用户反馈开展整理归纳，持续维护需求文档"],
        }],
    }, ensure_ascii=False)
    with patch("resume_optimizer.llm_enabled", return_value=True), patch(
        "resume_optimizer.call_llm_text", return_value=response,
    ), patch(
        "resume_optimizer.review_entailment_batch", return_value=[True],
    ):
        outcome = optimize_resume_with_provenance(resume, "产品经理")

    path = "experience[0].bullets[0]"
    assert outcome.resume.experience[0].bullets[0].startswith("负责围绕用户反馈")
    assert outcome.trusted_rewrites[path] == "负责收集用户反馈并维护需求文档"

    source = build_source_bundle(
        "工作经历\n甲公司｜产品经理\n负责收集用户反馈并维护需求文档",
        "",
        "",
    )
    gated, bindings, removed = enforce_resume_evidence(
        outcome.resume,
        source,
        trusted_rewrites=outcome.trusted_rewrites,
    )
    assert removed == []
    assert gated.experience[0].bullets == outcome.resume.experience[0].bullets
    binding = next(item for item in bindings if item.path == path)
    assert binding.source_claim == "负责收集用户反馈并维护需求文档"
    coverage, missing = measure_source_coverage(source, bindings)
    assert coverage == 1.0
    assert missing == []


def test_optimizer_regroups_one_record_into_a_lossless_achievement_sentence():
    from resume_optimizer import optimize_resume_with_provenance

    original_bullets = [
        "负责梳理客户需求",
        "通过10次用户访谈收集反馈",
        "输出需求优先级清单",
    ]
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "产品经理",
            "period": "2022.01-2023.01",
            "bullets": original_bullets,
        }],
    })
    combined = "负责梳理客户需求，通过10次用户访谈收集反馈并输出需求优先级清单"
    response = json.dumps({
        "experience": [{
            "index": 0,
            "bullets": [{
                "text": combined,
                "source_indices": [0, 1, 2],
            }],
        }],
    }, ensure_ascii=False)
    with patch("resume_optimizer.llm_enabled", return_value=True), patch(
        "resume_optimizer.call_llm_text", return_value=response,
    ), patch(
        "resume_optimizer.review_entailment_batch", return_value=[True],
    ):
        outcome = optimize_resume_with_provenance(resume, "产品经理")

    path = "experience[0].bullets[0]"
    assert outcome.resume.experience[0].bullets == [combined]
    assert outcome.trusted_rewrites[path] == "\n".join(original_bullets)

    source = build_source_bundle(
        "工作经历\n"
        "甲公司｜产品经理｜2022.01-2023.01\n"
        "负责梳理客户需求。\n"
        "通过10次用户访谈收集反馈。\n"
        "输出需求优先级清单。",
        "",
        "",
    )
    gated, bindings, removed = enforce_resume_evidence(
        outcome.resume,
        source,
        trusted_rewrites=outcome.trusted_rewrites,
    )
    assert removed == []
    assert gated.experience[0].bullets == [combined]
    binding = next(item for item in bindings if item.path == path)
    assert len(binding.block_ids) == 3


def test_optimizer_rejects_a_grouped_record_when_any_source_bullet_is_missing():
    from resume_optimizer import optimize_resume_with_provenance

    original_bullets = [
        "负责梳理客户需求",
        "通过10次用户访谈收集反馈",
        "输出需求优先级清单",
    ]
    resume = CanonicalResume.model_validate({
        "experience": [{"bullets": original_bullets}],
    })
    response = json.dumps({
        "experience": [{
            "index": 0,
            "bullets": [{
                "text": "负责梳理客户需求，通过10次用户访谈收集反馈",
                "source_indices": [0, 1],
            }],
        }],
    }, ensure_ascii=False)
    with patch("resume_optimizer.llm_enabled", return_value=True), patch(
        "resume_optimizer.call_llm_text", return_value=response,
    ):
        outcome = optimize_resume_with_provenance(resume, "产品经理")

    assert outcome.resume.experience[0].bullets == original_bullets
    assert outcome.trusted_rewrites == {}


def test_optimizer_keeps_unchanged_sibling_inside_an_atomic_grouped_record():
    from resume_optimizer import optimize_resume_with_provenance

    original_bullets = [
        "负责梳理客户需求通过10次用户访谈收集反馈",
        "协调研发和设计团队跟进需求落地",
        "输出需求优先级清单",
    ]
    resume = CanonicalResume.model_validate({
        "experience": [{"bullets": original_bullets}],
    })
    combined = "负责通过10次用户访谈收集反馈，梳理客户需求并输出需求优先级清单"
    response = json.dumps({
        "experience": [{
            "index": 0,
            "bullets": [
                {"text": combined, "source_indices": [0, 2]},
                {"text": original_bullets[1], "source_indices": [1]},
            ],
        }],
    }, ensure_ascii=False)
    with patch("resume_optimizer.llm_enabled", return_value=True), patch(
        "resume_optimizer.call_llm_text", return_value=response,
    ):
        outcome = optimize_resume_with_provenance(resume, "产品经理")

    assert outcome.resume.experience[0].bullets == [combined, original_bullets[1]]
    assert len(outcome.trusted_rewrites) == 2


def test_optimizer_reverts_high_risk_rewrite_when_semantic_review_rejects():
    from resume_optimizer import optimize_resume_with_provenance

    original = "参与课堂观察和作业批改"
    resume = CanonicalResume.model_validate({
        "experience": [{"bullets": [original]}],
    })
    response = json.dumps({
        "experience": [{
            "index": 0,
            "bullets": ["协助开展课堂观察记录与作业批改工作，支持日常教学运行"],
        }],
    }, ensure_ascii=False)
    with patch("resume_optimizer.llm_enabled", return_value=True), patch(
        "resume_optimizer.call_llm_text", return_value=response,
    ), patch(
        "resume_optimizer.review_entailment_batch", return_value=[False],
    ):
        outcome = optimize_resume_with_provenance(resume, "教师")

    assert outcome.resume.experience[0].bullets == [original]
    assert outcome.trusted_rewrites == {}
    assert outcome.semantic_rejected == 1


def test_semantic_review_parses_fenced_batch_and_reports_inconclusive_failure():
    from semantic_guard import review_entailment_batch

    response = """```json
    {"results": [
      {"index": 0, "pass": true, "reason": "同义改写"},
      {"index": 1, "pass": false, "reason": "责任升级"}
    ]}
    ```"""
    with patch("semantic_guard.llm_enabled", return_value=True), patch(
        "semantic_guard.remaining_request_seconds", return_value=60,
    ), patch("semantic_guard.call_llm_text", return_value=response):
        assert review_entailment_batch([("原文一", "改写一"), ("原文二", "改写二")]) == [True, False]

    with patch("semantic_guard.llm_enabled", return_value=True), patch(
        "semantic_guard.remaining_request_seconds", return_value=60,
    ), patch("semantic_guard.call_llm_text", return_value="not-json"):
        assert review_entailment_batch([("原文", "改写")]) is None


def test_semantic_review_respects_deadline_and_hard_fact_guard_runs_first():
    from resume_optimizer import optimize_resume_with_provenance
    from semantic_guard import review_entailment_batch

    with patch("semantic_guard.llm_enabled", return_value=True), patch(
        "semantic_guard.remaining_request_seconds", return_value=10,
    ), patch("semantic_guard.call_llm_text") as llm_call:
        assert review_entailment_batch([("原文", "改写")]) is None
        llm_call.assert_not_called()

    resume = CanonicalResume.model_validate({
        "experience": [{"bullets": ["负责用户反馈整理"]}],
    })
    response = json.dumps({
        "experience": [{
            "index": 0,
            "bullets": ["负责1000名用户反馈整理"],
        }],
    }, ensure_ascii=False)
    with patch("resume_optimizer.llm_enabled", return_value=True), patch(
        "resume_optimizer.call_llm_text", return_value=response,
    ), patch("resume_optimizer.review_entailment_batch") as reviewer:
        outcome = optimize_resume_with_provenance(resume, "产品经理")

    assert outcome.resume.experience[0].bullets == ["负责用户反馈整理"]
    reviewer.assert_not_called()


def test_long_resume_optimizer_batches_keep_global_record_indexes():
    from resume_optimizer import _build_optimizer_batches

    resume = CanonicalResume.model_validate({
        "experience": [
            {
                "organization": f"公司{index}",
                "role": "工程师",
                "bullets": [f"负责第{index}项已验证工作内容"],
            }
            for index in range(20)
        ],
    })
    batches = _build_optimizer_batches(resume)
    assert len(batches) >= 2
    indexes = [
        record["index"]
        for batch in batches
        for record in batch["experience"]
    ]
    assert indexes == list(range(20))


def test_long_source_composer_chunks_and_merges_without_losing_typed_sections():
    from resume_composer import _merge_drafts, _split_source_bundle

    source = build_source_bundle(
        "工作经历\n" + "\n".join(f"第{index}段真实经历" + "细节" * 60 for index in range(30)),
        "不要编造数字",
        "目标岗位：研究员",
    )
    chunks = _split_source_bundle(source, max_fact_chars=1000)
    assert len(chunks) >= 2
    assert all(any(block.source_type == "jd" for block in chunk.blocks) for chunk in chunks)

    merged = _merge_drafts([
        DraftResume.model_validate({
            "experience": [{"organization": "甲公司", "role": "工程师"}],
            "publications": ["论文A"],
        }),
        DraftResume.model_validate({
            "experience": [{"organization": "乙公司", "role": "研究员"}],
            "patents": ["专利B"],
            "additional_sections": {"专业会员": ["协会C"]},
        }),
    ])
    assert [item.organization for item in merged.experience] == ["甲公司", "乙公司"]
    assert merged.publications == ["论文A"]
    assert merged.patents == ["专利B"]
    assert merged.additional_sections == {"专业会员": ["协会C"]}


def test_visual_layout_report_uses_real_rendered_pdf(tmp_path):
    resume = {
        "meta": {"name": "李明", "target_role": "工艺工程师"},
        "summary": "材料科学与工程本科，具备新能源工艺实习和项目经历。",
        "experience": [{
            "company": "某新能源公司",
            "role": "工艺实习生",
            "period": "2025",
            "bullets": ["记录工艺参数。", "使用Minitab分析异常批次。"],
        }],
        "education": [{"school": "某工业大学", "degree": "本科", "major": "材料科学与工程"}],
        "projects": [],
        "skills": {"tools": ["Minitab"]},
    }
    pdf_path = tmp_path / "resume.pdf"
    render_pdf(resume, pdf_path, template="classic")
    report = analyze_pdf_layout(pdf_path)
    assert report["available"] is True
    assert report["page_count"] >= 1
    assert report["blank_pages"] == []
    assert report["first_page_has_core_section"] is True


def test_docx_section_heading_keeps_with_next(tmp_path):
    resume = {
        "meta": {"name": "李明"},
        "summary": "材料科学与工程本科。",
        "experience": [{"company": "某公司", "role": "实习生", "bullets": ["记录参数。"]}],
        "education": [],
        "projects": [],
        "skills": {},
    }
    output = tmp_path / "resume.docx"
    render_docx(resume, output, template="classic")
    from docx import Document
    doc = Document(output)
    assert abs(doc.sections[0].page_width.mm - 210) < 1
    assert abs(doc.sections[0].page_height.mm - 297) < 1
    heading = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "工作/实习经历")
    assert heading.paragraph_format.keep_with_next is True


def test_docx_conversion_uses_isolated_libreoffice_profile(tmp_path):
    docx_path = tmp_path / "resume.docx"
    pdf_path = tmp_path / "resume.pdf"
    docx_path.write_bytes(b"docx")

    def fake_run(args, **kwargs):
        outdir = Path(args[args.index("--outdir") + 1])
        (outdir / "resume.pdf").write_bytes(b"pdf")
        return type("Result", (), {"returncode": 0})()

    with patch("resume_renderer.shutil.which", return_value="/usr/bin/libreoffice"), patch(
        "resume_renderer.subprocess.run", side_effect=fake_run
    ) as run:
        assert _convert_docx_to_pdf(docx_path, pdf_path) is True

    args = run.call_args.args[0]
    profile_args = [value for value in args if value.startswith("-env:UserInstallation=file://")]
    assert len(profile_args) == 1
    assert pdf_path.read_bytes() == b"pdf"
