from __future__ import annotations

import uuid
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

import resume_renderer


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PACKAGE_R = "http://schemas.openxmlformats.org/package/2006/relationships"


def _resume() -> dict:
    return {
        "meta": {"name": "张三", "target_role": "产品经理"},
        "summary": "负责需求分析、方案设计与跨团队交付。",
        "experience": [
            {
                "company": "示例科技有限公司",
                "role": "高级产品经理",
                "period": "2021.03-至今",
                "bullets": ["推动三个核心模块按期上线。"],
            }
        ],
        "education": [],
        "projects": [],
        "skills": {},
    }


def _disable_layout_qa(monkeypatch) -> None:
    monkeypatch.setattr(
        resume_renderer,
        "inspect_docx_layout",
        lambda _path: {"available": False, "issues": []},
    )


def test_chinese_docx_declares_and_embeds_portable_font(tmp_path, monkeypatch):
    _disable_layout_qa(monkeypatch)
    original_font = b"OTTO" + bytes(range(64))
    monkeypatch.setattr(resume_renderer, "DEFAULT_DOC_EMBED_CJK_FONT", True)
    monkeypatch.setattr(
        resume_renderer,
        "_build_cjk_font_subset",
        lambda characters: (
            "Noto Sans CJK SC",
            original_font if "张" in characters and "•" in characters else b"",
        ),
    )

    output = tmp_path / "portable.docx"
    resume_renderer.render_docx(_resume(), output, template="classic")

    with zipfile.ZipFile(output) as archive:
        assert archive.testzip() is None
        names = set(archive.namelist())
        font_parts = [name for name in names if name.startswith("word/fonts/")]
        assert len(font_parts) == 1

        document = etree.fromstring(archive.read("word/document.xml"))
        for run in document.xpath("//w:r[.//w:t]", namespaces={"w": W}):
            text = "".join(run.xpath(".//w:t/text()", namespaces={"w": W}))
            if resume_renderer._contains_east_asian_text(text):
                fonts = run.xpath("./w:rPr/w:rFonts", namespaces={"w": W})
                assert len(fonts) == 1
                assert fonts[0].get(f"{{{W}}}eastAsia") == "Noto Sans CJK SC"
                assert fonts[0].get(f"{{{W}}}ascii") == "Noto Sans CJK SC"
                assert fonts[0].get(f"{{{W}}}hAnsi") == "Noto Sans CJK SC"

        numbering = etree.fromstring(archive.read("word/numbering.xml"))
        bullet_values = numbering.xpath(
            "//w:lvl[w:numFmt[@w:val='bullet']]/w:lvlText/@w:val",
            namespaces={"w": W},
        )
        assert "•" in bullet_values
        assert all("\uf0b7" not in value for value in bullet_values)

        font_table = etree.fromstring(archive.read("word/fontTable.xml"))
        embeds = font_table.xpath(
            "./w:font[@w:name='Noto Sans CJK SC']/w:embedRegular",
            namespaces={"w": W},
        )
        assert len(embeds) == 1
        relationship_id = embeds[0].get(f"{{{R}}}id")
        font_key = uuid.UUID(embeds[0].get(f"{{{W}}}fontKey").strip("{}"))

        relationships = etree.fromstring(archive.read("word/_rels/fontTable.xml.rels"))
        targets = {
            item.get("Id"): item.get("Target")
            for item in relationships.findall(f"{{{PACKAGE_R}}}Relationship")
        }
        assert targets[relationship_id] == font_parts[0].removeprefix("word/")
        embedded = archive.read(font_parts[0])
        assert resume_renderer._obfuscate_ooxml_font(embedded, font_key) == original_font

    editable_text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert "张三" in editable_text
    assert "推动三个核心模块按期上线" in editable_text


def test_chinese_docx_stays_valid_when_embedding_is_unavailable(tmp_path, monkeypatch):
    _disable_layout_qa(monkeypatch)
    monkeypatch.setattr(resume_renderer, "DEFAULT_DOC_EMBED_CJK_FONT", True)

    def unavailable(_characters):
        raise FileNotFoundError("test font missing")

    monkeypatch.setattr(resume_renderer, "_build_cjk_font_subset", unavailable)
    output = tmp_path / "fallback.docx"
    resume_renderer.render_docx(_resume(), output, template="classic")

    with zipfile.ZipFile(output) as archive:
        assert archive.testzip() is None
        assert not any(name.startswith("word/fonts/") for name in archive.namelist())
        document = etree.fromstring(archive.read("word/document.xml"))
        east_asia_fonts = document.xpath(
            "//w:r[.//w:t[contains(., '张')]]/w:rPr/w:rFonts/@w:eastAsia",
            namespaces={"w": W},
        )
        assert east_asia_fonts == ["Noto Sans CJK SC"]


def test_latin_only_docx_does_not_embed_cjk_font(tmp_path, monkeypatch):
    _disable_layout_qa(monkeypatch)
    monkeypatch.setattr(resume_renderer, "DEFAULT_DOC_EMBED_CJK_FONT", True)
    monkeypatch.setattr(
        resume_renderer,
        "_build_cjk_font_subset",
        lambda _characters: (_ for _ in ()).throw(AssertionError("must not subset")),
    )
    resume = {
        "meta": {"name": "Alex Smith"},
        "summary": "Product manager with enterprise delivery experience.",
        "experience": [],
        "education": [],
        "projects": [],
        "skills": {},
    }
    output = tmp_path / "latin.docx"
    resume_renderer.render_docx(resume, output, template="classic")

    with zipfile.ZipFile(output) as archive:
        assert not any(name.startswith("word/fonts/") for name in archive.namelist())


def test_cjk_portability_preserves_explicit_template_latin_font(tmp_path, monkeypatch):
    _disable_layout_qa(monkeypatch)
    monkeypatch.setattr(resume_renderer, "DEFAULT_DOC_EMBED_CJK_FONT", False)
    template = tmp_path / "template.docx"
    source = Document()
    paragraph = source.add_paragraph("个人总结")
    paragraph.runs[0].font.name = "Arial"
    source.add_paragraph("示例内容").runs[0].font.name = "Arial"
    source.save(template)

    output = tmp_path / "preserved.docx"
    resume_renderer.render_docx(_resume(), output, template=str(template))

    with zipfile.ZipFile(output) as archive:
        document = etree.fromstring(archive.read("word/document.xml"))
        fonts = document.xpath(
            "//w:r[.//w:t[contains(., '需求分析')]]/w:rPr/w:rFonts",
            namespaces={"w": W},
        )
        assert fonts
        assert all(node.get(f"{{{W}}}ascii") == "Arial" for node in fonts)
        assert all(node.get(f"{{{W}}}hAnsi") == "Arial" for node in fonts)
        assert all(node.get(f"{{{W}}}eastAsia") for node in fonts)
