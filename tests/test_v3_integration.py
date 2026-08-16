from io import BytesIO
import json
import threading

import pytest
from pydantic import ValidationError

from core.schemas import ResumeCopilotResponse
from core.v3.contracts import SourceAsset
from core.v3.document_graph import from_native_text
from core.v3.fact_graph import build_fact_graph
from core.v3.input_adapters import build_input_document_graph
from core.v3.jd_graph import build_requirement_graph
from core.v3.pipeline import run_v3_pipeline
from core.v3.planner import plan_resume
from core.v3.realizer import realize_plan
from core.v3.resume_adapter import frozen_to_resume_data
from core.v3.section_ontology import section_type
from core.v3.semantic_llm import compile_semantics
from core.v3.training_examples import build_training_records, maybe_write_training_trace
from core.v3.training_schema import (
    ConstrainedRealizerResponse,
    SCHEMA_FINGERPRINT,
    SCHEMA_VERSION,
    SemanticCompilationResponse,
    SemanticFactDecision,
    schema_fingerprint,
)


def _semantic_response_for_name_and_action(_model, _system, user_prompt, **_kwargs):
    request = json.loads(user_prompt)
    fact_id = request["candidates"][0]["candidate_fact_id"]
    return {
        "schema_version": SCHEMA_VERSION,
        "decisions": [{
            "candidate_fact_id": fact_id,
            "classification": "fact",
            "record_id": None,
            "atoms": [
                {
                    "quote": "张三",
                    "fact_type": "identity",
                    "destination_section": "contact",
                    "destination_field": "name",
                },
                {
                    "quote": "负责用户访谈",
                    "fact_type": "action",
                    "destination_section": "experience",
                    "destination_field": "bullet",
                },
            ],
            "context_spans": [],
        }],
    }


def test_frozen_schema_rejects_version_drift():
    with pytest.raises(ValidationError):
        SemanticCompilationResponse.model_validate({
            "schema_version": "resume_compiler_v3.1",
            "decisions": [],
        })


def test_frozen_schema_rejects_implicit_or_empty_model_outputs():
    with pytest.raises(ValidationError):
        SemanticCompilationResponse.model_validate({})
    with pytest.raises(ValidationError):
        SemanticCompilationResponse.model_validate({
            "schema_version": SCHEMA_VERSION,
            "decisions": [],
        })
    with pytest.raises(ValidationError):
        ConstrainedRealizerResponse.model_validate({
            "schema_version": SCHEMA_VERSION,
            "request_fact_ids": [],
            "claims": [],
        })


def test_frozen_schema_fingerprint_has_not_drifted():
    assert schema_fingerprint() == SCHEMA_FINGERPRINT


def test_frozen_schema_has_generic_degree_and_major_atoms():
    payload = SemanticCompilationResponse.model_validate({
        "schema_version": SCHEMA_VERSION,
        "decisions": [{
            "candidate_fact_id": "cv:fact:1",
            "classification": "fact",
            "record_id": None,
            "atoms": [
                {
                    "quote": "本科",
                    "fact_type": "degree",
                    "destination_section": "education",
                    "destination_field": "degree",
                },
                {
                    "quote": "临床医学",
                    "fact_type": "major",
                    "destination_section": "education",
                    "destination_field": "major",
                },
            ],
            "context_spans": [],
        }],
    })

    assert [atom.fact_type for atom in payload.decisions[0].atoms] == ["degree", "major"]


@pytest.mark.parametrize("classification", ["context", "intent", "instruction"])
def test_frozen_schema_requires_exact_spans_for_every_non_fact_classification(classification):
    with pytest.raises(ValidationError):
        SemanticFactDecision.model_validate({
            "candidate_fact_id": "query:fact:1",
            "classification": classification,
            "record_id": None,
            "atoms": [],
            "context_spans": [],
        })


def test_frozen_schema_has_no_unbounded_other_context_reason():
    with pytest.raises(ValidationError):
        SemanticFactDecision.model_validate({
            "candidate_fact_id": "query:fact:1",
            "classification": "context",
            "record_id": None,
            "atoms": [],
            "context_spans": [{"quote": "精通", "reason": "other"}],
        })


def test_semantic_compiler_creates_only_exact_source_atoms():
    text = "张三，负责用户访谈"
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    result = compile_semantics(
        fact_graph,
        llm_call=_semantic_response_for_name_and_action,
    )

    assert result.report.status == "success"
    assert [fact.text for fact in result.graph.eligible_facts()] == ["张三", "负责用户访谈"]
    assert all(
        "".join(span.quote(result.graph.documents) for span in fact.spans) == fact.text
        for fact in result.graph.facts
    )
    assert {fact.destination_section for fact in result.graph.eligible_facts()} == {"contact", "experience"}


def test_semantic_context_span_accounts_for_label_without_rendering_it():
    text = "字段标签：真实值"
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def response(_model, _system, user_prompt, **_kwargs):
        fact_id = json.loads(user_prompt)["candidates"][0]["candidate_fact_id"]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": fact_id,
                "classification": "fact",
                "record_id": None,
                "atoms": [{
                    "quote": "真实值",
                    "fact_type": "other",
                    "destination_section": "additional",
                    "destination_field": "item",
                }],
                "context_spans": [{"quote": "字段标签：", "reason": "label"}],
            }],
        }

    result = compile_semantics(fact_graph, llm_call=response)

    assert result.report.status == "success"
    assert [fact.text for fact in result.graph.eligible_facts()] == ["真实值"]
    context = [fact for fact in result.graph.facts if fact.fact_id in result.report.context_fact_ids]
    assert [fact.text for fact in context] == ["字段标签："]
    assert all(not fact.eligible for fact in context)


def test_top_level_key_value_label_is_structural_context_even_when_model_emits_full_line():
    text = "Professional Direction: Project Manager"
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def response(_model, _system, user_prompt, **_kwargs):
        candidate = json.loads(user_prompt)["candidates"][0]
        assert candidate["structural_context_spans"] == [{
            "quote": "Professional Direction: ",
            "reason": "label",
            "char_start": 0,
            "char_end": 24,
        }]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": candidate["candidate_fact_id"],
                "classification": "fact",
                "record_id": None,
                "atoms": [{
                    "quote": text,
                    "fact_type": "role",
                    "destination_section": "summary",
                    "destination_field": "summary",
                }],
                "context_spans": [],
            }],
        }

    result = compile_semantics(fact_graph, llm_call=response)

    assert result.report.fallback_fact_ids == ()
    assert [fact.text for fact in result.graph.eligible_facts()] == ["Project Manager"]
    contexts = [fact for fact in result.graph.facts if not fact.eligible]
    assert [fact.text for fact in contexts] == ["Professional Direction: "]
    assert any("structural_label_split_from_fact" in error for error in result.report.errors)


def test_top_level_key_value_label_atoms_are_recovered_without_duplicate_context():
    text = "Career Level: Junior"
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def response(_model, _system, user_prompt, **_kwargs):
        candidate = json.loads(user_prompt)["candidates"][0]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": candidate["candidate_fact_id"],
                "classification": "fact",
                "record_id": None,
                "atoms": [
                    {
                        "quote": "Career Level",
                        "fact_type": "other",
                        "destination_section": "summary",
                        "destination_field": "summary",
                    },
                    {
                        "quote": "Junior",
                        "fact_type": "other",
                        "destination_section": "summary",
                        "destination_field": "summary",
                    },
                ],
                "context_spans": [{"quote": ": ", "reason": "separator"}],
            }],
        }

    result = compile_semantics(fact_graph, llm_call=response)

    assert result.report.fallback_fact_ids == ()
    assert [fact.text for fact in result.graph.eligible_facts()] == ["Junior"]
    assert [fact.text for fact in result.graph.facts if not fact.eligible] == ["Career Level: "]


def test_record_local_key_value_label_cannot_become_a_standalone_bullet():
    text = "工作经历\n时期：2012年9月至2013年11月"
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def response(_model, _system, user_prompt, **_kwargs):
        candidate = json.loads(user_prompt)["candidates"][0]
        assert candidate["structural_context_spans"] == [{
            "quote": "时期：",
            "reason": "label",
            "char_start": 0,
            "char_end": 3,
        }]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": candidate["candidate_fact_id"],
                "classification": "fact",
                "record_id": candidate["locked_record_id"],
                "atoms": [
                    {
                        "quote": "时期：",
                        "fact_type": "other",
                        "destination_section": "experience",
                        "destination_field": "bullet",
                    },
                    {
                        "quote": "2012年9月至2013年11月",
                        "fact_type": "period",
                        "destination_section": "experience",
                        "destination_field": "period",
                    },
                ],
                "context_spans": [],
            }],
        }

    result = compile_semantics(fact_graph, llm_call=response)

    assert [fact.text for fact in result.graph.eligible_facts()] == [
        "2012年9月至2013年11月"
    ]
    assert result.graph.eligible_facts()[0].destination_field == "period"
    assert any(
        "structural_label_emitted_as_fact" in error
        for error in result.report.errors
    )


def test_short_uppercase_key_remains_a_factual_acronym():
    text = "工作经历\nCRM: 客户关系管理"
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def response(_model, _system, user_prompt, **_kwargs):
        candidate = json.loads(user_prompt)["candidates"][0]
        assert "structural_context_spans" not in candidate
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": candidate["candidate_fact_id"],
                "classification": "fact",
                "record_id": candidate["locked_record_id"],
                "atoms": [{
                    "quote": "CRM: 客户关系管理",
                    "fact_type": "skill",
                    "destination_section": "experience",
                    "destination_field": "bullet",
                }],
                "context_spans": [],
            }],
        }

    result = compile_semantics(fact_graph, llm_call=response)

    assert [fact.text for fact in result.graph.eligible_facts()] == [
        "CRM: 客户关系管理"
    ]


def test_semantic_context_cannot_hide_numeric_hard_anchor():
    text = "结果：提升30%"
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def response(_model, _system, user_prompt, **_kwargs):
        fact_id = json.loads(user_prompt)["candidates"][0]["candidate_fact_id"]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": fact_id,
                "classification": "fact",
                "record_id": None,
                "atoms": [{
                    "quote": "结果：",
                    "fact_type": "other",
                    "destination_section": "additional",
                    "destination_field": "item",
                }],
                "context_spans": [{"quote": "提升30%", "reason": "label"}],
            }],
        }

    result = compile_semantics(fact_graph, llm_call=response)

    assert result.report.status == "fallback"
    assert result.graph.eligible_facts()[0].text == text
    assert any("context_hides_hard_anchor" in error for error in result.report.errors)


def test_semantic_whole_candidate_placeholder_is_auditable_but_not_rendered():
    text = "[姓名]"
    graph = from_native_text(SourceAsset(
        source_id="query", source_type="query", filename="query.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def response(_model, _system, user_prompt, **_kwargs):
        fact_id = json.loads(user_prompt)["candidates"][0]["candidate_fact_id"]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": fact_id,
                "classification": "context",
                "record_id": None,
                "atoms": [],
                "context_spans": [{"quote": text, "reason": "placeholder"}],
            }],
        }

    result = compile_semantics(fact_graph, llm_call=response)

    assert result.report.status == "success"
    assert result.graph.eligible_facts() == []
    assert result.report.non_fact_ids == result.report.input_fact_ids
    context = [fact for fact in result.graph.facts if fact.fact_id in result.report.context_fact_ids]
    assert [fact.text for fact in context] == [text]
    assert context[0].classification == "ineligible"


def test_semantic_whole_candidate_context_cannot_hide_hard_anchor():
    text = "占位结果30%"
    graph = from_native_text(SourceAsset(
        source_id="query", source_type="query", filename="query.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def response(_model, _system, user_prompt, **_kwargs):
        fact_id = json.loads(user_prompt)["candidates"][0]["candidate_fact_id"]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": fact_id,
                "classification": "context",
                "record_id": None,
                "atoms": [],
                "context_spans": [{"quote": text, "reason": "placeholder"}],
            }],
        }

    result = compile_semantics(fact_graph, llm_call=response)

    assert result.report.status == "fallback"
    assert result.graph.eligible_facts() == []
    assert result.report.fail_closed_fact_ids == result.report.input_fact_ids
    assert any("non_fact_context_hides_hard_anchor" in error for error in result.report.errors)


def test_semantic_non_fact_context_must_cover_all_substantive_source_text():
    text = "以下是全部个人信息"
    graph = from_native_text(SourceAsset(
        source_id="query", source_type="query", filename="query.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def response(_model, _system, user_prompt, **_kwargs):
        fact_id = json.loads(user_prompt)["candidates"][0]["candidate_fact_id"]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": fact_id,
                "classification": "instruction",
                "record_id": None,
                "atoms": [],
                "context_spans": [{"quote": "以下是", "reason": "instruction"}],
            }],
        }

    result = compile_semantics(fact_graph, llm_call=response)

    assert result.report.status == "fallback"
    assert result.graph.eligible_facts() == []
    assert result.report.fail_closed_fact_ids == result.report.input_fact_ids
    assert any("non_fact_context_not_complete" in error for error in result.report.errors)


def test_semantic_query_fact_and_placeholder_are_split_without_leaking_placeholder():
    text = "姓名：张三 邮箱：[邮箱]"
    graph = from_native_text(SourceAsset(
        source_id="query", source_type="query", filename="query.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def response(_model, _system, user_prompt, **_kwargs):
        candidate = json.loads(user_prompt)["candidates"][0]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": candidate["candidate_fact_id"],
                "classification": "fact",
                "record_id": candidate["locked_record_id"],
                "atoms": [{
                    "quote": "张三",
                    "fact_type": "identity",
                    "destination_section": "contact",
                    "destination_field": "name",
                }],
                "context_spans": [
                    {"quote": "姓名：", "reason": "label"},
                    {"quote": " 邮箱：[邮箱]", "reason": "placeholder"},
                ],
            }],
        }

    result = compile_semantics(fact_graph, llm_call=response)

    assert result.report.status == "success"
    assert [fact.text for fact in result.graph.eligible_facts()] == ["张三"]
    assert "[邮箱]" not in "".join(fact.text for fact in result.graph.eligible_facts())


@pytest.mark.parametrize("placeholder", ["[公司]", "【公司"])
def test_semantic_placeholder_atom_cannot_discard_valid_role_and_period_siblings(
    placeholder: str,
):
    text = f"生产专员，{placeholder}2022年至今"
    fact_graph = build_fact_graph([from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=False,
    ))])

    def response(_model, _system, user_prompt, **_kwargs):
        candidate = json.loads(user_prompt)["candidates"][0]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": candidate["candidate_fact_id"],
                "classification": "fact",
                "record_id": candidate["locked_record_id"],
                "atoms": [
                    {
                        "quote": "生产专员",
                        "fact_type": "role",
                        "destination_section": "experience",
                        "destination_field": "role",
                    },
                    {
                        "quote": placeholder,
                        "fact_type": "organization",
                        "destination_section": "experience",
                        "destination_field": "organization",
                    },
                    {
                        "quote": "2022年至今",
                        "fact_type": "period",
                        "destination_section": "experience",
                        "destination_field": "period",
                    },
                ],
                "context_spans": [],
            }],
        }

    compiled = compile_semantics(fact_graph, llm_call=response)
    plan = plan_resume(compiled.graph, build_requirement_graph(""))
    frozen = realize_plan(plan, compiled.graph)
    resume_data = frozen_to_resume_data(frozen, compiled.graph)

    assert [fact.text for fact in compiled.graph.eligible_facts()] == [
        "生产专员", "2022年至今",
    ]
    assert any(
        fact.text == placeholder and not fact.eligible
        for fact in compiled.graph.facts
    )
    assert resume_data["experience"] == [{
        "role": "生产专员",
        "period": "2022年至今",
    }]
    assert placeholder not in json.dumps(resume_data, ensure_ascii=False)


def test_placeholder_atom_and_context_span_indexes_cannot_collide():
    text = "[公司]，生产专员"
    fact_graph = build_fact_graph([from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=False,
    ))])

    def response(_model, _system, user_prompt, **_kwargs):
        candidate = json.loads(user_prompt)["candidates"][0]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": candidate["candidate_fact_id"],
                "classification": "fact",
                "record_id": candidate["locked_record_id"],
                "atoms": [
                    {
                        "quote": "[公司]",
                        "fact_type": "organization",
                        "destination_section": "experience",
                        "destination_field": "organization",
                    },
                    {
                        "quote": "生产专员",
                        "fact_type": "role",
                        "destination_section": "experience",
                        "destination_field": "role",
                    },
                ],
                "context_spans": [{"quote": "，", "reason": "separator"}],
            }],
        }

    compiled = compile_semantics(fact_graph, llm_call=response)
    fact_ids = [fact.fact_id for fact in compiled.graph.facts]
    contexts = [
        fact for fact in compiled.graph.facts
        if fact.fact_id in compiled.report.context_fact_ids
    ]

    assert len(fact_ids) == len(set(fact_ids))
    assert [fact.text for fact in compiled.graph.eligible_facts()] == ["生产专员"]
    assert {fact.text for fact in contexts} == {"[公司]", "，"}
    assert any(":context:atom:0" in fact.fact_id for fact in contexts)
    assert any(":context:span:0" in fact.fact_id for fact in contexts)


def test_record_ids_are_unique_across_multiple_source_sections():
    text = "工作经历\n甲公司 产品经理 2020-2022\n负责访谈\n教育经历\n甲大学 本科 2016-2020"
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))

    fact_graph = build_fact_graph([graph])

    record_ids = [record.record_id for record in fact_graph.records]
    assert len(record_ids) >= 2
    assert len(record_ids) == len(set(record_ids))


def test_numbered_generic_query_headings_define_sections_without_becoming_facts():
    text = """1. 联系信息
张三
2. 总结或目标
十年跨国财务经验
3. 技能
财务分析
4. 专业经验
甲公司 2020年至今
财务经理
5. 教育背景
甲大学 2016年
工商管理硕士
6. 认证和执照
注册内部审计师
8. 奖项和荣誉
年度优秀员工
9. 志愿服务经历
社区财务委员会成员
10. 语言能力
法语流利
11. 兴趣爱好
可持续经济学"""
    graph = from_native_text(SourceAsset(
        source_id="query", source_type="query", filename="query.txt", text=text, native=True,
    ))

    fact_graph = build_fact_graph([graph])
    fact_texts = {fact.text for fact in fact_graph.facts}
    section_by_id = {section.section_id: section.section_type for section in fact_graph.sections}
    section_by_text = {
        fact.text: section_by_id.get(fact.section_id, "other")
        for fact in fact_graph.facts
    }

    assert not any(value in fact_texts for value in (
        "1. 联系信息", "2. 总结或目标", "3. 技能", "4. 专业经验",
        "5. 教育背景", "6. 认证和执照", "8. 奖项和荣誉",
    ))
    assert section_by_text["张三"] == "contact"
    assert section_by_text["十年跨国财务经验"] == "summary"
    assert section_by_text["财务分析"] == "skills"
    assert section_by_text["甲公司 2020年至今"] == "experience"
    assert section_by_text["甲大学 2016年"] == "education"
    assert section_by_text["注册内部审计师"] == "credentials"
    assert section_by_text["年度优秀员工"] == "awards"
    assert section_by_text["社区财务委员会成员"] == "activities"
    assert section_by_text["法语流利"] == "skills"
    assert section_by_text["可持续经济学"] == "additional"


def test_repeated_section_layout_ordinals_keep_the_same_section_type():
    assert section_type("工作经历 - II") == "experience"
    assert section_type("Work Experience 2") == "experience"

    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt",
        text="教育经历\n甲大学\n工作经历 - II\n负责用户访谈", native=True,
    ))
    fact_graph = build_fact_graph([graph])
    section_by_id = {
        section.section_id: section.section_type for section in fact_graph.sections
    }
    section_by_text = {
        fact.text: section_by_id.get(fact.section_id, "other")
        for fact in fact_graph.facts
    }

    assert "工作经历 - II" not in section_by_text
    assert section_by_text["负责用户访谈"] == "experience"


def test_semantic_compiler_preserves_only_same_destination_inter_atom_transport():
    text = "- Delivered results using structured workflows and clear communication"
    graph = from_native_text(SourceAsset(
        source_id="query", source_type="query", filename="query.txt",
        text="Professional Experience\n" + text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def response(_model, _system, user_prompt, **_kwargs):
        candidate = json.loads(user_prompt)["candidates"][0]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": candidate["candidate_fact_id"],
                "classification": "fact",
                "record_id": candidate["locked_record_id"],
                "atoms": [
                    {
                        "quote": "Delivered results",
                        "fact_type": "action",
                        "destination_section": "experience",
                        "destination_field": "bullet",
                    },
                    {
                        "quote": "structured workflows",
                        "fact_type": "method",
                        "destination_section": "experience",
                        "destination_field": "bullet",
                    },
                    {
                        "quote": "clear communication",
                        "fact_type": "method",
                        "destination_section": "experience",
                        "destination_field": "bullet",
                    },
                ],
                "context_spans": [],
            }],
        }

    compiled = compile_semantics(fact_graph, llm_call=response)
    plan = plan_resume(compiled.graph, build_requirement_graph(""))
    frozen = realize_plan(plan, compiled.graph)

    assert compiled.report.fallback_fact_ids == ()
    assert any(
        "implicit_transport_gap_preserved" in error
        for error in compiled.report.errors
    )
    assert [claim.text for claim in frozen.claims] == [
        "Delivered results using structured workflows and clear communication"
    ]


def test_semantic_compiler_still_rejects_unclaimed_leading_substance():
    text = "Consistently delivered results using structured workflows"
    graph = from_native_text(SourceAsset(
        source_id="query", source_type="query", filename="query.txt",
        text="Professional Experience\n" + text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def response(_model, _system, user_prompt, **_kwargs):
        candidate = json.loads(user_prompt)["candidates"][0]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": candidate["candidate_fact_id"],
                "classification": "fact",
                "record_id": candidate["locked_record_id"],
                "atoms": [
                    {
                        "quote": "delivered results",
                        "fact_type": "action",
                        "destination_section": "experience",
                        "destination_field": "bullet",
                    },
                    {
                        "quote": "structured workflows",
                        "fact_type": "method",
                        "destination_section": "experience",
                        "destination_field": "bullet",
                    },
                ],
                "context_spans": [],
            }],
        }

    compiled = compile_semantics(fact_graph, llm_call=response)

    assert compiled.graph.eligible_facts() == []
    assert compiled.report.fallback_fact_ids
    assert any(
        "substantive_source_not_covered" in error
        for error in compiled.report.errors
    )


def test_generic_document_title_is_structural_and_never_becomes_a_fact():
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt",
        text="个人简历\n张三\n志愿者经历\n社区服务", native=True,
    ))

    fact_graph = build_fact_graph([graph])

    assert "个人简历" not in {fact.text for fact in fact_graph.facts}
    assert "志愿者经历" not in {fact.text for fact in fact_graph.facts}
    assert {fact.text for fact in fact_graph.facts} == {"张三", "社区服务"}


def test_pure_record_field_labels_are_auditable_context_not_resume_facts():
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt",
        text=(
            "工作经历 - III\n"
            "职位：销售协调员\n"
            "组织：创新解决方案有限公司\n"
            "行业：\n"
            "职责：\n"
            "负责客户沟通"
        ),
        native=True,
    ))

    fact_graph = build_fact_graph([graph])
    by_text = {fact.text: fact for fact in fact_graph.facts}

    assert by_text["行业："].classification == "ineligible"
    assert by_text["职责："].classification == "ineligible"
    assert by_text["行业："].eligible is False
    assert by_text["职责："].eligible is False
    assert by_text["职位：销售协调员"].eligible is True
    assert by_text["组织：创新解决方案有限公司"].eligible is True
    assert by_text["负责客户沟通"].eligible is True


def test_record_local_achievement_label_does_not_split_experience_ownership():
    text = """4. 专业经验
甲公司 2020年至2022年
产品经理
主要成就：
- 建立库存管理制度
乙公司 2018年至2020年
运营经理
成就：
- 将处理时间减少30%"""
    graph = from_native_text(SourceAsset(
        source_id="query", source_type="query", filename="query.txt", text=text, native=True,
    ))

    fact_graph = build_fact_graph([graph])
    record_by_text = {fact.text: fact.record_id for fact in fact_graph.facts}

    assert len(fact_graph.records) == 2
    assert record_by_text["产品经理"] == record_by_text["- 建立库存管理制度"]
    assert record_by_text["产品经理"] == record_by_text["主要成就："]
    assert record_by_text["运营经理"] == record_by_text["- 将处理时间减少30%"]
    assert record_by_text["运营经理"] == record_by_text["成就："]
    assert record_by_text["产品经理"] != record_by_text["运营经理"]


def test_indentation_does_not_create_one_record_per_experience_line():
    text = """4. 专业经验
甲公司                                 2020年至2022年
   产品经理
   主要成就：
   - 建立库存管理制度
乙公司                                 2018年至2020年
   运营经理
   - 将处理时间减少30%"""
    graph = from_native_text(SourceAsset(
        source_id="query", source_type="query", filename="query.txt", text=text, native=True,
    ))

    fact_graph = build_fact_graph([graph])
    record_by_text = {fact.text: fact.record_id for fact in fact_graph.facts}

    assert len(fact_graph.records) == 2
    assert record_by_text["产品经理"] == record_by_text["- 建立库存管理制度"]
    assert record_by_text["运营经理"] == record_by_text["- 将处理时间减少30%"]
    assert record_by_text["产品经理"] != record_by_text["运营经理"]


def test_narrative_commas_do_not_split_a_record():
    text = """工作经历
甲公司 2020年至今
产品经理
负责访谈，输出需求优先级
主导上线，提升转化率30%"""
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))

    fact_graph = build_fact_graph([graph])
    record_ids = {fact.record_id for fact in fact_graph.facts if fact.text != "工作经历"}

    assert len(fact_graph.records) == 1
    assert record_ids == {fact_graph.records[0].record_id}


def test_implicit_query_experience_date_ranges_create_distinct_locked_records():
    text = """以下是我提供的全部个人信息。
2022年6月 - 至今
[公司]
财务总监
- 领导预算流程
2018年3月 - 2022年5月
[公司]
助理财务经理
- 制定内部控制框架
2008年6月 - 2018年2月
[公司]
高级财务分析师
- 监督财务运营
教育背景
2000 [学校]
经济学硕士"""
    graph = from_native_text(SourceAsset(
        source_id="query", source_type="query", filename="query.txt", text=text, native=True,
    ))

    fact_graph = build_fact_graph([graph])
    record_by_text = {fact.text: fact.record_id for fact in fact_graph.facts}

    experience_records = {
        record_by_text["2022年6月 - 至今"],
        record_by_text["2018年3月 - 2022年5月"],
        record_by_text["2008年6月 - 2018年2月"],
    }
    assert None not in experience_records
    assert len(experience_records) == 3
    assert record_by_text["财务总监"] == record_by_text["- 领导预算流程"]
    assert record_by_text["助理财务经理"] == record_by_text["- 制定内部控制框架"]
    assert record_by_text["高级财务分析师"] == record_by_text["- 监督财务运营"]
    assert len(experience_records | {record_by_text["2000 [学校]"]}) == 4


def test_short_month_year_ranges_split_plain_text_experience_records():
    text = """工作经历
科技解决方案有限公司 [城市] 9/21-2/23
IT支持专员
- 提供硬件支持
IT Consultants Co. [城市] 11/19-6/21
系统管理员
- 管理系统镜像
全球技术支持服务 [城市] 5/18-8/19
服务台协调员
- 解决技术问题
运营IT支持 [城市] 11/15-2/17
技术支持专员
- 管理用户访问权限"""
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))

    fact_graph = build_fact_graph([graph])
    record_by_text = {fact.text: fact.record_id for fact in fact_graph.facts}

    assert len(fact_graph.records) == 4
    assert record_by_text["IT支持专员"] == record_by_text["- 提供硬件支持"]
    assert record_by_text["系统管理员"] == record_by_text["- 管理系统镜像"]
    assert record_by_text["服务台协调员"] == record_by_text["- 解决技术问题"]
    assert record_by_text["技术支持专员"] == record_by_text["- 管理用户访问权限"]
    assert len({
        record_by_text["IT支持专员"],
        record_by_text["系统管理员"],
        record_by_text["服务台协调员"],
        record_by_text["技术支持专员"],
    }) == 4


def test_date_range_inside_narrative_does_not_create_a_record_boundary():
    text = """补充信息
在2021年至2022年期间参与系统升级
持续支持团队协作"""
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))

    fact_graph = build_fact_graph([graph])

    assert fact_graph.records == []
    assert all(fact.record_id is None for fact in fact_graph.facts)


def test_role_before_period_is_reassigned_from_previous_record():
    text = """工作经历
- **机械工长**
[公司], [城市]
2016年4月 – 2018年6月
- 实施安全审计
- **机械主管**
[公司], [城市]
2012年6月 – 2015年3月
- 管理施工运营
- **管道技术员**
[公司], [城市]
2009年8月 – 2012年5月
- 领导管道项目"""
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))

    fact_graph = build_fact_graph([graph])
    record_by_text = {fact.text: fact.record_id for fact in fact_graph.facts}

    assert len(fact_graph.records) == 3
    assert record_by_text["- **机械工长**"] == record_by_text["2016年4月 – 2018年6月"]
    assert record_by_text["- **机械主管**"] == record_by_text["2012年6月 – 2015年3月"]
    assert record_by_text["- **管道技术员**"] == record_by_text["2009年8月 – 2012年5月"]
    assert record_by_text["- 实施安全审计"] != record_by_text["- **机械主管**"]
    assert record_by_text["- 管理施工运营"] != record_by_text["- **管道技术员**"]


def test_blank_line_stops_implicit_record_prefix_lookbehind():
    text = """工作经验：17年零6个月

2022年6月 - 至今
财务总监
- 领导预算流程"""
    graph = from_native_text(SourceAsset(
        source_id="query", source_type="query", filename="query.txt", text=text, native=True,
    ))

    fact_graph = build_fact_graph([graph])
    record_by_text = {fact.text: fact.record_id for fact in fact_graph.facts}

    assert record_by_text["工作经验：17年零6个月"] is None
    assert record_by_text["2022年6月 - 至今"] is not None


def test_decimal_metric_is_not_a_date_or_record_boundary():
    text = """教育经历
机械技术工程学士学位
2011年毕业
平均绩点: 2.85
机械工程文凭"""
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))

    fact_graph = build_fact_graph([graph])
    record_by_text = {fact.text: fact.record_id for fact in fact_graph.facts}

    assert len(fact_graph.records) == 1
    assert record_by_text["平均绩点: 2.85"] == record_by_text["2011年毕业"]


def test_leading_middle_dot_bullets_do_not_create_one_record_per_line():
    text = """工作经历
公司报告解决方案部门 2012年3月-至今
指数与分析总监
·制定增长机会战略
·制定商业框架和运营蓝图
·利用与关键行业利益相关者的关系
Financial Insights 2008年2月-2012年2月
客户管理总监
·主导客户保留策略
·与主要客户建立关系"""
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))

    fact_graph = build_fact_graph([graph])
    record_by_text = {fact.text: fact.record_id for fact in fact_graph.facts}

    assert len(fact_graph.records) == 2
    assert record_by_text["·制定增长机会战略"] == record_by_text["指数与分析总监"]
    assert record_by_text["·利用与关键行业利益相关者的关系"] == record_by_text["指数与分析总监"]
    assert record_by_text["·主导客户保留策略"] == record_by_text["客户管理总监"]
    assert record_by_text["指数与分析总监"] != record_by_text["客户管理总监"]


def test_ocr_date_placeholder_range_keeps_preceding_header_as_new_record():
    text = """专业经历
甲公司，客户经理
2020年1月-2022年2月
负责客户维护
Market Data
Inc，高级客户关系经理
年月-年月
负责产品留存
"""
    fact_graph = build_fact_graph([from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="resume.txt",
        text=text, native=False,
    ))])
    record_by_text = {fact.text: fact.record_id for fact in fact_graph.facts}

    assert len(fact_graph.records) == 2
    assert record_by_text["Market Data"] == record_by_text["年月-年月"]
    assert record_by_text["Inc，高级客户关系经理"] == record_by_text["年月-年月"]
    assert record_by_text["负责客户维护"] != record_by_text["负责产品留存"]
    placeholder = next(fact for fact in fact_graph.facts if fact.text == "年月-年月")
    assert placeholder.record_id == record_by_text["Market Data"]
    assert placeholder.eligible is False
    assert placeholder.classification == "ineligible"


def test_date_placeholder_is_not_sent_to_semantic_model_or_rendered():
    text = """工作经历
客户经理，[公司]
年月-年月
负责客户维护"""
    fact_graph = build_fact_graph([from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="resume.txt",
        text=text, native=False,
    ))])
    seen: list[str] = []

    def exact_semantics(_model, _system, user_prompt, **_kwargs):
        candidates = json.loads(user_prompt)["candidates"]
        seen.extend(item["candidate_text"] for item in candidates)
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [
                {
                    "candidate_fact_id": item["candidate_fact_id"],
                    "classification": "fact",
                    "record_id": item["locked_record_id"],
                    "atoms": [{
                        "quote": item["candidate_text"],
                        "fact_type": "action",
                        "destination_section": "experience",
                        "destination_field": "bullet",
                    }],
                    "context_spans": [],
                }
                for item in candidates
            ],
        }

    compiled = compile_semantics(fact_graph, llm_call=exact_semantics)
    plan = plan_resume(compiled.graph, build_requirement_graph(""))
    frozen = realize_plan(plan, compiled.graph)
    resume_data = frozen_to_resume_data(frozen, compiled.graph)

    assert "年月-年月" not in seen
    assert all("年月-年月" not in claim.text for claim in frozen.claims)
    assert "年月-年月" not in json.dumps(resume_data, ensure_ascii=False)


def test_role_plus_anonymized_company_is_record_header_but_location_is_not():
    text = """工作经历
生产专员，[公司]2022年至今
负责设备维护
质量保证专员，【公司
[城市]
进行全面审计
"""
    fact_graph = build_fact_graph([from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="resume.txt",
        text=text, native=False,
    ))])
    record_by_text = {fact.text: fact.record_id for fact in fact_graph.facts}

    assert len(fact_graph.records) == 2
    assert record_by_text["质量保证专员，【公司"] == record_by_text["进行全面审计"]
    assert record_by_text["[城市]"] == record_by_text["质量保证专员，【公司"]
    assert record_by_text["负责设备维护"] != record_by_text["进行全面审计"]


def test_short_bulleted_degree_before_placeholder_moves_with_following_period():
    text = """教育经历
- 机械技术工程学士学位
[学校], [城市]
2011年毕业
平均绩点: 2.85"""
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))

    fact_graph = build_fact_graph([graph])
    record_by_text = {fact.text: fact.record_id for fact in fact_graph.facts}

    assert len(fact_graph.records) == 1
    assert record_by_text["- 机械技术工程学士学位"] == record_by_text["2011年毕业"]


def test_invalid_semantic_output_is_logged_and_source_falls_back_verbatim():
    text = "负责真实工作"
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def fabricated(_model, _system, user_prompt, **_kwargs):
        fact_id = json.loads(user_prompt)["candidates"][0]["candidate_fact_id"]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": fact_id,
                "classification": "fact",
                "record_id": None,
                "atoms": [{
                    "quote": "负责编造工作",
                    "fact_type": "action",
                    "destination_section": "experience",
                    "destination_field": "bullet",
                }],
                "context_spans": [],
            }],
        }

    result = compile_semantics(fact_graph, llm_call=fabricated)

    assert result.report.status == "fallback"
    assert result.graph.eligible_facts()[0].text == text
    assert any("atom_not_exact" in error for error in result.report.errors)


def test_semantic_compiler_retains_valid_sibling_when_one_decision_is_invalid():
    text = "张三\n负责用户访谈"
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def partially_invalid(_model, _system, user_prompt, **_kwargs):
        candidates = json.loads(user_prompt)["candidates"]
        first, second = candidates
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [
                {
                    "candidate_fact_id": first["candidate_fact_id"],
                    "classification": "fact",
                    "record_id": first["locked_record_id"],
                    "atoms": [{
                        "quote": first["candidate_text"],
                        "fact_type": "identity",
                        "destination_section": "contact",
                        "destination_field": "name",
                    }],
                    "context_spans": [],
                },
                {
                    "candidate_fact_id": second["candidate_fact_id"],
                    "classification": "instruction",
                    "record_id": second["locked_record_id"],
                    "atoms": [{
                        "quote": second["candidate_text"],
                        "fact_type": "action",
                        "destination_section": "experience",
                        "destination_field": "bullet",
                    }],
                    "context_spans": [],
                },
            ],
        }

    result = compile_semantics(fact_graph, llm_call=partially_invalid)

    assert result.report.status == "partial"
    assert result.report.response_batch_count == 1
    assert result.report.schema_valid_batch_count == 0
    assert result.report.raw_decision_count == 2
    assert result.report.valid_decision_count == 1
    assert result.report.invalid_decision_count == 1
    assert len(result.report.accepted_fact_ids) == 1
    assert len(result.report.fallback_fact_ids) == 1
    assert result.report.training_outputs[0] is not None
    assert [fact.text for fact in result.graph.eligible_facts()] == ["张三", "负责用户访谈"]


def test_semantic_compiler_recovers_valid_atoms_from_one_malformed_child():
    text = "- 主导现金流优化举措，促成30%的成本削减"
    graph = from_native_text(SourceAsset(
        source_id="query", source_type="query", filename="query.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def malformed_separator(_model, _system, user_prompt, **_kwargs):
        candidate = json.loads(user_prompt)["candidates"][0]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": candidate["candidate_fact_id"],
                "classification": "fact",
                "record_id": candidate["locked_record_id"],
                "atoms": [
                    {
                        "quote": "-",
                        "fact_type": "other",
                        "destination_section": "experience",
                        "destination_field": "separator",
                    },
                    {
                        "quote": "主导现金流优化举措",
                        "fact_type": "action",
                        "destination_section": "experience",
                        "destination_field": "bullet",
                    },
                    {
                        "quote": "促成30%的成本削减",
                        "fact_type": "result",
                        "destination_section": "experience",
                        "destination_field": "bullet",
                    },
                ],
                "context_spans": [],
            }],
        }

    result = compile_semantics(fact_graph, llm_call=malformed_separator)

    assert result.report.recovered_decision_count == 1
    assert result.report.invalid_atom_count == 1
    assert result.report.invalid_decision_count == 0
    assert result.report.fallback_fact_ids == ()
    assert [fact.text for fact in result.graph.eligible_facts()] == [
        "主导现金流优化举措",
        "促成30%的成本削减",
    ]


def test_deterministic_realizer_canonicalizes_claim_order_across_repeated_sections():
    from core.v3.contracts import CoverageLedger, FactGraph, FactUnit, NarrativeGroup, ResumePlan, SourceSpan
    from core.v3.realizer import realize_plan

    documents = {"cv": "甲乙丙"}
    facts = [
        FactUnit(
            fact_id=f"fact:{index}", source_id="cv", source_type="cv",
            text=text, spans=[SourceSpan(source_id="cv", char_start=index, char_end=index + 1)],
            destination_section=section, destination_field="item",
        )
        for index, (text, section) in enumerate((("甲", "experience"), ("乙", "skills"), ("丙", "experience")))
    ]
    graph = FactGraph(documents=documents, facts=facts)
    plan = ResumePlan(
        groups=[
            NarrativeGroup(group_id=f"group:{index}", section=section, fact_ids=[fact.fact_id])
            for index, (fact, section) in enumerate(zip(facts, ("experience", "skills", "experience")))
        ],
        ledger=CoverageLedger(
            eligible_fact_ids=[fact.fact_id for fact in facts],
            planned_fact_ids=[fact.fact_id for fact in facts],
        ),
    )

    frozen = realize_plan(plan, graph)

    assert [claim.text for claim in frozen.claims] == ["甲", "丙", "乙"]


def test_deterministic_realizer_preserves_source_punctuation_verbatim():
    from core.v3.contracts import CoverageLedger, FactGraph, FactUnit, NarrativeGroup, ResumePlan, SourceSpan
    from core.v3.realizer import realize_plan, validate_realized_claims

    documents = {"cv": "完成分析。输出报告；"}
    facts = [
        FactUnit(
            fact_id="fact:0", source_id="cv", source_type="cv", text="完成分析。",
            spans=[SourceSpan(source_id="cv", char_start=0, char_end=5)],
            destination_section="experience", destination_field="bullet",
        ),
        FactUnit(
            fact_id="fact:1", source_id="cv", source_type="cv", text="输出报告；",
            spans=[SourceSpan(source_id="cv", char_start=5, char_end=10)],
            destination_section="experience", destination_field="bullet",
        ),
    ]
    graph = FactGraph(documents=documents, facts=facts)
    plan = ResumePlan(
        groups=[NarrativeGroup(
            group_id="group:0", section="experience",
            fact_ids=[fact.fact_id for fact in facts],
        )],
        ledger=CoverageLedger(
            eligible_fact_ids=[fact.fact_id for fact in facts],
            planned_fact_ids=[fact.fact_id for fact in facts],
        ),
    )

    frozen = realize_plan(plan, graph)

    assert [claim.text for claim in frozen.claims] == ["完成分析。", "输出报告；"]
    assert validate_realized_claims(frozen, graph) == []


def test_deterministic_realizer_recomposes_atoms_from_one_source_sentence():
    from core.v3.contracts import CoverageLedger, FactGraph, FactUnit, NarrativeGroup, ResumePlan, SourceSpan
    from core.v3.realizer import realize_plan

    document = "- Delivered results using structured workflows and clear communication"
    pieces = ["Delivered results", "using structured workflows", "and clear communication"]
    facts = []
    for index, piece in enumerate(pieces):
        start = document.index(piece)
        facts.append(FactUnit(
            fact_id=f"fact:{index}", base_fact_id="transport:0",
            source_id="cv", source_type="cv", text=piece,
            spans=[SourceSpan(source_id="cv", char_start=start, char_end=start + len(piece))],
            destination_section="experience", destination_field="bullet",
        ))
    graph = FactGraph(documents={"cv": document}, facts=facts)
    plan = ResumePlan(
        groups=[NarrativeGroup(
            group_id="group:0", section="experience",
            fact_ids=[fact.fact_id for fact in facts],
        )],
        ledger=CoverageLedger(
            eligible_fact_ids=[fact.fact_id for fact in facts],
            planned_fact_ids=[fact.fact_id for fact in facts],
        ),
    )

    frozen = realize_plan(plan, graph)

    assert [claim.text for claim in frozen.claims] == [document.removeprefix("- ")]


def test_semantic_batches_run_concurrently_but_keep_training_order(monkeypatch):
    text = "负责用户访谈\n输出需求报告"
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="cv.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])
    barrier = threading.Barrier(2)
    monkeypatch.setenv("V3_SEMANTIC_BATCH_FACTS", "1")
    monkeypatch.setenv("V3_SEMANTIC_CONCURRENCY", "2")

    def response(_model, _system, user_prompt, **_kwargs):
        candidate = json.loads(user_prompt)["candidates"][0]
        barrier.wait(timeout=2)
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": candidate["candidate_fact_id"],
                "classification": "fact",
                "record_id": candidate["locked_record_id"],
                "atoms": [{
                    "quote": candidate["candidate_text"],
                    "fact_type": "action",
                    "destination_section": "experience",
                    "destination_field": "bullet",
                }],
                "context_spans": [],
            }],
        }

    result = compile_semantics(fact_graph, llm_call=response)

    assert result.report.status == "success"
    assert result.report.batch_count == 2
    assert [batch["candidates"][0]["candidate_text"] for batch in result.report.training_inputs] == [
        "负责用户访谈", "输出需求报告",
    ]
    assert [
        output["decisions"][0]["candidate_fact_id"]
        for output in result.report.training_outputs if output is not None
    ] == list(result.report.input_fact_ids)


def test_full_semantic_contract_fallback_skips_model_realizer():
    calls = {"realizer": 0}

    def invalid_semantic(*_args, **_kwargs):
        return {}

    def realizer(*_args, **_kwargs):
        calls["realizer"] += 1
        raise AssertionError("realizer must not run after complete semantic fallback")

    result = run_v3_pipeline(
        cv_text="完成分析。\n输出报告；",
        semantic_llm_call=invalid_semantic,
        realizer_llm_call=realizer,
    )

    assert result.semantic_report.status == "fallback"
    # Record-local contract: every unit is degraded, so no LLM call is made
    # and the whole resume stays on the exact deterministic path.
    assert result.realization_report.status == "deterministic"
    assert calls["realizer"] == 0
    assert result.quality_report["atomic_factuality"]["recall"] == 1.0


def test_partial_semantic_fallback_degrades_only_the_affected_unit():
    def partial_semantic(_model, _system, user_prompt, **_kwargs):
        candidates = json.loads(user_prompt)["candidates"]
        first, second = candidates
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [
                {
                    "candidate_fact_id": first["candidate_fact_id"],
                    "classification": "fact",
                    "record_id": first["locked_record_id"],
                    "atoms": [{
                        "quote": first["candidate_text"],
                        "fact_type": "action",
                        "destination_section": "experience",
                        "destination_field": "bullet",
                    }],
                    "context_spans": [],
                },
                {
                    "candidate_fact_id": second["candidate_fact_id"],
                    "classification": "fact",
                    "record_id": second["locked_record_id"],
                    "atoms": [{
                        "quote": "模型没有逐字复制原文",
                        "fact_type": "action",
                        "destination_section": "experience",
                        "destination_field": "bullet",
                    }],
                    "context_spans": [],
                },
            ],
        }

    def realizer(_model, _system, user_prompt, **_kwargs):
        request = json.loads(user_prompt)
        return {
            "schema_version": SCHEMA_VERSION,
            "request_fact_ids": request["request_fact_ids"],
            "units": [
                {
                    "unit_id": unit["unit_id"],
                    "claims": [
                        {
                            "claim_id": f"claim:{index}",
                            "section": fact["destination_section"],
                            "field": fact["destination_field"],
                            "text": fact["source_text"],
                            "fact_ids": [fact["fact_id"]],
                            "record_id": fact["record_id"],
                            "group_id": group["group_id"],
                        }
                        for group in unit["groups"]
                        for index, fact in enumerate(group["facts"])
                    ],
                }
                for unit in request["units"]
            ],
        }

    result = run_v3_pipeline(
        cv_text="负责用户访谈\n输出需求报告",
        semantic_llm_call=partial_semantic,
        realizer_llm_call=realizer,
    )

    assert result.semantic_report.status == "partial"
    assert result.semantic_report.fallback_fact_ids
    # R24 record-local contract: the unit holding the fallback fact restores
    # exact source sentences deterministically while the clean unit keeps LLM
    # realization; the whole-resume realizer is no longer shut down.
    assert result.realization_report.status == "partial"
    unit_status = {
        report["status"] for report in result.realization_report.unit_reports
    }
    assert "deterministic_degraded" in unit_status
    assert "llm" in unit_status
    assert result.quality_report["atomic_factuality"]["recall"] == 1.0
    assert result.quality_report["atomic_factuality"]["precision"] == 1.0


def test_v3_realizer_budget_admission_uses_exact_deterministic_fallback(monkeypatch):
    import server_runtime

    calls = {"realizer": 0}

    def realizer(*_args, **_kwargs):
        calls["realizer"] += 1
        raise AssertionError("realizer must not start below its declared budget")

    monkeypatch.setenv("V3_REALIZER_MIN_REMAINING_SECONDS", "240")
    monkeypatch.setattr(server_runtime, "remaining_request_seconds", lambda: 120.0)
    result = run_v3_pipeline(
        cv_text="张三，负责用户访谈",
        semantic_llm_call=_semantic_response_for_name_and_action,
        realizer_llm_call=realizer,
    )

    assert calls["realizer"] == 0
    assert result.realization_report.status == "budget_fallback"
    assert result.quality_report["atomic_factuality"]["precision"] == 1.0
    assert result.quality_report["atomic_factuality"]["recall"] == 1.0


def test_v3_adapter_never_emits_unknown_meta_fields_on_semantic_fallback():
    from core.v2_schemas import CanonicalResume

    result = run_v3_pipeline(
        query_text="姓名：[姓名]\n以下是我提供的全部个人信息。",
        semantic_llm_call=lambda *_args, **_kwargs: {},
    )

    assert "other_info" not in result.resume_data["meta"]
    rendered = json.dumps(result.resume_data, ensure_ascii=False)
    assert "[姓名]" not in rendered
    assert "以下是我提供的全部个人信息。" not in rendered
    assert result.semantic_report.fail_closed_fact_ids
    assert result.resume_data["framework"]["mode"] == "empty_profile"
    CanonicalResume.model_validate(result.resume_data)


def test_v3_adapter_keeps_record_overflow_inside_canonical_resume_contract():
    from core.v2_schemas import CanonicalResume
    from core.v3.contracts import FactGraph, FrozenResume, RealizedClaim
    from core.v3.resume_adapter import frozen_to_resume_data

    claims = [
        RealizedClaim(
            claim_id="education:school", section="education", field="school",
            text="甲大学", record_id="edu:1", group_id="education:1",
        ),
        RealizedClaim(
            claim_id="education:extra", section="education", field="item",
            text="获得校级奖学金", record_id="edu:1", group_id="education:1",
        ),
        RealizedClaim(
            claim_id="project:org:1", section="projects", field="organization",
            text="甲公司", record_id="project:1", group_id="projects:1",
        ),
        RealizedClaim(
            claim_id="project:org:2", section="projects", field="organization",
            text="乙部门", record_id="project:1", group_id="projects:1",
        ),
    ]
    frozen = FrozenResume(
        sections={
            "education": claims[:2],
            "projects": claims[2:],
        },
        claims=claims,
    )

    data = frozen_to_resume_data(frozen, FactGraph())

    assert data["education"] == [{"school": "甲大学"}]
    assert data["additional_sections"]["教育补充信息"] == ["获得校级奖学金"]
    assert data["projects"] == [{"company": "甲公司", "bullets": ["乙部门"]}]
    CanonicalResume.model_validate({
        **data,
        "projects": [{"organization": "甲公司", "bullets": ["乙部门"]}],
    })


def test_v3_adapter_keeps_unique_summary_fact_and_compacts_exact_redundancy():
    from core.v3.contracts import FactGraph, FrozenResume, RealizedClaim
    from core.v3.resume_adapter import frozen_to_resume_data

    values = [
        "Project Manager", "Junior", "1", "SaaS",
        "Project Manager with 1 years of experience in SaaS.",
    ]
    claims = [
        RealizedClaim(
            claim_id=f"summary:{index}", section="summary", field="summary",
            text=value, group_id=f"summary:{index}",
        )
        for index, value in enumerate(values)
    ]
    frozen = FrozenResume(sections={"summary": claims}, claims=claims)

    data = frozen_to_resume_data(frozen, FactGraph())

    assert data["summary"] == "Junior；Project Manager with 1 years of experience in SaaS."


def test_v3_adapter_never_emits_research_topic_on_experience_record():
    from core.resume_copilot_service import _canonical_resume_from_render_data
    from core.v3.contracts import FactGraph, FrozenResume, RealizedClaim
    from core.v3.resume_adapter import frozen_to_resume_data

    claims = [
        RealizedClaim(
            claim_id="experience:role", section="experience", field="role",
            text="临床项目总监", record_id="experience:1", group_id="experience:1",
        ),
        RealizedClaim(
            claim_id="experience:title", section="experience", field="title",
            text="神经科项目", record_id="experience:1", group_id="experience:1",
        ),
    ]
    frozen = FrozenResume(
        sections={"experience": claims},
        claims=claims,
    )

    data = frozen_to_resume_data(frozen, FactGraph())

    assert data["experience"] == [{
        "role": "临床项目总监",
        "bullets": ["神经科项目"],
    }]
    assert "topic" not in data["experience"][0]
    canonical = _canonical_resume_from_render_data(data)
    assert canonical.experience[0].role == "临床项目总监"
    assert canonical.experience[0].bullets == ["神经科项目"]


def test_v3_adapter_removes_only_public_markdown_and_list_presentation():
    from core.v3.contracts import FactGraph, FrozenResume, RealizedClaim
    from core.v3.resume_adapter import frozen_to_resume_data

    claims = [
        RealizedClaim(
            claim_id="experience:role", section="experience", field="role",
            text="- **机械工长**", record_id="experience:1", group_id="experience:1",
        ),
        RealizedClaim(
            claim_id="experience:bullet", section="experience", field="bullet",
            text="- 管理项目进度以满足截止日期和质量要求", record_id="experience:1",
            group_id="experience:1",
        ),
    ]
    frozen = FrozenResume(sections={"experience": claims}, claims=claims)

    data = frozen_to_resume_data(frozen, FactGraph())

    assert data["experience"] == [{
        "role": "机械工长",
        "bullets": ["管理项目进度以满足截止日期和质量要求"],
    }]


def test_v3_image_adapter_uses_hybrid_rapidocr_text_without_second_structure_call(
    monkeypatch,
):
    monkeypatch.setenv("LAYOUT_ORDER_ENGINE", "ppstructure_hybrid")

    graph = build_input_document_graph(
        b"transport-bytes-are-not-read-on-the-hybrid-fallback",
        filename="resume.png",
        fallback_text="工作经历\n产品经理\n负责用户访谈",
    )

    assert graph.extraction_engine == "ocr"
    assert graph.source_text == "工作经历\n产品经理\n负责用户访谈"
    assert graph.metadata["hybrid_layout"] is True
    assert graph.metadata["text_source"] == "PP-OCRv6"


def test_v3_hybrid_adapter_preserves_region_boundaries_but_removes_transport_marker(
    monkeypatch,
):
    from experimental_model_candidates import OCR_LAYOUT_REGION_SEPARATOR

    monkeypatch.setenv("LAYOUT_ORDER_ENGINE", "ppstructure_hybrid")
    graph = build_input_document_graph(
        b"transport",
        filename="resume.png",
        fallback_text=(
            "工作经历\n产品经理\n负责用户访谈\n"
            f"{OCR_LAYOUT_REGION_SEPARATOR}\n"
            "运营经理\n负责活动复盘"
        ),
    )

    assert OCR_LAYOUT_REGION_SEPARATOR not in graph.source_text
    body_nodes = [node for node in graph.nodes if node.kind != "heading"]
    assert [node.region_id for node in body_nodes] == [
        "hybrid-region:0",
        "hybrid-region:0",
        "hybrid-region:1",
        "hybrid-region:1",
    ]



def test_v3_hybrid_order_allows_cross_region_record_continuation(monkeypatch):
    from experimental_model_candidates import OCR_LAYOUT_REGION_SEPARATOR

    monkeypatch.setenv("LAYOUT_ORDER_ENGINE", "ppstructure_hybrid")
    marker = OCR_LAYOUT_REGION_SEPARATOR
    graph = build_input_document_graph(
        b"transport",
        filename="resume.png",
        fallback_text=(
            "工作经历\n甲公司 产品经理 2020年1月-2022年1月\n负责用户访谈\n"
            f"{marker}\n协调研发上线\n"
            f"{marker}\n质量经理，[公司]\n进行全面审计"
        ),
    )
    fact_graph = build_fact_graph([graph])
    record_by_text = {fact.text: fact.record_id for fact in fact_graph.facts}

    assert len(fact_graph.records) == 2
    assert record_by_text["协调研发上线"] == record_by_text["负责用户访谈"]
    assert record_by_text["质量经理，[公司]"] == record_by_text["进行全面审计"]
    assert record_by_text["协调研发上线"] != record_by_text["进行全面审计"]


def test_non_fact_decision_may_explain_exact_context_without_emitting_atoms():
    text = "请帮我优化简历"
    graph = from_native_text(SourceAsset(
        source_id="query", source_type="query", filename="query.txt", text=text, native=True,
    ))
    fact_graph = build_fact_graph([graph])

    def response(_model, _system, user_prompt, **_kwargs):
        candidates = json.loads(user_prompt)["candidates"]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [
                {
                    "candidate_fact_id": candidate["candidate_fact_id"],
                    "classification": "instruction",
                    "record_id": None,
                    "atoms": [],
                    "context_spans": [{
                        "quote": candidate["candidate_text"],
                        "reason": "instruction",
                    }],
                }
                for candidate in candidates
            ],
        }

    result = compile_semantics(fact_graph, llm_call=response)

    assert result.report.status == "success"
    assert not result.graph.eligible_facts()
    assert result.report.non_fact_ids


def test_end_to_end_fixed_schema_realizer_populates_renderer_schema():
    def realizer(_model, _system, user_prompt, **_kwargs):
        request = json.loads(user_prompt)
        units = []
        for unit in request["units"]:
            claims = []
            for index, group in enumerate(unit["groups"]):
                for fact in group["facts"]:
                    claims.append({
                        "claim_id": f"generated:{index}:{fact['fact_id']}",
                        "section": fact["destination_section"],
                        "field": fact["destination_field"],
                        "text": fact["source_text"],
                        "fact_ids": [fact["fact_id"]],
                        "record_id": fact["record_id"],
                        "group_id": group["group_id"],
                    })
            units.append({"unit_id": unit["unit_id"], "claims": claims})
        return {
            "schema_version": SCHEMA_VERSION,
            "request_fact_ids": request["request_fact_ids"],
            "units": units,
        }

    result = run_v3_pipeline(
        cv_text="张三，负责用户访谈",
        target_role="产品经理",
        semantic_llm_call=_semantic_response_for_name_and_action,
        realizer_llm_call=realizer,
    )

    assert result.semantic_report.status == "success"
    assert result.realization_report.status == "success"
    assert result.resume_data["meta"]["name"] == "张三"
    assert result.resume_data["meta"]["target_role"] == "产品经理"
    assert result.resume_data["experience"][0]["bullets"] == ["负责用户访谈"]
    assert result.quality_report["atomic_factuality"]["recall"] == 1.0
    assert result.quality_report["structural_invariants"]["violations"] == []


def test_v3_renderer_adapter_produces_editable_docx(tmp_path):
    from docx import Document
    from resume_renderer import export_resume_files

    result = run_v3_pipeline(
        cv_text="张三，负责用户访谈",
        target_role="产品经理",
        semantic_llm_call=_semantic_response_for_name_and_action,
        realizer_llm_call=lambda _model, _system, user_prompt, **_kwargs: {
            "schema_version": SCHEMA_VERSION,
            "request_fact_ids": (request := json.loads(user_prompt))["request_fact_ids"],
            "units": [
                {
                    "unit_id": unit["unit_id"],
                    "claims": [
                        {
                            "claim_id": f"claim:{index}",
                            "section": fact["destination_section"],
                            "field": fact["destination_field"],
                            "text": fact["source_text"],
                            "fact_ids": [fact["fact_id"]],
                            "record_id": fact["record_id"],
                            "group_id": group["group_id"],
                        }
                        for index, group in enumerate(unit["groups"])
                        for fact in group["facts"]
                    ],
                }
                for unit in request["units"]
            ],
        },
    )

    files = export_resume_files(
        resume_data=result.resume_data,
        output_dir=tmp_path,
        output_format="docx",
        template="classic",
    )
    document = Document(files["docx"])
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "张三" in text
    assert "负责用户访谈" in text
    assert files["pdf"] is None


def test_realizer_can_reuse_source_fact_in_optional_summary_without_ownership_error():
    def semantic(_model, _system, user_prompt, **_kwargs):
        fact_id = json.loads(user_prompt)["candidates"][0]["candidate_fact_id"]
        return {
            "schema_version": SCHEMA_VERSION,
            "decisions": [{
                "candidate_fact_id": fact_id,
                "classification": "fact",
                "record_id": None,
                "atoms": [{
                    "quote": "负责用户访谈",
                    "fact_type": "action",
                    "destination_section": "experience",
                    "destination_field": "bullet",
                }],
                "context_spans": [],
            }],
        }

    def realizer(_model, _system, user_prompt, **_kwargs):
        request = json.loads(user_prompt)
        fact_id = request["request_fact_ids"][0]
        unit = request["units"][0]
        group = unit["groups"][0]
        return {
            "schema_version": SCHEMA_VERSION,
            "request_fact_ids": request["request_fact_ids"],
            "units": [{
                "unit_id": unit["unit_id"],
                "claims": [{
                    "claim_id": "body",
                    "section": "experience",
                    "field": "bullet",
                    "text": "负责用户访谈",
                    "fact_ids": [fact_id],
                    "record_id": group["record_id"],
                    "group_id": group["group_id"],
                }],
            }],
            "summary_claims": [{
                "claim_id": "summary",
                "section": "summary",
                "field": "summary",
                "text": "具备负责用户访谈经验。",
                "fact_ids": [fact_id],
                "record_id": None,
                "group_id": "summary:profile",
            }],
        }

    result = run_v3_pipeline(
        cv_text="负责用户访谈",
        semantic_llm_call=semantic,
        realizer_llm_call=realizer,
    )

    assert result.realization_report.status == "success"
    assert result.resume_data["summary"] == "具备负责用户访谈经验。"
    assert result.resume_data["experience"][0]["bullets"] == ["负责用户访谈"]
    assert result.output.audit.ownership_errors == []


def test_jd_only_v3_is_structured_framework_without_candidate_facts():
    result = run_v3_pipeline(
        jd_text="招聘产品经理，要求用户研究",
        target_role="产品经理",
        use_llm=False,
    )

    assert result.output.plan.skeleton is True
    assert result.resume_data["framework"]["mode"] == "empty_profile"
    assert not result.output.graph.eligible_facts()
    rendered = json.dumps(result.resume_data, ensure_ascii=False)
    assert "用户研究" not in rendered
    assert "输出模块：结构化待填写简历框架" in result.reply_text
    assert "缺失信息明细" in result.reply_text
    assert "姓名" in result.reply_text


def test_docx_adapter_preserves_body_table_hierarchy_and_exact_spans():
    from docx import Document

    document = Document()
    document.add_heading("工作经历", level=1)
    document.add_paragraph("甲公司 产品经理 2020-2022")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "负责访谈"
    table.cell(0, 1).text = "输出报告"
    handle = BytesIO()
    document.save(handle)

    graph = build_input_document_graph(
        handle.getvalue(), filename="resume.docx", source_id="cv", source_type="cv",
    )

    assert graph.extraction_engine == "native_docx"
    assert any(node.metadata.get("table_row_id") for node in graph.nodes)
    assert all(
        "".join(span.quote(graph.documents()) for span in node.source_spans) == node.text
        for node in graph.nodes
    )


def test_docx_header_does_not_inherit_last_body_section():
    from docx import Document

    document = Document()
    document.sections[0].header.paragraphs[0].text = "header@example.com"
    document.add_heading("专业技能", level=1)
    document.add_paragraph("Python")
    handle = BytesIO()
    document.save(handle)
    graph = build_input_document_graph(handle.getvalue(), filename="resume.docx")

    fact_graph = build_fact_graph([graph])
    header_fact = next(fact for fact in fact_graph.facts if fact.text == "header@example.com")
    header_section = next(section for section in fact_graph.sections if section.section_id == header_fact.section_id)

    assert header_section.section_type == "other"


def test_image_adapter_consumes_raw_ppstructure_blocks(monkeypatch):
    import ppstructure_runtime

    monkeypatch.setenv("LAYOUT_ORDER_ENGINE", "ppstructure")
    monkeypatch.setattr(ppstructure_runtime, "extract_ppstructure_blocks", lambda *_args, **_kwargs: [{
        "page": 1,
        "block_id": 4,
        "block_order": 2,
        "block_label": "text",
        "block_content": "负责多列信息",
        "block_bbox": [1, 2, 30, 40],
        "region_id": "right",
        "custom": {"kept": True},
    }])

    graph = build_input_document_graph(b"image", filename="resume.png")

    assert graph.extraction_engine == "ppstructure"
    assert graph.nodes[0].region_id == "right"
    assert graph.nodes[0].metadata["custom"] == {"kept": True}


def test_native_pdf_keeps_page_bboxes_without_ocr():
    fitz = pytest.importorskip("fitz")
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Native PDF resume content with enough characters for extraction")
    content = document.tobytes()
    document.close()

    graph = build_input_document_graph(content, filename="resume.pdf")

    assert graph.extraction_engine == "native_pdf"
    assert graph.nodes and graph.nodes[0].page == 1
    assert graph.nodes[0].bbox is not None
    assert graph.metadata["page_engines"] == ["native_pdf"]


def test_scanned_pdf_page_uses_ppstructure_blocks(monkeypatch):
    fitz = pytest.importorskip("fitz")
    import ppstructure_runtime

    document = fitz.open()
    document.new_page()
    content = document.tobytes()
    document.close()
    monkeypatch.setattr(ppstructure_runtime, "extract_ppstructure_blocks", lambda *_args, **_kwargs: [{
        "block_id": 1,
        "block_order": 1,
        "block_label": "text",
        "block_content": "扫描简历内容",
        "block_bbox": [10, 20, 100, 40],
        "region_id": "scan-region",
    }])

    graph = build_input_document_graph(content, filename="scan.pdf")

    assert graph.extraction_engine == "ppstructure"
    assert graph.nodes[0].text == "扫描简历内容"
    assert graph.nodes[0].metadata["region_id"] == "scan-region"


def test_training_trace_is_off_by_default_and_records_are_versioned(monkeypatch, tmp_path):
    records = build_training_records(
        semantic_inputs=[{"schema_version": SCHEMA_VERSION}],
        semantic_outputs=[],
        semantic_status="fallback",
        semantic_errors=["invalid"],
        realizer_input=None,
        realizer_output=None,
        realizer_status="disabled",
        realizer_violations=[],
    )
    monkeypatch.setenv("V3_TRAINING_TRACE_DIR", str(tmp_path))
    monkeypatch.delenv("V3_TRAINING_TRACE_ENABLED", raising=False)

    assert records[0]["schema_version"] == SCHEMA_VERSION
    assert records[0]["schema_sha256"] == SCHEMA_FINGERPRINT
    assert maybe_write_training_trace(records) == ""
    assert list(tmp_path.iterdir()) == []


def test_training_records_preserve_failed_batch_alignment():
    records = build_training_records(
        semantic_inputs=[{"batch": 0}, {"batch": 1}],
        semantic_outputs=[None, {"decisions": []}],
        semantic_status="partial",
        semantic_errors=["batch:0:invalid"],
        realizer_input=None,
        realizer_output=None,
        realizer_status="disabled",
        realizer_violations=[],
    )

    assert [record["input"]["batch"] for record in records] == [0, 1]
    assert records[0]["output"] is None
    assert records[1]["output"] == {"decisions": []}


@pytest.mark.asyncio
async def test_service_environment_routes_to_v3_before_v2(monkeypatch):
    import resume_copilot_pipeline
    import resume_copilot_service

    context = resume_copilot_pipeline.PipelineContext(
        query_text="测试", scenario="scenario2", industry="other", user_stage="job_seeker",
    )
    response = ResumeCopilotResponse(
        files={}, reply_text="v3", scenario="scenario2", industry="other", user_stage="job_seeker",
    )
    called = {"v3": 0}

    async def ingest(**_kwargs):
        return context

    async def classify(ctx):
        return ctx

    async def v3(ctx, *, stage_render):
        called["v3"] += 1
        return response

    monkeypatch.setattr(resume_copilot_pipeline, "stage_ingest", ingest)
    monkeypatch.setattr(resume_copilot_pipeline, "stage_classify", classify)
    monkeypatch.setattr(resume_copilot_service, "_resume_copilot_v3_from_context", v3)
    monkeypatch.setenv("RESUME_PIPELINE_VERSION", "v3")

    result = await resume_copilot_service._resume_copilot_service_impl(
        query="测试", cv=None, cv_template=None, target_jd=None,
        target_jd_file=None, target_jd_url=None, jd_text=None, jd_url=None,
    )

    assert result.reply_text == "v3"
    assert called["v3"] == 1


@pytest.mark.asyncio
async def test_v3_ingest_builds_document_graph_once_and_captures_original_bytes(monkeypatch):
    from fastapi import UploadFile
    import resume_copilot_pipeline
    import v3.input_adapters

    raw = b"plain-resume-bytes"
    graph = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="resume.txt",
        text="张三\n负责访谈", native=True,
    ))
    calls = {"count": 0}

    def build(*_args, **_kwargs):
        calls["count"] += 1
        return graph

    monkeypatch.setattr(v3.input_adapters, "build_input_document_graph", build)
    monkeypatch.setenv("RESUME_PIPELINE_VERSION", "v3")
    upload = UploadFile(filename="resume.txt", file=BytesIO(raw))

    ctx = await resume_copilot_pipeline.stage_ingest(
        query="请优化", cv=upload, cv_template=None, target_jd=None,
        jd_text=None, target_jd_url=None, jd_url=None, target_jd_file=None,
    )

    assert calls["count"] == 1
    assert ctx.cv_content == raw
    assert ctx.cv_filename == "resume.txt"
    assert ctx.cv_document_graph is graph
    assert ctx.cv_text == graph.source_text


@pytest.mark.asyncio
async def test_v3_native_english_txt_bypasses_ocr_quality_gate(monkeypatch):
    from fastapi import UploadFile
    import resume_copilot_pipeline

    text = """ANONYMIZED SYNTHETIC PROFILE
Professional Direction: Project Manager
Education
High School
Skills
Process Improvement, Reporting, Asana
Experience Highlights
- Delivered results using structured workflows and clear communication
"""
    monkeypatch.setenv("RESUME_PIPELINE_VERSION", "v3")
    upload = UploadFile(filename="resume.txt", file=BytesIO(text.encode()))

    ctx = await resume_copilot_pipeline.stage_ingest(
        query="请优化", cv=upload, cv_template=None, target_jd=None,
        jd_text=None, target_jd_url=None, jd_url=None, target_jd_file=None,
    )

    assert ctx.cv_text == text
    assert ctx.cv_document_graph is not None
    assert ctx.cv_document_graph.extraction_engine == "text"
    assert ctx.ocr_quality is None
    assert ctx._low_ocr_quality is False
    assert not any(item.get("source") == "pdf_ocr_quality" for item in ctx.ocr_warnings)


def test_v3_ocr_quality_gate_uses_pdf_extraction_provenance():
    import resume_copilot_pipeline

    native_pdf = from_native_text(SourceAsset(
        source_id="cv", source_type="cv", filename="resume.pdf",
        text="Native PDF resume", native=True,
    ), "Native PDF resume", engine="native_pdf")
    native_pdf.metadata.update({"native": True, "page_engines": ["native_pdf"]})
    scanned_pdf = native_pdf.model_copy(update={
        "extraction_engine": "ppstructure",
        "metadata": {"native": False, "page_engines": ["ppstructure"]},
    })

    assert not resume_copilot_pipeline._cv_text_requires_ocr_quality_gate(
        "resume.pdf", native_pdf,
    )
    assert resume_copilot_pipeline._cv_text_requires_ocr_quality_gate(
        "resume.pdf", scanned_pdf,
    )
    assert resume_copilot_pipeline._cv_text_requires_ocr_quality_gate(
        "resume.png", native_pdf,
    )


@pytest.mark.asyncio
async def test_v3_hybrid_raster_ingest_extracts_text_before_building_graph(monkeypatch):
    from fastapi import UploadFile
    import resume_copilot_pipeline
    import v3.input_adapters

    raw = b"synthetic-image-transport"
    extracted = "工作经历\n产品经理\n负责用户访谈"
    calls = []

    def build(content, **kwargs):
        calls.append((content, kwargs["fallback_text"]))
        return from_native_text(SourceAsset(
            source_id="cv", source_type="cv", filename="resume.png",
            text=kwargs["fallback_text"], native=False,
        ), kwargs["fallback_text"], engine="ocr")

    monkeypatch.setattr(v3.input_adapters, "build_input_document_graph", build)
    monkeypatch.setattr(
        resume_copilot_pipeline, "extract_text_from_bytes",
        lambda _content, _filename: extracted,
    )
    monkeypatch.setenv("RESUME_PIPELINE_VERSION", "v3")
    monkeypatch.setenv("LAYOUT_ORDER_ENGINE", "ppstructure_hybrid")
    upload = UploadFile(filename="resume.png", file=BytesIO(raw))

    ctx = await resume_copilot_pipeline.stage_ingest(
        query="请优化", cv=upload, cv_template=None, target_jd=None,
        jd_text=None, target_jd_url=None, jd_url=None, target_jd_file=None,
    )

    assert calls == [(raw, extracted)]
    assert ctx.cv_document_graph is not None
    assert ctx.cv_text == extracted
