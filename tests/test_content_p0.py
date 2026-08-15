"""P0 regressions for factuality, source completeness and user-facing detail."""

from __future__ import annotations

import asyncio
from types import SimpleNamespace
from unittest.mock import patch

import pytest
from docx import Document
from docx.shared import Inches

from evidence_binding import (
    _flexible_literal,
    bind_resume_evidence,
    enforce_resume_evidence,
    measure_source_coverage,
)
from resume_composer import compose_from_query, compose_resume, compose_resume_with_outcome
from resume_copilot_pipeline import PipelineContext, _build_llm_reply, stage_classify
from resume_copilot_service import _collect_content_conflicts
from resume_renderer import render_docx
from resume_validator import check_required_fields
from source_adapter import build_source_bundle, candidate_blocks
from v2_pipeline import (
    _build_evidence_summary,
    _compact_canonical,
    _compact_identity_parts,
    _canonical_to_v1_format,
    _deterministic_fallback,
    _expand_optimizer_provenance,
    _ground_bullets,
    _ground_optimizer_output,
    _needs_optimizer,
    _quality_v2_presentation_cleanup,
    _recover_grounded_source_structure,
    _recover_missing_record_facts,
    _record_source_owners,
    _restore_attested_source_summary,
    _split_grounded_fact_bullet,
    run_v2_pipeline,
)
from v2_schemas import CanonicalResume, DraftResume, SourceBlock, SourceBundle


def _doc_text(path) -> str:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def test_quality_v2_cleanup_drops_duration_metadata_and_repairs_empty_metric_clause():
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "某制造企业",
            "role": "生产专员",
            "period": "2022年至今",
            "bullets": [
                "14个月",
                "领导小团队实施设备升级，实现了%的效率提升",
                "通过流程优化将处理时长缩短15%",
                "7个月内完成系统迁移",
            ],
        }],
    })

    cleaned = _quality_v2_presentation_cleanup(resume)

    assert cleaned.experience[0].bullets == [
        "领导小团队实施设备升级",
        "通过流程优化将处理时长缩短15%",
        "7个月内完成系统迁移",
    ]


def test_jd_only_request_reaches_framework_without_query_or_cv():
    ctx = PipelineContext(query_text="", cv_text="", jd_text="岗位：病理科技师", has_jd=True)
    classification = SimpleNamespace(
        target_role="病理科技师", industry="医疗", user_stage="job_seeker",
    )
    with patch("resume_copilot_pipeline.classify_resume_request", return_value=classification):
        result = asyncio.run(stage_classify(ctx))
    assert result.scenario == "scenario4"
    assert result.generation_text == ""

    with patch("v2_pipeline.compose_from_query", return_value=CanonicalResume.model_validate({
        "meta": {"target_role": "病理科技师"},
    })):
        generated = run_v2_pipeline("", "", "岗位：病理科技师")
    framework = generated.resume_dict["framework"]
    assert framework["mode"] == "empty_profile"
    assert [item["title"] for item in framework["sections"]] == [
        "基本信息", "个人总结", "教育经历", "工作/实习经历", "项目经历", "专业技能",
    ]


def test_query_embedded_jd_and_bare_role_are_not_candidate_facts():
    source = build_source_bundle(
        "",
        "请根据JD生成简历：负责用户调研与竞品分析\n数据分析师",
        "",
    )
    assert candidate_blocks(source) == []
    assert candidate_blocks(build_source_bundle("", "保留原学校", "")) == []


def test_exact_source_summary_is_restored_without_disclaimer_or_jd_text():
    source = build_source_bundle(
        "个人总结\n具备清晰的问题拆解和执行闭环能力，以真实岗位职责和结果为准。",
        "",
        "岗位要求：具备战略规划能力",
    )
    resume = CanonicalResume.model_validate({
        "meta": {"name": "张晨", "target_role": "产品经理"},
        "summary": "曾任星河科技产品经理。",
    })

    restored, added = _restore_attested_source_summary(resume, source)

    assert added == ["具备清晰的问题拆解和执行闭环能力"]
    assert "具备清晰的问题拆解和执行闭环能力" in restored.summary
    assert "以真实岗位职责和结果为准" not in restored.summary
    assert "战略规划能力" not in restored.summary


def test_attested_source_summary_is_not_duplicated_when_already_represented():
    source = build_source_bundle(
        "个人总结\n具备清晰的问题拆解和执行闭环能力。",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "summary": "具备清晰的问题拆解和执行闭环能力。",
    })

    restored, added = _restore_attested_source_summary(resume, source)

    assert restored is resume
    assert added == []


def test_structured_query_sections_remain_candidate_facts():
    source = build_source_bundle(
        "",
        "工作经历\n甲公司｜产品经理｜2022.01-2024.01\n负责用户调研并输出PRD",
        "",
    )
    facts = [block.text for block in candidate_blocks(source)]
    assert "甲公司｜产品经理｜2022.01-2024.01" in facts
    assert "负责用户调研并输出PRD" in facts
    assert "工作经历" not in facts


@pytest.mark.parametrize(
    "text",
    [
        "做过小学数学试讲和班级活动组织",
        "有小学语文试讲和课程设计经历",
        "在校期间做过校园社群活动",
    ],
)
def test_unscoped_portfolio_fact_is_not_promoted_to_employment(text):
    source = build_source_bundle("", text, "")
    factual = [block for block in candidate_blocks(source) if block.text == text]

    assert len(factual) == 1
    assert factual[0].section_hint == "projects"
    assert factual[0].record_id == "query:projects:0"


def test_role_narrative_does_not_manufacture_company_from_domain_modifier():
    assert _compact_identity_parts(
        "2016年7月至2025年5月做企业软件销售"
    ) == ("", "企业软件销售")


def test_dated_duty_sentence_is_content_not_a_role_field():
    original = "2019年7月至2025年5月负责培训机构排课、学员服务和数据台账"
    compacted = _compact_canonical(CanonicalResume.model_validate({
        "experience": [{
            "role": original,
            "period": "2019年7月至2025年5月",
            "bullets": ["负责排课、学员服务和数据台账"],
        }],
    }))

    assert compacted.experience[0].role == ""
    assert compacted.experience[0].bullets == [original]


@pytest.mark.parametrize(
    ("query", "section", "record", "expected_organization"),
    [
        (
            "2016年7月至2025年5月做企业软件销售，负责客户开发和方案演示",
            "experience",
            {"organization": "做企业软件", "role": "企业软件销售", "period": "2016年7月至2025年5月", "bullets": ["负责客户开发和方案演示"]},
            "",
        ),
        (
            "有小学语文试讲和课程设计经历",
            "projects",
            {"name": "小学语文试讲和课程设计", "organization": "有小学", "bullets": ["有小学语文试讲和课程设计经历"]},
            "",
        ),
        (
            "2020年7月至2025年6月在农商行做柜面和贷款资料审核",
            "experience",
            {"organization": "农商行", "role": "柜面和贷款资料审核", "period": "2020年7月至2025年6月", "bullets": []},
            "农商行",
        ),
    ],
)
def test_organization_field_requires_organization_grammar(
    query, section, record, expected_organization,
):
    source = build_source_bundle("", query, "")
    resume = CanonicalResume.model_validate({section: [record]})

    grounded, _, _ = enforce_resume_evidence(resume, source)

    assert getattr(grounded, section)[0].organization == expected_organization


@pytest.mark.parametrize(
    ("header", "organization", "role"),
    [
        ("XX实验小学语文老师", "XX实验小学", "语文老师"),
        ("XX教育培训机构助教", "XX教育培训机构", "助教"),
        ("华创新能源有限公司电池工艺实习生", "华创新能源有限公司", "电池工艺实习生"),
        ("辽宁省人民医院住院医师", "辽宁省人民医院", "住院医师"),
    ],
)
def test_compact_structured_header_binds_attested_organization(
    header: str,
    organization: str,
    role: str,
):
    source = build_source_bundle(
        f"工作经历\n{header}\n2022.01-2024.01\n负责日常工作并完成记录。",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": organization,
            "role": role,
            "period": "2022.01-2024.01",
            "bullets": ["负责日常工作并完成记录。"],
        }],
    })

    grounded, _, removed = enforce_resume_evidence(resume, source)

    assert grounded.experience[0].organization == organization
    assert grounded.experience[0].role == role
    assert removed == []


@pytest.mark.parametrize(
    ("token", "display"),
    [
        ("product_pm", "产品经理"),
        ("operations", "运营"),
        ("sales", "销售"),
    ],
)
def test_internal_target_taxonomy_is_not_rendered_in_summary(token, display):
    resume = CanonicalResume.model_validate({
        "meta": {"target_role": token},
        "skills": {"items": [{"name": "Excel", "category": "tool"}]},
    })

    assert f"求职方向为{display}" in _build_evidence_summary(resume)


def test_compact_query_project_survives_deterministic_fallback():
    query = (
        "姓名程洛，电话13210008017，邮箱chengluo@example.com，"
        "上海交通大学软件工程本科09-2022到06-2026，"
        "做过课程选课系统项目，负责需求分析和原型设计，技能SQL、Axure。"
    )

    resume = _deterministic_fallback("", query, "产品经理岗位")

    assert len(resume.projects) == 1
    assert resume.projects[0].name == "课程选课系统项目"
    assert resume.projects[0].bullets == ["负责需求分析和原型设计"]
    assert [item.name for item in resume.skills.items] == ["SQL", "Axure"]


def test_query_only_profile_keeps_dates_award_and_all_structured_facts():
    query = (
        "姓名：李然\n"
        "我在星河科技公司担任产品经理，2021.03-2024.06，"
        "负责用户调研、需求分析和产品方案设计，推动研发与测试按期上线。\n"
        "我参与校园二手交易平台项目，负责竞品分析、原型设计和数据复盘。\n"
        "我获得全国大学生创新创业大赛二等奖。\n"
        "技能：Axure、SQL、Excel\n"
        "教育经历：江南大学｜本科｜工业工程｜2017.09-2021.06"
    )

    resume = _compact_canonical(
        _deterministic_fallback("", query, "目标岗位：产品经理")
    )

    assert resume.meta.name == "李然"
    assert resume.meta.target_role == "产品经理"
    assert resume.education[0].school == "江南大学"
    assert resume.education[0].major == "工业工程"
    assert resume.experience[0].organization == "星河科技公司"
    assert resume.experience[0].role == "产品经理"
    assert resume.experience[0].period == "2021.03-2024.06"
    assert resume.experience[0].bullets == [
        "负责用户调研、需求分析和产品方案设计",
        "推动研发与测试按期上线",
    ]
    assert resume.projects[0].name == "校园二手交易平台项目"
    assert resume.projects[0].bullets == ["负责竞品分析、原型设计和数据复盘"]
    assert resume.awards == ["我获得全国大学生创新创业大赛二等奖"]
    assert [item.name for item in resume.skills.items] == ["Axure", "SQL", "Excel"]
    assert not any(title.startswith("待整理") for title in resume.additional_sections)


def test_duty_phrase_cannot_survive_as_role_or_summary_identity():
    resume = CanonicalResume.model_validate({
        "experience": [
            {
                "organization": "字节跳动",
                "role": "算法工程师",
                "period": "2019-2021",
                "bullets": ["负责模型开发"],
            },
            {
                "role": "单元测试",
                "period": "2017-2019",
                "bullets": [
                    "负责开发落地工作，熟练运用 TensorFlow、Caffe、Torch 等机器学习框架"
                ],
            },
        ],
    })

    compact = _compact_canonical(resume)

    assert compact.experience[1].role == ""
    assert "单元测试" in compact.experience[1].bullets
    assert "曾任字节跳动算法工程师、单元测试" not in compact.summary


def test_role_evidence_must_come_from_identity_context_across_industries():
    examples = (
        ("某科技公司", "测试工程师", "单元测试"),
        ("某实验学校", "语文教师", "课堂教学"),
        ("某人民医院", "住院医师", "患者诊疗"),
        ("某零售公司", "运营专员", "活动运营"),
    )
    for organization, valid_role, duty in examples:
        source = build_source_bundle(
            f"工作经历\n{organization}｜{valid_role}｜2022.01-2024.01\n负责{duty}并完成日常记录",
            "",
            "",
        )
        wrong = CanonicalResume.model_validate({
            "experience": [{
                "organization": organization,
                "role": duty,
                "period": "2022.01-2024.01",
                "bullets": [f"负责{duty}并完成日常记录"],
            }],
        })
        gated_wrong, wrong_bindings, _ = enforce_resume_evidence(wrong, source)
        assert gated_wrong.experience[0].role == ""
        assert gated_wrong.experience[0].bullets == [f"负责{duty}并完成日常记录"]
        assert not any(binding.path.endswith(".role") for binding in wrong_bindings)

        correct = CanonicalResume.model_validate({
            "experience": [{
                "organization": organization,
                "role": valid_role,
                "period": "2022.01-2024.01",
                "bullets": [f"负责{duty}并完成日常记录"],
            }],
        })
        gated_correct, correct_bindings, _ = enforce_resume_evidence(correct, source)
        assert gated_correct.experience[0].role == valid_role
        assert any(binding.path.endswith(".role") for binding in correct_bindings)


def test_record_ownership_graph_uses_joint_period_and_body_evidence():
    source = build_source_bundle(
        "工作经历\n"
        "甲公司｜产品经理｜2020.01-2022.01\n"
        "负责旧版需求梳理并输出需求清单\n"
        "甲公司｜产品经理｜2022.02-2024.01\n"
        "负责新版用户访谈并输出产品方案",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "experience": [
            {
                "organization": "甲公司",
                "role": "产品经理",
                "period": "2022.02-2024.01",
                "bullets": ["负责新版用户访谈并输出产品方案"],
            },
            {
                "organization": "甲公司",
                "role": "产品经理",
                "period": "2020.01-2022.01",
                "bullets": ["负责旧版需求梳理并输出需求清单"],
            },
        ],
    })

    owners = _record_source_owners(resume, source)

    assert owners[("experience", 0)].endswith(":1")
    assert owners[("experience", 1)].endswith(":0")


def test_record_ownership_graph_leaves_repeated_weak_identity_ambiguous():
    source = build_source_bundle(
        "工作经历\n"
        "甲公司｜产品经理｜2020.01-2022.01\n负责需求分析\n"
        "甲公司｜产品经理｜2022.02-2024.01\n负责需求分析",
        "",
        "",
    )
    ambiguous = CanonicalResume.model_validate({
        "experience": [{"organization": "甲公司", "role": "产品经理"}],
    })

    assert _record_source_owners(ambiguous, source) == {}


def test_record_ownership_graph_generalizes_across_typed_industry_records():
    source = build_source_bundle(
        "教育经历\n同济大学｜本科｜临床医学｜2014.09-2019.06\n"
        "同济大学｜硕士｜公共卫生｜2019.09-2022.06\n"
        "科研经历\n某医学中心｜慢病随访课题｜2021.01-2021.12\n完成随访数据复核\n"
        "某医学中心｜影像分析课题｜2022.01-2022.12\n完成影像标注复核\n"
        "校园经历\n学生会｜宣传部员｜2018.01-2018.12\n参与活动宣传\n"
        "学生会｜宣传部长｜2019.01-2019.12\n负责活动统筹\n"
        "项目经历\n门诊流程项目｜某医院｜项目成员｜2023.01-2023.06\n参与流程梳理\n"
        "病历质量项目｜某医院｜项目负责人｜2023.07-2023.12\n负责病历抽查",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "education": [
            {"school": "同济大学", "degree": "硕士", "major": "公共卫生", "period": "2019.09-2022.06"},
            {"school": "同济大学", "degree": "本科", "major": "临床医学", "period": "2014.09-2019.06"},
        ],
        "research": [
            {"institution": "某医学中心", "topic": "影像分析课题", "period": "2022.01-2022.12", "bullets": ["完成影像标注复核"]},
            {"institution": "某医学中心", "topic": "慢病随访课题", "period": "2021.01-2021.12", "bullets": ["完成随访数据复核"]},
        ],
        "activities": [
            {"organization": "学生会", "role": "宣传部长", "period": "2019.01-2019.12", "bullets": ["负责活动统筹"]},
            {"organization": "学生会", "role": "宣传部员", "period": "2018.01-2018.12", "bullets": ["参与活动宣传"]},
        ],
        "projects": [
            {"name": "病历质量项目", "organization": "某医院", "role": "项目负责人", "period": "2023.07-2023.12", "bullets": ["负责病历抽查"]},
            {"name": "门诊流程项目", "organization": "某医院", "role": "项目成员", "period": "2023.01-2023.06", "bullets": ["参与流程梳理"]},
        ],
    })

    owners = _record_source_owners(resume, source)

    expected_markers = {
        "education": ("公共卫生", "临床医学"),
        "research": ("影像分析课题", "慢病随访课题"),
        "activities": ("宣传部长", "宣传部员"),
        "projects": ("病历质量项目", "门诊流程项目"),
    }
    for section, markers in expected_markers.items():
        for index, marker in enumerate(markers):
            expected_record_id = next(
                block.record_id for block in candidate_blocks(source)
                if block.section_hint == section and marker in block.text
            )
            assert owners[(section, index)] == expected_record_id


def test_leading_body_phrase_cannot_bind_as_role_across_industries():
    examples = (
        ("某科技公司", "测试工程师", "单元测试", "负责覆盖核心模块"),
        ("某实验学校", "语文教师", "课堂教学", "负责课程设计"),
        ("某人民医院", "住院医师", "患者诊疗", "负责病历复核"),
        ("某零售公司", "运营专员", "活动运营", "负责排期复盘"),
    )
    for organization, valid_role, duty, detail in examples:
        source = build_source_bundle(
            f"工作经历\n{organization}｜{valid_role}｜2022.01-2024.01\n{duty}，{detail}",
            "",
            "",
        )
        wrong = CanonicalResume.model_validate({
            "experience": [{
                "organization": organization,
                "role": duty,
                "period": "2022.01-2024.01",
                "bullets": [f"{duty}，{detail}"],
            }],
        })

        gated_wrong, wrong_bindings, _ = enforce_resume_evidence(wrong, source)

        assert gated_wrong.experience[0].role == ""
        assert gated_wrong.experience[0].bullets == [f"{duty}，{detail}"]
        assert not any(binding.path.endswith(".role") for binding in wrong_bindings)


def test_flexible_evidence_literal_collapses_aligned_column_whitespace():
    literal = _flexible_literal(
        "RETAIL HUB INTERNATIONAL                                        2011年"
    )

    assert r"\ " not in literal
    assert "\\s*\\s*" not in literal
    assert literal.startswith(r"R\s*E\s*T\s*A\s*I\s*L")


def test_compact_role_before_duties_survives_with_identity_context():
    source = build_source_bundle(
        "工作经历\n某科技公司 产品经理，负责需求分析与版本规划",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "某科技公司",
            "role": "产品经理",
            "bullets": ["负责需求分析与版本规划"],
        }],
    })

    gated, bindings, _ = enforce_resume_evidence(resume, source)

    assert gated.experience[0].role == "产品经理"
    assert any(binding.path.endswith(".role") for binding in bindings)


def test_compound_tool_list_remains_one_coherent_bullet():
    source = "负责开发落地工作，熟练运用 TensorFlow、Caffe、Torch 等机器学习框架"

    assert _split_grounded_fact_bullet(source) == [source]


def test_audited_long_source_scaffold_does_not_repeat_llm_wording_pass():
    resume = CanonicalResume.model_validate({
        "experience": [{
            "role": "销售主管",
            "bullets": ["负责ERP销售策略并管理销售团队"],
        }],
    })

    assert _needs_optimizer(resume)
    assert not _needs_optimizer(resume, audited_source_scaffold=True)
    assert _needs_optimizer(
        resume,
        audited_source_scaffold=True,
        narrative_record_keys={("experience", 0)},
    )


def test_internal_raw_sections_are_not_public_resume_sections(tmp_path):
    canonical = CanonicalResume.model_validate({
        "meta": {"name": "候选人"},
        "additional_sections": {
            "待整理的原始信息": ["内部解析缓存，不应输出"],
            "专业会员": ["某专业协会会员"],
        },
    })
    payload = _canonical_to_v1_format(canonical)

    assert "待整理的原始信息" not in payload["additional_sections"]
    assert payload["additional_sections"] == {"专业会员": ["某专业协会会员"]}

    # Renderer also filters defensively for callers that bypass the V2 bridge.
    payload["additional_sections"]["待整理原始信息"] = ["仍然不应输出"]
    output = tmp_path / "additional-sections.docx"
    with patch("resume_renderer.inspect_docx_layout", return_value={"available": False, "issues": []}):
        render_docx(payload, output, template="minimal")
    text = _doc_text(output)
    assert "待整理原始信息" not in text
    assert "仍然不应输出" not in text
    assert "专业会员" in text
    assert "某专业协会会员" in text


def test_compact_query_project_survives_full_no_cv_pipeline_as_complete_action():
    query = (
        "姓名程洛，电话13210008017，邮箱chengluo@example.com，"
        "上海交通大学软件工程本科09-2022到06-2026，"
        "做过课程选课系统项目，负责需求分析和原型设计，技能SQL、Axure。"
    )

    with patch("v2_pipeline.compose_from_query", return_value=CanonicalResume()), patch(
        "v2_pipeline._needs_optimizer", return_value=False,
    ):
        result = run_v2_pipeline("", query, "产品经理岗位")

    assert len(result.resume.projects) == 1
    assert result.resume.projects[0].name == "课程选课系统项目"
    assert result.resume.projects[0].bullets == ["负责需求分析和原型设计"]
    assert "framework" not in result.resume_dict


def test_no_cv_long_candidate_query_uses_chunked_path_without_2000_char_cut():
    marker = "唯一证书：注册安全工程师"
    query = "工作经历\n甲公司｜工程师｜2020.01-2024.01\n" + ("负责现场巡检与问题闭环。" * 220) + "\n" + marker
    captured = []

    def fake_compose(source):
        captured.extend(block.text for block in candidate_blocks(source))
        return DraftResume(additional_sections={"证书": [marker]})

    with patch("resume_composer.llm_enabled", return_value=True), patch(
        "resume_composer.compose_resume", side_effect=fake_compose,
    ):
        result = compose_from_query(query, "目标岗位：安全工程师")
    assert marker in "\n".join(captured)
    assert result.additional_sections["证书"] == [marker]


def test_sparse_no_cv_profile_keeps_explicit_domain_and_target_role():
    query = "我是做智能硬件产品的，这是一个IoT智能硬件产品经理的岗位JD，帮我优化。"
    result = _deterministic_fallback("", query, "负责智能硬件产品规划")
    assert result.meta.target_role == "IoT智能硬件产品经理"
    assert any(item.name == "智能硬件产品" for item in result.skills.items)


def test_composer_retains_successful_result_when_another_fact_chunk_is_empty():
    chunks = [
        SourceBundle(blocks=[SourceBlock(block_id="resume_0", source_type="resume", text="甲公司｜工程师")]),
        SourceBundle(blocks=[SourceBlock(block_id="resume_1", source_type="resume", text="乙公司｜工程师")]),
    ]
    responses = [
        {"experience": [{"organization": "甲公司", "role": "工程师"}]},
        {},
    ]
    with patch("resume_composer.llm_enabled", return_value=True), patch(
        "resume_composer._split_source_bundle", return_value=chunks,
    ), patch("resume_composer.call_llm_typed", side_effect=responses):
        outcome = compose_resume_with_outcome(
            SourceBundle(blocks=chunks[0].blocks + chunks[1].blocks)
        )
    assert outcome.draft.experience[0].organization == "甲公司"
    assert outcome.completed_chunks == 1
    assert len(outcome.failed_chunks) == 1


def test_fabricated_bullet_with_short_shared_prefix_cannot_survive():
    source = "负责客户沟通、资料整理和流程跟进"
    resume = CanonicalResume.model_validate({
        "experience": [{
            "bullets": ["负责客户沟通，制定年度品牌战略并推动全国渠道增长"],
        }],
    })
    grounded = _ground_bullets(resume, source)
    rendered = "\n".join(grounded.experience[0].bullets)
    assert "年度品牌战略" not in rendered
    assert "全国渠道增长" not in rendered


def test_cross_company_splice_keeps_the_grounded_anchor_and_removes_foreign_fields():
    for source_text in (
        "工作经历\n甲公司｜产品经理｜2020.01-2022.01\n负责需求分析\n"
        "乙公司｜运营经理｜2022.02-2024.01\n负责活动运营",
        "工作经历\n甲公司\n产品经理\n负责需求分析\n乙公司\n运营经理\n负责活动运营",
    ):
        source = build_source_bundle(source_text, "", "")
        resume = CanonicalResume.model_validate({
            "experience": [{
                "organization": "甲公司",
                "role": "运营经理",
                "bullets": ["负责活动运营"],
            }],
        })
        gated, _bindings, removed = enforce_resume_evidence(resume, source)
        assert len(gated.experience) == 1
        assert gated.experience[0].organization == "甲公司"
        assert gated.experience[0].role == ""
        assert gated.experience[0].bullets == []
        assert "experience[0]" not in removed
        assert "experience[0].role" in removed
        assert "experience[0].bullets[0]" in removed


def test_long_ocr_line_coverage_detects_each_omitted_fact():
    source = build_source_bundle(
        "工作经历\n甲公司｜产品经理｜2020.01-2024.01\n"
        "负责客户沟通、用户调研、竞品分析、输出PRD、推动研发上线、分析数据复盘",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司", "role": "产品经理", "period": "2020.01-2024.01",
            "bullets": ["负责客户沟通"],
        }],
    })
    bindings = bind_resume_evidence(resume, source)
    coverage, missing = measure_source_coverage(source, bindings)
    assert coverage < 0.60
    assert len([item for item in missing if item.startswith("resume_2#u")]) >= 5


def test_lossless_fallback_preserves_all_records_and_section_types():
    education = "\n".join(
        f"第{i}学校｜专业{i}｜本科｜20{i:02d}.09-20{i + 4:02d}.06" for i in range(1, 6)
    )
    experience = "\n".join(
        f"第{i}公司｜工程师｜20{i:02d}.01-20{i:02d}.12\n负责事项{i}并完成交付{i}"
        for i in range(1, 8)
    )
    source = (
        f"教育经历\n{education}\n工作经历\n{experience}\n"
        "科研经历\n某大学实验室｜医学影像课题｜2023.01-2024.01\n完成实验记录\n"
        "校园经历\n学生会｜宣传部长｜2022.01-2022.12\n完成活动宣传\n"
        "项目经历\n病历质量项目｜2024.01-2024.06\n抽查病历并完成问题清单"
    )
    result = _deterministic_fallback(source, "", "")
    assert len(result.education) == 5
    assert len(result.experience) == 7
    assert len(result.research) == 1
    assert len(result.activities) == 1
    assert len(result.projects) == 1
    assert result.experience[-1].role == "工程师"
    assert "负责事项7" in result.experience[-1].bullets[0]


def test_fallback_structures_ocr_compact_teacher_history_without_leaking_directions():
    source = (
        "王宁简历\n"
        "华东师范大学硕士教育学 09-2017-06-2020\n"
        "09-2020- 06-2025上海市第一实验学校语文教师，"
        "负责初中语文授课、班级管理和校本教研。\n"
        "王宁|13710003003|wangning@example.com|3年经验|女\n"
        "候选人具备清晰的问题拆解和执行闭环能力，"
        "过往经历以真实岗位职责和结果为准。\n"
        "03-2019- 06-2019上海市育才中学教育实习，"
        "参与课堂观察和作业批改。"
    )

    result = _deterministic_fallback(
        source,
        "请优化为初中语文老师岗位，保留原学校。",
        "",
    )

    assert result.meta.name == "王宁"
    assert result.meta.target_role == "初中语文老师"
    assert result.education[0].school == "华东师范大学"
    assert result.education[0].major == "教育学"
    assert [item.organization for item in result.experience] == [
        "上海市第一实验学校",
        "上海市育才中学",
    ]
    bullets = "\n".join(
        bullet for item in result.experience for bullet in item.bullets
    )
    assert "班级管理和校本教研" in bullets
    assert "课堂观察和作业批改" in bullets
    assert "保留原学校" not in bullets
    assert "13710003003" not in bullets


def test_rejected_optimizer_wording_restores_original_bullet_instead_of_dropping_it():
    evidence = (
        "03-2019- 06-2019上海市育才中学教育实习，"
        "参与课堂观察和作业批改。"
    )
    original = CanonicalResume.model_validate({
        "experience": [{
            "organization": "上海市育才中学",
            "role": "教育实习",
            "period": "03-2019- 06-2019",
            "bullets": ["参与课堂观察和作业批改。"],
        }],
    })
    optimized = original.model_copy(deep=True)
    optimized.experience[0].bullets = [
        "协助开展课堂观察记录与作业批改工作，支持日常教学运行。"
    ]

    grounded = _ground_optimizer_output(original, optimized, evidence)

    assert grounded.experience[0].bullets == ["参与课堂观察和作业批改。"]


def test_trusted_optimizer_rewrite_still_passes_final_fact_grounder():
    evidence = "负责用户调研并输出分析报告"
    original = CanonicalResume.model_validate({
        "experience": [{"bullets": [evidence]}],
    })
    optimized = original.model_copy(deep=True)
    optimized.experience[0].bullets = [
        "负责用户调研并输出分析报告，推动全国业绩提升50%"
    ]

    grounded = _ground_optimizer_output(
        original,
        optimized,
        evidence,
        trusted_rewrites={"experience[0].bullets[0]": evidence},
    )

    assert grounded.experience[0].bullets == [evidence]


def test_clause_grounder_removes_only_fabricated_result_and_keeps_source_process():
    evidence = "通过问卷开展用户调研，覆盖200名用户。"
    original = CanonicalResume.model_validate({
        "experience": [{"bullets": [evidence]}],
    })
    optimized = original.model_copy(deep=True)
    optimized.experience[0].bullets = [
        "通过问卷开展用户调研，覆盖200名用户，提升转化率30%。"
    ]

    grounded = _ground_optimizer_output(original, optimized, evidence)
    output = grounded.experience[0].bullets[0]

    assert "通过问卷开展用户调研" in output
    assert "覆盖200名用户" in output
    assert "提升转化率" not in output
    assert "30%" not in output


def test_clause_grounder_accepts_one_bullet_supported_by_multiple_source_lines():
    evidence = (
        "负责梳理客户需求。\n"
        "使用半结构化访谈收集反馈。\n"
        "输出需求优先级清单。"
    )
    original = CanonicalResume.model_validate({
        "experience": [{"bullets": ["负责梳理客户需求。"]}],
    })
    optimized = original.model_copy(deep=True)
    optimized.experience[0].bullets = [
        "负责梳理客户需求，使用半结构化访谈收集反馈，输出需求优先级清单。"
    ]

    grounded = _ground_optimizer_output(original, optimized, evidence)

    assert grounded.experience[0].bullets == optimized.experience[0].bullets


def test_final_evidence_gate_repairs_only_unsupported_atomic_clause():
    source = build_source_bundle(
        "工作经历\n甲公司｜产品经理｜2022.01-2024.01\n"
        "通过问卷开展用户调研，覆盖200名用户。",
        "",
        "岗位要求：提升转化率30%",
    )
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "产品经理",
            "period": "2022.01-2024.01",
            "bullets": [
                "通过问卷开展用户调研，覆盖200名用户，提升转化率30%。",
            ],
        }],
    })

    gated, bindings, removed = enforce_resume_evidence(resume, source)

    assert gated.experience[0].bullets == [
        "通过问卷开展用户调研，覆盖200名用户。",
    ]
    assert "30%" not in gated.experience[0].bullets[0]
    assert "experience[0].bullets[0]" not in removed
    binding = next(
        item for item in bindings
        if item.path == "experience[0].bullets[0]"
    )
    assert len(binding.fact_ids) == 2


def test_exact_compound_claim_auto_binds_multiple_facts_in_one_record():
    source = build_source_bundle(
        "工作经历\n甲公司｜产品经理｜2022.01-2024.01\n"
        "负责梳理客户需求。\n通过10次用户访谈收集反馈。\n"
        "输出需求优先级清单。",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "产品经理",
            "period": "2022.01-2024.01",
            "bullets": [
                "负责梳理客户需求，通过10次用户访谈收集反馈，输出需求优先级清单。",
            ],
        }],
    })

    bindings = bind_resume_evidence(resume, source)
    binding = next(
        item for item in bindings
        if item.path == "experience[0].bullets[0]"
    )

    assert binding.mode == "rewritten"
    assert len(binding.block_ids) == 3
    assert len(binding.fact_ids) == 3
    assert {item.record_id for item in source.fact_units if item.fact_id in binding.fact_ids} == {
        "resume:experience:0",
    }


def test_atomic_repair_restores_unique_complete_sentence_from_same_record():
    source = build_source_bundle(
        "工作经历\n甲公司｜产品经理｜2022.01-2024.01\n"
        "负责客户沟通并输出需求清单。\n"
        "乙公司｜运营专员｜2020.01-2021.01\n负责渠道活动执行。",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "产品经理",
            "period": "2022.01-2024.01",
            "bullets": ["围绕客户沟通开展全国渠道增长工作。"],
        }],
    })

    gated, _bindings, removed = enforce_resume_evidence(resume, source)

    assert gated.experience[0].bullets == ["负责客户沟通并输出需求清单。"]
    assert "全国渠道增长" not in gated.experience[0].bullets[0]
    assert "experience[0].bullets[0]" not in removed


def test_grouped_rewrite_keeps_every_source_block_in_provenance_and_coverage():
    source = build_source_bundle(
        "工作经历\n"
        "甲公司｜产品经理｜2022.01-2023.01\n"
        "负责梳理客户需求。\n"
        "通过10次用户访谈收集反馈。\n"
        "输出需求优先级清单。",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "产品经理",
            "period": "2022.01-2023.01",
            "bullets": [
                "负责梳理客户需求，通过10次用户访谈收集反馈，输出需求优先级清单。",
            ],
        }],
    })
    path = "experience[0].bullets[0]"
    provenance = (
        "负责梳理客户需求。\n"
        "通过10次用户访谈收集反馈。\n"
        "输出需求优先级清单。"
    )

    bindings = bind_resume_evidence(
        resume,
        source,
        trusted_rewrites={path: provenance},
    )

    binding = next(item for item in bindings if item.path == path)
    assert binding.source_claim == provenance
    assert len(binding.block_ids) == 3
    coverage, missing = measure_source_coverage(
        source,
        bindings,
        allow_distributed=True,
    )
    assert coverage == 1.0
    assert missing == []


def test_grouped_optimizer_provenance_expands_atom_sources_without_duplicates():
    compound = (
        "负责梳理客户需求；通过10次用户访谈收集反馈；"
        "输出需求优先级清单。"
    )
    before = CanonicalResume.model_validate({
        "experience": [{
            "bullets": [
                "负责梳理客户需求",
                "通过10次用户访谈收集反馈",
                "输出需求优先级清单",
            ],
        }],
    })
    atom_provenance = {
        f"experience[0].bullets[{index}]": compound
        for index in range(3)
    }

    expanded = _expand_optimizer_provenance(
        "experience[0].bullets[0]",
        "\n".join(before.experience[0].bullets),
        before,
        atom_provenance,
    )

    assert expanded == compound


def test_fact_ledger_recovers_missing_body_only_into_its_source_record():
    source = build_source_bundle(
        "工作经历\n"
        "甲公司｜产品经理｜2022.01-2023.01\n"
        "负责需求分析并输出PRD。\n"
        "负责通过10次用户访谈整理需求优先级。\n"
        "乙公司｜运营专员｜2023.02-2024.01\n"
        "负责活动策划与执行。",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "experience": [
            {
                "organization": "甲公司",
                "role": "产品经理",
                "period": "2022.01-2023.01",
                "bullets": ["负责需求分析并输出PRD。"],
            },
            {
                "organization": "乙公司",
                "role": "运营专员",
                "period": "2023.02-2024.01",
                "bullets": ["负责活动策划与执行。"],
            },
        ],
    })
    before_bindings = bind_resume_evidence(resume, source)
    before_coverage, _ = measure_source_coverage(
        source,
        before_bindings,
        allow_distributed=True,
    )

    recovered, stats, changed_paths = _recover_missing_record_facts(
        resume,
        source,
        before_bindings,
    )
    after_bindings = bind_resume_evidence(recovered, source)
    after_coverage, _ = measure_source_coverage(
        source,
        after_bindings,
        allow_distributed=True,
    )

    assert stats.appended_bullets == 1
    assert stats.expanded_bullets == 0
    assert changed_paths == {"experience[0].bullets[1]"}
    assert "负责通过10次用户访谈整理需求优先级" in recovered.experience[0].bullets
    assert recovered.experience[1].bullets == ["负责活动策划与执行。"]
    assert after_coverage > before_coverage


def test_fact_ledger_recovers_rest_of_compact_ocr_line_after_short_prefix():
    source = build_source_bundle(
        "工作经历\n甲公司｜产品经理｜2020.01-2024.01\n"
        "负责客户沟通、用户调研、竞品分析、输出PRD、推动研发上线、分析数据复盘",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "产品经理",
            "period": "2020.01-2024.01",
            "bullets": ["负责客户沟通"],
        }],
    })
    before_bindings = bind_resume_evidence(resume, source)

    recovered, stats, _ = _recover_missing_record_facts(
        resume,
        source,
        before_bindings,
    )
    after_bindings = bind_resume_evidence(recovered, source)
    coverage, missing = measure_source_coverage(
        source,
        after_bindings,
        allow_distributed=True,
    )

    assert stats.appended_bullets == 0
    assert stats.expanded_bullets == 1
    assert len(recovered.experience[0].bullets) == 1
    assert "推动研发上线" in recovered.experience[0].bullets[-1]
    assert "分析数据复盘" in recovered.experience[0].bullets[-1]
    assert coverage == 1.0
    assert missing == []


def test_fact_ledger_uses_period_to_disambiguate_repeated_employer_and_role():
    source = build_source_bundle(
        "工作经历\n"
        "甲公司｜运营专员｜2020.01-2021.01\n"
        "负责活动策划。\n"
        "负责通过问卷复盘活动。\n"
        "甲公司｜运营专员｜2022.01-2023.01\n"
        "负责社群运营。",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "experience": [
            {
                "organization": "甲公司",
                "role": "运营专员",
                "period": "2020.01-2021.01",
                "bullets": ["负责活动策划。"],
            },
            {
                "organization": "甲公司",
                "role": "运营专员",
                "period": "2022.01-2023.01",
                "bullets": ["负责社群运营。"],
            },
        ],
    })
    bindings = bind_resume_evidence(resume, source)

    recovered, stats, _ = _recover_missing_record_facts(resume, source, bindings)

    assert stats.appended_bullets == 1
    assert "负责通过问卷复盘活动" in recovered.experience[0].bullets
    assert recovered.experience[1].bullets == ["负责社群运营。"]


def test_grounded_structure_recovers_one_complete_missing_record_only():
    source = build_source_bundle(
        "工作经历\n"
        "甲公司｜产品经理｜2020.01-2021.01\n"
        "负责需求分析并输出PRD。\n"
        "乙公司｜运营专员｜2021.02-2023.01\n"
        "负责活动策划并完成执行复盘。",
        "",
        "",
    )
    current = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "产品经理",
            "period": "2020.01-2021.01",
            "bullets": ["负责需求分析并输出PRD。"],
        }],
    })
    fallback = CanonicalResume.model_validate({
        "experience": [
            current.experience[0].model_dump(),
            {
                "organization": "乙公司",
                "role": "运营专员",
                "period": "2021.02-2023.01",
                "bullets": ["负责活动策划并完成执行复盘。"],
            },
        ],
    })

    recovered, stats = _recover_grounded_source_structure(
        current, fallback, source,
    )
    gated, _, removed = enforce_resume_evidence(recovered, source)

    assert stats.appended_records == 1
    assert [item.organization for item in gated.experience] == ["甲公司", "乙公司"]
    assert gated.experience[1].period == "2021.02-2023.01"
    assert removed == []


@pytest.mark.parametrize(
    ("organization", "role", "period", "bullet"),
    [
        ("甲科技公司", "产品经理", "2021.01-2023.01", "负责用户调研并输出需求优先级清单。"),
        ("乙中学", "语文教师", "2020.09-2024.06", "负责初中语文授课并跟踪阶段测评结果。"),
        ("丙医院", "住院医师", "2019.07-2023.07", "负责门诊初步问诊并完成患者随访记录。"),
        ("丁零售公司", "运营专员", "2022.03-2024.03", "负责会员活动策划并协调门店完成落地。"),
    ],
)
def test_complete_record_recovery_generalizes_across_industries(
    organization: str,
    role: str,
    period: str,
    bullet: str,
):
    source = build_source_bundle(
        f"工作经历\n{organization}｜{role}｜{period}\n{bullet}",
        "",
        "",
    )
    fallback = CanonicalResume.model_validate({
        "experience": [{
            "organization": organization,
            "role": role,
            "period": period,
            "bullets": [bullet],
        }],
    })

    recovered, stats = _recover_grounded_source_structure(
        CanonicalResume(), fallback, source,
    )
    gated, _, removed = enforce_resume_evidence(recovered, source)

    assert stats.appended_records == 1
    assert len(gated.experience) == 1
    assert gated.experience[0].organization == organization
    assert gated.experience[0].role == role
    assert gated.experience[0].bullets == [bullet]
    assert removed == []


def test_grounded_structure_does_not_create_incomplete_or_ambiguous_record():
    source = build_source_bundle(
        "工作经历\n甲公司｜产品经理\n负责需求分析并输出PRD。",
        "",
        "",
    )
    fallback = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "产品经理",
            "bullets": ["负责需求分析并输出PRD。"],
        }],
    })

    recovered, stats = _recover_grounded_source_structure(
        CanonicalResume(), fallback, source,
    )

    assert recovered.experience == []
    assert stats.appended_records == 0


def test_grounded_structure_recovers_education_skills_certificates_and_awards():
    source = build_source_bundle(
        "教育经历\n北京大学｜公共卫生硕士｜2018.09-2021.06\n"
        "专业技能\n流行病学调查\n统计分析\n"
        "证书与资质\n医师资格证书\n"
        "荣誉奖项\n优秀住院医师",
        "",
        "",
    )
    current = CanonicalResume.model_validate({
        "education": [{
            "school": "北京大学",
            "degree": "硕士",
            "period": "2018.09-2021.06",
        }],
        "skills": {"items": [{"name": "统计分析", "category": "other"}]},
    })
    fallback = CanonicalResume.model_validate({
        "education": [{
            "school": "北京大学",
            "degree": "硕士",
            "major": "公共卫生",
            "period": "2018.09-2021.06",
        }],
        "skills": {"items": [
            {"name": "流行病学调查", "category": "domain"},
            {"name": "统计分析", "category": "other"},
        ]},
        "certifications": ["医师资格证书"],
        "awards": ["优秀住院医师"],
    })

    recovered, stats = _recover_grounded_source_structure(
        current, fallback, source,
    )
    gated, _, removed = enforce_resume_evidence(recovered, source)

    assert gated.education[0].major == "公共卫生"
    assert {item.name for item in gated.skills.items} == {"流行病学调查", "统计分析"}
    assert gated.certifications == ["医师资格证书"]
    assert gated.awards == ["优秀住院医师"]
    assert stats.filled_fields == 1
    assert stats.appended_values == 3
    assert removed == []


def test_quality_v2_structure_recovery_fills_fields_without_importing_content(
    monkeypatch,
):
    monkeypatch.setenv("PIPELINE_PROFILE", "quality_v2")
    source = build_source_bundle(
        "工作经历\n甲公司｜产品经理｜2022.01-2024.01\n"
        "负责需求分析。\n输出需求清单。\n"
        "专业技能\nPython\n荣誉奖项\n优秀员工",
        "",
        "",
    )
    current = CanonicalResume.model_validate({
        "experience": [{
            "organization": "",
            "role": "产品经理",
            "period": "2022.01-2024.01",
            "bullets": ["负责需求分析。"],
        }],
    })
    fallback = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "产品经理",
            "period": "2022.01-2024.01",
            "bullets": ["负责需求分析。", "输出需求清单。"],
        }],
        "skills": {"items": [{"name": "Python", "category": "language"}]},
        "awards": ["优秀员工"],
    })

    recovered, stats = _recover_grounded_source_structure(
        current, fallback, source,
    )

    assert recovered.experience[0].organization == "甲公司"
    assert recovered.experience[0].bullets == ["负责需求分析。"]
    assert recovered.skills.items == []
    assert recovered.awards == []
    assert stats.filled_fields == 1
    assert stats.appended_bullets == 0
    assert stats.appended_values == 0
    assert stats.appended_records == 0


def test_fact_ledger_does_not_turn_a_short_role_header_into_a_bullet():
    source = build_source_bundle(
        "项目经历\n增长项目\n负责人\n负责活动策划与执行。",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "projects": [{
            "name": "增长项目",
            "role": "负责人",
            "bullets": ["负责活动策划与执行。"],
        }],
    })
    bindings = bind_resume_evidence(resume, source)

    recovered, _, _ = _recover_missing_record_facts(resume, source, bindings)

    assert recovered.projects[0].bullets == ["负责活动策划与执行。"]


def test_grounded_original_summary_keeps_unique_fact_and_stays_complete():
    text = (
        "个人总结\n拥有8年三甲医院临床经验，完成急重症诊疗与病例复核。\n"
        "工作经历\n甲医院｜主治医师｜2016.01-2024.01\n完成急重症诊疗与病例复核"
    )
    source = build_source_bundle(text, "", "")
    resume = CanonicalResume.model_validate({
        "summary": "拥有8年三甲医院临床经验，完成急重症诊疗与病例复核。",
        "experience": [{
            "organization": "甲医院", "role": "主治医师", "period": "2016.01-2024.01",
            "bullets": ["完成急重症诊疗与病例复核"],
        }],
    })
    gated, _bindings, _removed = enforce_resume_evidence(resume, source)
    compacted = _compact_canonical(gated)
    assert "8年三甲医院临床经验" in compacted.summary
    assert len(compacted.summary) <= 260
    assert compacted.summary.endswith("。")


def test_compaction_merges_same_period_duplicate_recovery_rows():
    resume = CanonicalResume.model_validate({
        "experience": [
            {
                "organization": "",
                "role": "助理财务经理",
                "period": "2018年3月 - 2022年5月",
                "bullets": ["制定并实施内部控制框架和会计政策"],
            },
            {
                "organization": "",
                "role": "财务经理",
                "period": "2018年3月-2022年5月",
                "bullets": [
                    "制定并实施内部控制框架和会计政策。",
                    "管理短期和长期财务规划及预算编制。",
                ],
            },
        ],
    })

    compacted = _compact_canonical(resume)

    assert len(compacted.experience) == 1
    assert compacted.experience[0].role == "助理财务经理"
    assert len(compacted.experience[0].bullets) == 2


def test_compaction_keeps_distinct_concurrent_roles_separate():
    resume = CanonicalResume.model_validate({
        "experience": [
            {
                "organization": "甲医院",
                "role": "主治医师",
                "period": "2020年至2022年",
                "bullets": ["负责门诊诊疗"],
            },
            {
                "organization": "乙大学",
                "role": "讲师",
                "period": "2020年至2022年",
                "bullets": ["承担临床课程教学"],
            },
        ],
    })

    assert len(_compact_canonical(resume).experience) == 2


def test_compaction_prefers_complete_certification_and_drops_year_fragment():
    resume = CanonicalResume.model_validate({
        "certifications": [
            "监管合规证书 - 2011",
            "- 监管合规",
            "- 2011",
        ],
    })

    assert _compact_canonical(resume).certifications == ["监管合规证书 - 2011"]


@pytest.mark.parametrize(
    "period",
    ["209年-至今", "2609年3月-至今", "2014年16月-至今", "年月-年月"],
)
def test_compaction_clears_impossible_non_education_periods(period):
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "工程师",
            "period": period,
            "bullets": ["负责系统开发"],
        }],
    })

    compacted = _compact_canonical(resume)

    assert len(compacted.experience) == 1
    assert compacted.experience[0].period == ""
    assert compacted.experience[0].organization == "甲公司"
    assert compacted.experience[0].bullets == ["负责系统开发"]


def test_compaction_keeps_valid_non_education_periods_and_education_future_date():
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "工程师",
            "period": "2020.03-2024.12",
            "bullets": ["负责系统开发"],
        }],
        "education": [{
            "school": "乙大学",
            "degree": "本科",
            "period": "预计2030年毕业",
        }],
    })

    compacted = _compact_canonical(resume)

    assert compacted.experience[0].period == "2020.03-2024.12"
    assert compacted.education[0].period == "预计2030年毕业"


def test_compaction_reclassifies_metric_fragment_role_as_content():
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "实现10%的效率提升",
            "period": "2020年至2022年",
            "bullets": [],
        }],
    })

    compacted = _compact_canonical(resume)

    assert compacted.experience[0].role == ""
    assert compacted.experience[0].bullets == ["实现10%的效率提升"]


def test_compaction_reclassifies_narrative_identity_fragments_as_content():
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "从欧洲、中东及非洲区客户经理晋升，取得公司最高留存率",
            "role": "确保产品留存无损失",
            "period": "2020年至2022年",
            "bullets": [],
        }],
    })

    compacted = _compact_canonical(resume)

    assert compacted.experience[0].organization == ""
    assert compacted.experience[0].role == ""
    assert compacted.experience[0].bullets == ["确保产品留存无损失"]


def test_compaction_reclassifies_relational_duty_sentence_as_content():
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "",
            "role": "与市场数据部门的主要联络人，确保产品留存无损失",
            "period": "2020年至2022年",
            "bullets": [],
        }],
    })

    compacted = _compact_canonical(resume)

    assert compacted.experience[0].role == ""
    assert compacted.experience[0].bullets == [
        "与市场数据部门的主要联络人，确保产品留存无损失",
    ]


def test_compaction_drops_generic_name_and_invalid_year_credential_fragment():
    resume = CanonicalResume.model_validate({
        "meta": {"name": "个人"},
        "certifications": ["实用护理师资格证书", "（2664)"],
    })

    compacted = _compact_canonical(resume)

    assert compacted.meta.name == ""
    assert compacted.certifications == ["实用护理师资格证书"]


def test_compaction_prunes_recovered_highlight_duplicates_but_keeps_unique_fact():
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "甲公司",
            "role": "产品经理",
            "bullets": ["负责需求分析", "推动版本上线"],
        }],
        "skills": {"items": [{"name": "招聘、培养和激励高绩效团队"}]},
        "additional_sections": {"经历亮点": [
            "负责需求分析；推动版本上线",
            "撸长招聘、培养和激励高绩效团队",
            "此前担任：副总裁，投资策略（收入1.1亿美元>副总裁，全球客户管理",
            "与关键决策者培养关系，顾问及导师 年月-年月",
            "指数与分析总监。负责：",
            "独立建立跨区域合作机制",
        ]},
    })

    compacted = _compact_canonical(resume)

    assert compacted.additional_sections == {
        "经历亮点": ["独立建立跨区域合作机制"],
    }


def test_work_experience_cannot_bind_to_calendar_year_suffix():
    source = build_source_bundle(
        "教育背景\n某大学｜2023年9月-至今\n软件工程硕士",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "meta": {"work_experience": "3年"},
        "education": [{
            "school": "某大学",
            "degree": "软件工程硕士",
            "period": "2023年9月-至今",
        }],
    })

    gated, _bindings, removed = enforce_resume_evidence(resume, source)

    assert gated.meta.work_experience == ""
    assert "meta.work_experience" in removed


def test_placeholder_company_token_cannot_support_residual_organization():
    source = build_source_bundle(
        "工作经历\n[公司] Technologies, [国家]\n4G网络顾问\n2018年至2019年\n负责网络优化",
        "",
        "",
    )
    resume = CanonicalResume.model_validate({
        "experience": [{
            "organization": "Technologies",
            "role": "4G网络顾问",
            "period": "2018年至2019年",
            "bullets": ["负责网络优化"],
        }],
    })

    gated, _bindings, _removed = enforce_resume_evidence(resume, source)

    assert gated.experience[0].organization == ""


def test_summary_formats_machine_role_slug_and_bare_seniority_for_people():
    resume = CanonicalResume.model_validate({
        "meta": {"target_role": "product_pm", "work_experience": "4年"},
        "experience": [{
            "organization": "甲公司",
            "role": "产品经理",
            "bullets": ["负责需求分析和版本规划"],
        }],
    })

    compacted = _compact_canonical(resume)

    assert "求职方向为产品经理" in compacted.summary
    assert "Product PM" not in compacted.summary
    assert "4年经验" in compacted.summary
    assert "product_pm" not in compacted.summary


def test_summary_keeps_role_achievement_skills_and_education_before_generic_prose():
    resume = CanonicalResume.model_validate({
        "meta": {"target_role": "数据分析师", "work_experience": "4年"},
        "summary": (
            "具备良好的沟通能力和团队协作意识。"
            "工作认真负责并保持持续学习。"
            "能够适应不同业务环境并推进任务。"
        ),
        "experience": [{
            "organization": "甲公司",
            "role": "数据分析师",
            "period": "2020.01-2024.01",
            "bullets": ["分析120万条交易数据并输出经营看板，支持季度复盘。"],
        }],
        "education": [{"school": "乙大学", "major": "统计学", "degree": "本科"}],
        "skills": {"items": [
            {"name": "SQL", "category": "language"},
            {"name": "Python", "category": "language"},
            {"name": "Tableau", "category": "tool"},
        ]},
    })

    summary = _build_evidence_summary(resume)

    assert "曾任甲公司数据分析师" in summary
    assert "120万条交易数据" in summary
    assert "SQL、Python、Tableau" in summary
    assert "乙大学" in summary and "统计学" in summary
    assert len(summary) <= 260
    assert summary.endswith("。")


def test_new_required_rules_do_not_require_seniority_and_count_campus():
    project_profile = {
        "meta": {"name": "李明", "phone": "13812345678", "email": "li@example.com"},
        "summary": "求职方向为产品经理。",
        "education": [{"school": "某大学", "degree": "本科", "major": "管理", "period": "2020-2024"}],
        "projects": [{"name": "调研项目", "period": "2023-2024", "bullets": ["完成用户调研报告"]}],
        "skills": {"others": ["用户调研"]},
    }
    fields = {item.field for item in check_required_fields(project_profile, user_stage="experienced")}
    assert "meta.work_experience" not in fields

    campus_profile = {
        **project_profile,
        "projects": [],
        "campus_experience": [{
            "company": "学生会", "role": "宣传部长", "period": "2022-2023",
            "bullets": ["完成活动宣传"],
        }],
    }
    fields = {item.field for item in check_required_fields(campus_profile, user_stage="student")}
    assert "experience/projects/campus" not in fields


def test_production_conflicts_and_llm_reply_always_include_confirmation_and_framework_modules():
    resume_data = {
        "meta": {"target_role": "产品经理"},
        "summary": "求职方向为产品经理。",
        "experience": [
            {"company": "甲公司", "role": "产品经理", "period": "2020.01-2022.12", "bullets": ["完成需求文档"]},
            {"company": "乙公司", "role": "产品经理", "period": "2022.01-2023.12", "bullets": ["完成产品上线"]},
        ],
    }
    conflicts = [item.model_dump() for item in _collect_content_conflicts(resume_data, "产品经理")]
    assert any("甲公司" in item["description"] and "乙公司" in item["description"] for item in conflicts)

    framework = {
        "framework": {
            "mode": "empty_profile",
            "sections": [
                {"title": "基本信息"}, {"title": "个人总结"}, {"title": "教育经历"},
                {"title": "工作/实习经历"}, {"title": "项目经历"}, {"title": "专业技能"},
            ],
        },
    }
    with patch("resume_copilot_pipeline.ENABLE_LLM_REPLY", True), patch(
        "resume_copilot_pipeline.llm_enabled", return_value=True,
    ), patch("resume_copilot_pipeline.call_llm_text", return_value="我先给你整理了一版。"):
        reply = _build_llm_reply(
            audit_report={}, score=0, missing_fields=[], changes=[],
            resume_data=framework, conflicts=conflicts,
            direction="根据目标JD搭建待填写结构", framework_mode=True,
        )
    assert "待填写框架" in reply
    assert "已生成模块：基本信息、个人总结、教育经历、工作/实习经历、项目经历、专业技能" in reply
    assert "需要确认的时间或内容冲突" in reply
    assert "甲公司" in reply and "乙公司" in reply


def test_minimal_and_custom_docx_keep_cross_industry_sections(tmp_path):
    payload = {
        "meta": {"name": "李明"},
        "summary": "求职方向为临床研究。",
        "education": [{"school": "某医科大学", "degree": "硕士", "major": "临床医学", "period": "2020-2023"}],
        "research": [{"company": "某实验室", "role": "影像研究", "period": "2022-2023", "bullets": ["完成影像标注"]}],
        "campus_experience": [{"company": "志愿协会", "role": "志愿者", "period": "2021", "bullets": ["完成义诊支持"]}],
        "skills": {"domains": ["临床研究"]},
        "publications": ["医学影像研究，第一作者，2023"],
    }
    minimal = tmp_path / "minimal.docx"
    with patch("resume_renderer.inspect_docx_layout", return_value={"available": False, "issues": []}):
        render_docx(payload, minimal, template="minimal")
    minimal_text = _doc_text(minimal)
    for expected in ("某医科大学", "某实验室", "志愿协会", "临床研究", "医学影像研究"):
        assert expected in minimal_text

    template = tmp_path / "custom-template.docx"
    template_doc = Document()
    template_doc.add_paragraph("{{name}}")
    template_doc.save(template)
    custom = tmp_path / "custom.docx"
    with patch("resume_renderer.inspect_docx_layout", return_value={"available": False, "issues": []}):
        render_docx(payload, custom, template=str(template))
    custom_text = _doc_text(custom)
    for expected in ("某医科大学", "某实验室", "志愿协会", "临床研究", "医学影像研究"):
        assert expected in custom_text


def test_static_docx_template_preserves_style_package_and_replaces_sample_body(tmp_path):
    template = tmp_path / "static-style-template.docx"
    template_doc = Document()
    template_doc.styles["Normal"].font.name = "Arial"
    template_doc.sections[0].top_margin = Inches(1.05)
    template_doc.sections[0].left_margin = Inches(0.92)
    template_doc.sections[0].header.paragraphs[0].text = "CUSTOM TEMPLATE HEADER"
    template_doc.sections[0].footer.paragraphs[0].text = "CUSTOM TEMPLATE FOOTER"
    template_doc.add_paragraph("示例候选人内容（应替换）")
    template_doc.save(template)

    payload = {
        "meta": {"name": "李明", "target_role": "产品经理"},
        "summary": "具有真实的产品项目经历。",
        "experience": [{
            "company": "甲公司",
            "role": "产品经理",
            "period": "2022-2024",
            "bullets": ["负责需求分析并推动版本上线"],
        }],
        "education": [],
        "projects": [],
        "skills": {},
    }
    output = tmp_path / "from-static-style.docx"
    with patch("resume_renderer.inspect_docx_layout", return_value={"available": False, "issues": []}):
        render_docx(payload, output, template=str(template))

    rendered = Document(output)
    text = _doc_text(output)
    assert "李明" in text
    assert "甲公司" in text
    assert "示例候选人内容（应替换）" not in text
    assert rendered.sections[0].header.paragraphs[0].text == "CUSTOM TEMPLATE HEADER"
    assert rendered.sections[0].footer.paragraphs[0].text == "CUSTOM TEMPLATE FOOTER"
    assert abs(rendered.sections[0].top_margin.inches - 1.05) < 0.01
    assert abs(rendered.sections[0].left_margin.inches - 0.92) < 0.01
    assert rendered.styles["Normal"].font.name == "Arial"


def test_custom_template_keeps_section_order_and_heading_styles(tmp_path):
    template = tmp_path / "sectioned-template.docx"
    template_doc = Document()
    template_doc.add_paragraph("个人总结", style="Heading 2")
    template_doc.add_paragraph("示例总结")
    template_doc.add_paragraph("工作经历", style="Heading 2")
    template_doc.add_paragraph("示例工作内容")
    template_doc.save(template)

    payload = {
        "meta": {"name": "李明", "target_role": "产品经理"},
        "summary": "负责真实产品需求分析与版本推进。",
        "experience": [{
            "company": "甲公司",
            "role": "产品经理",
            "period": "2022-2024",
            "bullets": ["完成用户调研并推动版本上线"],
        }],
        "education": [],
        "projects": [],
        "skills": {},
    }
    output = tmp_path / "sectioned-output.docx"
    with patch("resume_renderer.inspect_docx_layout", return_value={"available": False, "issues": []}):
        render_docx(payload, output, template=str(template))

    rendered = Document(output)
    texts = [paragraph.text for paragraph in rendered.paragraphs if paragraph.text.strip()]
    summary_heading = texts.index("个人总结")
    work_heading = texts.index("工作经历")
    assert summary_heading < texts.index("负责真实产品需求分析与版本推进。") < work_heading
    assert work_heading < texts.index("完成用户调研并推动版本上线")
    assert "示例总结" not in texts
    assert "示例工作内容" not in texts
    assert rendered.paragraphs[summary_heading].style.name == "Heading 2"
